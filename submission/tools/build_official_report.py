#!/usr/bin/env python3
"""Build RouteContract's report from the organizer's retained DOCX template."""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import shutil
import tempfile
import urllib.parse
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from .report_ooxml import mark_data_rows_cannot_split
    from .report_content_contract import (
        StrictJsonError,
        count_reader_facing_evidence_ids,
        decode_strict_json,
        materialize_external_evidence,
    )
except ImportError:  # Direct script execution puts this directory on sys.path.
    from report_ooxml import mark_data_rows_cannot_split
    from report_content_contract import (
        StrictJsonError,
        count_reader_facing_evidence_ids,
        decode_strict_json,
        materialize_external_evidence,
    )


EXPECTED_TEMPLATE_SHA256 = (
    "937679bac40cbfaced3457530c232c9d190a74f6b5d67c58b4bc33014a579195"
)
FONT_NAME = "Malgun Gothic"
BODY_FONT_PT = 10
REPORT_IMAGE_WIDTH_INCHES = 4.15
# The organizer's supplemental guide caps the prioritized summary at ten rows.
# Attachment 1 table entries are report body text, so they use the organizer's
# required Malgun Gothic 10 pt rather than a compact exception.
SBOM_FONT_PT = BODY_FONT_PT
SBOM_LINE_SPACING_PT = 11.2
SBOM_MAX_ROWS = 10
CONTENT_MAX_BYTES = 1024 * 1024
PLACEHOLDER_RE = re.compile(r"\[\[[^\]]+\]\]")
CORE_PROPERTY_NAMESPACES = {
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
}
PUBLIC_DOCUMENT_IDENTITY = "RouteContract project"
WORDPROCESSINGML_NAMESPACE = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)
REVISION_IDENTIFIER_ELEMENT_TAGS = frozenset(
    f"{{{WORDPROCESSINGML_NAMESPACE}}}{local_name}"
    for local_name in ("rsids", "rsidRoot", "rsid")
)
UPSTREAM_ISSUE_38456_URL = "https://github.com/apache/shardingsphere/issues/38456"
EXTERNAL_LINK_ALIASES = {
    "[결과 Issue]": "result_issue_url",
    "[활성화 기록]": "activation_record_url",
    "[모집 기록]": "recruitment_record_url",
    "[검증 프로토콜]": "protocol_issue_url",
}


def load_document_dependencies() -> None:
    """Load report-rendering dependencies only after input gates have passed."""
    global PillowImage
    global Document
    global WD_CELL_VERTICAL_ALIGNMENT
    global WD_ALIGN_PARAGRAPH
    global OxmlElement
    global qn
    global RELATIONSHIP_TYPE
    global Inches
    global Pt
    global RGBColor
    global Run
    global etree

    from PIL import Image as PillowImage
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt, RGBColor
    from docx.text.run import Run
    from lxml import etree


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--content", type=Path, required=True)
    parser.add_argument(
        "--assets-dir",
        type=Path,
        help="PNG directory; defaults to an assets directory beside --content",
    )
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument(
        "--strict-final",
        action="store_true",
        help="refuse to build while any [[submission gate]] remains",
    )
    return parser.parse_args()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_content(path: Path, *, strict: bool) -> dict[str, Any]:
    try:
        if path.stat().st_size > CONTENT_MAX_BYTES:
            raise ValueError("report content exceeds the 1048576-byte safety limit")
        raw = path.read_bytes()
        if len(raw) > CONTENT_MAX_BYTES:
            raise ValueError("report content exceeds the 1048576-byte safety limit")
    except ValueError:
        raise
    except OSError:
        raise ValueError("report content input is unavailable") from None
    try:
        data = decode_strict_json(raw, maximum_bytes=CONTENT_MAX_BYTES)
    except StrictJsonError:
        raise ValueError("report content must be valid UTF-8 strict JSON") from None
    if not isinstance(data, dict):
        raise ValueError("report content must be an object")
    required = {
        "metadata",
        "assets",
        "project_intro",
        "background",
        "environment",
        "architecture",
        "features",
        "effects",
        "other",
        "external_evidence",
        "sbom",
    }
    missing = required.difference(data)
    unexpected = set(data).difference(required)
    if missing or unexpected:
        raise ValueError(
            "content keys do not match schema; "
            f"missing_count={len(missing)}, unexpected_count={len(unexpected)}"
        )
    rows = data["sbom"]
    if not isinstance(rows, list) or not 1 <= len(rows) <= SBOM_MAX_ROWS:
        raise ValueError(
            f"the official Attachment 1 must contain 1 to {SBOM_MAX_ROWS} prioritized rows"
        )
    expected_row_keys = {"name", "version", "license", "url", "purpose"}
    for index, row in enumerate(rows, start=1):
        if not isinstance(row, dict) or set(row) != expected_row_keys:
            raise ValueError(
                f"SBOM row {index} must contain exactly: "
                + ", ".join(sorted(expected_row_keys))
            )
        if any(not isinstance(row[key], str) or not row[key].strip() for key in row):
            raise ValueError(f"SBOM row {index} values must be non-empty strings")
    return materialize_external_evidence(data, allow_placeholders=not strict)


def iter_strings(value: Any) -> Iterable[str]:
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for child in value.values():
            yield from iter_strings(child)
    elif isinstance(value, list):
        for child in value:
            yield from iter_strings(child)


def discovered_https_targets(data: dict[str, Any]) -> set[str]:
    targets: set[str] = set()
    for text in iter_strings(data):
        for raw in re.findall(r"https://\S+", text):
            target = raw.rstrip(".,;:!?)]}'\"")
            parsed = urllib.parse.urlparse(target)
            if parsed.scheme == "https" and parsed.netloc:
                targets.add(target)
    return targets


def validate_submission_gates(data: dict[str, Any], strict: bool) -> None:
    placeholder_count = sum(_unresolved_gate_count(text) for text in iter_strings(data))
    if strict and placeholder_count:
        raise ValueError(
            f"unresolved final-submission gates (count={placeholder_count})"
        )
    evidence_id_count = count_reader_facing_evidence_ids(data)
    if strict and evidence_id_count:
        raise ValueError(
            "reader-facing report content contains audit evidence IDs "
            f"(count={evidence_id_count})"
        )


def _unresolved_gate_count(text: str) -> int:
    complete = list(PLACEHOLDER_RE.finditer(text))
    fragments = PLACEHOLDER_RE.sub("", text)
    return len(complete) + fragments.count("[[") + fragments.count("]]")


def resolve_assets(data: dict[str, Any], assets_dir: Path) -> dict[str, Path]:
    expected = {"architecture", "baseline_candidate"}
    specs = data["assets"]
    if set(specs) != expected:
        raise ValueError(f"assets must contain exactly: {', '.join(sorted(expected))}")
    root = assets_dir.resolve()
    resolved: dict[str, Path] = {}
    for key in sorted(expected):
        filename = Path(specs[key]["filename"])
        if filename.is_absolute() or ".." in filename.parts:
            raise ValueError(f"asset filename must stay under --assets-dir: {filename}")
        path = (root / filename).resolve()
        try:
            path.relative_to(root)
        except ValueError as error:
            raise ValueError(f"asset escapes --assets-dir: {path}") from error
        if not path.is_file():
            raise FileNotFoundError(f"required report asset is missing: {path}")
        with PillowImage.open(path) as image:
            if image.format != "PNG" or image.size != (1200, 675):
                raise ValueError(
                    f"asset must be a QA-approved 1200x675 PNG: {path} "
                    f"(format={image.format}, size={image.size})"
                )
        resolved[key] = path
    return resolved


def table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def find_table(document: Document, predicate):
    for table in document.tables:
        if predicate(table):
            return table
    raise ValueError("required official-template table was not found")


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def strip_guide_page(document: Document) -> None:
    guide = find_table(document, lambda table: "결과보고서 작성 안내" in table_text(table))
    remove_element(guide._tbl)


def strip_ai_attachment(document: Document) -> None:
    ai_title = find_table(
        document,
        lambda table: "AI 모델 활용 및 라이선스 기술 명세서" in table_text(table),
    )
    body = document.element.body
    children = list(body.iterchildren())
    ai_index = children.index(ai_title._tbl)

    start = ai_index
    while start > 0:
        previous = children[start - 1]
        if previous.tag == qn("w:p") and not "".join(previous.itertext()).strip():
            start -= 1
            continue
        break

    for element in children[start:]:
        if element.tag != qn("w:sectPr"):
            remove_element(element)


def set_run_font(run, size_pt: float, *, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT_NAME)


def append_hyperlink(paragraph, display: str, target: str, size_pt: float) -> None:
    relationship_id = paragraph.part.relate_to(
        target, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    run = Run(run_element, paragraph)
    run.add_text(display)
    set_run_font(run, size_pt, bold=False)
    run.font.color.rgb = RGBColor(5, 99, 193)
    run.underline = True


def append_text_with_hyperlinks(
    paragraph,
    text: str,
    size_pt: float,
    *,
    targets: Iterable[str] = (),
    aliases: dict[str, str] | None = None,
) -> None:
    candidates = [(target, target) for target in targets if target in text]
    candidates.append(("#38456", UPSTREAM_ISSUE_38456_URL))
    if aliases:
        candidates.extend(
            (display, target)
            for display, target in aliases.items()
            if display in text
        )
    matches: list[tuple[int, int, str, str]] = []
    for display, target in candidates:
        start = text.find(display)
        while start >= 0:
            matches.append((start, start + len(display), display, target))
            start = text.find(display, start + len(display))
    matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    cursor = 0
    for start, end, display, target in matches:
        if start < cursor:
            continue
        if start > cursor:
            run = paragraph.add_run(text[cursor:start])
            set_run_font(run, size_pt, bold=False)
        append_hyperlink(paragraph, display, target, size_pt)
        cursor = end
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size_pt, bold=False)


def set_paragraph_rhythm(paragraph, *, alignment=None, after_pt: float = 1.5) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = 1.0
    fmt.keep_together = False
    fmt.keep_with_next = False
    fmt.widow_control = True


def clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_rhythm(paragraph, after_pt=0)


def set_plain_cell(
    cell,
    text: str,
    *,
    size_pt: float = BODY_FONT_PT,
    bold: bool = False,
    alignment=None,
    hyperlink_target: str | None = None,
) -> None:
    if alignment is None:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    set_paragraph_rhythm(paragraph, alignment=alignment, after_pt=0)
    if hyperlink_target is None:
        run = paragraph.add_run(text)
        set_run_font(run, size_pt, bold=bold)
    else:
        append_hyperlink(paragraph, text, hyperlink_target, size_pt)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, width in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def set_block_cell(
    cell,
    blocks: list[dict[str, str]],
    *,
    image_path: Path | None = None,
    image_caption: str | None = None,
    image_after_block: int | None = None,
    hyperlink_targets: Iterable[str] = (),
    hyperlink_aliases: dict[str, str] | None = None,
) -> None:
    clear_cell(cell)
    first_paragraph = True

    def next_paragraph():
        nonlocal first_paragraph
        paragraph = cell.paragraphs[0] if first_paragraph else cell.add_paragraph()
        first_paragraph = False
        return paragraph

    def append_image() -> None:
        if image_path is None or image_caption is None:
            return
        picture_paragraph = next_paragraph()
        set_paragraph_rhythm(
            picture_paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=1.0
        )
        picture_paragraph.paragraph_format.keep_with_next = True
        inline_shape = picture_paragraph.add_run().add_picture(
            str(image_path), width=Inches(REPORT_IMAGE_WIDTH_INCHES)
        )
        # python-docx does not expose picture alt text through its public API.
        # Set it on the drawing properties so the evidence figures remain
        # understandable to screen-reader users and survive deterministic rebuilds.
        inline_shape._inline.docPr.set("descr", image_caption)
        inline_shape._inline.docPr.set("title", image_path.stem)
        caption_paragraph = next_paragraph()
        set_paragraph_rhythm(
            caption_paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=2.0
        )
        caption_paragraph.paragraph_format.keep_together = True
        caption_run = caption_paragraph.add_run(image_caption)
        set_run_font(caption_run, BODY_FONT_PT, bold=False)
        caption_run.italic = True

    if image_path is not None and image_after_block is None:
        append_image()

    for index, block in enumerate(blocks):
        paragraph = next_paragraph()
        set_paragraph_rhythm(paragraph, after_pt=2.0 if index < len(blocks) - 1 else 0)
        paragraph.paragraph_format.keep_together = True
        lead = block.get("lead", "").strip()
        text = block.get("text", "").strip()
        if lead:
            lead_run = paragraph.add_run(f"{lead}: ")
            set_run_font(lead_run, BODY_FONT_PT, bold=True)
        append_text_with_hyperlinks(
            paragraph,
            text,
            BODY_FONT_PT,
            targets=hyperlink_targets,
            aliases=hyperlink_aliases,
        )
        if image_path is not None and image_after_block == index:
            append_image()
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def remove_fixed_row_heights(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        for height in list(tr_pr.findall(qn("w:trHeight"))):
            tr_pr.remove(height)


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def mark_effects_row_cannot_split(table, row_index: int = 11) -> None:
    """Keep the compact expectations/application row on one report page."""
    mark_data_rows_cannot_split([table.rows[row_index]])


def mark_environment_row_cannot_split(table, row_index: int = 7) -> None:
    """Keep the compact environment row from straddling report pages."""
    mark_data_rows_cannot_split([table.rows[row_index]])


def strip_cloned_row_identity(row_element) -> None:
    """Drop revision/session identifiers that must not be duplicated by a clone."""
    for element in row_element.iter():
        for attribute in list(element.attrib):
            local_name = attribute.rsplit("}", 1)[-1]
            if local_name.startswith("rsid") or local_name in {"paraId", "textId"}:
                del element.attrib[attribute]


def clone_row_element(row):
    cloned = deepcopy(row._tr)
    strip_cloned_row_identity(cloned)
    return cloned


def insert_row_after(table, row):
    """Insert a structural copy immediately after ``row`` and return it."""
    matches = [index for index, candidate in enumerate(table.rows) if candidate._tr is row._tr]
    if len(matches) != 1:
        raise ValueError("report row must occur exactly once before continuation insertion")
    row._tr.addnext(clone_row_element(row))
    return table.rows[matches[0] + 1]


def fill_feature_report_rows(
    table,
    blocks: list[dict[str, str]],
    *,
    image_path: Path,
    image_caption: str,
    hyperlink_targets: Iterable[str] = (),
) -> None:
    """Place installation in its own unsplittable continuation row."""
    split_indexes = [
        index for index, block in enumerate(blocks) if block.get("lead") == "설치·릴리스"
    ]
    if split_indexes != [6]:
        raise ValueError("report feature blocks must place 설치·릴리스 at index 6")

    original_row_index = 9
    continuation_row = insert_row_after(table, table.rows[original_row_index])
    set_block_cell(
        table.cell(original_row_index, 1),
        blocks[: split_indexes[0]],
        image_path=image_path,
        image_caption=image_caption,
        image_after_block=2,
        hyperlink_targets=hyperlink_targets,
    )
    set_plain_cell(continuation_row.cells[0], "")
    set_block_cell(
        continuation_row.cells[1],
        blocks[split_indexes[0] :],
        hyperlink_targets=hyperlink_targets,
    )
    mark_data_rows_cannot_split([continuation_row])


def fill_other_report_rows(
    table,
    blocks: list[dict[str, str]],
    *,
    original_row_index: int = 11,
    hyperlink_targets: Iterable[str] = (),
    hyperlink_aliases: dict[str, str] | None = None,
) -> None:
    """Put the roadmap and later evidence blocks in one unsplittable row."""
    split_indexes = [
        index
        for index, block in enumerate(blocks)
        if block.get("lead") == "품질관리·발전 로드맵"
    ]
    if split_indexes != [1]:
        raise ValueError(
            "report 기타 blocks must place 품질관리·발전 로드맵 at index 1"
        )
    if original_row_index != len(table.rows) - 1:
        raise ValueError("report 기타 source row must be the final row before continuation")

    copy_row(table, table.rows[original_row_index])
    continuation_row_index = original_row_index + 1

    set_block_cell(
        table.cell(original_row_index, 1),
        blocks[: split_indexes[0]],
        hyperlink_targets=hyperlink_targets,
        hyperlink_aliases=hyperlink_aliases,
    )
    set_plain_cell(table.cell(continuation_row_index, 0), "")
    set_block_cell(
        table.cell(continuation_row_index, 1),
        blocks[split_indexes[0] :],
        hyperlink_targets=hyperlink_targets,
        hyperlink_aliases=hyperlink_aliases,
    )
    mark_data_rows_cannot_split([table.rows[continuation_row_index]])


def fill_header_and_metadata(document: Document, data: dict[str, Any]) -> None:
    metadata = find_table(
        document,
        lambda table: len(table.rows) == 3 and table.cell(0, 0).text.strip() == "항    목",
    )
    values = data["metadata"]
    set_plain_cell(metadata.cell(1, 1), values["team_name"])
    set_plain_cell(metadata.cell(1, 3), values["team_size"], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_plain_cell(metadata.cell(2, 1), values["division"], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_plain_cell(metadata.cell(2, 3), values["task_type"], alignment=WD_ALIGN_PARAGRAPH.CENTER)


def fill_main_report(
    document: Document, data: dict[str, Any], assets: dict[str, Path]
) -> None:
    main = find_table(
        document,
        lambda table: "프로젝트 개요" in table_text(table) and "개발배경 및 목적" in table_text(table),
    )
    values = data["metadata"]
    hyperlink_targets = discovered_https_targets(data)
    evidence = data["external_evidence"]
    hyperlink_aliases = {
        display: evidence[key]
        for display, key in EXTERNAL_LINK_ALIASES.items()
        if isinstance(evidence.get(key), str)
        and evidence[key].startswith("https://")
    }
    set_plain_cell(main.cell(1, 1), values["project_name"], bold=True)
    set_plain_cell(
        main.cell(2, 1), values["repository_url"],
        hyperlink_target=values["repository_url"],
    )
    set_plain_cell(
        main.cell(3, 1), values["video_url"],
        hyperlink_target=values["video_url"],
    )
    set_block_cell(
        main.cell(4, 1), data["project_intro"], hyperlink_targets=hyperlink_targets
    )
    set_block_cell(
        main.cell(6, 1), data["background"], hyperlink_targets=hyperlink_targets
    )
    set_block_cell(
        main.cell(7, 1), data["environment"], hyperlink_targets=hyperlink_targets
    )
    set_block_cell(
        main.cell(8, 1),
        data["architecture"],
        image_path=assets["architecture"],
        image_caption=data["assets"]["architecture"]["caption"],
        hyperlink_targets=hyperlink_targets,
    )
    fill_feature_report_rows(
        main,
        data["features"],
        image_path=assets["baseline_candidate"],
        image_caption=data["assets"]["baseline_candidate"]["caption"],
        hyperlink_targets=hyperlink_targets,
    )
    set_block_cell(
        main.cell(11, 1), data["effects"], hyperlink_targets=hyperlink_targets
    )
    fill_other_report_rows(
        main,
        data["other"],
        original_row_index=12,
        hyperlink_targets=hyperlink_targets,
        hyperlink_aliases=hyperlink_aliases,
    )
    remove_fixed_row_heights(main)
    mark_environment_row_cannot_split(main)
    mark_effects_row_cannot_split(main)


def copy_row(table, row) -> None:
    table._tbl.append(clone_row_element(row))


def fill_sbom(document: Document, rows: list[dict[str, str]]) -> None:
    sbom = find_table(
        document,
        lambda table: len(table.columns) == 6 and table.cell(0, 0).text.strip() == "번호",
    )
    while len(sbom.rows) < len(rows) + 1:
        copy_row(sbom, sbom.rows[-1])
    while len(sbom.rows) > len(rows) + 1:
        remove_element(sbom.rows[-1]._tr)

    mark_repeat_header(sbom.rows[0])
    keys = ("name", "version", "license", "url", "purpose")
    for index, item in enumerate(rows, start=1):
        set_plain_cell(
            sbom.cell(index, 0),
            str(index),
            size_pt=SBOM_FONT_PT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for column, key in enumerate(keys, start=1):
            alignment = WD_ALIGN_PARAGRAPH.LEFT
            if key in {"version", "license"}:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_plain_cell(
                sbom.cell(index, column),
                item[key],
                size_pt=SBOM_FONT_PT,
                alignment=alignment,
                hyperlink_target=item[key] if key == "url" else None,
            )
        for cell in sbom.rows[index].cells:
            set_cell_margins(cell, top=20, start=40, bottom=20, end=40)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = Pt(SBOM_LINE_SPACING_PT)
    mark_data_rows_cannot_split(sbom.rows[1:])
    remove_fixed_row_heights(sbom)


def style_retained_paragraphs(document: Document) -> None:
    for paragraph in document.paragraphs:
        if "필요 시, 행을 추가" in paragraph.text:
            remove_element(paragraph._p)


def assert_geometry(document: Document) -> None:
    if len(document.sections) != 2:
        raise ValueError(f"expected 2 official sections, found {len(document.sections)}")
    portrait, landscape = document.sections
    actual = (
        round(portrait.page_width.inches, 2),
        round(portrait.page_height.inches, 2),
        round(portrait.left_margin.inches, 2),
        round(portrait.right_margin.inches, 2),
        round(portrait.top_margin.inches, 2),
        round(portrait.bottom_margin.inches, 2),
        round(landscape.page_width.inches, 2),
        round(landscape.page_height.inches, 2),
        round(landscape.left_margin.inches, 2),
        round(landscape.right_margin.inches, 2),
        round(landscape.top_margin.inches, 2),
        round(landscape.bottom_margin.inches, 2),
    )
    expected = (8.27, 11.69, 1.18, 1.18, 1.38, 1.18, 11.69, 8.27, 0.79, 0.79, 0.98, 0.98)
    if actual != expected:
        raise ValueError(f"official section geometry changed: {actual!r}")


def row_cannot_split(row) -> bool:
    tr_pr = row._tr.trPr
    return tr_pr is not None and tr_pr.find(
        f"{{{WORDPROCESSINGML_NAMESPACE}}}cantSplit"
    ) is not None


def assert_main_report_table_structure(table) -> None:
    if len(table.rows) != 14:
        raise ValueError(
            f"expected exactly 14 main report rows, found {len(table.rows)}"
        )
    expected_labels = {
        9: "프로젝트주요기능",
        10: "",
        11: "기대효과및활용분야",
        12: "기타",
        13: "",
    }
    actual_labels = {
        index: re.sub(r"\s+", "", table.cell(index, 0).text)
        for index in expected_labels
    }
    if actual_labels != expected_labels:
        raise ValueError("main report continuation row labels changed")
    if not table.cell(10, 1).text.strip().startswith("설치·릴리스:"):
        raise ValueError("main report installation continuation changed")
    missing_cannot_split = [
        index for index in (7, 10, 11, 13) if not row_cannot_split(table.rows[index])
    ]
    if missing_cannot_split:
        raise ValueError(
            "main report rows lost cannot-split protection "
            f"(count={len(missing_cannot_split)})"
        )
    fixed_height_count = sum(
        1
        for row in table.rows
        for child in row._tr.iter()
        if child.tag == f"{{{WORDPROCESSINGML_NAMESPACE}}}trHeight"
    )
    if fixed_height_count:
        raise ValueError(
            "main report contains fixed row heights "
            f"(count={fixed_height_count})"
        )


def assert_output_scope(document: Document) -> None:
    all_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [table_text(table) for table in document.tables]
    )
    forbidden = ("결과보고서 작성 안내", "AI 모델 활용 및 라이선스 기술 명세서")
    leaked = [text for text in forbidden if text in all_text]
    if leaked:
        raise ValueError(f"removed official sections leaked into output: {leaked}")
    if len(document.tables) != 5:
        raise ValueError(f"expected exactly 5 retained official tables, found {len(document.tables)}")
    if len(document.inline_shapes) != 2:
        raise ValueError(f"expected exactly 2 report images, found {len(document.inline_shapes)}")
    main = find_table(
        document,
        lambda table: "프로젝트 개요" in table_text(table)
        and "개발배경 및 목적" in table_text(table),
    )
    assert_main_report_table_structure(main)
    assert_geometry(document)


def sanitize_core_properties(source: bytes) -> bytes:
    """Remove organizer/person metadata while keeping a deterministic project identity."""
    root = etree.fromstring(source)
    creator = root.find("dc:creator", namespaces=CORE_PROPERTY_NAMESPACES)
    last_modified_by = root.find("cp:lastModifiedBy", namespaces=CORE_PROPERTY_NAMESPACES)
    revision = root.find("cp:revision", namespaces=CORE_PROPERTY_NAMESPACES)
    if creator is None or last_modified_by is None or revision is None:
        raise ValueError("official template core properties are missing required fields")
    creator.text = PUBLIC_DOCUMENT_IDENTITY
    last_modified_by.text = PUBLIC_DOCUMENT_IDENTITY
    revision.text = "1"
    for field in ("dcterms:created", "dcterms:modified"):
        element = root.find(field, namespaces=CORE_PROPERTY_NAMESPACES)
        if element is not None:
            root.remove(element)
    return etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )


def assert_sanitized_core_properties(output: Path) -> None:
    with ZipFile(output) as package:
        root = etree.fromstring(package.read("docProps/core.xml"))
    creator = root.findtext("dc:creator", namespaces=CORE_PROPERTY_NAMESPACES)
    last_modified_by = root.findtext("cp:lastModifiedBy", namespaces=CORE_PROPERTY_NAMESPACES)
    revision = root.findtext("cp:revision", namespaces=CORE_PROPERTY_NAMESPACES)
    created = root.find("dcterms:created", namespaces=CORE_PROPERTY_NAMESPACES)
    modified = root.find("dcterms:modified", namespaces=CORE_PROPERTY_NAMESPACES)
    if (creator, last_modified_by, revision) != (
        PUBLIC_DOCUMENT_IDENTITY,
        PUBLIC_DOCUMENT_IDENTITY,
        "1",
    ) or created is not None or modified is not None:
        raise ValueError("output DOCX core properties were not privacy-sanitized")


def is_story_part(name: str) -> bool:
    return (
        name == "word/document.xml"
        or re.fullmatch(r"word/header\d+\.xml", name) is not None
        or re.fullmatch(r"word/footer\d+\.xml", name) is not None
        or name in {"word/footnotes.xml", "word/endnotes.xml"}
    )


def is_revision_identifier_part(name: str) -> bool:
    """Return whether a Word part may carry revision-session identifiers."""
    return is_story_part(name) or name in {"word/settings.xml", "word/styles.xml"}


def sanitize_story_part(source: bytes) -> bytes:
    """Remove Word revision-session attributes and elements from a privacy part."""
    root = etree.fromstring(source)
    for element in root.iter():
        for attribute in list(element.attrib):
            if attribute.startswith(f"{{{WORDPROCESSINGML_NAMESPACE}}}rsid"):
                del element.attrib[attribute]
    for element in list(root.iter()):
        if element.tag not in REVISION_IDENTIFIER_ELEMENT_TAGS:
            continue
        parent = element.getparent()
        if parent is None:
            raise ValueError(
                "revision identifier cannot be the root of a Word privacy part"
            )
        parent.remove(element)
    return etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )


def assert_sanitized_story_properties(output: Path) -> None:
    with ZipFile(output) as package:
        leaked_count = 0
        for name in package.namelist():
            if not is_revision_identifier_part(name):
                continue
            root = etree.fromstring(package.read(name))
            leaked_count += sum(
                1
                for element in root.iter()
                for attribute in element.attrib
                if attribute.startswith(f"{{{WORDPROCESSINGML_NAMESPACE}}}rsid")
            )
            leaked_count += sum(
                1
                for element in root.iter()
                if element.tag in REVISION_IDENTIFIER_ELEMENT_TAGS
            )
    if leaked_count:
        raise ValueError(
            "output DOCX contains Word revision session identifiers "
            f"(count={leaked_count})"
        )


def restore_preserve_only_package_parts(template: Path, output: Path) -> None:
    """Preserve organizer parts except document/image relationships and media."""
    with tempfile.NamedTemporaryFile(
        prefix=f".{output.stem}-", suffix=".docx", dir=output.parent, delete=False
    ) as stream:
        temporary = Path(stream.name)
    try:
        with ZipFile(template) as reference, ZipFile(output) as generated, ZipFile(
            temporary, "w", compression=ZIP_DEFLATED
        ) as rebuilt:
            reference_names = set(reference.namelist())
            generated_names = set(generated.namelist())
            removed = reference_names - generated_names
            added = generated_names - reference_names
            unexpected_added = {name for name in added if not name.startswith("word/media/")}
            if removed or unexpected_added:
                raise ValueError(
                    "python-docx changed the official package part set: "
                    f"unexpected_added_count={len(unexpected_added)}, "
                    f"removed_count={len(removed)}"
                )
            editable = {
                "[Content_Types].xml",
                "word/document.xml",
                "word/_rels/document.xml.rels",
                "docProps/core.xml",
            }
            for item in generated.infolist():
                if item.filename in editable or item.filename.startswith("word/media/"):
                    # python-docx stamps rewritten package parts with the build
                    # time. Normalize those ZIP headers so identical inputs
                    # produce a byte-identical report across repeated builds.
                    item.date_time = (1980, 1, 1, 0, 0, 0)
                    content = generated.read(item.filename)
                    if item.filename == "docProps/core.xml":
                        content = sanitize_core_properties(content)
                    output_item = item
                else:
                    reference_item = reference.getinfo(item.filename)
                    content = reference.read(item.filename)
                    output_item = reference_item
                if is_revision_identifier_part(item.filename):
                    content = sanitize_story_part(content)
                rebuilt.writestr(output_item, content)
        os.replace(temporary, output)
    finally:
        if temporary.exists():
            temporary.unlink()


def main() -> None:
    args = parse_args()
    template = args.template.resolve()
    content_path = args.content.resolve()
    assets_dir = (
        args.assets_dir.resolve() if args.assets_dir else (content_path.parent / "assets").resolve()
    )
    output = args.output.resolve()

    data = load_content(content_path, strict=args.strict_final)
    validate_submission_gates(data, args.strict_final)

    actual_hash = sha256(template)
    if actual_hash != EXPECTED_TEMPLATE_SHA256:
        raise ValueError(
            "official template SHA-256 mismatch: "
            f"expected {EXPECTED_TEMPLATE_SHA256}, got {actual_hash}"
        )

    load_document_dependencies()
    assets = resolve_assets(data, assets_dir)
    output.parent.mkdir(parents=True, exist_ok=True)
    if output == template:
        raise ValueError("output must not overwrite the retained official template")
    shutil.copy2(template, output)

    document = Document(output)
    strip_guide_page(document)
    strip_ai_attachment(document)
    fill_header_and_metadata(document, data)
    fill_main_report(document, data, assets)
    fill_sbom(document, data["sbom"])
    style_retained_paragraphs(document)
    assert_output_scope(document)
    document.save(output)
    restore_preserve_only_package_parts(template, output)
    assert_sanitized_core_properties(output)
    assert_sanitized_story_properties(output)

    verified = Document(output)
    assert_output_scope(verified)
    print(f"built={output}")
    print(f"sha256={sha256(output)}")


if __name__ == "__main__":
    main()
