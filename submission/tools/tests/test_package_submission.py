from __future__ import annotations

import base64
import hashlib
import importlib.util
import json
import os
import re
import subprocess
import sys
import tempfile
import unicodedata
import unittest
import urllib.parse
import xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import call, patch
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from submission.tools import build_official_report


SCRIPT = Path(__file__).resolve().parents[1] / "package_submission.py"
REPORT_BUILDER_SCRIPT = SCRIPT.parent / "build_official_report.py"
REPOSITORY_ROOT = Path(__file__).resolve().parents[3]
SPEC = importlib.util.spec_from_file_location("package_submission", SCRIPT)
assert SPEC is not None and SPEC.loader is not None
package_submission = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(package_submission)
TEST_CURRENT_UTC = datetime(2026, 8, 14, 0, 0, 0, tzinfo=timezone.utc)


def digest(character: str = "a") -> str:
    return character * 64


def valid_license_reviews() -> list[dict[str, object]]:
    return [
        {
            "action": "resolve or renew the OCI license review before expiry",
            "componentName": "mysql",
            "componentVersion": "8.4.11",
            "expires": "2099-01-01",
            "owner": "test maintainers",
            "purl": "pkg:oci/mysql@sha256%3A" + "b" * 64,
            "rationaleCode": "MYSQL_OCI_PACKAGE_LICENSE_CONCLUSION_INCOMPLETE",
            "reviewedAt": "2026-08-13",
            "scope": "test-container",
            "status": "manual-review-required",
        },
        {
            "action": "resolve the redistribution NOTICE review before expiry",
            "componentName": "jts-io-common",
            "componentVersion": "1.19.0",
            "expires": "2099-01-01",
            "owner": "test maintainers",
            "purl": "pkg:maven/org.locationtech.jts.io/jts-io-common@1.19.0",
            "rationaleCode": "JTS_IO_COMMON_REDISTRIBUTION_NOTICE_TREATMENT_UNCONFIRMED",
            "reviewedAt": "2026-08-13",
            "scope": "test-runtime",
            "status": "manual-review-required",
        },
    ]


def valid_vulnerability_exceptions() -> list[dict[str, object]]:
    coordinates = (
        ("OSV-003", "GHSA-c2rv-hwqm-wjpg", "pkg:maven/org.apache.calcite/calcite-core@1.40.0", "1.42.0", "MODERATE"),
    )
    return [
        {
            "advisory": advisory,
            "exceptionId": exception_id,
            "expires": "2099-01-01",
            "fixedVersion": fixed_version,
            "owner": "test maintainers",
            "purl": purl,
            "rationaleCode": "SHARDINGSPHERE_5_5_3_TEST_GRAPH",
            "reviewedAt": "2026-08-12",
            "scope": "aggregate-test-only",
            "severity": severity,
        }
        for exception_id, advisory, purl, fixed_version, severity in coordinates
    ]


def findings_for(exceptions: list[dict[str, object]]) -> list[dict[str, object]]:
    return [
        {
            "action": "time-bounded reviewed exception; re-evaluate by expiry",
            "advisory": item["advisory"],
            "exceptionExpires": item["expires"],
            "exceptionId": item["exceptionId"],
            "fixedVersion": item["fixedVersion"],
            "owner": item["owner"],
            "purl": item["purl"],
            "rationaleCode": item["rationaleCode"],
            "reachabilityEvidence": {
                "exampleProfile": True,
                "publishedProfile": False,
                "publishedRuntime": False,
            },
            "reviewedAt": item["reviewedAt"],
            "scope": item["scope"],
            "severity": item["severity"],
        }
        for item in exceptions
    ]


def valid_test_summary(revision: str) -> str:
    suites = {
        "io.github.ym0506.routecontract.RouteContractTest": 18,
        "io.github.ym0506.routecontract.example.DataSourceProxyComparisonMySqlTest": 1,
        "io.github.ym0506.routecontract.example.FailureBoundaryMySqlTest": 1,
        "io.github.ym0506.routecontract.example.ObservedExecutionRegressionCorpusMySqlTest": 7,
        "io.github.ym0506.routecontract.example.OperationCorrelationMySqlTest": 5,
        "io.github.ym0506.routecontract.internal.ShardingSphere553PreflightTest": 3,
        "io.github.ym0506.routecontract.manifest.ObservedExecutionManifestTest": 17,
    }
    lines = [
        "format=routecontract-test-summary-v1",
        f"revision={revision}",
        "suite_count=7",
        "test_count=52",
        "failure_count=0",
        "error_count=0",
        "skipped_count=0",
    ]
    lines.extend(
        f"suite={suite}|tests={suites[suite]}|failures=0|errors=0|skipped=0"
        for suite in sorted(suites)
    )
    return "\n".join(lines) + "\n"


def valid_manifest() -> dict:
    return {
        "schema_version": 5,
        "official_notice_url": "https://osscontest.kr/notice/39",
        "submission_identity": {
            "receipt_number": "123",
            "team_name": "홍길동전",
            "registered_project_name": "RouteContract",
            "team_size": "1명",
            "division": "학생",
            "task_type": "자유과제",
        },
        "project": {
            "slug": "routecontract",
            "repository_url": "https://github.com/example-owner/routecontract",
            "commit": "1" * 40,
            "tag": "v0.1.0",
            "ci_run_url": (
                "https://github.com/example-owner/routecontract/actions/runs/123456"
            ),
            "release_url": (
                "https://github.com/example-owner/routecontract/releases/tag/v0.1.0"
            ),
        },
        "report": {
            "docx_sha256": digest("2"),
            "pdf_sha256": digest("3"),
            "runtime_ai_attachment": "not_applicable",
        },
        "video": {
            "youtube_url": "https://www.youtube.com/watch?v=abcdefghijk",
            "title": "RouteContract demo",
            "duration_seconds": 173.0,
            "local_file_sha256": digest("4"),
            "external_evidence_branch": "rc_only",
            "caption_contract": {
                "schema_version": 1,
                "source_path": "submission/video-caption-cues.json",
                "source_sha256": (
                    "19560225c18ca8156a13760e8412e464"
                    "62383df5e80a92bc2dd7d4615a1f0158"
                ),
            },
        },
        "release_evidence": {
            "workflow_artifact_id": 987654,
            "workflow_artifact_sha256": digest("a"),
            "source_archive_filename": "routecontract-0.1.0-source.zip",
            "source_archive_sha256": digest("5"),
            "aggregate_sbom_json_sha256": digest("6"),
            "aggregate_sbom_xml_sha256": digest("7"),
            "signature_filenames": [],
        },
        "participant_attestations": {
            "registration_matches_report": True,
            "single_entry_per_participant_confirmed": True,
            "duplicate_benefit_status_reviewed": True,
            "ai_assistance_scope_confirmed": True,
            "core_behavior_boundaries_artifacts_and_dependency_roles_reviewed_and_explainable": True,
            "report_free_text_contains_no_external_evidence_claims": True,
            "report_free_text_privacy_reviewed": True,
            "public_external_evidence_history_and_maintainer_edits_reviewed": True,
            "source_and_dependency_licenses_reviewed": True,
            "final_pdf_visual_qa_completed": True,
            "final_local_video_actual_screen_caption_watchthrough_completed": True,
            "final_public_video_frame_audio_caption_equivalence_review_completed": True,
            "five_year_public_repository_visibility_obligation_if_selected_accepted": True,
            "owner_voice_ai_assistance_disclosed_and_participant_reviewed": True,
            "maintenance_order_and_period_confirmed": True,
            "origin_and_prior_work_statement_confirmed": True,
        },
        "duplicate_benefit_confirmation": {
            "status": "not_applicable",
            "sha256": None,
        },
    }


def valid_report_content(branch: str = "rc_only") -> dict:
    content_path = REPOSITORY_ROOT / "submission" / "report-content.ko.json"
    content = json.loads(content_path.read_text(encoding="utf-8"))
    manifest = valid_manifest()
    content["metadata"].update(
        {
            "team_name": manifest["submission_identity"]["team_name"],
            "team_size": manifest["submission_identity"]["team_size"],
            "division": manifest["submission_identity"]["division"],
            "task_type": manifest["submission_identity"]["task_type"],
            "project_name": manifest["submission_identity"]["registered_project_name"],
            "repository_url": manifest["project"]["repository_url"],
            "video_url": manifest["video"]["youtube_url"],
        }
    )

    def resolve_placeholders(value: object) -> object:
        if isinstance(value, str):
            if value == package_submission.REPORT_CONTENT_CONTRACT.EXTERNAL_EVIDENCE_SUMMARY_MARKER:
                return value
            return package_submission.PLACEHOLDER_RE.sub("resolved evidence", value)
        if isinstance(value, list):
            return [resolve_placeholders(child) for child in value]
        if isinstance(value, dict):
            return {key: resolve_placeholders(child) for key, child in value.items()}
        return value

    content = resolve_placeholders(content)
    assert isinstance(content, dict)
    external = {
        "branch": branch,
        "final_stable_tag": manifest["project"]["tag"],
        "tested_tag": manifest["project"]["tag"],
        "qualified_result_count": 1,
        "result_issue_url": f"{manifest['project']['repository_url']}/issues/42",
        "activation_record_url": None,
        "recruitment_record_url": None,
        "protocol_issue_url": f"{manifest['project']['repository_url']}/issues/9",
        "cutoff_utc": "2026-08-01T06:00:00Z",
    }
    if branch in {"rc_only", "zero"}:
        external["tested_tag"] = "v0.1.0-rc1"
        external["activation_record_url"] = (
            f"{manifest['project']['repository_url']}/blob/{'b' * 40}/docs/evidence/"
            "independent-rc-activation-v0.1.0-rc1.json"
        )
        external["recruitment_record_url"] = (
            f"{manifest['project']['repository_url']}/issues/9#issuecomment-777"
        )
    if branch == "zero":
        external["qualified_result_count"] = 0
        external["result_issue_url"] = None
    content["external_evidence"] = external
    return content


def valid_ffprobe_video() -> dict:
    return {
        "format": {
            "duration": "173.000000",
            "tags": {
                "encoder": "Lavf test fixture",
            },
        },
        "streams": [
            {
                "index": 0,
                "codec_type": "video",
                "width": 1920,
                "height": 1080,
                "disposition": {
                    "default": 1,
                    "attached_pic": 0,
                    "still_image": 0,
                },
                "tags": {
                    "language": "und",
                    "handler_name": "VideoHandler",
                },
            },
        ],
    }


def valid_youtube_probe() -> dict:
    return {
        "id": "abcdefghijk",
        "duration": 173,
        "availability": "public",
        "live_status": "not_live",
        "age_limit": 0,
        "formats": [
            {
                "format_id": "140",
                "height": None,
                "vcodec": "none",
                "acodec": "mp4a.40.2",
                "url": "https://example.invalid/audio",
            },
            {
                "format_id": "137",
                "height": 1080,
                "vcodec": "avc1.640028",
                "acodec": "none",
                "url": "https://example.invalid/video",
                "has_drm": False,
            },
        ],
    }


class SubmissionClaimTextTest(unittest.TestCase):
    def test_develop_kim_provenance_is_consistent_across_public_docs(self) -> None:
        origin = (REPOSITORY_ROOT / "ORIGIN_AND_PRIOR_WORK.md").read_text(
            encoding="utf-8"
        )
        competitive = (
            REPOSITORY_ROOT / "docs" / "competitive-analysis.md"
        ).read_text(encoding="utf-8")
        evidence_matrix = (
            REPOSITORY_ROOT / "docs" / "evidence-matrix.md"
        ).read_text(encoding="utf-8")

        self.assertIn("`Develop-KIM`은 프로젝트 소유자의 계정이 아니므로", origin)
        for public_doc in (competitive, evidence_matrix):
            self.assertIn("`Develop-KIM` is not the participant's account", public_doc)
            self.assertNotIn("Until participant ownership", public_doc)
            self.assertNotIn("Until the participant's ownership", public_doc)

    def test_ai_assistance_uses_bounded_owner_understanding_contract(self) -> None:
        disclosure = (REPOSITORY_ROOT / "AI_ASSISTANCE.md").read_text(
            encoding="utf-8"
        )
        submission_readme = (
            REPOSITORY_ROOT / "submission" / "README.md"
        ).read_text(encoding="utf-8")
        self.assertNotIn(
            "reviewing and understanding every submitted source line and dependency",
            disclosure,
        )
        for required in (
            "reviewing the final submitted diff",
            "core behavior, generated evidence",
            "safety boundaries, and direct-dependency roles",
        ):
            self.assertIn(required, disclosure)
        for required in (
            "core_behavior_boundaries_artifacts_and_dependency_roles_reviewed_and_explainable",
            "purpose and non-proof limits",
            "runtime, compile-only, test, build/audit and report-tool dependencies",
            "not a claim of line-by-line authorship",
        ):
            self.assertIn(required, submission_readme)

    def test_release_toolchain_and_javadoc_classifier_disclosure_are_pinned(
        self,
    ) -> None:
        workflows = {
            path.name: path.read_text(encoding="utf-8")
            for path in (
                REPOSITORY_ROOT / ".github" / "workflows" / "ci.yml",
                REPOSITORY_ROOT
                / ".github"
                / "workflows"
                / "dependency-submission.yml",
                REPOSITORY_ROOT
                / ".github"
                / "workflows"
                / "release-evidence.yml",
            )
        }
        for name, workflow in workflows.items():
            with self.subTest(workflow=name):
                self.assertIn("java-version: '17.0.20+101'", workflow)
                self.assertNotRegex(workflow, r"java-version: [\"']17[\"']")
                self.assertIn(
                    "actions/setup-java@b6effb05e454b25005698d916606bdc6ffcbf961",
                    workflow,
                )
                self.assertIn(
                    "grep -Fxq 'IMPLEMENTOR_VERSION=\"Temurin-17.0.20.1+1\"' "
                    '"${JAVA_HOME}/release"',
                    workflow,
                )
                self.assertIn(
                    "grep -Fxq 'JAVA_RUNTIME_VERSION=\"17.0.20.1+1\"' "
                    '"${JAVA_HOME}/release"',
                    workflow,
                )
        self.assertIn("java -fullversion", workflows["ci.yml"])
        self.assertIn("java -fullversion", workflows["release-evidence.yml"])

        third_party = (REPOSITORY_ROOT / "THIRD_PARTY.md").read_text(
            encoding="utf-8"
        )
        sbom = (REPOSITORY_ROOT / "docs" / "sbom.md").read_text(encoding="utf-8")
        notice = (REPOSITORY_ROOT / "NOTICE").read_text(encoding="utf-8")
        readme = (REPOSITORY_ROOT / "README.md").read_text(encoding="utf-8")
        normalized_third_party = " ".join(third_party.split())
        normalized_sbom = " ".join(sbom.split())
        for required in (
            "Eclipse Temurin/OpenJDK standard-doclet 17.0.20.1+1",
            "GPL-2.0-only WITH Classpath-exception-2.0",
            "jQuery 3.7.1",
            "jQuery UI 1.14.1",
            "`legal/`",
            "not present in the main library JAR or its runtime dependency graph",
            "https://github.com/adoptium/temurin17-binaries/releases/tag/jdk-17.0.20.1%2B1",
            "https://github.com/openjdk/jdk17u/tree/jdk-17.0.20.1%2B1/src/jdk.javadoc",
            "https://github.com/jquery/jquery/tree/3.7.1",
            "https://github.com/jquery/jquery-ui/tree/1.14.1",
            "main JAR accepts only `.class` entry paths under the exact RouteContract package namespace plus an exact metadata allowlist",
            "sources JAR applies the same path policy to `.java` entries and requires every declared package to match its path",
            "source ZIP rejects JTS/Mahout-named files and package paths, every compiled `.class`",
            "do not determine the semantic origin of renamed or copied source/class bytes",
            "does not configure shading or dependency embedding",
            "do not prove the absence or semantic provenance of renamed, relocated, transformed, or copied bytes",
            "owner's source/provenance review",
            "published POM declares no direct JTS or Mahout dependency",
            "cannot contain a Maven parent or relocation",
            "upstream metadata gap remains disclosed",
        ):
            self.assertIn(required, normalized_third_party)
        for required in (
            "Gradle dependency-profile BOMs",
            "not a shipped-file inventory",
            "Javadoc classifier",
            "do not add them as runtime or direct dependency components",
            "non-bundled distribution boundary is distinct from the unresolved upstream metadata gap",
            "do not determine the semantic origin of renamed or copied source/class bytes",
            "reopens if those payload invariants change or a JTS/Mahout published dependency enters the release",
        ):
            self.assertIn(required, normalized_sbom)
        for unsupported in (
            "only first-party class/source entries",
            "proves the absence of JTS/Mahout source/class bytes",
        ):
            self.assertNotIn(unsupported, normalized_third_party)
            self.assertNotIn(unsupported, normalized_sbom)
        for required in (
            "Javadoc classifier",
            "17.0.20.1+1",
            "legal/",
            "THIRD_PARTY.md",
        ):
            self.assertIn(required, notice)
        self.assertIn("Javadoc classifier", readme)
        self.assertIn("THIRD_PARTY.md", readme)

    def test_storyboard_pins_caption_first_nonmisleading_cuts(self) -> None:
        storyboard = (
            REPOSITORY_ROOT / "submission" / "video-storyboard.md"
        ).read_text(encoding="utf-8")
        for required in (
            "`0:00.000–2:53.000`의 모든 프레임",
            "실제 terminal·browser·source",
            "실제 명령을 입력하는 순간",
            "로그아웃한 공개 GitHub의 실제 URL",
            "final revision의 tracked file path",
            "제목·로고 화면, 슬라이드, 요약판, 모의 terminal",
            "정지 screenshot 삽입, 검은 전환 화면",
            "구간 전환은 실제 화면에서",
            "실제 실행 · 대기 구간 8×",
            "쉬운 한국어",
            "`submission/video-caption-cues.json`이 cue 시각·문구·분기의 유일한 원본",
            "generated/reference-only",
            "JSON과 byte-for-byte 같아야 한다",
            "RouteContract는 JDBC 실행 기록을",
            "ShardingSphere의 기능 결과는 같아도",
            "관측된 실행 시도는 1회→2회",
            "방금 실행한 실제 MySQL 결과입니다",
            "CI에 연결하면 exit 1",
            "승인 기록은 자동으로 바뀌지 않습니다",
            "입력값은 저장하지 않습니다",
            "사람이 승인한 기준과 비교합니다",
            "성능·거래 완료를 판단하지 않습니다",
            "이 구간은 실제 GitHub Actions 실패 화면이 아니다",
            "이 local task 자체가 required CI check이거나",
            "실제 PR을 막았다고 말하지 않는다",
            "machine-readable test/verifier 출력에서 실제로 일치한",
            "고정 요약문은 출력하지",
            "ROUTECONTRACT_MANIFEST_DEMO businessResult=UNCHANGED",
            "ROUTECONTRACT_FILE_CI_DEMO approvedAttempts=1 candidateAttempts=2",
            "ROUTECONTRACT_FINGERPRINT_DRIFT_DEMO businessResult=UNCHANGED",
            "BUILD FAILED in <실제 Gradle 소요 시간>",
            "verified_child_exit     0",
            "verified_child_exit     1",
            "바로 이어 final revision의 tracked approved/candidate JSON과 expected diff",
            "제한된 duration-only 전체 줄과 일치할 때만",
            "actual source",
            "actual browser",
            "complete route plan을 판정하지 않습니다",
            "transaction commit",
            "다음 공통 증거가 실제로 공개된 뒤에만 녹화한다",
            'RouteSnapshot snapshot = RouteContract.capture("orders.find", () -> {',
            "Order actual = repository.find(userId);",
            "assertEquals(expectedOrderId, actual.id());",
            "hasExactlyObservedPhysicalAttempts(1)",
            "독립 검증은 공개 양식으로 받습니다",
            "stable 외부 검증 미확보",
            "없는 결과는 만들지 않습니다",
            "실제 사람·독립성을 자동 증명하지 않는다",
            "Tab 자동완성",
            "선택한 분기의 결정적 공개 화면 하나",
            "같은 화면에서 8초 동안",
        ):
            self.assertIn(required, storyboard)

        for stale_visual_or_claim in (
            "RouteContract가 CI에서 차단",
            "CI gate는 exit 1",
            "실제 non-zero CI gate",
            "고정 화면 한 장",
            "한 장의 검증 카드",
            "Release 검증 카드",
            "설치 좌표 카드",
            "기존 테스트에 붙이는 실제 API 카드",
            "actual final-revision output",
            "[MYSQL BASELINE -> CANDIDATE]",
            "[SAME-BUDGET FINGERPRINT DRIFT]",
            "[INTENTIONAL CI GATE]",
            "demoMeaning             expected violation verified",
            "Maven Central",
            "Apache ShardingSphere PR #39535",
            "Task A",
        ):
            self.assertNotIn(stale_visual_or_claim, storyboard)

        timeline_headings = list(
            re.finditer(
                r"^## [0-9]:[0-9]{2}–[0-9]:[0-9]{2} — .+$",
                storyboard,
                flags=re.MULTILINE,
            )
        )
        self.assertEqual(10, len(timeline_headings))
        timeline_end = storyboard.index("## 고정 자막·YouTube 문안")
        for index, heading in enumerate(timeline_headings):
            body_end = (
                timeline_headings[index + 1].start()
                if index + 1 < len(timeline_headings)
                else timeline_end
            )
            section = storyboard[heading.end() : body_end]
            with self.subTest(heading=heading.group(0)):
                self.assertIn("실제 화면 흐름:", section)
                self.assertRegex(
                    section,
                    r"\*\*실제 (?:terminal|browser|source) 화면:\*\*",
                )

        install_section = storyboard.split(
            "## 1:00–1:22 — 실제 공개 Release와 실제 사용 source", 1
        )[1].split(
            "## 1:22–1:40 — 같은 횟수의 구조 변화도 실제 terminal에서 확인",
            1,
        )[0]
        self.assertIn("실제 GitHub Release, Actions run, tracked source", install_section)
        self.assertIn("**실제 browser 화면:**", install_section)
        self.assertIn("**실제 source 화면:**", install_section)
        self.assertIn("별도 화면에 재입력하지 않는다", install_section)
        self.assertNotIn("complete route plan을 증명", install_section)
        self.assertNotIn("transaction commit을 증명", install_section)

        ci_section = storyboard.split(
            "## 0:46–1:00 — 로컬 intentional-red 경로를 실제 terminal에서 실행", 1
        )[1].split(
            "## 1:00–1:22 — 실제 공개 Release와 실제 사용 source", 1
        )[0]
        self.assertIn("0:46.000–0:51.000", ci_section)
        self.assertIn("0:51.000–0:54.000", ci_section)
        self.assertIn("0:58.000", ci_section)
        self.assertIn("이 구간에는 source나 다른 화면을 끼우지 않는다", ci_section)
        self.assertIn("0:12–0:46에서 이미 보여 줬으므로 반복하지 않는다", ci_section)
        self.assertNotIn("**바로 이어지는 실제 source 화면:**", ci_section)

        stable_section = storyboard.split(
            "## 2:07–2:25 — 실제 공개 GitHub 화면에서 안정판 증거 확인", 1
        )[1].split(
            "## 2:25–2:34 — 실제 공개 Issue 화면에서 외부 결과 확인", 1
        )[0]
        for required in (
            "제출 revision full SHA를 가리키는 annotated stable `v0.1.0` tag의 peeled commit",
            "merge PR에서 ruleset-required",
            "같은 final SHA의 main-push `Java 17 / MySQL integration / SBOM` success",
            "logged-out GitHub commit page",
            "merge PR의 Checks tab",
            "final main-push Actions run으로",
            "PR-only라 main push 증거로 세지 않는다",
            "2:07.000–2:15.000 (8초 dwell)",
            "2:15.000–2:19.500 (4.5초 dwell)",
            "2:19.500–2:25.000 (5.5초 dwell)",
            "3초 이상 머문다",
            "scroll 없이 머문다",
            "1:00–1:22의 실제",
            "이 18초에 반복해 넘기지 않는다",
            "별도 요약 화면에 옮겨 적지 않는다",
        ):
            self.assertIn(required, stable_section)
        for duplicated_release_proof in (
            "exact tag SHA의 successful `Release evidence` run",
            "immutable Release의 JAR·sources·Javadoc·POM·SBOM·`SHA256SUMS`",
            "빈 Maven repository에서 Release asset을 검증한 consumer 결과",
        ):
            self.assertNotIn(duplicated_release_proof, stable_section)

    def test_storyboard_pins_readable_caption_corpus_and_timeline(self) -> None:
        storyboard = (
            REPOSITORY_ROOT / "submission" / "video-storyboard.md"
        ).read_text(encoding="utf-8")

        def seconds(value: str) -> float:
            minutes, raw_seconds = value.split(":", 1)
            return int(minutes) * 60 + float(raw_seconds)

        caption_rows = re.findall(
            r"^\| (공통|zero|rc_only) \| ([0-9]:[0-9]{2}\.[0-9]{3}) "
            r"\| ([0-9]:[0-9]{2}\.[0-9]{3}) \| ([^|\n]+) \| ([^|\n]+) \|$",
            storyboard,
            flags=re.MULTILINE,
        )
        expected_rows = [
            ("공통", "0:00.500", "0:05.200", "RouteContract는 JDBC 실행 기록을", "승인본과 비교합니다"),
            ("공통", "0:05.700", "0:11.500", "ShardingSphere의 기능 결과는 같아도", "관측된 실행 시도는 1회→2회"),
            ("공통", "0:12.500", "0:19.000", "방금 실행한 실제 MySQL 결과입니다", "명령과 종료 상태를 함께 봅니다"),
            ("공통", "0:19.500", "0:27.000", "승인된 기록은 실행 한 번입니다", "변경된 기록은 두 번입니다"),
            ("공통", "0:27.500", "0:35.000", "기능 결과는 그대로 한 행입니다", "달라진 것은 내부 실행 모습입니다"),
            ("공통", "0:35.500", "0:44.500", "정한 한도를 넘자 두 위반을 냈습니다", "자동 승인하지 않고 검토를 요구합니다"),
            ("공통", "0:46.500", "0:52.500", "CI에 연결하면 exit 1", "의도한 실패로 빌드를 멈춥니다"),
            ("공통", "0:53.000", "0:59.500", "승인 기록은 자동으로 바뀌지 않습니다", "사람이 차이를 본 뒤에만 바꿉니다"),
            ("공통", "1:00.500", "1:07.000", "공개 배포 파일과 검사값을 확인합니다", "빈 저장소에 직접 설치합니다"),
            ("공통", "1:07.500", "1:14.500", "설치한 파일로 실제 MySQL을 실행합니다", "통과한 공개 기록을 직접 봅니다"),
            ("공통", "1:15.000", "1:21.500", "기존 기능 테스트를 그대로 감쌉니다", "새 기록은 승인본과 비교됩니다"),
            ("공통", "1:22.500", "1:28.500", "실행 횟수는 그대로 한 번입니다", "그래도 기록의 모양은 달라졌습니다"),
            ("공통", "1:29.000", "1:35.000", "입력값은 저장하지 않습니다", "자료형 개수만 한 개에서 두 개로 바뀝니다"),
            ("공통", "1:35.500", "1:39.500", "횟수만 같아도 승인하지 않습니다", "SQL 뜻은 판단하지 않습니다"),
            ("공통", "1:40.500", "1:47.000", "실제 MySQL 여덟 사례를 스무 번씩", "모두 같은 기록으로 되풀이했습니다"),
            ("공통", "1:47.500", "1:54.500", "동시에 실행한 20쌍은 섞이지 않았습니다", "실제 호출의 겹침은 측정하지 않았습니다"),
            ("공통", "1:55.500", "2:01.500", "기록을 한 작업별로 묶고", "사람이 승인한 기준과 비교합니다"),
            ("공통", "2:02.000", "2:06.500", "의도치 않은 차이는 CI 실패", "승인 기준은 자동으로 바뀌지 않습니다"),
            ("공통", "2:07.500", "2:15.000", "제출 코드와 안정판이 같은 코드인지", "공개 이력에서 직접 확인합니다"),
            ("공통", "2:15.500", "2:24.500", "코드 변경 검사와 main 검사 결과를", "실제 공개 화면에서 확인합니다"),
            ("zero", "2:25.500", "2:33.500", "독립 검증은 공개 양식으로 받습니다", "없는 결과는 만들지 않습니다"),
            ("rc_only", "2:25.500", "2:33.500", "정해진 양식의 RC 결과 접수는 1건", "자기 확인 진술이며 안정판 검증은 아닙니다"),
            ("공통", "2:34.500", "2:41.000", "검증 범위는 5.5.3 동기 실행", "성능·거래 완료를 판단하지 않습니다"),
            ("공통", "2:41.500", "2:48.000", "기능 결과가 같아도 hook 보고 실행 시도는", "한 번에서 두 번으로 달라질 수 있습니다"),
            ("공통", "2:48.500", "2:52.500", "새 기록을 승인본과 비교해", "CI에서 의도치 않은 차이를 멈춥니다"),
        ]
        self.assertEqual(expected_rows, caption_rows)
        self.assertEqual(23, sum(branch == "공통" for branch, *_ in caption_rows))
        self.assertEqual(1, sum(branch == "zero" for branch, *_ in caption_rows))
        self.assertEqual(1, sum(branch == "rc_only" for branch, *_ in caption_rows))

        for branch, start_text, end_text, line_one, line_two in caption_rows:
            with self.subTest(branch=branch, start=start_text, text=line_one):
                start = seconds(start_text)
                end = seconds(end_text)
                duration = end - start
                self.assertGreaterEqual(duration, 4.0)
                self.assertLessEqual(duration, 9.0)
                self.assertLessEqual(len(line_one.strip()), 34)
                self.assertLessEqual(len(line_two.strip()), 34)
                visible = re.sub(r"\s+", "", line_one + line_two)
                self.assertLessEqual(len(visible) / duration, 8.0)

        caption_corpus = "\n".join(
            line_one + "\n" + line_two
            for _branch, _start, _end, line_one, line_two in caption_rows
        )
        for internal_term in (
            "RCM201",
            "RCM202",
            "RCM301",
            "RCM302",
            "checksum",
            "attestation",
            "Maven",
            "Release",
            "capture",
            "approved",
            "candidate",
            "fingerprint",
            "parameter",
            "alias",
            "caller-operation",
            "callback",
            "datasource-proxy",
            "SBOM",
            "PR",
            "commit",
            "stable",
            "cutoff",
            "manifest",
        ):
            with self.subTest(internal_caption_term=internal_term):
                self.assertNotIn(internal_term, caption_corpus)

        for branch in ("zero", "rc_only"):
            selected = [
                (seconds(start), seconds(end))
                for row_branch, start, end, _line_one, _line_two in caption_rows
                if row_branch in {"공통", branch}
            ]
            self.assertEqual(selected, sorted(selected))
            self.assertEqual(0.5, selected[0][0])
            self.assertEqual(172.5, selected[-1][1])
            for previous, current in zip(selected, selected[1:]):
                self.assertGreaterEqual(current[0] - previous[1], 0.5)

        heading_matches = re.findall(
            r"^## (\d+):(\d{2})–(\d+):(\d{2}) — (.+)$",
            storyboard,
            flags=re.MULTILINE,
        )
        self.assertEqual(10, len(heading_matches))
        timeline = [
            (int(sm) * 60 + int(ss), int(em) * 60 + int(es))
            for sm, ss, em, es, _label in heading_matches
        ]
        self.assertEqual(0, timeline[0][0])
        self.assertEqual(173, timeline[-1][1])
        for previous, current in zip(timeline, timeline[1:]):
            self.assertEqual(previous[1], current[0])

        self.assertIn("audio stream이 정확히 0개", storyboard)
        self.assertIn("빈 audio track", storyboard)
        self.assertIn("burned-in caption", storyboard)
        self.assertNotIn("내레이션:", storyboard)
        self.assertNotIn("Yuna 180 wpm", storyboard)

    def test_video_delivery_docs_separate_machine_and_owner_qc(self) -> None:
        readme = (REPOSITORY_ROOT / "submission" / "README.md").read_text(
            encoding="utf-8"
        )
        storyboard = (
            REPOSITORY_ROOT / "submission" / "video-storyboard.md"
        ).read_text(encoding="utf-8")

        for required in (
            "at least 1920x1080",
            "exactly zero audio",
            "170 through 175 seconds inclusive",
            "at least 20 decoded frames per second on average",
            "a complete selected-stream decode",
            "post-decode file-hash recheck",
            "public, non-live, age-unrestricted",
            "downloadable 1080p",
            "checks all reported formats",
            "cannot grade whether the",
            "burned-in captions match the visible real execution",
            "owner must still watch the checksummed local file",
            "`ffmpeg` for a full decode",
            "download the entire public video",
        ):
            self.assertIn(required, readme)
        for required in (
            "1920×1080 이상",
            "audio stream이 정확히 0개",
            "공개·non-live·연령 제한 없음",
            "1080p 이상 format",
            "실제 화면·burned-in caption 일치와 화면 가독성",
        ):
            self.assertIn(required, storyboard)

    def test_public_docs_keep_build_consumer_and_contest_boundaries_explicit(self) -> None:
        contributing = (REPOSITORY_ROOT / "CONTRIBUTING.md").read_text(encoding="utf-8")
        releasing = (REPOSITORY_ROOT / "RELEASING.md").read_text(encoding="utf-8")
        readme_ko = (REPOSITORY_ROOT / "README.md").read_text(encoding="utf-8")
        readme_en = (REPOSITORY_ROOT / "README.en.md").read_text(encoding="utf-8")
        competitive = (
            REPOSITORY_ROOT / "docs" / "competitive-analysis.md"
        ).read_text(encoding="utf-8")
        specification = (
            REPOSITORY_ROOT / "docs" / "specification.md"
        ).read_text(encoding="utf-8")
        submission_readme = (
            REPOSITORY_ROOT / "submission" / "README.md"
        ).read_text(encoding="utf-8")
        storyboard = (
            REPOSITORY_ROOT / "submission" / "video-storyboard.md"
        ).read_text(encoding="utf-8")
        release_evidence = (
            REPOSITORY_ROOT / ".github" / "workflows" / "release-evidence.yml"
        ).read_text(encoding="utf-8")
        ci_workflow = (
            REPOSITORY_ROOT / ".github" / "workflows" / "ci.yml"
        ).read_text(encoding="utf-8")

        for required in (
            "validateOfficialCycloneDxSbom",
            "same-checkout packaging evidence",
            "Synthetic, non-sensitive test values are allowed",
            "scripts/verify-release-assets-consumer.sh",
        ):
            self.assertIn(required, contributing)
        self.assertNotIn("repository ownership, signing and", releasing)
        self.assertIn("v0.1 packaging gate requires no signature assets", releasing)
        self.assertIn(
            "asserted-absent task-specific Gradle user home before any\n  "
            "project-local Python or Gradle command",
            releasing,
        )
        self.assertIn("intentional contract-gate\n  child exit `1`", releasing)
        self.assertIn("혼합 자동화", readme_ko)
        self.assertTrue(readme_ko.startswith("# RouteContract for ShardingSphere-JDBC\n"))
        self.assertIn("사용하거나 도입을 평가하는 Java 개발자·팀", readme_ko)
        self.assertIn("`1 → 2` 자체를 성능 결함으로 단정", readme_ko)
        self.assertIn(
            "MyBatis·JPA·Hibernate별 end-to-end 호환성을 검증했다는 뜻은 아닙니다",
            readme_ko,
        )
        self.assertNotIn("위 Quick Start의 Jackson 2 BOM", readme_ko)
        self.assertLess(readme_ko.index("## Quick Start"), readme_ko.index("## 가장 작은 사용 예"))
        self.assertLess(readme_ko.index("## 가장 작은 사용 예"), readme_ko.index("## 검증된 핵심 시나리오"))
        self.assertLess(readme_ko.index("## 검증된 핵심 시나리오"), readme_ko.index("## 기존 도구와의 정확한 차이"))
        self.assertLess(readme_ko.index("## 기존 도구와의 정확한 차이"), readme_ko.index("## 코드·공개 증거 경계"))
        self.assertLess(readme_ko.index("## v0.1 지원 범위"), readme_ko.index("## 의존성·Release 호환성 상세"))
        self.assertIn(
            "[검증 증거 매트릭스](docs/evidence-matrix.md)", readme_ko
        )
        self.assertIn(
            "datasource-proxy도 충분히 신뢰할 수 있는 직접 구현 대안", readme_ko
        )
        self.assertIn("모든 물리 data source wrapper 없이", readme_ko)
        self.assertIn("docs/empirical-comparison.md", readme_ko)
        self.assertIn("Mixed automation", readme_en)
        self.assertTrue(readme_en.startswith("# RouteContract for ShardingSphere-JDBC\n"))
        self.assertIn("developers and teams using or evaluating", readme_en)
        self.assertIn("does not label `1 → 2` itself as a performance defect", readme_en)
        self.assertIn(
            "not verified end-to-end compatibility with each of MyBatis, JPA, and Hibernate",
            readme_en,
        )
        self.assertNotIn("BOM shown in Quick Start", readme_en)
        self.assertLess(readme_en.index("## Quick Start"), readme_en.index("## Smallest usage example"))
        self.assertLess(readme_en.index("## Smallest usage example"), readme_en.index("## Verified core scenarios"))
        self.assertLess(readme_en.index("## Verified core scenarios"), readme_en.index("## Precise comparison with existing tools"))
        self.assertLess(readme_en.index("## Precise comparison with existing tools"), readme_en.index("## Code and public-evidence boundaries"))
        self.assertLess(readme_en.index("## v0.1 support boundary"), readme_en.index("## Dependency and Release compatibility details"))
        driver_coordinate = (
            'testImplementation("org.apache.shardingsphere:'
            'shardingsphere-jdbc:5.5.3")'
        )
        for readme in (readme_ko, readme_en):
            self.assertEqual(2, readme.count(driver_coordinate))
            for block in re.findall(r"```groovy\n(.*?)```", readme, re.DOTALL):
                if "routecontract-shardingsphere-5.5:0.1.0" not in block:
                    continue
                self.assertLess(block.index("jackson-bom:2.18.9"), block.index(driver_coordinate))
                self.assertLess(
                    block.index(driver_coordinate),
                    block.index("routecontract-shardingsphere-5.5:0.1.0"),
                )
        self.assertIn(
            "[Verification evidence matrix](docs/evidence-matrix.md)",
            readme_en,
        )
        self.assertIn("RouteContract v0.1 implemented surface", competitive)
        self.assertIn("evidence status at its declared snapshot", competitive)
        self.assertIn(
            "deterministic structural added/removed attempt\n  signatures",
            specification,
        )
        self.assertIn("does not establish SQL semantic equivalence", specification)
        self.assertNotIn("semantic added/removed signatures", specification)
        for readme, heading in (
            (readme_ko, "## 공개 Release 자산을 registry 없이 사용하기"),
            (readme_en, "## Consume public Release assets without a registry"),
        ):
            release_section = readme.split(heading, 1)[1].split("\n## ", 1)[0]
            ordered = (
                "gh release download v0.1.0",
                "python3 scripts/install-release-assets.py",
                "exclusiveContent",
                'maven { url = uri("/absolute/path/to/routecontract-maven") }',
                'testImplementation(platform("com.fasterxml.jackson:jackson-bom:2.18.9"))',
                'testImplementation("org.apache.shardingsphere:shardingsphere-jdbc:5.5.3")',
                'testImplementation("io.github.ym0506.routecontract:routecontract-shardingsphere-5.5:0.1.0")',
            )
            self.assertTrue(all(value in release_section for value in ordered))
            self.assertEqual(
                sorted(release_section.index(value) for value in ordered),
                [release_section.index(value) for value in ordered],
            )
        release_quick_start = "      - name: Verify public Quick Start"
        ci_quick_start = "      - name: Verify the documented real-MySQL Quick Start"
        ci_package_job = (
            "  build-and-sbom:\n"
            "    name: Java 17 / MySQL integration / SBOM\n"
            "    runs-on: ubuntu-24.04"
        )
        release_package_job = (
            "  build-release-evidence:\n"
            "    name: Build release evidence\n"
            "    runs-on: ubuntu-24.04"
        )
        test_and_build = "      - name: Test and build evidence"
        test_summary = "      - name: Create revision-bound privacy-minimized test summary"
        python_setup = "      - name: Set up report-package Python"
        python_install = (
            "      - name: Install checksum-locked report-package test dependencies"
        )
        package_tests = "      - name: Test fail-closed contest packaging rules"
        for workflow, quick_start, package_job in (
            (ci_workflow, ci_quick_start, ci_package_job),
            (release_evidence, release_quick_start, release_package_job),
        ):
            self.assertIn(package_job, workflow)
            self.assertEqual(
                1,
                workflow.count(
                    "--requirement submission/report-package-ci-requirements.txt"
                ),
            )
            self.assertNotIn(
                "--requirement submission/report-builder-requirements.txt", workflow
            )
            self.assertEqual(
                1,
                workflow.count(
                    "uses: actions/setup-python@"
                    "5fda3b95a4ea91299a34e894583c3862153e4b97 # v7.0.0"
                ),
            )
            self.assertEqual(1, workflow.count("python-version: '3.12.14'"))
            self.assertEqual(2, workflow.count("architecture: x64"))
            self.assertEqual(1, workflow.count("check-latest: false"))
            self.assertIn("${RUNNER_TEMP}/routecontract-report-test-venv", workflow)
            self.assertEqual(1, workflow.count('test ! -e "${report_venv}"'))
            self.assertEqual(1, workflow.count('python -m venv "${report_venv}"'))
            self.assertEqual(1, workflow.count('"${report_venv}/bin/python"'))
            self.assertIn("--only-binary=:all: --no-deps", workflow)
            self.assertEqual(1, workflow.count("--require-hashes"))
            self.assertEqual(1, workflow.count("--isolated"))
            self.assertEqual(1, workflow.count("--index-url https://pypi.org/simple"))
            self.assertIn("sys.version_info[:3] == (3, 12, 14)", workflow)
            self.assertIn('platform.python_implementation() == "CPython"', workflow)
            self.assertIn('platform.machine() == "x86_64"', workflow)
            self.assertEqual(
                2,
                workflow.count(
                    '"${RUNNER_TEMP}/routecontract-report-test-venv/bin/python"'
                ),
            )
            self.assertNotIn("python3 -m unittest", workflow)
            self.assertLess(workflow.index(python_setup), workflow.index(python_install))
            self.assertLess(workflow.index(python_install), workflow.index(package_tests))
            self.assertEqual(1, workflow.count("routecontract-quickstart-gradle-home"))
            quick_start_offset = workflow.index(quick_start)
            next_step_offset = workflow.find("\n      - name:", quick_start_offset + 1)
            self.assertNotEqual(-1, next_step_offset)
            quick_start_block = workflow[quick_start_offset:next_step_offset]
            status_check = (
                'test -z "$(git status --porcelain=v1 --untracked-files=all '
                '--ignore-submodules=none)"'
            )
            ignored_check = 'test -z "$(git clean -ndx)"'
            home_assignment = (
                'quickstart_gradle_home="${RUNNER_TEMP}/'
                'routecontract-quickstart-gradle-home"'
            )
            absent_home_check = 'test ! -e "${quickstart_gradle_home}"'
            quick_start_invocation = (
                'GRADLE_USER_HOME="${quickstart_gradle_home}" '
                "./scripts/quickstart-demo.sh"
            )
            for command in (
                status_check,
                ignored_check,
                home_assignment,
                absent_home_check,
                quick_start_invocation,
            ):
                self.assertIn(command, quick_start_block)
            self.assertLess(
                quick_start_block.index(status_check),
                quick_start_block.index(ignored_check),
            )
            self.assertLess(
                quick_start_block.index(ignored_check),
                quick_start_block.index(home_assignment),
            )
            self.assertLess(
                quick_start_block.index(home_assignment),
                quick_start_block.index(absent_home_check),
            )
            self.assertLess(
                quick_start_block.index(absent_home_check),
                quick_start_block.index(quick_start_invocation),
            )
            self.assertEqual(
                1,
                workflow.count(
                    'git status --porcelain=v1 --untracked-files=all '
                    "--ignore-submodules=none"
                ),
            )
            self.assertEqual(1, workflow.count("git clean -ndx"))
            full_gradle_validation = (
                "./gradlew --no-daemon --no-build-cache clean check"
            )
            self.assertEqual(1, workflow.count(full_gradle_validation))
            self.assertLess(workflow.index(quick_start), workflow.index(python_setup))
            self.assertLess(workflow.index(quick_start), workflow.index(python_install))
            self.assertLess(
                workflow.index(quick_start),
                workflow.index(full_gradle_validation),
            )
        self.assertEqual(1, release_evidence.count("./scripts/quickstart-demo.sh"))
        self.assertIn(release_quick_start, release_evidence)
        release_tag_binding = (
            "      - name: Require tag revision and public main to match"
        )
        release_version_binding = (
            "      - name: Require tag and project versions to match"
        )
        self.assertLess(
            release_evidence.index(release_tag_binding),
            release_evidence.index(release_quick_start),
        )
        self.assertLess(
            release_evidence.index(release_quick_start),
            release_evidence.index(release_version_binding),
        )
        self.assertLess(
            release_evidence.index(release_quick_start),
            release_evidence.index(test_and_build),
        )
        self.assertLess(release_evidence.index(test_and_build), release_evidence.index(test_summary))
        self.assertIn(
            "Do not invent, append, or combine a detail that the application does not display",
            submission_readme.replace("\n", " "),
        )
        self.assertIn("high-confidence lexical privacy scanner", submission_readme)
        self.assertIn("it is a heuristic", submission_readme)
        self.assertIn("report_free_text_privacy_reviewed=true", submission_readme)
        self.assertIn(
            "--content submission/private/report-content.final.ko.json \\\n"
            "  --assets-dir submission/assets \\\n"
            "  --output submission/draft/routecontract-result-report-final.docx \\\n"
            "  --strict-final",
            submission_readme,
        )
        self.assertNotIn(
            "--content submission/report-content.ko.json \\\n+  --output submission/draft/routecontract-result-report-final.docx",
            submission_readme,
        )
        self.assertEqual(
            "[[APPLICATION_TASK_TYPE]]",
            json.loads(
                (REPOSITORY_ROOT / "submission" / "report-content.ko.json").read_text(
                    encoding="utf-8"
                )
            )["metadata"]["task_type"],
        )
        self.assertEqual(
            "[[APPLICATION_TASK_TYPE]]",
            json.loads(
                (
                    REPOSITORY_ROOT
                    / "submission"
                    / "package-manifest.example.json"
                ).read_text(encoding="utf-8")
            )["submission_identity"]["task_type"],
        )
        self.assertIn("공개 오픈소스 시연 영상의 결과 우선", storyboard)
        self.assertNotIn("금상", storyboard)
        self.assertNotIn("은상", storyboard)

    def test_august_20_notice_41_rubric_has_exact_second_round_weights(self) -> None:
        matrix = (REPOSITORY_ROOT / "docs" / "evidence-matrix.md").read_text(
            encoding="utf-8"
        )
        self.assertIn("Rubric source checked: 2026-08-21 KST", matrix)
        self.assertIn(
            "[August 20 judging notice](https://osscontest.kr/notice/41)",
            matrix,
        )

        section = matrix.split("## 2차 평가(발표): 70 points", 1)[1].split(
            "The second round applies", 1
        )[0]
        rows = re.findall(r"^\| ([^|]+?) \| (\d+) \|", section, flags=re.MULTILINE)
        self.assertEqual(7, len(rows))
        self.assertEqual(7, len({name for name, _ in rows}))
        self.assertEqual(
            {
                "작품발표(PT)": 10,
                "활용성": 15,
                "작품 데모(완성도)": 10,
                "커뮤니티 확장 가능성": 5,
                "오픈소스SW 적절성": 15,
                "기능테스트": 10,
                "라이선스 검증": 5,
            },
            {name: int(points) for name, points in rows},
        )
        self.assertEqual(70, sum(int(points) for _, points in rows))

    def test_second_round_evidence_mappings_track_application_and_oss_proof(
        self,
    ) -> None:
        matrix = (REPOSITORY_ROOT / "docs" / "evidence-matrix.md").read_text(
            encoding="utf-8"
        )
        section = matrix.split("## 2차 평가(발표): 70 points", 1)[1].split(
            "The second round applies", 1
        )[0]

        self.assertRegex(
            section,
            r"\| 활용성 \| 15 \| E04, E05, E06, E08, E09, E10 \|",
        )
        self.assertRegex(
            section,
            r"\| 오픈소스SW 적절성 \| 15 \| "
            r"E02, E03, E06, E08, E09, E11, E12 \|",
        )
        functional_row = next(
            line for line in section.splitlines() if line.startswith("| 기능테스트 |")
        )
        self.assertIn("final stable exact-revision", functional_row)
        self.assertIn("clean-clone Quick Start", functional_row)
        self.assertNotIn("independent-user", functional_row)
        oss_row = next(
            line
            for line in section.splitlines()
            if line.startswith("| 오픈소스SW 적절성 |")
        )
        self.assertIn("owner license/NOTICE", oss_row)
        self.assertNotIn("external consumption", oss_row)

    def test_independent_rc_protocol_uses_twelve_asset_supply_chain_contract(self) -> None:
        protocol = (REPOSITORY_ROOT / "docs/independent-install-study.md").read_text(
            encoding="utf-8"
        )
        issue_form = (
            REPOSITORY_ROOT / ".github/ISSUE_TEMPLATE/independent-rc1-install.yml"
        ).read_text(encoding="utf-8")
        self.assertIn("exactly eleven project payloads", protocol)
        self.assertIn("supply-chain-evidence.json", protocol)
        self.assertIn("exactly twelve uploaded assets", protocol)
        self.assertIn("Never mix different RC filenames", protocol)
        self.assertIn("cannot authenticate an earlier RC's JAR bytes", protocol)
        self.assertIn("workflow artifact identity/digest", protocol)
        self.assertIn("exact 17-file flat allowlist", protocol)
        self.assertIn("the 11 payload names", issue_form)
        self.assertIn("12/12 verify-asset=pass", issue_form)
        self.assertIn("flat files=17", issue_form)
        self.assertIn(
            "Direct project assets: 12 files (the tagged 11-payload allowlist + SHA256SUMS)",
            issue_form,
        )
        for stale in (
            "exactly ten project payloads",
            "11/11 verify-asset=pass",
            "tagged 10-payload allowlist",
        ):
            self.assertNotIn(stale, protocol + issue_form)

    def test_storyboard_requires_two_exclusive_external_result_branches(
        self,
    ) -> None:
        storyboard = (
            REPOSITORY_ROOT / "submission" / "video-storyboard.md"
        ).read_text(encoding="utf-8")

        for required in (
            "**rc-only-result 분기:**",
            "정해진 양식의 RC 결과 접수 1건",
            "참가자의 자기 확인 진술",
            "stable 검증·adoption 아님",
            "실제 result Issue만 보여 준다",
            "**0-result 분기:**",
            "독립 검증은 공개 양식으로 받습니다",
            "stable 외부 검증 미확보",
            "실제 Discussion #28만 보여 준다",
            "activation·모집·프로토콜 링크는 보고서와 영상 설명에 남기고",
            "package manifest의 `video.external_evidence_branch`",
            "실제 browser 화면·burned-in caption은 같은 분기를 사용",
            "`rc_only` ↔ `rc-only-result`, `zero` ↔ `0-result`",
            "실제 Discussion #28",
            "같은 화면에서 8초 동안 전환·scroll 없이 머문다",
            "다른 분기의 문구나 화면은 한 프레임도 넣지 않는다",
            "게시하지 않았다면 browser 화면과 자막에서 제외",
            "final-stable-result 분기는 fail-closed",
            "실제 사람·독립성·채택·endorsement를 추정하지 않는다",
            "실제 사람·독립성을 자동 증명하지 않는다",
        ):
            self.assertIn(required, storyboard)

        for unconditional_claim in (
            "**qualified-result 분기:**",
            "비작성자가 직접 남긴 첫 quick-start 결과 Issue",
            "최종 revision의 CI와 checksummed Release, 비작성자의 첫 설치 "
            "결과, upstream 질문을 모두 공개 링크로 남겼습니다.",
            "- [ ] 실제 비작성자의 첫 quick-start 결과가 본인 계정으로 "
            "공개되어 있다.",
            "비작성자 독립 clean-install 첫 결과 1건을 adoption과 구분해 링크했습니다.",
            "모든 acceptance criteria",
            "qualified Issue",
            "Issue #9 form의 14개 필수 self-attestation",
            "REST/GraphQL은 현재 editor·last edit·retained body edit·title rename",
            "maintainer 수정·삭제·은폐·이전·누락",
        ):
            self.assertNotIn(unconditional_claim, storyboard)

    def test_report_uses_generated_external_evidence_slot(self) -> None:
        report = json.loads(
            (REPOSITORY_ROOT / "submission" / "report-content.ko.json").read_text(
                encoding="utf-8"
            )
        )
        external = report["external_evidence"]
        self.assertEqual(
            set(package_submission.REPORT_CONTENT_CONTRACT.EXTERNAL_EVIDENCE_KEYS),
            set(external),
        )
        self.assertEqual(
            "https://github.com/ym0506/routecontract/issues/9",
            external["protocol_issue_url"],
        )
        slot = [item for item in report["other"] if item["lead"] == "외부 검증"]
        self.assertEqual(
            [package_submission.REPORT_CONTENT_CONTRACT.EXTERNAL_EVIDENCE_SUMMARY_MARKER],
            [item["text"] for item in slot],
        )
        community = next(
            item["text"]
            for item in report["other"]
            if item["lead"] == "품질관리·발전 로드맵"
        )

        self.assertIn("외부 결과는 링크로만 보고한다", community)
        self.assertNotIn("비작성자 독립 설치 결과", community)
        self.assertNotIn("qualified 결과", community)
        self.assertNotIn("외부 검증 미확보", community)
        self.assertNotIn(
            "외부 설치 기록·license/security 재검토를 완료", community
        )


class TaggedIssueFormAllowlistTest(unittest.TestCase):
    FORM_ROOT = REPOSITORY_ROOT / ".github/ISSUE_TEMPLATE"
    APPROVED = {
        "independent-rc1-install.yml": (
            "0f4afc4ac098e0ee425704168f045352b3e2a77f856a0ae7438a9f93d955e583"
        ),
        "independent-rc2-install.yml": (
            "518c4102b9a0f7725b46b825ad5952263b3418bdb07b0164c54a037d902e7f8a"
        ),
    }

    def test_accepts_reviewed_rc1_and_rc2_bytes(self) -> None:
        self.assertEqual(
            self.APPROVED,
            package_submission.APPROVED_ISSUE_FORM_SHA256_BY_FILENAME,
        )
        for filename, expected_sha256 in self.APPROVED.items():
            with self.subTest(filename=filename):
                data = (self.FORM_ROOT / filename).read_bytes()
                self.assertEqual(expected_sha256, hashlib.sha256(data).hexdigest())
                package_submission._validate_tagged_issue_form_bytes(data, filename)

    def test_rejects_tampering_cross_version_bytes_and_unreviewed_rc3(self) -> None:
        rc1 = (self.FORM_ROOT / "independent-rc1-install.yml").read_bytes()
        rc2 = (self.FORM_ROOT / "independent-rc2-install.yml").read_bytes()
        for data, filename in (
            (rc2 + b"\n", "independent-rc2-install.yml"),
            (rc1, "independent-rc2-install.yml"),
            (rc2, "independent-rc1-install.yml"),
        ):
            with self.subTest(filename=filename, data_sha256=hashlib.sha256(data).hexdigest()), (
                self.assertRaisesRegex(
                    package_submission.GateError,
                    "reviewed version-specific form",
                )
            ):
                package_submission._validate_tagged_issue_form_bytes(data, filename)

        with self.assertRaisesRegex(
            package_submission.GateError,
            "outside the contest package allowlist",
        ):
            package_submission._validate_tagged_issue_form_bytes(
                rc2.replace(b"0.1.0-rc2", b"0.1.0-rc3"),
                "independent-rc3-install.yml",
            )


class ReportExternalEvidenceContractTest(unittest.TestCase):
    def materialize(self, branch: str) -> dict:
        return package_submission.validate_and_materialize_report_content(
            valid_report_content(branch),
            valid_manifest(),
            current_utc=TEST_CURRENT_UTC,
        )

    @staticmethod
    def summary(content: dict) -> str:
        return next(
            item["text"] for item in content["other"] if item["lead"] == "외부 검증"
        )

    def test_final_stable_branch_is_fail_closed_without_distinct_protocol(self) -> None:
        with self.assertRaisesRegex(
            package_submission.GateError, "distinct reviewed stable form and protocol"
        ):
            self.materialize("final_stable")

    def test_rc_only_branch_names_rc_and_denies_stable_validation_and_adoption(
        self,
    ) -> None:
        summary = self.summary(self.materialize("rc_only"))
        self.assertEqual(
            "cutoff 2026-08-01T06:00:00Z까지 exact RC v0.1.0-rc1의 Task A 형식 "
            "요건을 충족한 API-visible 공개 self-attestation 1건 [결과 Issue]을 "
            "확인했다. 이는 실제 사람·비공개 독립성·stable 검증·채택·endorsement를 "
            "자동 증명하지 않는다. API로 복원할 수 없는 변경 이력은 자동 검증 범위 "
            "밖이다. [활성화 기록]·[모집 기록]·[검증 프로토콜]",
            summary,
        )

    def test_zero_branch_requires_honest_zero_and_no_result_issue(self) -> None:
        summary = self.summary(self.materialize("zero"))
        self.assertEqual(
            "cutoff 2026-08-01T06:00:00Z까지 exact RC v0.1.0-rc1 공개 모집 "
            "[모집 기록]과 [검증 프로토콜]을 운영했으나 API-visible Task A 형식 "
            "요건을 충족한 공개 결과는 0건이었다. 따라서 독립 외부 설치·채택·stable "
            "검증을 주장하지 않는다. API로 복원할 수 없는 변경 이력은 자동 검증 "
            "범위 밖이다. [활성화 기록]",
            summary,
        )

    def test_rejects_reader_facing_evidence_ids_in_every_owner_overlay(self) -> None:
        markers = ("E09", "e09", "Ｅ０９", "Ｅ09", "E０9")
        for path in package_submission.REPORT_CONTENT_CONTRACT.OWNER_FREE_TEXT_OVERLAY_STRING_PATHS:
            for marker in markers:
                content = valid_report_content("zero")
                parent = content
                for component in path[:-1]:
                    parent = parent[component]
                value = f"owner-reviewed prose {marker}"
                parent[path[-1]] = value
                with self.subTest(path=path, marker=marker), self.assertRaisesRegex(
                    package_submission.GateError,
                    r"reader-facing audit evidence IDs \(count=1\)",
                ) as caught:
                    package_submission.validate_and_materialize_report_content(
                        content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                    )
                self.assertNotIn(value, str(caught.exception))

    def test_reader_facing_evidence_id_counter_is_normalized_and_bounded(self) -> None:
        counter = package_submission.REPORT_CONTENT_CONTRACT.count_reader_facing_evidence_ids
        for marker in ("E09", "e09", "Ｅ０９", "Ｅ09", "E０9", "E\u200b09", "ThinkPad E14"):
            with self.subTest(marker=marker):
                self.assertEqual(1, counter(marker))
        for allowed in ("E15", "AE09", "E09Z"):
            with self.subTest(allowed=allowed):
                self.assertEqual(0, counter(allowed))

    def test_report_sentence_keeps_only_the_branch_necessary_public_url(self) -> None:
        rc_summary = self.summary(self.materialize("rc_only"))
        zero_summary = self.summary(self.materialize("zero"))

        self.assertIn("[결과 Issue]", rc_summary)
        self.assertIn("[모집 기록]", rc_summary)
        self.assertIn("[모집 기록]", zero_summary)
        for summary in (rc_summary, zero_summary):
            self.assertIn("[활성화 기록]", summary)
            self.assertIn("[검증 프로토콜]", summary)
            self.assertNotIn(
                "/docs/evidence/independent-rc-activation-v0.1.0-rc1.json",
                summary,
            )
            self.assertIn("활성화", summary)

    def test_rejects_branch_specific_permalink_omissions_or_leaks(self) -> None:
        mutations = (
            ("rc_only", {"activation_record_url": None}, "activation-record permalink"),
            ("rc_only", {"recruitment_record_url": None}, "Issue #9 comment"),
            ("zero", {"recruitment_record_url": None}, "Issue #9 comment"),
        )
        for branch, mutation, message in mutations:
            content = valid_report_content(branch)
            content["external_evidence"].update(mutation)
            with self.subTest(branch=branch, mutation=mutation), self.assertRaisesRegex(
                package_submission.GateError, message
            ):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

    def test_rejects_future_cutoff_against_explicit_validation_utc(self) -> None:
        content = valid_report_content("zero")
        content["external_evidence"]["cutoff_utc"] = "2026-08-15T00:00:00Z"

        with self.assertRaisesRegex(package_submission.GateError, "later than current UTC"):
            package_submission.validate_and_materialize_report_content(
                content, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )

    def test_rejects_unicode_issue_number(self) -> None:
        unicode_issue = valid_report_content("rc_only")
        unicode_issue["external_evidence"]["result_issue_url"] = (
            "https://github.com/example-owner/routecontract/issues/²"
        )
        with self.assertRaisesRegex(
            package_submission.GateError, "canonical public Issue URL"
        ):
            package_submission.validate_and_materialize_report_content(
                unicode_issue, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )

    def test_rejects_change_outside_documented_private_overlay(self) -> None:
        mutations = (
            lambda content: content["features"][0].__setitem__(
                "text", content["features"][0]["text"] + " 변경"
            ),
            lambda content: content["other"].append(
                {
                    "lead": "추가",
                    "text": "최종 안정판이 외부에서 정상 동작함을 확인했다.",
                }
            ),
            lambda content: content["sbom"][0].__setitem__("version", "99.0"),
        )
        for mutate in mutations:
            content = valid_report_content("rc_only")
            mutate(content)
            with self.subTest(mutate=mutate), self.assertRaisesRegex(
                package_submission.GateError,
                "canonical closed source outside documented private-overlay fields",
            ):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

    def test_accepts_each_documented_private_overlay_path(self) -> None:
        paths = package_submission.REPORT_CONTENT_CONTRACT.PRIVATE_OVERLAY_STRING_PATHS
        self.assertEqual(
            paths,
            (
                *package_submission.REPORT_CONTENT_CONTRACT.STRUCTURED_PRIVATE_OVERLAY_STRING_PATHS,
                *package_submission.REPORT_CONTENT_CONTRACT.OWNER_FREE_TEXT_OVERLAY_STRING_PATHS,
            ),
        )
        self.assertEqual(
            {
                ("environment", 2, "text"),
                ("features", 4, "text"),
                ("features", 6, "text"),
                ("other", 4, "text"),
                ("other", 5, "text"),
                ("other", 8, "text"),
            },
            set(
                package_submission.REPORT_CONTENT_CONTRACT.OWNER_FREE_TEXT_OVERLAY_STRING_PATHS
            ),
        )
        for path in paths:
            content = valid_report_content("zero")
            parent = content
            for component in path[:-1]:
                parent = parent[component]
            parent[path[-1]] = f"resolved private overlay {path!r}"
            with self.subTest(path=path):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

    def test_rejects_sensitive_visible_private_overlay_values_without_echoing(self) -> None:
        cases = (
            (("environment", 2, "text"), "macOS /Users/alice/private-project", "LOCAL_PATH"),
            (("features", 4, "text"), "CI token=super-secret-value", "CREDENTIAL_OR_TOKEN"),
            (("features", 4, "text"), "AWS_SECRET_ACCESS_KEY=super-secret-value", "CREDENTIAL_OR_TOKEN"),
            (("features", 4, "text"), "access AKIAIOSFODNN7EXAMPLE", "CREDENTIAL_OR_TOKEN"),
            (
                ("features", 4, "text"),
                "JWT eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiIxMjM0NTY3ODkwIn0.SflKxwRJSMeKKF2QT4fwpMeJf36POk6yJV_adQssw5c",
                "CREDENTIAL_OR_TOKEN",
            ),
            (("features", 6, "text"), "jdbc:mysql://localhost:3306/private", "JDBC_URL"),
            (("other", 4, "text"), "contact alice@example.com", "EMAIL_ADDRESS"),
            (("other", 4, "text"), "연락 개발자@예시.한국", "EMAIL_ADDRESS"),
            (("other", 8, "text"), "SELECT secret FROM customer", "RAW_SQL"),
            (("other", 8, "text"), "SEL/**/ECT secret FR/**/OM customer", "RAW_SQL"),
            (("other", 8, "text"), "SELECT/**/secret/**/FROM customer", "RAW_SQL"),
            (("other", 8, "text"), "INSERT/**/INTO customer VALUES (1)", "RAW_SQL"),
            (("other", 8, "text"), "UPDATE/**/customer/**/SET status = 1", "RAW_SQL"),
            (("other", 8, "text"), "alice@\nexample.com", "EMAIL_ADDRESS"),
            (("other", 8, "text"), "pass\tword=hunter2", "CREDENTIAL_OR_TOKEN"),
            (("other", 8, "text"), "jd\r\nbc:mysql://db", "JDBC_URL"),
            (
                ("other", 8, "text"),
                "SEL-- comment\nECT secret FR-- comment\nOM customer",
                "RAW_SQL",
            ),
            (
                ("other", 8, "text"),
                "SEL# comment\nECT secret FR# comment\nOM customer",
                "RAW_SQL",
            ),
            (
                ("other", 8, "text"),
                "SELECT-- comment\nsecret-- comment\nFROM customer",
                "RAW_SQL",
            ),
            (("other", 8, "text"), "database is 10.0.0.8:3306", "PRIVATE_TOPOLOGY"),
            (("other", 8, "text"), "database is [::1]:3306", "PRIVATE_TOPOLOGY"),
            (("other", 8, "text"), "phone 010-1234-5678", "KOREAN_PHONE_NUMBER"),
            (("other", 8, "text"), "phone 010 1234 5678", "KOREAN_PHONE_NUMBER"),
            (("other", 8, "text"), "phone 010--1234---5678", "KOREAN_PHONE_NUMBER"),
            (("other", 8, "text"), "identity 020506-3123456", "KOREAN_RESIDENT_NUMBER"),
            (("other", 8, "text"), "-----BEGIN PRIVATE KEY-----", "PRIVATE_KEY"),
            (("other", 8, "text"), "device 00:1A:2B:3C:4D:5E", "MAC_ADDRESS"),
            (("other", 8, "text"), "workspace /var/folders/ab/private", "LOCAL_PATH"),
            (("other", 8, "text"), "workspace /private/tmp/private", "LOCAL_PATH"),
            (("other", 8, "text"), "workspace /workspace/private", "LOCAL_PATH"),
            (("other", 8, "text"), "workspace /workspace", "LOCAL_PATH"),
            (("other", 8, "text"), "workspace ~/private", "LOCAL_PATH"),
            (("other", 8, "text"), "workspace ~alice/private", "LOCAL_PATH"),
            (("other", 8, "text"), "workspace ~alice", "LOCAL_PATH"),
        )
        for path, value, category in cases:
            content = valid_report_content("zero")
            parent = content
            for component in path[:-1]:
                parent = parent[component]
            parent[path[-1]] = value
            with self.subTest(path=path, category=category):
                with self.assertRaisesRegex(
                    package_submission.GateError, f"privacy-safe.*{category}"
                ) as caught:
                    package_submission.validate_and_materialize_report_content(
                        content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                    )
                self.assertNotIn(value, str(caught.exception))

    def test_report_builder_cli_rejects_privacy_bypasses_without_echoing(self) -> None:
        cases = (
            ("PRIVATECANARY alice@\nexample.com", "EMAIL_ADDRESS"),
            ("PRIVATECANARY pass\tword=hunter2", "CREDENTIAL_OR_TOKEN"),
            ("PRIVATECANARY phone 010--1234---5678", "KOREAN_PHONE_NUMBER"),
            ("PRIVATECANARY workspace /workspace", "LOCAL_PATH"),
            (
                "PRIVATECANARY SEL-- comment\nECT secret FR-- comment\nOM customer",
                "RAW_SQL",
            ),
            ("PRIVATECANARY Cafe\u0301", "COMBINING_OR_ENCLOSING_MARK"),
        )
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            template = root / "template.docx"
            template.write_bytes(b"not reached")
            for index, (value, category) in enumerate(cases):
                content = valid_report_content("zero")
                content["other"][8]["text"] = value
                content_path = root / f"content-{index}.json"
                content_path.write_text(json.dumps(content), encoding="utf-8")
                process = subprocess.run(
                    [
                        sys.executable,
                        str(REPORT_BUILDER_SCRIPT),
                        "--template",
                        str(template),
                        "--content",
                        str(content_path),
                        "--output",
                        str(root / f"report-{index}.docx"),
                        "--strict-final",
                    ],
                    cwd=REPOSITORY_ROOT,
                    check=False,
                    capture_output=True,
                    text=True,
                )
                with self.subTest(category=category):
                    self.assertNotEqual(0, process.returncode)
                    self.assertIn(category, process.stderr)
                    self.assertNotIn(value, process.stderr)
                    self.assertNotIn("PRIVATECANARY", process.stderr)

    def test_rejects_unicode_obfuscation_and_disallowed_controls(self) -> None:
        cases = (
            ("alice＠example．com", "EMAIL_ADDRESS"),
            ("ＳＥＬＥＣＴ secret ＦＲＯＭ customer", "RAW_SQL"),
            ("ｐａｓｓｗｏｒｄ＝hunter2", "CREDENTIAL_OR_TOKEN"),
            ("ｊｄｂｃ：ｍｙｓｑｌ：／／10.0.0.8/private", "JDBC_URL"),
            ("／Ｕｓｅｒｓ／alice／private", "LOCAL_PATH"),
            (f"password{chr(0xFE0F)}=hunter2", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"pass{chr(0x034F)}word=hunter2", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0x200B)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0x2060)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0x115F)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0x1160)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0x3164)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0xFFA0)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0xE0100)}text", "DEFAULT_IGNORABLE_OR_FILLER"),
            (f"safe{chr(0x0338)}text", "COMBINING_OR_ENCLOSING_MARK"),
            (f"safe{chr(0x0488)}text", "COMBINING_OR_ENCLOSING_MARK"),
            (f"safe{chr(0xD800)}text", "CONTROL_OR_FORMAT_CHARACTER"),
            (f"safe{chr(0xE000)}text", "CONTROL_OR_FORMAT_CHARACTER"),
            (f"safe{chr(0x0378)}text", "CONTROL_OR_FORMAT_CHARACTER"),
            (f"safe{chr(0x000B)}text", "CONTROL_OR_FORMAT_CHARACTER"),
            ("Cafe\u0301", "COMBINING_OR_ENCLOSING_MARK"),
        )
        for value, category in cases:
            content = valid_report_content("zero")
            content["other"][8]["text"] = value
            with self.subTest(category=category):
                with self.assertRaisesRegex(
                    package_submission.GateError, f"privacy-safe.*{category}"
                ) as caught:
                    package_submission.validate_and_materialize_report_content(
                        content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                    )
                self.assertNotIn(value, str(caught.exception))

    def test_rejects_owner_marker_before_placeholder_error_can_echo_it(self) -> None:
        value = "[[password is super-secret-value]]"
        content = valid_report_content("zero")
        content["other"][8]["text"] = value

        with self.assertRaisesRegex(
            package_submission.GateError, "privacy-safe.*UNRESOLVED_MARKER"
        ) as caught:
            package_submission.validate_and_materialize_report_content(
                content, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )
        self.assertNotIn(value, str(caught.exception))
        self.assertNotIn("super-secret-value", str(caught.exception))

    def test_placeholder_errors_never_echo_values_in_memory_visible_or_cli(self) -> None:
        marker = "[[AWS_SECRET_ACCESS_KEY=canary-secret-value]]"
        with self.assertRaisesRegex(package_submission.GateError, r"count=1") as caught:
            package_submission.reject_placeholders({"owner": marker}, "manifest")
        self.assertNotIn(marker, str(caught.exception))
        self.assertNotIn("canary-secret-value", str(caught.exception))

        valid_visible = (
            "개발 보조 AI "
            + package_submission.NO_RUNTIME_AI_DISCLOSURE
            + " SBOM(소프트웨어 자재명세서)"
        )
        for label, docx_text, pdf_text in (
            ("DOCX", marker, valid_visible),
            ("PDF", valid_visible, marker),
        ):
            with self.subTest(label=label), self.assertRaisesRegex(
                package_submission.GateError, rf"{label}.*count=1"
            ) as caught:
                package_submission.validate_report_text_contract(docx_text, pdf_text)
            self.assertNotIn(marker, str(caught.exception))
            self.assertNotIn("canary-secret-value", str(caught.exception))

        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            manifest = root / "manifest.json"
            manifest_value = valid_manifest()
            manifest_value["submission_identity"]["team_name"] = marker
            manifest.write_text(json.dumps(manifest_value), encoding="utf-8")
            dummy = root / "input.bin"
            dummy.write_bytes(b"x")
            evidence = root / "evidence"
            evidence.mkdir()
            process = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--manifest",
                    str(manifest),
                    "--template",
                    str(dummy),
                    "--content",
                    str(dummy),
                    "--report-pdf",
                    str(dummy),
                    "--video-file",
                    str(dummy),
                    "--release-evidence-dir",
                    str(evidence),
                    "--release-evidence-artifact",
                    str(dummy),
                    "--output",
                    str(root / "output"),
                    "--repository-root",
                    str(REPOSITORY_ROOT),
                ],
                cwd=REPOSITORY_ROOT,
                check=False,
                capture_output=True,
                text=True,
            )
            self.assertEqual(1, process.returncode)
            self.assertIn("count=1", process.stderr)
            self.assertNotIn(marker, process.stderr)
            self.assertNotIn("canary-secret-value", process.stderr)

    def test_report_builder_placeholder_errors_never_echo_values_direct_or_cli(
        self,
    ) -> None:
        marker = "[[AWS_SECRET_ACCESS_KEY=builder-canary-secret]]"
        with self.assertRaisesRegex(ValueError, r"count=1") as caught:
            build_official_report.validate_submission_gates(
                {"metadata": {"team_name": marker}}, strict=True
            )
        self.assertNotIn(marker, str(caught.exception))
        self.assertNotIn("builder-canary-secret", str(caught.exception))

        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            content = valid_report_content("zero")
            content["metadata"]["team_name"] = marker
            content_path = root / "report-content.json"
            content_path.write_text(json.dumps(content), encoding="utf-8")
            dummy_template = root / "template.docx"
            dummy_template.write_bytes(b"not reached")
            process = subprocess.run(
                [
                    sys.executable,
                    str(REPORT_BUILDER_SCRIPT),
                    "--template",
                    str(dummy_template),
                    "--content",
                    str(content_path),
                    "--output",
                    str(root / "report.docx"),
                    "--strict-final",
                ],
                cwd=REPOSITORY_ROOT,
                check=False,
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(0, process.returncode)
            self.assertIn("count=1", process.stderr)
            self.assertNotIn(marker, process.stderr)
            self.assertNotIn("builder-canary-secret", process.stderr)

    def test_report_builder_strict_final_rejects_normalized_evidence_ids(self) -> None:
        markers = ("E09", "e09", "Ｅ０９", "Ｅ09", "E０9", "E\u200b09")
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            template = root / "template.docx"
            template.write_bytes(b"not reached")
            for index, marker in enumerate(markers):
                direct = {"metadata": {"team_name": marker}}
                with self.subTest(marker=marker, path="direct"), self.assertRaisesRegex(
                    ValueError, r"reader-facing report content contains audit evidence IDs \(count=1\)"
                ) as caught:
                    build_official_report.validate_submission_gates(direct, strict=True)
                self.assertNotIn(marker, str(caught.exception))

                content = valid_report_content("zero")
                content["environment"][2]["text"] = marker
                content_path = root / f"content-{index}.json"
                content_path.write_text(
                    json.dumps(content, ensure_ascii=False), encoding="utf-8"
                )
                process = subprocess.run(
                    [
                        sys.executable,
                        str(REPORT_BUILDER_SCRIPT),
                        "--template",
                        str(template),
                        "--content",
                        str(content_path),
                        "--output",
                        str(root / f"report-{index}.docx"),
                        "--strict-final",
                    ],
                    cwd=REPOSITORY_ROOT,
                    check=False,
                    capture_output=True,
                    text=True,
                )
                with self.subTest(marker=marker, path="cli"):
                    self.assertNotEqual(0, process.returncode)
                    expected_error = (
                        "DEFAULT_IGNORABLE_OR_FILLER"
                        if "\u200b" in marker
                        else "audit evidence IDs (count=1)"
                    )
                    self.assertIn(expected_error, process.stderr)
                    self.assertNotIn(marker, process.stderr)

    def test_placeholder_fragments_are_counted_without_echo_in_all_report_gates(
        self,
    ) -> None:
        valid_visible = (
            "개발 보조 AI "
            + package_submission.NO_RUNTIME_AI_DISCLOSURE
            + " SBOM(소프트웨어 자재명세서)"
        )
        for marker in (
            "[[PRIVATECANARYSECRETXYZ",
            "PRIVATECANARYSECRETXYZ]]",
        ):
            with self.subTest(marker=marker), self.assertRaisesRegex(
                package_submission.GateError, r"count=1"
            ) as caught:
                package_submission.reject_placeholders({"owner": marker}, "manifest")
            self.assertNotIn(marker, str(caught.exception))
            self.assertNotIn("PRIVATECANARYSECRETXYZ", str(caught.exception))

            with self.subTest(marker=marker), self.assertRaisesRegex(
                ValueError, r"count=1"
            ) as caught:
                build_official_report.validate_submission_gates(
                    {"metadata": {"team_name": marker}}, strict=True
                )
            self.assertNotIn(marker, str(caught.exception))
            self.assertNotIn("PRIVATECANARYSECRETXYZ", str(caught.exception))

            for label, docx_text, pdf_text in (
                ("DOCX", marker, valid_visible),
                ("PDF", valid_visible, marker),
            ):
                with self.subTest(marker=marker, label=label), self.assertRaisesRegex(
                    package_submission.GateError, rf"{label}.*count=1"
                ) as caught:
                    package_submission.validate_report_text_contract(
                        docx_text, pdf_text
                    )
                self.assertNotIn(marker, str(caught.exception))
                self.assertNotIn("PRIVATECANARYSECRETXYZ", str(caught.exception))

            with tempfile.TemporaryDirectory() as raw:
                root = Path(raw)
                manifest_path = root / "manifest.json"
                manifest = valid_manifest()
                manifest["submission_identity"]["team_name"] = marker
                manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
                content_path = root / "report-content.json"
                content = valid_report_content("zero")
                content["metadata"]["team_name"] = marker
                content_path.write_text(json.dumps(content), encoding="utf-8")
                dummy = root / "input.bin"
                dummy.write_bytes(b"x")
                evidence = root / "evidence"
                evidence.mkdir()
                package_process = subprocess.run(
                    [
                        sys.executable,
                        str(SCRIPT),
                        "--manifest",
                        str(manifest_path),
                        "--template",
                        str(dummy),
                        "--content",
                        str(dummy),
                        "--report-pdf",
                        str(dummy),
                        "--video-file",
                        str(dummy),
                        "--release-evidence-dir",
                        str(evidence),
                        "--release-evidence-artifact",
                        str(dummy),
                        "--output",
                        str(root / "package-output"),
                        "--repository-root",
                        str(REPOSITORY_ROOT),
                    ],
                    cwd=REPOSITORY_ROOT,
                    check=False,
                    capture_output=True,
                    text=True,
                )
                self.assertNotEqual(0, package_process.returncode)
                self.assertIn("count=1", package_process.stderr)
                self.assertNotIn(marker, package_process.stderr)
                self.assertNotIn("PRIVATECANARYSECRETXYZ", package_process.stderr)

                template = root / "template.docx"
                template.write_bytes(b"not reached")
                builder_process = subprocess.run(
                    [
                        sys.executable,
                        str(REPORT_BUILDER_SCRIPT),
                        "--template",
                        str(template),
                        "--content",
                        str(content_path),
                        "--output",
                        str(root / "report.docx"),
                        "--strict-final",
                    ],
                    cwd=REPOSITORY_ROOT,
                    check=False,
                    capture_output=True,
                    text=True,
                )
                self.assertNotEqual(0, builder_process.returncode)
                self.assertIn("count=1", builder_process.stderr)
                self.assertNotIn(marker, builder_process.stderr)
                self.assertNotIn("PRIVATECANARYSECRETXYZ", builder_process.stderr)

    def test_exact_key_errors_and_clis_report_counts_without_key_values(self) -> None:
        canary = "PRIVATE_CANARY_UNEXPECTED_KEY"
        direct_cases = (
            lambda: package_submission.REPORT_CONTENT_CONTRACT._require_exact_keys(
                {canary: True}, {"expected"}, "external evidence"
            ),
            lambda: package_submission.require_exact_keys(
                {canary: True}, {"expected"}, "manifest"
            ),
        )
        for call in direct_cases:
            with self.subTest(call=call), self.assertRaisesRegex(
                (ValueError, package_submission.GateError),
                r"missing_count=1, unexpected_count=1",
            ) as caught:
                call()
            self.assertNotIn(canary, str(caught.exception))

        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            content = valid_report_content("zero")
            content[canary] = True
            content_path = root / "report-content.json"
            content_path.write_text(json.dumps(content), encoding="utf-8")
            with self.assertRaisesRegex(
                ValueError, r"missing_count=0, unexpected_count=1"
            ) as caught:
                build_official_report.load_content(content_path, strict=False)
            self.assertNotIn(canary, str(caught.exception))

            array_path = root / "array.json"
            array_path.write_text(json.dumps([canary]), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "must be an object") as caught:
                build_official_report.load_content(array_path, strict=False)
            self.assertNotIn(canary, str(caught.exception))

            template = root / "template.docx"
            template.write_bytes(b"not reached")
            builder_process = subprocess.run(
                [
                    sys.executable,
                    str(REPORT_BUILDER_SCRIPT),
                    "--template",
                    str(template),
                    "--content",
                    str(content_path),
                    "--output",
                    str(root / "report.docx"),
                ],
                cwd=REPOSITORY_ROOT,
                check=False,
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(0, builder_process.returncode)
            self.assertIn("unexpected_count=1", builder_process.stderr)
            self.assertNotIn(canary, builder_process.stderr)

            manifest = valid_manifest()
            manifest[canary] = True
            manifest_path = root / "manifest.json"
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            dummy = root / "input.bin"
            dummy.write_bytes(b"x")
            evidence = root / "evidence"
            evidence.mkdir()
            package_process = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--manifest",
                    str(manifest_path),
                    "--template",
                    str(dummy),
                    "--content",
                    str(dummy),
                    "--report-pdf",
                    str(dummy),
                    "--video-file",
                    str(dummy),
                    "--release-evidence-dir",
                    str(evidence),
                    "--release-evidence-artifact",
                    str(dummy),
                    "--output",
                    str(root / "package-output"),
                    "--repository-root",
                    str(REPOSITORY_ROOT),
                ],
                cwd=REPOSITORY_ROOT,
                check=False,
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(0, package_process.returncode)
            self.assertIn("unexpected_count=1", package_process.stderr)
            self.assertNotIn(canary, package_process.stderr)

    def test_visible_content_mismatch_errors_never_echo_owner_values(self) -> None:
        canary = "PRIVATECANARYSECRETXYZ"
        content = valid_report_content("zero")
        cases = (
            ("missing", "anchor", "anchor", ["anchor", canary], []),
            (
                "value-order",
                "anchor" + canary,
                "anchor" + canary[::-1],
                ["anchor", canary],
                [],
            ),
            (
                "row-order",
                "anchor" + canary,
                "anchor" + canary[::-1],
                ["anchor"],
                [canary],
            ),
        )
        for label, docx_text, pdf_text, values, rows in cases:
            with self.subTest(label=label), patch.object(
                package_submission, "visible_report_values", return_value=values
            ), patch.object(
                package_submission, "visible_report_block_rows", return_value=rows
            ), patch.object(
                package_submission, "external_evidence_summary", return_value="anchor"
            ), self.assertRaisesRegex(package_submission.GateError, r"count=1") as caught:
                package_submission.validate_report_visible_content(
                    docx_text, pdf_text, content
                )
            self.assertNotIn(canary, str(caught.exception))

    def test_accepts_privacy_safe_owner_and_public_link_overlays(self) -> None:
        content = valid_report_content("zero")
        content["environment"][2]["text"] = "Apple silicon, macOS, JDK 17, Docker Desktop"
        content["features"][4]["text"] = (
            "52 tests passed: https://github.com/example-owner/routecontract/actions/runs/123456"
        )
        content["features"][6]["text"] = (
            "Release: https://github.com/example-owner/routecontract/releases/tag/v0.1.0"
        )
        content["other"][4]["text"] = (
            "JTS is non-bundled; the upstream metadata gap remains disclosed."
        )
        content["other"][5]["text"] = (
            "Revision and public evidence links were checked against the final manifest."
        )
        content["other"][8]["text"] = (
            "실패 경계를 좁히고 재현 fixture와 검증 계약을 직접 설명할 수 있게 됐다. "
            "AWS SDK와 JWT 검토 용어, 공개 IPv6 https://[2606:4700:4700::1111]/dns-query, "
            "https://github.com/apache/shardingsphere/issues/38456, "
            "org.apache.shardingsphere:shardingsphere-jdbc:5.5.3, "
            "v0.1.0 / Java 17 / MySQL 8.4.11도 값 없이 기록했다."
        )
        package_submission.validate_and_materialize_report_content(
            content, valid_manifest(), current_utc=TEST_CURRENT_UTC
        )

    def test_rejects_rc_result_collapsed_into_final_stable_branch(self) -> None:
        content = valid_report_content("final_stable")
        content["external_evidence"]["tested_tag"] = "v0.1.0-rc1"

        with self.assertRaisesRegex(
            package_submission.GateError, "distinct reviewed stable form and protocol"
        ):
            package_submission.validate_and_materialize_report_content(
                content, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )

    def test_rejects_rc_only_branch_without_an_exact_rc(self) -> None:
        content = valid_report_content("rc_only")
        content["external_evidence"]["tested_tag"] = "v0.1.0"

        with self.assertRaisesRegex(package_submission.GateError, "exact vMAJOR"):
            package_submission.validate_and_materialize_report_content(
                content, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )

    def test_rejects_rc_from_a_different_stable_version(self) -> None:
        content = valid_report_content("rc_only")
        content["external_evidence"]["tested_tag"] = "v0.2.0-rc1"

        with self.assertRaisesRegex(package_submission.GateError, "RC of the exact final"):
            package_submission.validate_and_materialize_report_content(
                content, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )

    def test_rejects_unanchored_positive_result_count(self) -> None:
        for branch in ("rc_only",):
            content = valid_report_content(branch)
            content["external_evidence"]["qualified_result_count"] = 2
            with self.subTest(branch=branch), self.assertRaisesRegex(
                package_submission.GateError, "exactly one qualified"
            ):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

    def test_rejects_zero_branch_with_a_result_or_positive_count(self) -> None:
        for mutation in (
            {"qualified_result_count": 1},
            {"result_issue_url": "https://github.com/example-owner/routecontract/issues/42"},
        ):
            content = valid_report_content("zero")
            content["external_evidence"].update(mutation)
            with self.subTest(mutation=mutation), self.assertRaisesRegex(
                package_submission.GateError, "count=0 and result_issue_url=null"
            ):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

    def test_rejects_protocol_issue_substituted_for_participant_result(self) -> None:
        content = valid_report_content("rc_only")
        content["external_evidence"]["result_issue_url"] = content[
            "external_evidence"
        ]["protocol_issue_url"]

        with self.assertRaisesRegex(package_submission.GateError, "cannot substitute"):
            package_submission.validate_and_materialize_report_content(
                content, valid_manifest(), current_utc=TEST_CURRENT_UTC
            )

    def test_rejects_mismatched_final_tag_and_manual_summary_text(self) -> None:
        mismatched = valid_report_content("rc_only")
        mismatched["external_evidence"]["final_stable_tag"] = "v0.1.1"
        manual = valid_report_content("rc_only")
        self.summary(manual)
        next(item for item in manual["other"] if item["lead"] == "외부 검증")[
            "text"
        ] = "RC 결과를 최종 결과로 직접 입력"

        for content, message in (
            (mismatched, "does not match the final package tag"),
            (manual, "must remain the generated structured-evidence marker"),
        ):
            with self.subTest(message=message), self.assertRaisesRegex(
                package_submission.GateError, message
            ):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

    def test_rejects_unresolved_or_partial_structured_branch(self) -> None:
        source = json.loads(
            (REPOSITORY_ROOT / "submission" / "report-content.ko.json").read_text(
                encoding="utf-8"
            )
        )
        partial = deepcopy(source)
        partial["external_evidence"]["branch"] = "rc_only"

        for content in (source, partial):
            with self.subTest(branch=content["external_evidence"]["branch"]), self.assertRaisesRegex(
                package_submission.GateError, "unresolved structured fields"
            ):
                package_submission.validate_and_materialize_report_content(
                    content, valid_manifest(), current_utc=TEST_CURRENT_UTC
                )

class PublicJsonTransportTest(unittest.TestCase):
    URL = "https://api.github.com/repos/example-owner/routecontract"

    @staticmethod
    def nested_json(depth: int, shape: str) -> bytes:
        value = b"0"
        for index in range(depth):
            container = shape if shape != "mixed" else ("array" if index % 2 else "object")
            value = b"[" + value + b"]" if container == "array" else b'{"a":' + value + b"}"
        return value

    @staticmethod
    def node_json(node_count: int, shape: str) -> bytes:
        if node_count < 1:
            raise AssertionError("node_count must include the root")
        if shape == "array":
            value: object = [0] * (node_count - 1)
        elif shape == "object":
            value = {f"k{index}": 0 for index in range(node_count - 1)}
        elif shape == "mixed":
            if node_count < 3:
                raise AssertionError("mixed JSON needs a list, object, and scalar")
            value = [{"nested": 0}, *([0] * (node_count - 3))]
        else:
            raise AssertionError(f"unknown JSON shape: {shape}")
        return json.dumps(value, separators=(",", ":")).encode("utf-8")

    @staticmethod
    def activation_record_payload(raw: bytes) -> tuple[dict[str, object], str, str]:
        expected_url = (
            "https://github.com/example-owner/routecontract/blob/"
            + "a" * 40
            + "/docs/evidence/independent-rc-activation-v0.1.0-rc2.json"
        )
        expected_path = (
            "docs/evidence/independent-rc-activation-v0.1.0-rc2.json"
        )
        return (
            {
                "type": "file",
                "path": expected_path,
                "sha": "b" * 40,
                "html_url": expected_url,
                "encoding": "base64",
                "content": base64.b64encode(raw).decode("ascii"),
                "size": len(raw),
            },
            expected_url,
            expected_path,
        )

    def test_object_list_and_page_helpers_reject_duplicate_keys_at_any_depth(self) -> None:
        cases = (
            (
                "object-top-level",
                package_submission.request_json,
                b'{"CANARY_SECRET":1,"CANARY_SECRET":2}',
                "bytes",
            ),
            (
                "object-nested",
                package_submission.request_json,
                b'{"outer":{"CANARY_SECRET":1,"CANARY_SECRET":2}}',
                "bytes",
            ),
            (
                "list-nested",
                package_submission.request_json_list,
                b'[{"outer":{"CANARY_SECRET":1,"CANARY_SECRET":2}}]',
                "bytes",
            ),
            (
                "page-nested",
                package_submission.request_json_list_page,
                b'[{"outer":{"CANARY_SECRET":1,"CANARY_SECRET":2}}]',
                "page",
            ),
        )
        for label, helper, raw, transport in cases:
            patcher = (
                patch.object(
                    package_submission,
                    "request_bytes_with_headers",
                    return_value=(raw, {}),
                )
                if transport == "page"
                else patch.object(package_submission, "request_bytes", return_value=raw)
            )
            with self.subTest(label=label), patcher, self.assertRaises(
                package_submission.GateError
            ) as caught:
                helper(self.URL)
            self.assertNotIn("CANARY_SECRET", str(caught.exception))

    def test_public_json_helpers_reject_nonfinite_and_malformed_utf8_generically(self) -> None:
        cases = (
            (package_submission.request_json, b'{"id":NaN}', "bytes"),
            (package_submission.request_json_list, b'[{"id":Infinity}]', "bytes"),
            (package_submission.request_json_list_page, b'[{"id":-Infinity}]', "page"),
            (package_submission.request_json, b'{"id":1e999}', "bytes"),
            (package_submission.request_json, b'{"id":"\xff"}', "bytes"),
            (
                package_submission.request_json,
                self.nested_json(
                    package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
                    + 1,
                    "object",
                ),
                "bytes",
            ),
        )
        for helper, raw, transport in cases:
            patcher = (
                patch.object(
                    package_submission,
                    "request_bytes_with_headers",
                    return_value=(raw, {}),
                )
                if transport == "page"
                else patch.object(package_submission, "request_bytes", return_value=raw)
            )
            with self.subTest(raw=raw), patcher, self.assertRaisesRegex(
                package_submission.GateError, "strict JSON"
            ):
                helper(self.URL)

    def test_public_json_helpers_accept_strict_object_list_and_page(self) -> None:
        self.assertEqual(8_000_000, package_submission.MAX_PUBLIC_JSON_RESPONSE_BYTES)
        self.assertEqual(
            package_submission.MAX_PUBLIC_JSON_RESPONSE_BYTES,
            package_submission._decode_public_contents_file.__kwdefaults__["maximum_size"],
        )
        with patch.object(
            package_submission, "request_bytes", return_value=b'{"id":1}'
        ):
            self.assertEqual({"id": 1}, package_submission.request_json(self.URL))
        with patch.object(
            package_submission, "request_bytes", return_value=b'[{"id":1}]'
        ):
            self.assertEqual([{"id": 1}], package_submission.request_json_list(self.URL))
        with patch.object(
            package_submission,
            "request_bytes_with_headers",
            return_value=(b'[{"id":1}]', {"link": ["next"]}),
        ):
            self.assertEqual(
                ([{"id": 1}], ["next"]),
                package_submission.request_json_list_page(self.URL),
            )

    def test_shared_strict_decoder_enforces_iterative_depth_and_node_budgets(self) -> None:
        contract = package_submission.REPORT_CONTENT_CONTRACT
        maximum_depth = contract.STRICT_JSON_MAX_CONTAINER_DEPTH
        maximum_nodes = contract.STRICT_JSON_MAX_NODES

        for shape in ("array", "object", "mixed"):
            accepted = self.nested_json(maximum_depth, shape)
            rejected = self.nested_json(maximum_depth + 1, shape).replace(
                b"0", b'"CANARY_DEPTH_SECRET"', 1
            )
            with self.subTest(shape=shape, boundary="max"):
                contract.decode_strict_json(accepted)
            with self.subTest(shape=shape, boundary="max-plus-one"), self.assertRaises(
                contract.StrictJsonError
            ) as caught:
                contract.decode_strict_json(rejected)
            self.assertIsNone(caught.exception.__cause__)
            self.assertNotIn("CANARY_DEPTH_SECRET", str(caught.exception))

        accepted_elements = maximum_nodes - 1
        accepted = b"[" + b"0," * (accepted_elements - 1) + b"0]"
        rejected = accepted[:-1] + b",0]"
        self.assertEqual(accepted_elements, len(contract.decode_strict_json(accepted)))
        with self.assertRaises(contract.StrictJsonError) as caught:
            contract.decode_strict_json(rejected)
        self.assertIsNone(caught.exception.__cause__)

        for depth in (1_050, 1_200):
            with self.subTest(parser_depth=depth), self.assertRaises(
                contract.StrictJsonError
            ) as caught:
                contract.decode_strict_json(self.nested_json(depth, "mixed"))
            self.assertIsNone(caught.exception.__cause__)

    def test_shared_strict_decoder_enforces_explicit_byte_budget(self) -> None:
        contract = package_submission.REPORT_CONTENT_CONTRACT
        self.assertEqual({}, contract.decode_strict_json(b"{}", maximum_bytes=2))
        with self.assertRaises(contract.StrictJsonError) as caught:
            contract.decode_strict_json(b" {}", maximum_bytes=2)
        self.assertIsNone(caught.exception.__cause__)
        for invalid_limit in (-1, True, 1.5):
            with self.subTest(invalid_limit=invalid_limit), self.assertRaises(
                contract.StrictJsonError
            ) as caught:
                contract.decode_strict_json(b"{}", maximum_bytes=invalid_limit)
            self.assertIsNone(caught.exception.__cause__)

    def test_shared_and_activation_decoders_have_exact_integer_and_tree_parity(
        self,
    ) -> None:
        contract = package_submission.REPORT_CONTENT_CONTRACT
        activation = package_submission.RC_ACTIVATION_RECORD_VALIDATOR
        self.assertEqual(64, contract.STRICT_JSON_MAX_CONTAINER_DEPTH)
        self.assertEqual(100_000, contract.STRICT_JSON_MAX_NODES)
        self.assertEqual(1_000, contract.STRICT_JSON_MAX_INTEGER_DIGITS)
        self.assertEqual(
            contract.STRICT_JSON_MAX_CONTAINER_DEPTH,
            activation.MAX_JSON_NESTING_DEPTH,
        )
        self.assertEqual(contract.STRICT_JSON_MAX_NODES, activation.MAX_JSON_NODE_COUNT)
        self.assertEqual(
            contract.STRICT_JSON_MAX_INTEGER_DIGITS,
            activation.MAX_JSON_INTEGER_DIGITS,
        )
        self.assertEqual(
            package_submission.MAX_PUBLIC_ACTIVATION_RECORD_BYTES,
            activation.MAX_RECORD_BYTES,
        )

        decoders = (contract.decode_strict_json, activation._decode_strict_json)
        accepted_integer = b"-" + b"9" * contract.STRICT_JSON_MAX_INTEGER_DIGITS
        rejected_integer = accepted_integer + b"9"
        for decoder in decoders:
            with self.subTest(decoder=decoder.__module__, boundary="integer-max"):
                self.assertEqual(int(accepted_integer), decoder(accepted_integer))
            with self.subTest(
                decoder=decoder.__module__, boundary="integer-max-plus-one"
            ), self.assertRaises(ValueError) as caught:
                decoder(rejected_integer)
            self.assertIsNone(caught.exception.__cause__)

            encoded = b'{"ok":true}'
            self.assertEqual(
                {"ok": True}, decoder(encoded, maximum_bytes=len(encoded))
            )
            with self.assertRaises(ValueError) as caught:
                decoder(encoded, maximum_bytes=len(encoded) - 1)
            self.assertIsNone(caught.exception.__cause__)

        for shape in ("array", "object", "mixed"):
            accepted = self.node_json(contract.STRICT_JSON_MAX_NODES, shape)
            rejected = self.node_json(contract.STRICT_JSON_MAX_NODES + 1, shape)
            for decoder in decoders:
                with self.subTest(
                    decoder=decoder.__module__, shape=shape, boundary="nodes-max"
                ):
                    decoder(accepted)
                with self.subTest(
                    decoder=decoder.__module__,
                    shape=shape,
                    boundary="nodes-max-plus-one",
                ), self.assertRaises(ValueError) as caught:
                    decoder(rejected)
                self.assertIsNone(caught.exception.__cause__)

    def test_public_activation_record_uses_the_strict_json_decoder(self) -> None:
        valid_raw = b'{"schemaVersion":2}'
        payload, expected_url, expected_path = self.activation_record_payload(valid_raw)
        record, observed_raw, blob_sha = package_submission._decode_activation_record(
            payload, expected_url, expected_path
        )
        self.assertEqual({"schemaVersion": 2}, record)
        self.assertEqual(valid_raw, observed_raw)
        self.assertEqual("b" * 40, blob_sha)

        maximum_depth = (
            package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
        )
        boundary_raw = self.nested_json(maximum_depth, "object")
        payload, expected_url, expected_path = self.activation_record_payload(boundary_raw)
        boundary_record, observed_raw, _ = package_submission._decode_activation_record(
            payload, expected_url, expected_path
        )
        self.assertIsInstance(boundary_record, dict)
        self.assertEqual(boundary_raw, observed_raw)

        canary = "CANARY_PRIVATE_DUPLICATE_KEY"
        deep = b"[" * 10_000 + b"0" + b"]" * 10_000
        post_parse_deep = b'{"a":' * 2_000 + b"0" + b"}" * 2_000
        huge_integer = b'{"value":' + b"9" * 5_000 + b"}"
        malformed = (
            b'{"value":NaN}',
            b'{"value":Infinity}',
            b'{"value":-Infinity}',
            b'{"value":1e999}',
            self.nested_json(maximum_depth + 1, "object").replace(
                b"0", ('"' + canary + '"').encode(), 1
            ),
            self.nested_json(1_050, "object"),
            self.nested_json(1_200, "mixed"),
            deep,
            post_parse_deep,
            huge_integer,
            '{"value":1}'.encode("utf-16"),
            ('{"' + canary + '":1,"' + canary + '":2}').encode(),
        )
        for raw in malformed:
            payload, expected_url, expected_path = self.activation_record_payload(raw)
            with self.subTest(raw_prefix=raw[:24]), self.assertRaisesRegex(
                package_submission.GateError, "strict JSON"
            ) as caught:
                package_submission._decode_activation_record(
                    payload, expected_url, expected_path
                )
            self.assertIsNone(caught.exception.__cause__)
            self.assertNotIn(canary, str(caught.exception))

    def test_public_activation_record_enforces_the_exact_one_mibibyte_raw_limit(
        self,
    ) -> None:
        maximum = package_submission.MAX_PUBLIC_ACTIVATION_RECORD_BYTES
        accepted_raw = b"{}" + b" " * (maximum - 2)
        payload, expected_url, expected_path = self.activation_record_payload(
            accepted_raw
        )
        record, observed_raw, _ = package_submission._decode_activation_record(
            payload, expected_url, expected_path
        )
        self.assertEqual({}, record)
        self.assertEqual(accepted_raw, observed_raw)

        rejected_raw = accepted_raw + b" "
        payload, expected_url, expected_path = self.activation_record_payload(
            rejected_raw
        )
        with self.assertRaisesRegex(
            package_submission.GateError, "bounded public ordinary file"
        ) as caught:
            package_submission._decode_activation_record(
                payload, expected_url, expected_path
            )
        self.assertIsNone(caught.exception.__cause__)

    def test_tracked_public_activation_record_is_accepted_by_the_contents_decoder(
        self,
    ) -> None:
        expected_path = "docs/evidence/independent-rc-activation-v0.1.0-rc2.json"
        raw = (REPOSITORY_ROOT / expected_path).read_bytes()
        expected_url = (
            "https://github.com/ym0506/routecontract/blob/"
            + "a" * 40
            + f"/{expected_path}"
        )
        payload = {
            "type": "file",
            "path": expected_path,
            "sha": "b" * 40,
            "html_url": expected_url,
            "encoding": "base64",
            "content": base64.b64encode(raw).decode("ascii"),
            "size": len(raw),
        }
        record, observed_raw, _ = package_submission._decode_activation_record(
            payload, expected_url, expected_path
        )
        self.assertEqual("v0.1.0-rc2", record["tag"])
        self.assertEqual(raw, observed_raw)

    def test_activation_record_duplicate_key_cli_traceback_does_not_echo_key(self) -> None:
        canary = "CANARY_PRIVATE_DUPLICATE_KEY"
        script = f"""
import base64
from submission.tools import package_submission as module
raw = b'{{\"{canary}\":1,\"{canary}\":2}}'
url = 'https://github.com/example-owner/routecontract/blob/' + 'a' * 40 + '/record.json'
payload = {{'type':'file','path':'record.json','sha':'b'*40,'html_url':url,'encoding':'base64','content':base64.b64encode(raw).decode('ascii'),'size':len(raw)}}
module._decode_activation_record(payload, url, 'record.json')
"""
        process = subprocess.run(
            [sys.executable, "-c", script],
            cwd=REPOSITORY_ROOT,
            check=False,
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(0, process.returncode)
        self.assertIn("strict JSON", process.stderr)
        self.assertNotIn(canary, process.stderr)

    def test_private_and_supply_loaders_share_the_strict_json_contract(self) -> None:
        canary = "CANARY_PRIVATE_DUPLICATE_KEY"
        deep = b"[" * 10_000 + b"0" + b"]" * 10_000
        malformed = (
            ('{"' + canary + '":1,"' + canary + '":2}').encode(),
            b'{"value":NaN}',
            b'{"value":Infinity}',
            b'{"value":-Infinity}',
            b'{"value":1e999}',
            b'{"value":' + b"9" * 5_000 + b"}",
            self.nested_json(
                package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
                + 1,
                "object",
            ),
            deep,
            '{"value":1}'.encode("utf-16"),
            b"\xff",
        )
        with tempfile.TemporaryDirectory() as raw_root:
            root = Path(raw_root)
            path = root / "input.json"
            for raw in malformed:
                path.write_bytes(raw)
                loaders = (
                    (
                        "private-manifest",
                        lambda: package_submission.load_json(path, "package manifest"),
                    ),
                    (
                        "private-report",
                        lambda: package_submission.load_json(path, "report content"),
                    ),
                    (
                        "supply-evidence",
                        lambda: package_submission.load_strict_json(
                            path, "supply-chain evidence"
                        ),
                    ),
                    (
                        "builder-report",
                        lambda: build_official_report.load_content(path, strict=False),
                    ),
                )
                for label, loader in loaders:
                    with self.subTest(label=label, raw_prefix=raw[:24]), self.assertRaises(
                        (package_submission.GateError, ValueError)
                    ) as caught:
                        loader()
                    self.assertIsNone(caught.exception.__cause__)
                    self.assertNotIn(canary, str(caught.exception))

            path.write_bytes(b" " * (1024 * 1024 + 1))
            oversized_loaders = (
                lambda: package_submission.load_json(path, "package manifest"),
                lambda: package_submission.load_strict_json(
                    path, "supply-chain evidence"
                ),
                lambda: build_official_report.load_content(path, strict=False),
            )
            for loader in oversized_loaders:
                with self.subTest(loader=loader), self.assertRaisesRegex(
                    (package_submission.GateError, ValueError), "safety limit"
                ) as caught:
                    loader()
                self.assertIsNone(caught.exception.__cause__)

            path.write_text('{"ok":true}', encoding="utf-8")
            self.assertEqual(
                {"ok": True}, package_submission.load_json(path, "package manifest")
            )
            self.assertEqual(
                {"ok": True},
                package_submission.load_strict_json(path, "supply-chain evidence"),
            )
            report = valid_report_content("zero")
            path.write_text(json.dumps(report), encoding="utf-8")
            observed_report = build_official_report.load_content(path, strict=False)
            self.assertEqual(report["metadata"], observed_report["metadata"])
            self.assertEqual("zero", observed_report["external_evidence"]["branch"])

            path.write_bytes(
                self.nested_json(
                    package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
                    + 1,
                    "object",
                )
            )
            with patch.object(
                package_submission.REPORT_CONTENT_CONTRACT,
                "CANONICAL_REPORT_SOURCE",
                path,
            ), self.assertRaisesRegex(ValueError, "canonical report source") as caught:
                package_submission.REPORT_CONTENT_CONTRACT._canonical_report_source()
            self.assertIsNone(caught.exception.__cause__)

            maximum = (
                package_submission.REPORT_CONTENT_CONTRACT.CANONICAL_REPORT_SOURCE_MAX_BYTES
            )
            accepted_source = b"{}" + b" " * (maximum - 2)
            path.write_bytes(accepted_source)
            with patch.object(
                package_submission.REPORT_CONTENT_CONTRACT,
                "CANONICAL_REPORT_SOURCE",
                path,
            ):
                self.assertEqual(
                    {},
                    package_submission.REPORT_CONTENT_CONTRACT._canonical_report_source(),
                )

            path.write_bytes(accepted_source + b" ")
            with patch.object(
                package_submission.REPORT_CONTENT_CONTRACT,
                "CANONICAL_REPORT_SOURCE",
                path,
            ), self.assertRaisesRegex(ValueError, "canonical report source") as caught:
                package_submission.REPORT_CONTENT_CONTRACT._canonical_report_source()
            self.assertIsNone(caught.exception.__cause__)

    def test_private_and_supply_loader_cli_errors_never_echo_values(self) -> None:
        canary = "CANARY_PRIVATE_DUPLICATE_KEY"
        with tempfile.TemporaryDirectory() as raw_root:
            root = Path(raw_root)
            path = root / "input.json"
            path.write_text(
                '{"' + canary + '":1,"' + canary + '":2}', encoding="utf-8"
            )
            calls = (
                (
                    "package-manifest",
                    "from submission.tools import package_submission as m; "
                    f"m.load_json(__import__('pathlib').Path({str(path)!r}), 'package manifest')",
                ),
                (
                    "package-report",
                    "from submission.tools import package_submission as m; "
                    f"m.load_json(__import__('pathlib').Path({str(path)!r}), 'report content')",
                ),
                (
                    "supply-evidence",
                    "from submission.tools import package_submission as m; "
                    f"m.load_strict_json(__import__('pathlib').Path({str(path)!r}), 'supply-chain evidence')",
                ),
                (
                    "builder-report",
                    "from submission.tools import build_official_report as m; "
                    f"m.load_content(__import__('pathlib').Path({str(path)!r}), strict=False)",
                ),
            )
            for label, command in calls:
                process = subprocess.run(
                    [sys.executable, "-c", command],
                    cwd=REPOSITORY_ROOT,
                    check=False,
                    capture_output=True,
                    text=True,
                )
                with self.subTest(label=label):
                    self.assertNotEqual(0, process.returncode)
                    self.assertNotIn(canary, process.stderr)

            duplicate = '{"' + canary + '":1,"' + canary + '":2}'
            manifest_path = root / "manifest.json"
            report_path = root / "report.json"
            dummy = root / "input.bin"
            dummy.write_bytes(b"x")
            evidence = root / "evidence"
            evidence.mkdir()
            package_arguments = [
                sys.executable,
                str(SCRIPT),
                "--manifest",
                str(manifest_path),
                "--template",
                str(dummy),
                "--content",
                str(report_path),
                "--report-pdf",
                str(dummy),
                "--video-file",
                str(dummy),
                "--release-evidence-dir",
                str(evidence),
                "--release-evidence-artifact",
                str(dummy),
                "--output",
                str(root / "package-output"),
                "--repository-root",
                str(REPOSITORY_ROOT),
            ]
            for surface in ("manifest", "report"):
                manifest_path.write_text(
                    duplicate if surface == "manifest" else json.dumps(valid_manifest()),
                    encoding="utf-8",
                )
                report_path.write_text(
                    duplicate if surface == "report" else json.dumps(valid_report_content("zero")),
                    encoding="utf-8",
                )
                process = subprocess.run(
                    package_arguments,
                    cwd=REPOSITORY_ROOT,
                    check=False,
                    capture_output=True,
                    text=True,
                )
                with self.subTest(cli="package", surface=surface):
                    self.assertNotEqual(0, process.returncode)
                    self.assertIn("strict UTF-8 JSON", process.stderr)
                    self.assertNotIn(canary, process.stderr)

            report_path.write_text(duplicate, encoding="utf-8")
            builder_process = subprocess.run(
                [
                    sys.executable,
                    str(REPORT_BUILDER_SCRIPT),
                    "--template",
                    str(dummy),
                    "--content",
                    str(report_path),
                    "--output",
                    str(root / "report.docx"),
                ],
                cwd=REPOSITORY_ROOT,
                check=False,
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(0, builder_process.returncode)
            self.assertIn("strict JSON", builder_process.stderr)
            self.assertNotIn(canary, builder_process.stderr)

    def test_graphql_command_failure_and_timeout_do_not_echo_output(self) -> None:
        canary = "CANARY_GRAPHQL_SECRET"
        failures = (
            subprocess.CompletedProcess(
                ["/safe/gh"], 1, stdout=canary, stderr=canary
            ),
            subprocess.TimeoutExpired(["/safe/gh"], 60, output=canary, stderr=canary),
        )
        for failure in failures:
            side_effect = failure if isinstance(failure, BaseException) else None
            return_value = None if side_effect is not None else failure
            with self.subTest(kind=type(failure).__name__), patch.object(
                package_submission,
                "require_safe_github_cli_release_verification",
                return_value="/safe/gh",
            ), patch.object(
                package_submission.subprocess,
                "run",
                return_value=return_value,
                side_effect=side_effect,
            ) as invoked, self.assertRaises(package_submission.GateError) as caught:
                package_submission.request_graphql_issue("owner", "repo", 7)
            self.assertNotIn(canary, str(caught.exception))
            self.assertIsNone(caught.exception.__cause__)
            self.assertEqual(60, invoked.call_args.kwargs["timeout"])

    def test_graphql_decoder_rejects_nonfinite_deep_and_oversized_output(self) -> None:
        malformed = (
            '{"data":{"repository":{"id":1e999}}}',
            self.nested_json(
                package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
                + 1,
                "object",
            ).decode(),
            " " * (package_submission.MAX_JSON_TOOL_OUTPUT_BYTES + 1),
        )
        for output in malformed:
            with self.subTest(length=len(output)), patch.object(
                package_submission,
                "require_safe_github_cli_release_verification",
                return_value="/safe/gh",
            ), patch.object(
                package_submission,
                "run",
                return_value=output,
            ), self.assertRaisesRegex(package_submission.GateError, "invalid JSON") as caught:
                package_submission.request_graphql_issue("owner", "repo", 7)
            self.assertIsNone(caught.exception.__cause__)


class PublicExternalEvidenceTest(unittest.TestCase):
    RECORD_COMMIT = "b" * 40
    TAG_COMMIT = "c" * 40
    RECORD_BLOB = "d" * 40
    RECORD_TREE = "e" * 40
    FORM_BLOB = "f" * 40
    PROTOCOL_BLOB = "7" * 40
    README_BLOB = "6" * 40
    TAG_TREE = "a" * 40
    TAG_OBJECT = "9" * 40
    RUN_ID = 1
    ARTIFACT_ID = 2
    ACTIVATION_PULL_NUMBER = 88
    REPOSITORY_ID = 1300192
    REPOSITORY_NODE = "R_kgDOExample"

    def setUp(self) -> None:
        self.manifest = package_submission.validate_manifest(valid_manifest())
        self.repository_url = self.manifest["project"]["repository_url"]
        self.api_base = "https://api.github.com/repos/example-owner/routecontract"
        self.git_show_patcher = patch.object(
            package_submission, "_git_show_bytes", side_effect=self.tagged_bytes
        )
        self.git_show_patcher.start()
        self.addCleanup(self.git_show_patcher.stop)
        self.graphql_issues_by_number: dict[int, dict] = {}
        self.graphql_mutate = None
        self.graphql_patcher = patch.object(
            package_submission,
            "request_graphql_issue",
            side_effect=self.fake_graphql_issue,
        )
        self.graphql_patcher.start()
        self.addCleanup(self.graphql_patcher.stop)
        self.activation_graphql_mutate = None
        self.activation_graphql_patcher = patch.object(
            package_submission,
            "request_graphql_activation_pull",
            side_effect=self.fake_graphql_activation_pull,
        )
        self.activation_graphql_patcher.start()
        self.addCleanup(self.activation_graphql_patcher.stop)
        self.activation_pull_patcher = patch.object(
            package_submission,
            "request_json_list",
            side_effect=self.fake_activation_pull_requests,
        )
        self.activation_pull_patcher.start()
        self.addCleanup(self.activation_pull_patcher.stop)
        artifact_names = (
            *self.public_assets(),
            *package_submission.RC_ACTIVATION_RECORD_VALIDATOR.WORKFLOW_ONLY_FILES,
        )
        self.artifact_download_patcher = patch.object(
            package_submission,
            "download_and_validate_public_rc_workflow_artifact",
            return_value={name: digest("7") for name in sorted(artifact_names)},
        )
        self.artifact_download = self.artifact_download_patcher.start()
        self.addCleanup(self.artifact_download_patcher.stop)

    @staticmethod
    def tagged_bytes(_commit: str, path: str) -> bytes:
        return (REPOSITORY_ROOT / path).read_bytes()

    def fake_graphql_issue(self, _owner: str, _repository: str, number: int) -> dict:
        rest = deepcopy(self.graphql_issues_by_number[number])
        user = rest["user"]
        issue = {
            "id": rest["node_id"],
            "number": rest["number"],
            "url": rest["html_url"],
            "title": rest["title"],
            "body": rest["body"],
            "createdAt": rest["created_at"],
            "updatedAt": rest["updated_at"],
            "authorAssociation": rest["author_association"],
            "author": {
                "__typename": "User",
                "login": user["login"],
                "id": user["node_id"],
                "databaseId": user["id"],
            },
            "editor": None,
            "lastEditedAt": None,
            "includesCreatedEdit": False,
            "userContentEdits": {
                "totalCount": 0,
                "nodes": [],
                "pageInfo": {"hasNextPage": False},
            },
            "timelineItems": {
                "totalCount": 0,
                "nodes": [],
                "pageInfo": {"hasNextPage": False},
            },
        }
        payload = {
            "data": {
                "repository": {
                    "id": self.REPOSITORY_NODE,
                    "nameWithOwner": "example-owner/routecontract",
                    "issue": issue,
                }
            }
        }
        if self.graphql_mutate is not None:
            self.graphql_mutate(payload)
        return payload

    def fake_graphql_activation_pull(
        self, owner: str, repository: str, number: int
    ) -> dict:
        rest = self.activation_pull_request()
        payload = {
            "data": {
                "repository": {
                    "id": self.REPOSITORY_NODE,
                    "nameWithOwner": f"{owner}/{repository}",
                    "pullRequest": {
                        "id": rest["node_id"],
                        "databaseId": rest["id"],
                        "number": number,
                        "url": rest["html_url"],
                        "state": "MERGED",
                        "merged": True,
                        "mergedAt": rest["merged_at"],
                        "baseRefName": "main",
                        "baseRepository": {
                            "id": self.REPOSITORY_NODE,
                            "nameWithOwner": f"{owner}/{repository}",
                        },
                        "mergeCommit": {"oid": self.RECORD_COMMIT},
                    },
                }
            }
        }
        if self.activation_graphql_mutate is not None:
            self.activation_graphql_mutate(payload)
        return payload

    @staticmethod
    def public_assets(version: str = "0.1.0-rc1") -> list[str]:
        return [
            f"routecontract-{version}-source.zip",
            f"routecontract-shardingsphere-5.5-{version}.jar",
            f"routecontract-shardingsphere-5.5-{version}-sources.jar",
            f"routecontract-shardingsphere-5.5-{version}-javadoc.jar",
            "routecontract-shardingsphere-5.5.pom",
            "routecontract-shardingsphere-5.5-cyclonedx.json",
            "routecontract-shardingsphere-5.5-cyclonedx.xml",
            "routecontract-aggregate-cyclonedx.json",
            "routecontract-aggregate-cyclonedx.xml",
            "supply-chain-evidence.json",
            "test-summary.txt",
            "SHA256SUMS",
        ]

    def payload_digests(self) -> dict[str, str]:
        return {
            name: hashlib.sha256(name.encode("utf-8")).hexdigest()
            for name in self.public_assets()
            if name != package_submission.CHECKSUMS_NAME
        }

    def checksum_bytes(self) -> bytes:
        return "".join(
            f"{checksum}  {name}\n"
            for name, checksum in sorted(self.payload_digests().items())
        ).encode("utf-8")

    def activation_record(self) -> dict:
        form = "independent-rc1-install.yml"
        return {
            "issueFormFilename": form,
            "issueFormPermalink": (
                f"{self.repository_url}/blob/{self.TAG_COMMIT}/"
                f".github/ISSUE_TEMPLATE/{form}"
            ),
            "issueFormUrl": f"{self.repository_url}/issues/new?template={form}",
            "publicAssets": self.public_assets(),
            "releaseEvidence": {
                "artifactDigest": f"sha256:{digest('d')}",
                "artifactFileCount": 17,
                "artifactId": self.ARTIFACT_ID,
                "headSha": self.TAG_COMMIT,
                "runId": self.RUN_ID,
                "runUrl": f"{self.repository_url}/actions/runs/{self.RUN_ID}",
            },
            "releaseImmutability": {"enabled": True, "enforcedByOwner": True},
            "releaseState": {"draft": False, "immutable": True, "prerelease": True},
            "releaseUrl": f"{self.repository_url}/releases/tag/v0.1.0-rc1",
            "repository": self.repository_url,
            "schemaVersion": 2,
            "sha256sumsSha256": hashlib.sha256(self.checksum_bytes()).hexdigest(),
            "tag": "v0.1.0-rc1",
            "tagCommit": self.TAG_COMMIT,
            "taggedProtocolUrl": (
                f"{self.repository_url}/blob/v0.1.0-rc1/docs/independent-install-study.md"
            ),
            "taggedReadmeUrl": f"{self.repository_url}/blob/v0.1.0-rc1/README.md",
        }

    def activation_marker(self) -> str:
        return (
            "ROUTECONTRACT_RC_ACTIVATION_VERIFIED tag=v0.1.0-rc1 "
            f"commit={self.TAG_COMMIT} run={self.RUN_ID} "
            f"artifact={self.ARTIFACT_ID} assets=12"
        )

    def eligible_issue(self, content: dict, number: int = 42) -> dict:
        evidence = content["external_evidence"]
        body_lines = [f"Exact tag tested: {evidence['tested_tag']}"]
        if evidence["branch"] in {"rc_only", "zero"}:
            body_lines.extend(
                (
                    self.activation_marker(),
                    f"ACTIVATION_RECORD_PERMALINK {evidence['activation_record_url']}",
                    "PUBLIC_RECRUITMENT_RECORD_PERMALINK "
                    f"{evidence['recruitment_record_url']}",
                )
            )
        body_lines.extend(
            f"- [x] {label}"
            for label in package_submission.RESULT_ISSUE_REQUIRED_CHECKBOXES
        )
        body_lines.extend(
            (
                "### Task A first outcome — exact-tag Quick Start",
                "UNASSISTED_PASS",
            )
        )
        return {
            "repository_url": self.api_base,
            "html_url": f"{self.repository_url}/issues/{number}",
            "number": number,
            "title": f"[independent-install] clean attempt {number}",
            "node_id": f"I_kwDOExample{number}",
            "user": {
                "id": 10_000 + number,
                "node_id": f"U_kgDOExample{number}",
                "login": f"independent-tester-{number}",
                "type": "User",
            },
            "author_association": "NONE",
            "body": "\n".join(body_lines),
            "labels": [{"name": "evidence"}, {"name": "community"}],
            "created_at": "2026-07-23T00:00:00Z",
            "updated_at": "2026-07-23T00:01:00Z",
        }

    def unrelated_issue(self, number: int) -> dict:
        return {
            "repository_url": self.api_base,
            "html_url": f"{self.repository_url}/issues/{number}",
            "number": number,
            "title": f"ordinary bug {number}",
            "node_id": f"I_kwDOOrdinary{number}",
            "user": {
                "id": 20_000 + number,
                "node_id": f"U_kgDOOrdinary{number}",
                "login": "ordinary-user",
                "type": "User",
            },
            "author_association": "NONE",
            "body": "unrelated report",
            "labels": [{"name": "bug"}],
            "created_at": "2026-07-23T00:00:00Z",
            "updated_at": "2026-07-23T00:01:00Z",
        }

    def recruitment_comment(self, content: dict) -> dict:
        evidence = content["external_evidence"]
        return {
            "html_url": evidence["recruitment_record_url"],
            "issue_url": f"{self.api_base}/issues/9",
            "user": {"login": "example-owner", "type": "User"},
            "author_association": "OWNER",
            "body": "\n".join(
                (
                    "ROUTECONTRACT_PUBLIC_RECRUITMENT_OPEN tag=v0.1.0-rc1",
                    self.activation_marker(),
                    f"ACTIVATION_RECORD_PERMALINK {evidence['activation_record_url']}",
                )
            ),
            "created_at": "2026-07-21T00:00:00Z",
            "updated_at": "2026-07-21T00:05:00Z",
        }

    def release_assets(self, tested_tag: str) -> list[dict]:
        checksum_sha = hashlib.sha256(self.checksum_bytes()).hexdigest()
        assets: list[dict] = []
        for asset_id, name in enumerate(self.public_assets(), start=100):
            asset_sha = (
                checksum_sha
                if name == package_submission.CHECKSUMS_NAME
                else self.payload_digests()[name]
            )
            assets.append(
                {
                    "id": asset_id,
                    "name": name,
                    "size": len(self.checksum_bytes()) if name == "SHA256SUMS" else 123,
                    "state": "uploaded",
                    "digest": f"sha256:{asset_sha}",
                    "url": f"{self.api_base}/releases/assets/{asset_id}",
                    "browser_download_url": (
                        f"{self.repository_url}/releases/download/{tested_tag}/{name}"
                    ),
                    "created_at": "2026-07-20T02:00:00Z",
                    "updated_at": "2026-07-20T02:01:00Z",
                }
            )
        return assets

    def activation_pull_request(self) -> dict:
        return {
            "id": 8800,
            "node_id": "PR_kwDOActivation88",
            "number": self.ACTIVATION_PULL_NUMBER,
            "html_url": f"{self.repository_url}/pull/{self.ACTIVATION_PULL_NUMBER}",
            "state": "closed",
            "merged": True,
            "merged_at": "2026-07-20T05:00:00Z",
            "merge_commit_sha": self.RECORD_COMMIT,
            "base": {
                "ref": "main",
                "repo": {"full_name": "example-owner/routecontract"},
            },
        }

    def fake_activation_pull_requests(self, url: str) -> list[dict]:
        self.assertEqual(
            f"{self.api_base}/commits/{self.RECORD_COMMIT}/pulls?per_page=100",
            url,
        )
        return [deepcopy(self.activation_pull_request())]

    def fake_json(
        self,
        content: dict,
        *,
        record: dict | None = None,
        issues: list[dict] | None = None,
        mutate=None,
    ):
        evidence = content["external_evidence"]
        activation = deepcopy(record or self.activation_record())
        issues_by_number = {
            issue["number"]: deepcopy(issue) for issue in (issues or [])
        }
        self.graphql_issues_by_number = deepcopy(issues_by_number)

        def respond(url: str) -> dict:
            if url == self.api_base:
                value = {
                    "id": self.REPOSITORY_ID,
                    "node_id": self.REPOSITORY_NODE,
                    "full_name": "example-owner/routecontract",
                    "private": False,
                    "archived": False,
                }
            elif "/contents/docs/evidence/" in url:
                raw = json.dumps(activation, separators=(",", ":")).encode("utf-8")
                value = {
                    "type": "file",
                    "path": (
                        "docs/evidence/independent-rc-activation-v0.1.0-rc1.json"
                    ),
                    "sha": self.RECORD_BLOB,
                    "html_url": evidence["activation_record_url"],
                    "encoding": "base64",
                    "content": base64.b64encode(raw).decode("ascii"),
                    "size": len(raw),
                }
            elif "/contents/.github/ISSUE_TEMPLATE/independent-rc1-install.yml" in url:
                raw = self.tagged_bytes(self.TAG_COMMIT, ".github/ISSUE_TEMPLATE/independent-rc1-install.yml")
                value = {
                    "type": "file",
                    "path": ".github/ISSUE_TEMPLATE/independent-rc1-install.yml",
                    "sha": self.FORM_BLOB,
                    "html_url": activation["issueFormPermalink"],
                    "encoding": "base64",
                    "content": base64.b64encode(raw).decode("ascii"),
                    "size": len(raw),
                }
            elif "/contents/docs/independent-install-study.md" in url:
                raw = self.tagged_bytes(self.TAG_COMMIT, "docs/independent-install-study.md")
                value = {
                    "type": "file",
                    "path": "docs/independent-install-study.md",
                    "sha": self.PROTOCOL_BLOB,
                    "html_url": activation["taggedProtocolUrl"],
                    "encoding": "base64",
                    "content": base64.b64encode(raw).decode("ascii"),
                    "size": len(raw),
                }
            elif "/contents/README.md" in url:
                raw = self.tagged_bytes(self.TAG_COMMIT, "README.md")
                value = {
                    "type": "file",
                    "path": "README.md",
                    "sha": self.README_BLOB,
                    "html_url": activation["taggedReadmeUrl"],
                    "encoding": "base64",
                    "content": base64.b64encode(raw).decode("ascii"),
                    "size": len(raw),
                }
            elif url == f"{self.api_base}/commits/main":
                value = {"sha": self.manifest["project"]["commit"]}
            elif f"/commits/{self.RECORD_COMMIT}" in url:
                value = {
                    "sha": self.RECORD_COMMIT,
                    "parents": [{"sha": self.TAG_COMMIT}],
                    "files": [
                        {
                            "filename": (
                                "docs/evidence/"
                                "independent-rc-activation-v0.1.0-rc1.json"
                            ),
                            "status": "added",
                        }
                    ],
                    "commit": {
                        "tree": {"sha": self.RECORD_TREE},
                        "author": {"date": "2026-07-20T04:45:00Z"},
                        "committer": {"date": "2026-07-20T04:50:00Z"},
                    },
                }
            elif url == f"{self.api_base}/pulls/{self.ACTIVATION_PULL_NUMBER}":
                value = self.activation_pull_request()
            elif f"/git/trees/{self.RECORD_TREE}" in url:
                value = {
                    "sha": self.RECORD_TREE,
                    "truncated": False,
                    "tree": [
                        {
                            "path": (
                                "docs/evidence/"
                                "independent-rc-activation-v0.1.0-rc1.json"
                            ),
                            "mode": "100644",
                            "type": "blob",
                            "sha": self.RECORD_BLOB,
                        }
                    ],
                }
            elif url == f"{self.api_base}/compare/{self.RECORD_COMMIT}...main":
                value = {
                    "status": "ahead",
                    "ahead_by": 3,
                    "behind_by": 0,
                    "base_commit": {"sha": self.RECORD_COMMIT},
                    "merge_base_commit": {"sha": self.RECORD_COMMIT},
                }
            elif f"/commits/{self.TAG_COMMIT}" in url:
                value = {
                    "sha": self.TAG_COMMIT,
                    "commit": {"tree": {"sha": self.TAG_TREE}},
                }
            elif f"/git/trees/{self.TAG_TREE}" in url:
                value = {
                    "sha": self.TAG_TREE,
                    "truncated": False,
                    "tree": [
                        {
                            "path": ".github/ISSUE_TEMPLATE/independent-rc1-install.yml",
                            "mode": "100644",
                            "type": "blob",
                            "sha": self.FORM_BLOB,
                        },
                        {
                            "path": "docs/independent-install-study.md",
                            "mode": "100644",
                            "type": "blob",
                            "sha": self.PROTOCOL_BLOB,
                        },
                        {
                            "path": "README.md",
                            "mode": "100644",
                            "type": "blob",
                            "sha": self.README_BLOB,
                        },
                    ],
                }
            elif "/git/ref/tags/v0.1.0-rc1" in url:
                value = {
                    "ref": "refs/tags/v0.1.0-rc1",
                    "object": {
                        "type": "tag",
                        "sha": self.TAG_OBJECT,
                        "url": f"{self.api_base}/git/tags/{self.TAG_OBJECT}",
                    },
                }
            elif url == f"{self.api_base}/git/tags/{self.TAG_OBJECT}":
                value = {
                    "sha": self.TAG_OBJECT,
                    "tag": "v0.1.0-rc1",
                    "object": {"type": "commit", "sha": self.TAG_COMMIT},
                }
            elif url == f"{self.api_base}/actions/runs/{self.RUN_ID}":
                value = {
                    "id": self.RUN_ID,
                    "html_url": activation["releaseEvidence"]["runUrl"],
                    "head_sha": self.TAG_COMMIT,
                    "head_branch": "v0.1.0-rc1",
                    "event": "push",
                    "status": "completed",
                    "conclusion": "success",
                    "name": "Release evidence",
                    "path": ".github/workflows/release-evidence.yml",
                    "repository": {"full_name": "example-owner/routecontract"},
                    "created_at": "2026-07-20T00:00:00Z",
                    "updated_at": "2026-07-20T01:00:00Z",
                }
            elif url == f"{self.api_base}/actions/artifacts/{self.ARTIFACT_ID}":
                value = {
                    "id": self.ARTIFACT_ID,
                    "name": f"routecontract-release-evidence-{self.TAG_COMMIT}",
                    "digest": activation["releaseEvidence"]["artifactDigest"],
                    "expired": False,
                    "size_in_bytes": 456,
                    "workflow_run": {
                        "id": self.RUN_ID,
                        "head_sha": self.TAG_COMMIT,
                        "head_branch": "v0.1.0-rc1",
                    },
                    "created_at": "2026-07-20T00:20:00Z",
                    "updated_at": "2026-07-20T00:30:00Z",
                }
            elif "/releases/tags/v0.1.0-rc1" in url:
                value = {
                    "draft": False,
                    "prerelease": True,
                    "immutable": True,
                    "tag_name": "v0.1.0-rc1",
                    "html_url": f"{self.repository_url}/releases/tag/v0.1.0-rc1",
                    "created_at": "2026-07-20T01:30:00Z",
                    "published_at": "2026-07-20T03:00:00Z",
                    "updated_at": "2026-07-20T04:00:00Z",
                    "assets": self.release_assets("v0.1.0-rc1"),
                }
            elif "/issues/comments/777" in url:
                value = self.recruitment_comment(content)
            elif "/releases/tags/v0.1.0" in url:
                value = {
                    "draft": False,
                    "prerelease": False,
                    "immutable": True,
                    "tag_name": "v0.1.0",
                    "html_url": f"{self.repository_url}/releases/tag/v0.1.0",
                    "created_at": "2026-07-20T00:00:00Z",
                    "published_at": "2026-07-21T00:00:00Z",
                    "updated_at": "2026-07-21T00:05:00Z",
                }
            elif "/issues/" in url:
                number = int(url.rsplit("/", 1)[1])
                if number not in issues_by_number:
                    raise AssertionError(f"unexpected direct Issue request: {url}")
                value = deepcopy(issues_by_number[number])
            else:
                raise AssertionError(f"unexpected public-evidence request: {url}")
            value = deepcopy(value)
            if mutate is not None:
                mutate(url, value)
            return value

        return respond

    @staticmethod
    def fake_pages(
        pages: dict[int, list[dict]],
        links: dict[int, list[str]] | None = None,
    ):
        headers = links or {}

        def respond(url: str) -> tuple[list[dict], list[str]]:
            query = urllib.parse.parse_qs(urllib.parse.urlparse(url).query)
            page = int(query["page"][0]) if "page" in query else 1
            return deepcopy(pages.get(page, [])), deepcopy(headers.get(page, []))

        return respond

    def materialized(self, branch: str) -> dict:
        return package_submission.validate_and_materialize_report_content(
            valid_report_content(branch),
            self.manifest,
            current_utc=TEST_CURRENT_UTC,
        )

    def test_both_enabled_branches_bind_to_public_github_evidence(self) -> None:
        for branch in ("rc_only", "zero"):
            content = self.materialized(branch)
            issue = self.eligible_issue(content)
            page_one = [issue] if branch != "zero" else [self.unrelated_issue(17)]
            with self.subTest(branch=branch), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, issues=[issue]),
            ), patch.object(
                package_submission, "request_json_list_page",
                side_effect=self.fake_pages({1: page_one, 2: []}),
            ) as listed, patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ):
                metadata = package_submission.validate_public_external_evidence(
                    content, self.manifest
                )
            self.assertEqual(branch, metadata["branch"])
            self.assertGreaterEqual(listed.call_count, 1)
            if branch == "rc_only":
                self.assertIn("activation_record_commit", metadata)
                self.assertIn("recruitment_effective_at", metadata)
                self.assertEqual(17, metadata["activation_artifact_file_count"])
                self.assertTrue(metadata["activation_artifact_raw_osv_absent"])
                self.assertEqual("independent-tester-42", metadata["result_issue_author"])
            else:
                self.assertIn("activation_record_commit", metadata)
                self.assertIn("recruitment_effective_at", metadata)
                self.assertEqual(17, metadata["activation_artifact_file_count"])
                self.assertTrue(metadata["activation_artifact_raw_osv_absent"])
                self.assertIsNone(metadata["result_issue_url"])
                self.assertEqual(0, metadata["qualified_result_count"])
        self.assertEqual(2, self.artifact_download.call_count)

    def test_rejects_nonexact_activation_schema_and_version_derived_form(self) -> None:
        for mutation in (
            {"schemaVersion": 1},
            {"schemaVersion": 2.0},
            {"issueFormFilename": "independent-rc-install.yml"},
            {
                "issueFormPermalink": (
                    f"{self.repository_url}/blob/{self.TAG_COMMIT}/.github/"
                    "ISSUE_TEMPLATE/independent-rc-install.yml"
                )
            },
            {
                "issueFormUrl": (
                    f"{self.repository_url}/issues/new?template=independent-rc2-install.yml"
                )
            },
            {"releaseState": {"draft": 0, "immutable": 1, "prerelease": 1}},
            {"releaseState": {"draft": False, "immutable": True, "prerelease": None}},
        ):
            content = self.materialized("rc_only")
            record = self.activation_record()
            record.update(mutation)
            with self.subTest(mutation=mutation), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, record=record),
            ), patch.object(
                package_submission, "request_json_list_page",
                side_effect=self.fake_pages({1: [], 2: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaisesRegex(package_submission.GateError, "exact activated RC identity"):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

        content = self.materialized("rc_only")
        record = self.activation_record()
        record["releaseEvidence"]["artifactFileCount"] = 17.0
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, record=record),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaisesRegex(package_submission.GateError, "exact activated RC identity"):
            package_submission.validate_public_external_evidence(content, self.manifest)

        for key, value in (
            ("releaseEvidence", ["not-an-object"]),
            ("releaseImmutability", "not-an-object"),
        ):
            record = self.activation_record()
            record[key] = value
            with self.subTest(key=key), patch.object(
                package_submission,
                "request_json",
                side_effect=self.fake_json(content, record=record),
            ), patch.object(
                package_submission,
                "request_json_list_page",
                side_effect=self.fake_pages({1: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaisesRegex(package_submission.GateError, "malformed nested"):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

    def test_rejects_broken_rc_run_artifact_tree_main_asset_or_checksum(self) -> None:
        content = self.materialized("rc_only")
        issue = self.eligible_issue(content)

        def wrong_tree(_url: str, value: dict) -> None:
            if value.get("sha") == self.RECORD_TREE and "tree" in value:
                value["tree"][0]["mode"] = "120000"

        def wrong_main(url: str, value: dict) -> None:
            if "/compare/" in url:
                value["merge_base_commit"]["sha"] = "8" * 40

        def wrong_main_head(url: str, value: dict) -> None:
            if url == f"{self.api_base}/commits/main":
                value["sha"] = "8" * 40

        def activation_behind_main(url: str, value: dict) -> None:
            if "/compare/" in url:
                value["behind_by"] = 1

        def missing_archived(url: str, value: dict) -> None:
            if url == self.api_base:
                value.pop("archived")

        def failed_run(url: str, value: dict) -> None:
            if url.endswith(f"/actions/runs/{self.RUN_ID}"):
                value["conclusion"] = "failure"

        def wrong_artifact(url: str, value: dict) -> None:
            if url.endswith(f"/actions/artifacts/{self.ARTIFACT_ID}"):
                value["digest"] = f"sha256:{digest('e')}"

        def mutable_release(url: str, value: dict) -> None:
            if "/releases/tags/v0.1.0-rc1" in url:
                value["immutable"] = False

        def starter_asset(url: str, value: dict) -> None:
            if "/releases/tags/v0.1.0-rc1" in url:
                value["assets"][0]["state"] = "starter"

        def missing_tagged_protocol(url: str, value: dict) -> None:
            if url.endswith(f"/git/trees/{self.TAG_TREE}?recursive=1"):
                value["tree"] = [
                    entry
                    for entry in value["tree"]
                    if entry["path"] != "docs/independent-install-study.md"
                ]

        def corrupt_form_bytes(url: str, value: dict) -> None:
            if "/contents/.github/ISSUE_TEMPLATE/independent-rc1-install.yml" in url:
                raw = (
                    b"name: Independent RC installation\n"
                    b'title: "[independent-install] "\n'
                    b"labels: [evidence, community]\n"
                    b"description: |\n"
                    b"  body: []\n"
                    b"  Issue-form source: <record issueFormPermalink>\n"
                    b"body: []\n"
                )
                value["content"] = base64.b64encode(raw).decode("ascii")
                value["size"] = len(raw)

        def post_cutoff_record_date(url: str, value: dict) -> None:
            if url.endswith(f"/commits/{self.RECORD_COMMIT}"):
                value["commit"]["committer"]["date"] = "2026-08-02T00:00:00Z"

        def post_cutoff_pull_merge(url: str, value: dict) -> None:
            if url.endswith(f"/pulls/{self.ACTIVATION_PULL_NUMBER}"):
                value["merged_at"] = "2026-08-02T00:00:00Z"

        for label, mutator in (
            ("tree", wrong_tree),
            ("main", wrong_main),
            ("main-head", wrong_main_head),
            ("activation-behind-main", activation_behind_main),
            ("repository", missing_archived),
            ("run", failed_run),
            ("artifact", wrong_artifact),
            ("release", mutable_release),
            ("asset", starter_asset),
            ("tagged-protocol", missing_tagged_protocol),
            ("form-structural-decoy", corrupt_form_bytes),
            ("record-date-after-cutoff", post_cutoff_record_date),
            ("activation-pull-after-cutoff", post_cutoff_pull_merge),
        ):
            with self.subTest(label=label), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, issues=[issue], mutate=mutator),
            ), patch.object(
                package_submission, "request_json_list_page",
                side_effect=self.fake_pages({1: [issue], 2: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(content, self.manifest)

        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[issue]),
        ), patch.object(
            package_submission, "request_json_list", return_value=[]
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [issue], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaisesRegex(package_submission.GateError, "main pull request"):
            package_submission.validate_public_external_evidence(
                content, self.manifest
            )

        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[issue]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [issue], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=b"corrupt\n"
        ), self.assertRaises(package_submission.GateError):
            package_submission.validate_public_external_evidence(content, self.manifest)

    def test_recruitment_requires_exact_owner_lines_and_final_edit_time(self) -> None:
        content = self.materialized("rc_only")
        issue = self.eligible_issue(content)

        def bot(url: str, value: dict) -> None:
            if "/issues/comments/777" in url:
                value["user"]["type"] = "Bot"

        def member(url: str, value: dict) -> None:
            if "/issues/comments/777" in url:
                value["author_association"] = "MEMBER"

        def partial_line(url: str, value: dict) -> None:
            if "/issues/comments/777" in url:
                value["body"] = value["body"].replace(
                    "ROUTECONTRACT_PUBLIC_RECRUITMENT_OPEN tag=v0.1.0-rc1",
                    "ROUTECONTRACT_PUBLIC_RECRUITMENT_OPEN tag=v0.1.0-rc1 extra",
                )

        def edited_after_cutoff(url: str, value: dict) -> None:
            if "/issues/comments/777" in url:
                value["updated_at"] = "2026-08-02T00:00:00Z"

        def hidden_comment(url: str, value: dict) -> None:
            if "/issues/comments/777" in url:
                value["body"] = f"<!--\n{value['body']}\n-->"

        def activation_time_equal(url: str, value: dict) -> None:
            if "/issues/comments/777" in url:
                value["created_at"] = "2026-07-20T04:00:00Z"

        for label, mutator in (
            ("bot", bot),
            ("association", member),
            ("partial", partial_line),
            ("updated", edited_after_cutoff),
            ("hidden-comment", hidden_comment),
            ("equal-activation-second", activation_time_equal),
        ):
            with self.subTest(label=label), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, issues=[issue], mutate=mutator),
            ), patch.object(
                package_submission, "request_json_list_page",
                side_effect=self.fake_pages({1: [issue], 2: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(content, self.manifest)

    def test_activation_pull_association_accepts_null_list_merge_sha_only_when_direct_is_exact(
        self,
    ) -> None:
        """GitHub's anonymous commit/pulls response can omit a known squash SHA."""
        content = self.materialized("zero")
        associated = self.activation_pull_request()
        associated["merge_commit_sha"] = None
        with patch.object(
            package_submission,
            "request_json_list",
            return_value=[associated],
        ), patch.object(
            package_submission,
            "request_json",
            side_effect=self.fake_json(content),
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            metadata = package_submission.validate_public_external_evidence(
                content, self.manifest
            )

        self.assertEqual(
            self.RECORD_COMMIT,
            metadata["activation_pull_request"]["merge_commit_sha"],
        )
        self.assertEqual(
            self.RECORD_COMMIT,
            metadata["activation_pull_request"]["graphql_verified"][
                "merge_commit_sha"
            ],
        )

    def test_activation_pull_association_accepts_null_direct_merge_sha_only_when_list_is_exact(
        self,
    ) -> None:
        """GitHub can instead omit the SHA only from the direct PR response."""
        content = self.materialized("zero")

        def omit_direct_sha(url: str, payload: dict) -> None:
            if url == f"{self.api_base}/pulls/{self.ACTIVATION_PULL_NUMBER}":
                payload["merge_commit_sha"] = None

        with patch.object(
            package_submission,
            "request_json",
            side_effect=self.fake_json(content, mutate=omit_direct_sha),
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            metadata = package_submission.validate_public_external_evidence(
                content, self.manifest
            )

        self.assertEqual(
            self.RECORD_COMMIT,
            metadata["activation_pull_request"]["merge_commit_sha"],
        )

    def test_activation_pull_association_rejects_optional_sha_or_direct_identity_drift(
        self,
    ) -> None:
        content = self.materialized("zero")

        def direct_mutator(field: str, value):
            def mutate(url: str, payload: dict) -> None:
                if url != f"{self.api_base}/pulls/{self.ACTIVATION_PULL_NUMBER}":
                    return
                if field == "base_ref":
                    payload["base"]["ref"] = value
                elif field == "base_repository":
                    payload["base"]["repo"]["full_name"] = value
                else:
                    payload[field] = value

            return mutate

        cases = (
            ("id", 9999),
            ("node_id", "PR_drifted"),
            ("number", self.ACTIVATION_PULL_NUMBER + 1),
            ("html_url", f"{self.repository_url}/pull/999"),
            ("state", "open"),
            ("merged", False),
            ("merged_at", None),
            ("merge_commit_sha", "8" * 40),
            ("base_ref", "next"),
            ("base_repository", "foreign-owner/routecontract"),
        )
        for field, value in cases:
            associated = self.activation_pull_request()
            associated["merge_commit_sha"] = None
            with self.subTest(field=field), patch.object(
                package_submission,
                "request_json_list",
                return_value=[associated],
            ), patch.object(
                package_submission,
                "request_json",
                side_effect=self.fake_json(
                    content, mutate=direct_mutator(field, value)
                ),
            ), patch.object(
                package_submission,
                "request_json_list_page",
                side_effect=self.fake_pages({1: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

    def test_direct_main_compare_requires_exact_typed_ancestor_counts(self) -> None:
        content = self.materialized("zero")

        def comparison_mutator(changes: dict[str, object], missing: str | None = None):
            def mutate(url: str, payload: dict) -> None:
                if "/compare/" not in url:
                    return
                if missing is not None:
                    payload.pop(missing)
                payload.update(changes)

            return mutate

        cases = (
            ("missing-ahead", {}, "ahead_by"),
            ("missing-behind", {}, "behind_by"),
            ("bool-ahead", {"ahead_by": True}, None),
            ("bool-behind", {"behind_by": False}, None),
            ("float-ahead", {"ahead_by": 1.0}, None),
            ("float-behind", {"behind_by": 0.0}, None),
            ("identical-with-ahead", {"status": "identical", "ahead_by": 1}, None),
            ("ahead-with-zero", {"status": "ahead", "ahead_by": 0}, None),
            ("behind-positive", {"behind_by": 1}, None),
        )
        for label, changes, missing in cases:
            with self.subTest(label=label), patch.object(
                package_submission,
                "request_json",
                side_effect=self.fake_json(
                    content, mutate=comparison_mutator(changes, missing)
                ),
            ), patch.object(
                package_submission,
                "request_json_list_page",
                side_effect=self.fake_pages({1: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaisesRegex(
                package_submission.GateError, "ancestor of final public main"
            ):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

        associated = self.activation_pull_request()
        associated["merge_commit_sha"] = "8" * 40
        with patch.object(
            package_submission,
            "request_json_list",
            return_value=[associated],
        ), patch.object(
            package_submission,
            "request_json",
            side_effect=self.fake_json(content),
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaises(package_submission.GateError):
            package_submission.validate_public_external_evidence(content, self.manifest)

    def test_activation_pull_association_rejects_ambiguous_or_foreign_candidates(
        self,
    ) -> None:
        content = self.materialized("zero")
        first = self.activation_pull_request()
        first["merge_commit_sha"] = None
        second = deepcopy(first)
        second.update(
            {
                "id": 8900,
                "node_id": "PR_kwDOActivation89",
                "number": 89,
                "html_url": f"{self.repository_url}/pull/89",
            }
        )
        ordinary_json = self.fake_json(content)

        def ambiguous_json(url: str) -> dict:
            if url == f"{self.api_base}/pulls/89":
                direct = deepcopy(second)
                direct["merge_commit_sha"] = self.RECORD_COMMIT
                return direct
            return ordinary_json(url)

        def ambiguous_graphql(owner: str, repository: str, number: int) -> dict:
            payload = self.fake_graphql_activation_pull(owner, repository, number)
            if number == 89:
                pull = payload["data"]["repository"]["pullRequest"]
                pull.update(
                    {
                        "id": second["node_id"],
                        "databaseId": second["id"],
                        "url": second["html_url"],
                    }
                )
            return payload

        for label, associated in (
            ("ambiguous", [first, second]),
            (
                "foreign",
                [
                    {
                        **first,
                        "base": {
                            "ref": "main",
                            "repo": {"full_name": "foreign-owner/routecontract"},
                        },
                    }
                ],
            ),
        ):
            with self.subTest(label=label), patch.object(
                package_submission,
                "request_json_list",
                return_value=associated,
            ), patch.object(
                package_submission,
                "request_json",
                side_effect=ambiguous_json,
            ), patch.object(
                package_submission,
                "request_json_list_page",
                side_effect=self.fake_pages({1: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), patch.object(
                package_submission,
                "request_graphql_activation_pull",
                side_effect=ambiguous_graphql,
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

    def test_activation_pull_accepts_both_rest_shas_null_only_with_exact_graphql(
        self,
    ) -> None:
        content = self.materialized("zero")
        associated = self.activation_pull_request()
        associated["merge_commit_sha"] = None

        def omit_direct_sha(url: str, payload: dict) -> None:
            if url == f"{self.api_base}/pulls/{self.ACTIVATION_PULL_NUMBER}":
                payload["merge_commit_sha"] = None

        with patch.object(
            package_submission, "request_json_list", return_value=[associated]
        ), patch.object(
            package_submission,
            "request_json",
            side_effect=self.fake_json(content, mutate=omit_direct_sha),
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            metadata = package_submission.validate_public_external_evidence(
                content, self.manifest
            )

        self.assertEqual(
            self.RECORD_COMMIT,
            metadata["activation_pull_request"]["graphql_verified"][
                "merge_commit_sha"
            ],
        )

    def test_activation_pull_rejects_missing_merge_sha_key_but_accepts_explicit_null(
        self,
    ) -> None:
        content = self.materialized("zero")

        associated_missing = self.activation_pull_request()
        associated_missing.pop("merge_commit_sha")
        with self.subTest(surface="associated"), patch.object(
            package_submission, "request_json_list", return_value=[associated_missing]
        ), patch.object(
            package_submission, "request_json", side_effect=self.fake_json(content)
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaises(package_submission.GateError):
            package_submission.validate_public_external_evidence(content, self.manifest)

        def remove_direct_sha(url: str, payload: dict) -> None:
            if url == f"{self.api_base}/pulls/{self.ACTIVATION_PULL_NUMBER}":
                payload.pop("merge_commit_sha")

        associated_null = self.activation_pull_request()
        associated_null["merge_commit_sha"] = None
        with self.subTest(surface="direct"), patch.object(
            package_submission, "request_json_list", return_value=[associated_null]
        ), patch.object(
            package_submission,
            "request_json",
            side_effect=self.fake_json(content, mutate=remove_direct_sha),
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaises(package_submission.GateError):
            package_submission.validate_public_external_evidence(content, self.manifest)

    def test_activation_pull_rejects_noninteger_rest_and_graphql_identities(self) -> None:
        content = self.materialized("zero")

        for field in ("id", "number"):
            for value in (True, 1.0, 0, -1):
                associated = self.activation_pull_request()
                associated[field] = value
                with self.subTest(surface="associated", field=field, value=value), patch.object(
                    package_submission, "request_json_list", return_value=[associated]
                ), patch.object(
                    package_submission, "request_json", side_effect=self.fake_json(content)
                ), patch.object(
                    package_submission,
                    "request_json_list_page",
                    side_effect=self.fake_pages({1: []}),
                ), patch.object(
                    package_submission, "request_bytes", return_value=self.checksum_bytes()
                ), self.assertRaises(package_submission.GateError):
                    package_submission.validate_public_external_evidence(
                        content, self.manifest
                    )

                def mutate_direct(url: str, payload: dict) -> None:
                    if url == f"{self.api_base}/pulls/{self.ACTIVATION_PULL_NUMBER}":
                        payload[field] = value

                with self.subTest(surface="direct", field=field, value=value), patch.object(
                    package_submission,
                    "request_json",
                    side_effect=self.fake_json(content, mutate=mutate_direct),
                ), patch.object(
                    package_submission,
                    "request_json_list_page",
                    side_effect=self.fake_pages({1: []}),
                ), patch.object(
                    package_submission, "request_bytes", return_value=self.checksum_bytes()
                ), self.assertRaises(package_submission.GateError):
                    package_submission.validate_public_external_evidence(
                        content, self.manifest
                    )

        pull_path = ("data", "repository", "pullRequest")
        for field in ("databaseId", "number"):
            for value in (True, 1.0, 0, -1):
                def mutate_graphql(payload: dict, field=field, value=value) -> None:
                    payload["data"]["repository"]["pullRequest"][field] = value

                self.activation_graphql_mutate = mutate_graphql
                try:
                    with self.subTest(surface="graphql", field=field, value=value), patch.object(
                        package_submission, "request_json", side_effect=self.fake_json(content)
                    ), patch.object(
                        package_submission,
                        "request_json_list_page",
                        side_effect=self.fake_pages({1: []}),
                    ), patch.object(
                        package_submission, "request_bytes", return_value=self.checksum_bytes()
                    ), self.assertRaises(package_submission.GateError):
                        package_submission.validate_public_external_evidence(
                            content, self.manifest
                        )
                finally:
                    self.activation_graphql_mutate = None

        strict_fields = (
            ((*pull_path, "merged"), 1),
            ((*pull_path, "url"), self.ACTIVATION_PULL_NUMBER),
            ((*pull_path, "state"), 1),
            ((*pull_path, "mergedAt"), 1),
            (("data", "repository", "id"), 1),
            ((*pull_path, "baseRepository", "id"), 1),
        )
        for path, value in strict_fields:
            def mutate_graphql(payload: dict, path=path, value=value) -> None:
                target = payload
                for component in path[:-1]:
                    target = target[component]
                target[path[-1]] = value

            self.activation_graphql_mutate = mutate_graphql
            try:
                with self.subTest(path=path), patch.object(
                    package_submission, "request_json", side_effect=self.fake_json(content)
                ), patch.object(
                    package_submission,
                    "request_json_list_page",
                    side_effect=self.fake_pages({1: []}),
                ), patch.object(
                    package_submission, "request_bytes", return_value=self.checksum_bytes()
                ), self.assertRaises(package_submission.GateError):
                    package_submission.validate_public_external_evidence(
                        content, self.manifest
                    )
            finally:
                self.activation_graphql_mutate = None

    def test_activation_pull_rejects_graphql_partial_or_identity_drift(self) -> None:
        content = self.materialized("zero")

        def mutate(path: tuple[str, ...], value):
            def apply(payload: dict) -> None:
                target = payload
                for key in path[:-1]:
                    target = target[key]
                target[path[-1]] = value

            return apply

        pull_path = ("data", "repository", "pullRequest")
        cases = (
            (("data", "repository"), None),
            (("data", "unexpected"), {}),
            (("data", "repository", "unexpected"), True),
            (("data", "repository", "id"), "R_wrong"),
            (("data", "repository", "nameWithOwner"), "other/repo"),
            ((*pull_path, "id"), "PR_wrong"),
            ((*pull_path, "databaseId"), 9999),
            ((*pull_path, "number"), 89),
            ((*pull_path, "url"), f"{self.repository_url}/pull/89"),
            ((*pull_path, "state"), "CLOSED"),
            ((*pull_path, "merged"), False),
            ((*pull_path, "mergedAt"), None),
            ((*pull_path, "baseRefName"), "next"),
            ((*pull_path, "baseRepository", "id"), "R_wrong"),
            ((*pull_path, "baseRepository", "nameWithOwner"), "other/repo"),
            ((*pull_path, "mergeCommit"), None),
            ((*pull_path, "mergeCommit", "oid"), "8" * 40),
        )
        for path, value in cases:
            self.activation_graphql_mutate = mutate(path, value)
            with self.subTest(path=path), patch.object(
                package_submission,
                "request_json",
                side_effect=self.fake_json(content),
            ), patch.object(
                package_submission,
                "request_json_list_page",
                side_effect=self.fake_pages({1: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )
        self.activation_graphql_mutate = None
        with patch.object(
            package_submission,
            "request_graphql_activation_pull",
            side_effect=package_submission.GateError("GraphQL unavailable"),
        ), patch.object(
            package_submission,
            "request_json",
            side_effect=self.fake_json(content),
        ), patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=self.fake_pages({1: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaises(package_submission.GateError):
            package_submission.validate_public_external_evidence(content, self.manifest)

    def test_result_requires_checked_lines_user_type_exact_tag_and_cutoff(self) -> None:
        cases: list[tuple[str, str, object]] = []
        rc_content = self.materialized("rc_only")
        unchecked = self.eligible_issue(rc_content)
        required = package_submission.RESULT_ISSUE_REQUIRED_CHECKBOXES[0]
        unchecked["body"] = unchecked["body"].replace(
            f"- [x] {required}", f"- [ ] {required}"
        )
        cases.append(("unchecked", "rc_only", unchecked))

        duplicate_state = self.eligible_issue(rc_content)
        duplicate_state["body"] += f"\n- [ ] {required}"
        cases.append(("checked-and-unchecked", "rc_only", duplicate_state))

        for index, label in enumerate(
            package_submission.RESULT_ISSUE_REQUIRED_CHECKBOXES
        ):
            omitted = self.eligible_issue(rc_content, number=100 + index)
            omitted["body"] = omitted["body"].replace(f"- [x] {label}\n", "")
            cases.append((f"omitted-required-{index}", "rc_only", omitted))

        bot = self.eligible_issue(rc_content)
        bot["user"]["type"] = "Bot"
        cases.append(("bot", "rc_only", bot))

        unknown = self.eligible_issue(rc_content)
        unknown["author_association"] = "UNKNOWN"
        cases.append(("association", "rc_only", unknown))

        owner = self.eligible_issue(rc_content)
        owner["user"] = {"login": "example-owner", "type": "User"}
        owner["author_association"] = "OWNER"
        cases.append(("owner", "rc_only", owner))

        after_cutoff = self.eligible_issue(rc_content)
        after_cutoff["updated_at"] = "2026-08-02T00:00:00Z"
        cases.append(("post-cutoff", "rc_only", after_cutoff))

        equal_recruitment = self.eligible_issue(rc_content)
        equal_recruitment["created_at"] = "2026-07-21T00:05:00Z"
        cases.append(("equal-recruitment-second", "rc_only", equal_recruitment))

        inline_comment = self.eligible_issue(rc_content)
        for label in package_submission.RESULT_ISSUE_REQUIRED_CHECKBOXES:
            inline_comment["body"] = inline_comment["body"].replace(
                f"- [x] {label}", f"- [<!-- -->x] {label}"
            )
        cases.append(("inline-comment-checkbox-forgery", "rc_only", inline_comment))

        for wrapper_label, prefix, suffix in (
            ("html-comment", "<!--\n", "\n-->"),
            ("fenced-code", "```text\n", "\n```"),
            ("blockquote", "> ", ""),
            ("space-tab-indented-code", " \t", ""),
            ("nested-under-negation", "  ", ""),
        ):
            hidden = self.eligible_issue(rc_content)
            if wrapper_label == "blockquote":
                hidden["body"] = "\n".join(
                    f"> {line}" for line in hidden["body"].splitlines()
                )
            elif wrapper_label == "space-tab-indented-code":
                hidden["body"] = "\n".join(
                    f" \t{line}" for line in hidden["body"].splitlines()
                )
            elif wrapper_label == "nested-under-negation":
                hidden["body"] = "- These statements are NOT true:\n" + "\n".join(
                    f"  {line}" for line in hidden["body"].splitlines()
                )
            else:
                hidden["body"] = f"{prefix}{hidden['body']}{suffix}"
            cases.append((wrapper_label, "rc_only", hidden))

        for label, branch, issue in cases:
            content = rc_content
            with self.subTest(label=label), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, issues=[issue]),
            ), patch.object(
                package_submission, "request_json_list_page",
                side_effect=self.fake_pages({1: [issue], 2: []}),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(content, self.manifest)

    def test_graphql_requires_exact_unedited_opener_authored_issue(self) -> None:
        content = self.materialized("rc_only")
        issue = self.eligible_issue(content)

        def mutate(path: tuple[str, ...], value) -> None:
            def apply(payload: dict) -> None:
                parent = payload
                for component in path[:-1]:
                    parent = parent[component]
                parent[path[-1]] = value

            self.graphql_mutate = apply

        cases = (
            (("data", "repository", "id"), "R_wrong"),
            (("data", "repository", "issue", "body"), "changed"),
            (("data", "repository", "issue", "title"), "changed"),
            (("data", "repository", "issue", "author", "id"), "U_wrong"),
            (("data", "repository", "issue", "author", "__typename"), "Bot"),
            (
                ("data", "repository", "issue", "editor"),
                {"__typename": "User", "login": "maintainer"},
            ),
            (("data", "repository", "issue", "lastEditedAt"), "2026-07-23T00:02:00Z"),
            (("data", "repository", "issue", "includesCreatedEdit"), True),
            (("data", "repository", "issue", "userContentEdits", "totalCount"), 1),
            (("data", "repository", "issue", "userContentEdits", "totalCount"), False),
            (("data", "repository", "issue", "timelineItems", "totalCount"), 1),
            (("data", "repository", "issue", "timelineItems", "totalCount"), False),
        )
        for path, value in cases:
            mutate(path, value)
            try:
                with self.subTest(path=path), patch.object(
                    package_submission,
                    "request_json",
                    side_effect=self.fake_json(content, issues=[issue]),
                ), patch.object(
                    package_submission,
                    "request_json_list_page",
                    side_effect=self.fake_pages({1: [issue]}),
                ), patch.object(
                    package_submission, "request_bytes", return_value=self.checksum_bytes()
                ), self.assertRaisesRegex(
                    package_submission.GateError, "edited, renamed, or not the exact"
                ):
                    package_submission.validate_public_external_evidence(
                        content, self.manifest
                    )
            finally:
                self.graphql_mutate = None

    def test_graphql_transport_rejects_duplicate_or_partial_envelopes(self) -> None:
        self.graphql_patcher.stop()
        valid = '{"data":{"repository":null}}'
        with patch.object(
            package_submission,
            "require_safe_github_cli_release_verification",
            return_value="/usr/local/bin/gh",
        ), patch.object(package_submission, "run", return_value=valid) as invoked:
            self.assertEqual(
                {"data": {"repository": None}},
                package_submission.request_graphql_issue("owner", "repo", 7),
            )
        command = invoked.call_args.args[0]
        self.assertIn("--hostname", command)
        self.assertIn("github.com", command)
        self.assertIn("number=7", command)

        for label, raw in (
            ("duplicate", '{"data":{},"data":{}}'),
            ("errors", '{"data":{},"errors":[]}'),
            ("extensions", '{"data":{},"extensions":{}}'),
            ("null-data", '{"data":null}'),
            ("array-data", '{"data":[]}'),
        ):
            with self.subTest(label=label), patch.object(
                package_submission,
                "require_safe_github_cli_release_verification",
                return_value="/usr/local/bin/gh",
            ), patch.object(
                package_submission, "run", return_value=raw
            ), self.assertRaises(package_submission.GateError):
                package_submission.request_graphql_issue("owner", "repo", 7)

    def test_zero_enumerates_and_rejects_an_eligible_result(self) -> None:
        content = self.materialized("zero")
        issue = self.eligible_issue(content)
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[issue]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [issue], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaisesRegex(package_submission.GateError, "qualified result count"):
            package_submission.validate_public_external_evidence(content, self.manifest)

    def test_direct_collector_rejects_boolean_or_float_result_counts(self) -> None:
        for branch, value in (("zero", False), ("rc_only", True), ("rc_only", 1.0)):
            content = self.materialized(branch)
            content["external_evidence"]["qualified_result_count"] = value
            with self.subTest(branch=branch, value=value), self.assertRaisesRegex(
                package_submission.GateError, "typed"
            ):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

    def test_follows_page_two_and_rejects_foreign_or_partial_pagination(self) -> None:
        content = self.materialized("rc_only")
        issue = self.eligible_issue(content)
        first_page = [self.unrelated_issue(number) for number in range(1000, 1100)]
        earliest = datetime(2026, 7, 21, 0, 5, 0, tzinfo=timezone.utc)
        initial_url = package_submission._issue_initial_url(self.api_base, earliest)
        initial_query = urllib.parse.parse_qsl(
            urllib.parse.urlparse(initial_url).query
        )
        next_url = (
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            + urllib.parse.urlencode(
                [
                    *initial_query,
                    ("page", "2"),
                    ("after", "opaque-cursor-token"),
                ]
            )
        )
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[issue]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages(
                {1: first_page, 2: [issue]},
                {
                    1: [
                        f'<{next_url}>; rel="next"',
                        f'<{next_url}>; rel="last"',
                    ]
                },
            ),
        ) as listed, patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            metadata = package_submission.validate_public_external_evidence(
                content, self.manifest
            )
        self.assertEqual([42], metadata["enumeration_eligible_issue_numbers"])
        self.assertEqual(2, metadata["enumeration_pages"])
        self.assertEqual(
            [initial_url, next_url],
            [call.args[0] for call in listed.call_args_list],
        )
        self.assertIn(("since", "2026-07-21T00:04:59Z"), initial_query)

        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[issue]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages(
                {1: [], 2: [issue]},
                {1: [f'<{next_url}>; rel="next"']},
            ),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            empty_first_page = package_submission.validate_public_external_evidence(
                content, self.manifest
            )
        self.assertEqual([42], empty_first_page["enumeration_eligible_issue_numbers"])
        self.assertEqual(2, empty_first_page["enumeration_pages"])

        for label, page, header in (
            (
                "foreign",
                first_page,
                '<https://example.invalid/issues?page=2>; rel="next"',
            ),
            ("partial", first_page, f'<{next_url}>; rel='),
        ):
            with self.subTest(label=label), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, issues=[issue]),
            ), patch.object(
                package_submission, "request_json_list_page",
                side_effect=self.fake_pages(
                    {1: page, 2: [issue], 3: []}, {1: [header]}
                ),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaises(package_submission.GateError):
                package_submission.validate_public_external_evidence(content, self.manifest)

    def test_rejects_missing_next_when_link_advertises_a_later_last_page(self) -> None:
        earliest = datetime(2026, 7, 21, 0, 5, 0, tzinfo=timezone.utc)
        initial_url = package_submission._issue_initial_url(self.api_base, earliest)
        initial_query = urllib.parse.parse_qsl(
            urllib.parse.urlparse(initial_url).query
        )
        later_url = (
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            + urllib.parse.urlencode([*initial_query, ("page", "2")])
        )
        for branch in ("rc_only", "zero"):
            content = self.materialized(branch)
            issue = self.eligible_issue(content)
            with self.subTest(branch=branch), patch.object(
                package_submission, "request_json",
                side_effect=self.fake_json(content, issues=[issue]),
            ), patch.object(
                package_submission, "request_json_list_page",
                return_value=([], [f'<{later_url}>; rel="last"']),
            ), patch.object(
                package_submission, "request_bytes", return_value=self.checksum_bytes()
            ), self.assertRaisesRegex(package_submission.GateError, "omits next"):
                package_submission.validate_public_external_evidence(
                    content, self.manifest
                )

    def test_rejects_next_link_that_skips_numeric_pages(self) -> None:
        content = self.materialized("zero")
        earliest = datetime(2026, 7, 21, 0, 5, 0, tzinfo=timezone.utc)
        initial_query = urllib.parse.parse_qsl(
            urllib.parse.urlparse(
                package_submission._issue_initial_url(self.api_base, earliest)
            ).query
        )
        skipped = (
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            + urllib.parse.urlencode([*initial_query, ("page", "999")])
        )
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[]),
        ), patch.object(
            package_submission, "request_json_list_page",
            return_value=([], [f'<{skipped}>; rel="next"']),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaisesRegex(package_submission.GateError, "skips a numeric page"):
            package_submission.validate_public_external_evidence(
                content, self.manifest
            )

    def test_rejects_two_eligible_results_instead_of_selecting_one(self) -> None:
        content = self.materialized("rc_only")
        first = self.eligible_issue(content, 42)
        second = self.eligible_issue(content, 43)
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(content, issues=[first, second]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [first, second], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ), self.assertRaisesRegex(package_submission.GateError, "qualified result count"):
            package_submission.validate_public_external_evidence(content, self.manifest)

    def test_contextual_form_issues_do_not_poison_zero_or_one_qualified_result(self) -> None:
        rc_content = self.materialized("rc_only")
        eligible = self.eligible_issue(rc_content, 42)
        deviation = self.eligible_issue(rc_content, 43)
        deviation["body"] = deviation["body"].replace(
            "UNASSISTED_PASS", "PROTOCOL_DEVIATION"
        )
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(rc_content, issues=[eligible, deviation]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [eligible, deviation], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            metadata = package_submission.validate_public_external_evidence(
                rc_content, self.manifest
            )
        self.assertEqual(1, metadata["qualified_result_count"])
        self.assertEqual([42], metadata["enumeration_eligible_issue_numbers"])

        zero_content = self.materialized("zero")
        not_run = self.eligible_issue(zero_content, 44)
        not_run["body"] = not_run["body"].replace("UNASSISTED_PASS", "NOT_RUN")
        with patch.object(
            package_submission, "request_json",
            side_effect=self.fake_json(zero_content, issues=[not_run]),
        ), patch.object(
            package_submission, "request_json_list_page",
            side_effect=self.fake_pages({1: [not_run], 2: []}),
        ), patch.object(
            package_submission, "request_bytes", return_value=self.checksum_bytes()
        ):
            metadata = package_submission.validate_public_external_evidence(
                zero_content, self.manifest
            )
        self.assertEqual(0, metadata["qualified_result_count"])
        self.assertEqual([], metadata["enumeration_eligible_issue_numbers"])

    def test_transport_preserves_repeated_link_fields_and_rejects_invalid_utf8(self) -> None:
        class Headers:
            @staticmethod
            def raw_items():
                return iter((('Link', '<https://api.github.com/a>; rel=next'),
                             ('lInK', '<https://api.github.com/b>; rel="last"')))

        class Response:
            headers = Headers()

            def __enter__(self):
                return self

            def __exit__(self, *_args):
                return False

            @staticmethod
            def read(_limit: int | None = None) -> bytes:
                return b"[]"

        with patch.object(
            package_submission.urllib.request, "urlopen", return_value=Response()
        ), patch.object(package_submission, "verified_tls_context", return_value=None):
            data, headers = package_submission.request_bytes_with_headers(
                "https://api.github.com/example", limit=100
            )
        self.assertEqual(b"[]", data)
        self.assertEqual(
            [
                '<https://api.github.com/a>; rel=next',
                '<https://api.github.com/b>; rel="last"',
            ],
            headers["link"],
        )
        self.assertEqual(
            {"next": "https://api.github.com/a", "last": "https://api.github.com/b"},
            package_submission._parse_link_headers(headers["link"]),
        )
        with patch.object(
            package_submission,
            "request_bytes_with_headers",
            return_value=(b"\xff", {}),
        ), self.assertRaisesRegex(package_submission.GateError, "did not return JSON"):
            package_submission.request_json_list_page("https://api.github.com/example")

    def test_link_traversal_accepts_pr_only_empty_and_full_terminal_pages(self) -> None:
        content = self.materialized("zero")
        evidence = content["external_evidence"]
        earliest = datetime(2026, 7, 21, 0, 5, 0, tzinfo=timezone.utc)
        cutoff = datetime(2026, 8, 1, 6, 0, 0, tzinfo=timezone.utc)
        fixed_query = urllib.parse.parse_qsl(
            urllib.parse.urlparse(
                package_submission._issue_initial_url(self.api_base, earliest)
            ).query
        )

        def next_url(page: int) -> str:
            return (
                f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
                + urllib.parse.urlencode(
                    [*fixed_query, ("page", str(page)), ("after", f"cursor-{page}")]
                )
            )

        pull_request = {
            "number": 900,
            "repository_url": self.api_base,
            "pull_request": {"url": "https://api.github.com/pulls/900"},
        }
        with patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=[
                ([pull_request], [f'<{next_url(2)}>; REL=Next']),
                ([], []),
            ],
        ):
            eligible, metadata = package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )
        self.assertEqual([], eligible)
        self.assertEqual(2, metadata["enumeration_pages"])

        terminal_full_page = [
            {
                "number": number,
                "repository_url": self.api_base,
                "pull_request": {"url": f"https://api.github.com/pulls/{number}"},
            }
            for number in range(1000, 1100)
        ]
        with patch.object(
            package_submission,
            "request_json_list_page",
            return_value=(terminal_full_page, []),
        ):
            _, terminal_metadata = package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )
        self.assertEqual(1, terminal_metadata["enumeration_pages"])

    def test_exact_page_cap_accepts_terminal_and_rejects_another_next(self) -> None:
        content = self.materialized("zero")
        evidence = content["external_evidence"]
        earliest = datetime(2026, 7, 21, 0, 5, 0, tzinfo=timezone.utc)
        cutoff = datetime(2026, 8, 1, 6, 0, 0, tzinfo=timezone.utc)
        fixed_query = urllib.parse.parse_qsl(
            urllib.parse.urlparse(
                package_submission._issue_initial_url(self.api_base, earliest)
            ).query
        )

        def next_url(page: int) -> str:
            return (
                f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
                + urllib.parse.urlencode(
                    [*fixed_query, ("page", str(page)), ("after", f"cursor-{page}")]
                )
            )

        def pages(*, next_after_cap: bool):
            page = 0

            def respond(_url: str):
                nonlocal page
                page += 1
                if page < package_submission.MAX_ISSUE_ENUMERATION_PAGES:
                    return [], [f'<{next_url(page + 1)}>; rel="next"']
                if next_after_cap:
                    return [], [f'<{next_url(page + 1)}>; rel="next"']
                return [], []

            return respond

        with patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=pages(next_after_cap=False),
        ):
            _, metadata = package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )
        self.assertEqual(100, metadata["enumeration_pages"])

        with patch.object(
            package_submission,
            "request_json_list_page",
            side_effect=pages(next_after_cap=True),
        ), self.assertRaisesRegex(package_submission.GateError, "safety page limit"):
            package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )

    def test_rejects_malformed_foreign_duplicate_and_looping_link_targets(self) -> None:
        content = self.materialized("zero")
        evidence = content["external_evidence"]
        earliest = datetime(2026, 7, 21, 0, 5, 0, tzinfo=timezone.utc)
        cutoff = datetime(2026, 8, 1, 6, 0, 0, tzinfo=timezone.utc)
        initial = package_submission._issue_initial_url(self.api_base, earliest)
        parsed = urllib.parse.urlparse(initial)
        query = parsed.query
        invalid = (
            f"https://api.github.com/repositories/999/issues?{query}&page=2",
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            f"{query}&state=all&page=2",
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            f"{query}&after=",
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            f"{query}&unexpected=value",
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            f"{query}&page=2#fragment",
        )
        for target in invalid:
            with self.subTest(target=target), self.assertRaises(
                package_submission.GateError
            ):
                package_submission._validated_issue_operation_key(
                    target, self.api_base, self.REPOSITORY_ID, earliest
                )
        for headers in (
            [""],
            [
                f'<{initial}>; rel="next"',
                f'<{initial}>; rel=Next',
            ],
        ):
            with self.subTest(headers=headers), self.assertRaises(
                package_submission.GateError
            ):
                package_submission._parse_link_headers(headers)

        with patch.object(
            package_submission,
            "request_json_list_page",
            return_value=([], [f'<{initial}>; rel="next"']),
        ), self.assertRaisesRegex(package_submission.GateError, "loop"):
            package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )

        foreign_last = '<https://example.invalid/issues?page=9>; rel="last"'
        with patch.object(
            package_submission,
            "request_json_list_page",
            return_value=([], [foreign_last]),
        ), self.assertRaisesRegex(package_submission.GateError, "foreign"):
            package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )

        reverse_next = (
            f"https://api.github.com/repositories/{self.REPOSITORY_ID}/issues?"
            f"{query}&before=reverse-cursor"
        )
        with patch.object(
            package_submission,
            "request_json_list_page",
            return_value=([], [f'<{reverse_next}>; rel="next"']),
        ), self.assertRaisesRegex(package_submission.GateError, "reverse cursor"):
            package_submission._enumerate_result_issues(
                evidence,
                self.manifest,
                self.api_base,
                self.REPOSITORY_ID,
                self.REPOSITORY_NODE,
                earliest,
                cutoff,
                self.activation_marker(),
            )



class PublicRcWorkflowArtifactBindingTest(unittest.TestCase):
    TAG_COMMIT = "c" * 40
    RUN_ID = 11
    ARTIFACT_ID = 22

    @staticmethod
    def public_assets() -> tuple[str, ...]:
        version = "0.1.0-rc1"
        return (
            f"routecontract-{version}-source.zip",
            f"routecontract-shardingsphere-5.5-{version}.jar",
            f"routecontract-shardingsphere-5.5-{version}-sources.jar",
            f"routecontract-shardingsphere-5.5-{version}-javadoc.jar",
            "routecontract-shardingsphere-5.5.pom",
            "routecontract-shardingsphere-5.5-cyclonedx.json",
            "routecontract-shardingsphere-5.5-cyclonedx.xml",
            "routecontract-aggregate-cyclonedx.json",
            "routecontract-aggregate-cyclonedx.xml",
            "supply-chain-evidence.json",
            "test-summary.txt",
            "SHA256SUMS",
        )

    def public_payloads(self) -> dict[str, bytes]:
        return {
            name: f"release-byte::{name}\n".encode("utf-8")
            for name in self.public_assets()
            if name != package_submission.CHECKSUMS_NAME
        }

    def declared(self) -> dict[str, str]:
        return {
            name: hashlib.sha256(data).hexdigest()
            for name, data in self.public_payloads().items()
        }

    def checksum_bytes(self) -> bytes:
        return "".join(
            f"{value}  {name}\n" for name, value in sorted(self.declared().items())
        ).encode("ascii")

    def members(self) -> dict[str, bytes]:
        result = self.public_payloads()
        result[package_submission.CHECKSUMS_NAME] = self.checksum_bytes()
        result.update(
            {
                name: f"workflow-only::{name}\n".encode("utf-8")
                for name in package_submission.RC_ACTIVATION_RECORD_VALIDATOR.WORKFLOW_ONLY_FILES
            }
        )
        return result

    @staticmethod
    def write_zip(path: Path, members: dict[str, bytes]) -> None:
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
            for name, data in members.items():
                archive.writestr(name, data)

    def record(self, artifact_zip: Path) -> dict:
        return {
            "tag": "v0.1.0-rc1",
            "tagCommit": self.TAG_COMMIT,
            "issueFormFilename": "independent-rc1-install.yml",
            "publicAssets": list(self.public_assets()),
            "sha256sumsSha256": hashlib.sha256(self.checksum_bytes()).hexdigest(),
            "releaseEvidence": {
                "artifactDigest": f"sha256:{package_submission.sha256(artifact_zip)}",
                "artifactFileCount": 17,
                "artifactId": self.ARTIFACT_ID,
                "headSha": self.TAG_COMMIT,
                "runId": self.RUN_ID,
                "runUrl": "https://github.com/example-owner/routecontract/actions/runs/11",
            },
        }

    def test_accepts_exact_17_flat_files_and_binds_all_release_bytes(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            artifact_zip = Path(raw) / "artifact.zip"
            source_members = self.members()
            self.write_zip(artifact_zip, source_members)
            observed = package_submission.validate_public_rc_workflow_artifact_archive(
                artifact_zip,
                artifact_zip.stat().st_size,
                self.record(artifact_zip),
                self.declared(),
            )

        expected_names = set(self.public_assets()) | set(
            package_submission.RC_ACTIVATION_RECORD_VALIDATOR.WORKFLOW_ONLY_FILES
        )
        self.assertEqual(17, len(observed))
        self.assertEqual(expected_names, set(observed))
        self.assertNotIn("osv-raw.json", observed)
        for name in self.public_assets():
            self.assertEqual(
                hashlib.sha256(source_members[name]).hexdigest(), observed[name]
            )

    def test_rejects_digest_size_member_bytes_raw_osv_and_nonflat_names(self) -> None:
        mutations = {
            "release-byte-mismatch": lambda members: members.__setitem__(
                self.public_assets()[0], b"different public release bytes"
            ),
            "missing": lambda members: members.pop("environment.txt"),
            "raw-osv": lambda members: (
                members.pop("environment.txt"),
                members.__setitem__("osv-raw.json", b'{"results":[]}\n'),
            ),
            "nonflat": lambda members: (
                members.pop("environment.txt"),
                members.__setitem__("nested/environment.txt", b"unsafe\n"),
            ),
        }
        for label, mutate in mutations.items():
            with self.subTest(label=label), tempfile.TemporaryDirectory() as raw:
                artifact_zip = Path(raw) / "artifact.zip"
                members = self.members()
                mutate(members)
                self.write_zip(artifact_zip, members)
                with self.assertRaises(package_submission.GateError):
                    package_submission.validate_public_rc_workflow_artifact_archive(
                        artifact_zip,
                        artifact_zip.stat().st_size,
                        self.record(artifact_zip),
                        self.declared(),
                    )

        with tempfile.TemporaryDirectory() as raw:
            artifact_zip = Path(raw) / "artifact.zip"
            self.write_zip(artifact_zip, self.members())
            record = self.record(artifact_zip)
            record["releaseEvidence"]["artifactDigest"] = f"sha256:{digest('0')}"
            with self.assertRaisesRegex(package_submission.GateError, "digest"):
                package_submission.validate_public_rc_workflow_artifact_archive(
                    artifact_zip,
                    artifact_zip.stat().st_size,
                    record,
                    self.declared(),
                )
            with self.assertRaisesRegex(package_submission.GateError, "size"):
                package_submission.validate_public_rc_workflow_artifact_archive(
                    artifact_zip,
                    artifact_zip.stat().st_size + 1,
                    self.record(artifact_zip),
                    self.declared(),
                )

    def test_rejects_a_symlink_member_even_with_the_exact_name_allowlist(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            artifact_zip = Path(raw) / "artifact.zip"
            with ZipFile(artifact_zip, "w", compression=ZIP_DEFLATED) as archive:
                for name, data in self.members().items():
                    if name == "environment.txt":
                        info = ZipInfo(name)
                        info.external_attr = (
                            package_submission.stat.S_IFLNK | 0o777
                        ) << 16
                        archive.writestr(info, b"target")
                    else:
                        archive.writestr(name, data)
            with self.assertRaisesRegex(package_submission.GateError, "flat regular"):
                package_submission.validate_public_rc_workflow_artifact_archive(
                    artifact_zip,
                    artifact_zip.stat().st_size,
                    self.record(artifact_zip),
                    self.declared(),
                )

    def test_download_uses_safe_gh_exact_artifact_endpoint_and_declared_size(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            source = Path(raw) / "source.zip"
            self.write_zip(source, self.members())
            record = self.record(source)
            manifest = package_submission.validate_manifest(valid_manifest())

            def fake_download(
                gh: str,
                repository_root: Path,
                endpoint: str,
                destination: Path,
                expected_size: int,
                *,
                accept: str,
            ) -> None:
                self.assertEqual("/safe/bin/gh", gh)
                self.assertEqual(REPOSITORY_ROOT, repository_root)
                self.assertEqual(
                    "repos/example-owner/routecontract/actions/artifacts/22/zip",
                    endpoint,
                )
                self.assertEqual(source.stat().st_size, expected_size)
                self.assertEqual("application/vnd.github+json", accept)
                destination.write_bytes(source.read_bytes())

            with patch.object(
                package_submission,
                "require_safe_github_cli_release_verification",
                return_value="/safe/bin/gh",
            ), patch.object(
                package_submission.RC_ACTIVATION_RECORD_VALIDATOR,
                "_download_gh_file",
                side_effect=fake_download,
            ) as downloaded:
                observed = (
                    package_submission.download_and_validate_public_rc_workflow_artifact(
                        manifest, record, source.stat().st_size, self.declared()
                    )
                )

        self.assertEqual(1, downloaded.call_count)
        self.assertEqual(17, len(observed))


class PublicExternalSnapshotRevalidationTest(unittest.TestCase):
    @staticmethod
    def snapshot() -> dict:
        return {
            "branch": "rc_only",
            "qualified_result_count": 1,
            "result_issue_url": "https://github.com/example-owner/routecontract/issues/42",
            "result_issue_body_sha256": digest("1"),
            "result_issue_labels": ["community", "evidence"],
            "result_issue_updated_at": "2026-07-23T00:01:00+00:00",
            "recruitment_body_sha256": digest("2"),
            "activation_run": {"updated_at": "2026-07-20T01:00:00+00:00"},
            "activation_artifact_id": 2,
            "activation_artifact_digest": f"sha256:{digest('3')}",
            "activation_artifact_size_bytes": 456,
            "activation_artifact_raw_osv_absent": True,
            "activation_artifact_member_digests": {"a": digest("4")},
            "activation_artifact": {"size_in_bytes": 456},
            "activation_release": {
                "updated_at": "2026-07-20T04:00:00+00:00",
                "assets": [{"name": "a", "digest": f"sha256:{digest('4')}"}],
            },
        }

    def test_unchanged_second_observation_passes_and_reuses_artifact_binding(self) -> None:
        initial = self.snapshot()
        with patch.object(
            package_submission,
            "validate_public_external_evidence",
            return_value=deepcopy(initial),
        ) as collected:
            observed = package_submission.revalidate_public_external_evidence(
                initial, {"external_evidence": {}}, valid_manifest()
            )
        self.assertEqual(initial, observed)
        self.assertIs(initial, collected.call_args.kwargs["artifact_binding_cache"])

    def test_each_public_claim_mutation_fails_before_final_archive(self) -> None:
        mutation_paths = (
            ("result_issue_body_sha256",),
            ("result_issue_labels",),
            ("result_issue_updated_at",),
            ("recruitment_body_sha256",),
            ("activation_run", "updated_at"),
            ("activation_artifact", "size_in_bytes"),
            ("activation_release", "updated_at"),
            ("activation_release", "assets"),
        )
        for path in mutation_paths:
            initial = self.snapshot()
            changed = deepcopy(initial)
            parent = changed
            for component in path[:-1]:
                parent = parent[component]
            current = parent[path[-1]]
            parent[path[-1]] = [*current, "changed"] if isinstance(current, list) else f"{current}-changed"
            with self.subTest(path=path), patch.object(
                package_submission,
                "validate_public_external_evidence",
                return_value=changed,
            ), self.assertRaisesRegex(
                package_submission.GateError, "changed between packaging observations"
            ):
                package_submission.revalidate_public_external_evidence(
                    initial, {"external_evidence": {}}, valid_manifest()
                )

        source = (REPOSITORY_ROOT / "submission" / "tools" / "package_submission.py").read_text(
            encoding="utf-8"
        )
        main_source = source[source.index("def main()") :]
        public_requery = main_source.index("revalidate_public_evidence(")
        external_requery = main_source.index("revalidate_public_external_evidence(")
        archive = main_source.index("build_upload_zip(")
        publish = main_source.index("os.replace(staging, output)")
        self.assertLess(public_requery, external_requery)
        self.assertLess(external_requery, archive)
        self.assertLess(archive, publish)

    def test_malformed_second_observation_fails_before_canonical_comparison(self) -> None:
        mutations = (
            ("missing", lambda value: value.pop("qualified_result_count")),
            ("null", lambda value: value.__setitem__("qualified_result_count", None)),
            ("float", lambda value: value.__setitem__("qualified_result_count", 1.0)),
        )
        for label, mutate in mutations:
            initial = self.snapshot()
            observed = deepcopy(initial)
            mutate(observed)
            with self.subTest(label=label), patch.object(
                package_submission,
                "validate_public_external_evidence",
                return_value=observed,
            ), patch.object(
                package_submission,
                "canonical_external_snapshot_bytes",
                side_effect=AssertionError("canonical comparison must not run"),
            ), self.assertRaisesRegex(
                package_submission.GateError, "malformed"
            ):
                package_submission.revalidate_public_external_evidence(
                    initial, {"external_evidence": {}}, valid_manifest()
                )

    def test_complete_release_ci_video_snapshot_is_recollected_exactly(self) -> None:
        initial = {
            "repository_full_name": "example-owner/routecontract",
            "commit": "a" * 40,
            "ci_run_id": 11,
            "ci_conclusion": "success",
            "workflow_artifact_id": 22,
            "workflow_artifact_sha256": digest("3"),
            "release_id": 33,
            "release_tag": "v0.1.0",
            "release_immutable": True,
            "youtube_video_id": "video",
            "youtube_title": "title",
            "youtube_duration_seconds": 172.0,
            "youtube_availability": "public",
            "youtube_live_status": "not_live",
            "youtube_age_limit": 0,
            "youtube_max_video_height": 1080,
        }
        arguments = (
            valid_manifest(),
            {"duration_seconds": 172.0},
            {"public_release_assets": {}},
            Path("evidence"),
            REPOSITORY_ROOT,
        )
        with patch.object(
            package_submission, "validate_public_evidence", return_value=deepcopy(initial)
        ) as collected:
            self.assertEqual(
                initial,
                package_submission.revalidate_public_evidence(initial, *arguments),
            )
        collected.assert_called_once_with(*arguments)

        for key in (
            "commit",
            "ci_conclusion",
            "workflow_artifact_sha256",
            "release_immutable",
            "youtube_availability",
            "youtube_duration_seconds",
        ):
            changed = deepcopy(initial)
            changed[key] = f"changed-{key}"
            with self.subTest(key=key), patch.object(
                package_submission, "validate_public_evidence", return_value=changed
            ), self.assertRaisesRegex(
                package_submission.GateError,
                "release/CI/video state changed between packaging observations",
            ):
                package_submission.revalidate_public_evidence(initial, *arguments)

    def test_public_collector_direct_call_fail_closes_stable_branch(self) -> None:
        with patch.object(package_submission, "request_json") as requested, self.assertRaisesRegex(
            package_submission.GateError, "supports only rc_only or zero"
        ):
            package_submission.validate_public_external_evidence(
                valid_report_content("final_stable"), valid_manifest()
            )
        requested.assert_not_called()


class ManifestTest(unittest.TestCase):
    def test_default_builder_python_uses_canonical_current_interpreter(self) -> None:
        argv = [
            "package_submission.py",
            "--manifest",
            "m",
            "--template",
            "t",
            "--content",
            "c",
            "--report-pdf",
            "p",
            "--video-file",
            "v",
            "--release-evidence-dir",
            "e",
            "--release-evidence-artifact",
            "a",
            "--output",
            "o",
        ]
        with patch.object(sys, "argv", argv):
            args = package_submission.parse_args()

        self.assertEqual(Path(sys.executable).resolve(), args.builder_python)

    def test_accepts_standard_venv_final_python_symlink(self) -> None:
        with tempfile.TemporaryDirectory(dir=SCRIPT.parents[2]) as raw:
            root = Path(raw)
            target = root / "python-target"
            target.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            target.chmod(0o755)
            venv = root / "venv"
            (venv / "bin").mkdir(parents=True)
            interpreter = venv / "bin" / "python"
            interpreter.symlink_to(target)

            self.assertEqual(
                interpreter,
                package_submission.require_python_interpreter(
                    interpreter, "report builder Python"
                ),
            )

    def test_rejects_symlink_in_builder_python_parent_path(self) -> None:
        with tempfile.TemporaryDirectory(dir=SCRIPT.parents[2]) as raw:
            root = Path(raw)
            real = root / "real"
            real.mkdir()
            interpreter = real / "python"
            interpreter.write_text("#!/bin/sh\nexit 0\n", encoding="utf-8")
            interpreter.chmod(0o755)
            linked_parent = root / "linked"
            linked_parent.symlink_to(real, target_is_directory=True)

            with self.assertRaisesRegex(package_submission.GateError, "symlink component"):
                package_submission.require_python_interpreter(
                    linked_parent / "python", "report builder Python"
                )

    def test_accepts_complete_manifest(self) -> None:
        checked = package_submission.validate_manifest(valid_manifest())
        self.assertEqual("example-owner", checked["github_owner"])
        basename = "2026 오픈소스 개발자대회 결과보고서_123(홍길동전)"
        self.assertEqual(
            {
                "basename": basename,
                "docx": f"{basename}.docx",
                "pdf": f"{basename}.pdf",
                "zip": f"{basename}.zip",
            },
            checked["official_submission_filenames"],
        )

    def test_rejects_legacy_or_non_integer_manifest_schema_version(self) -> None:
        for version in (1, 2, 3, 4, 4.0, 5.0, True, "4", "5", None):
            manifest = valid_manifest()
            manifest["schema_version"] = version
            with self.subTest(version=version), self.assertRaisesRegex(
                package_submission.GateError,
                r"^manifest\.schema_version must be 5$",
            ):
                package_submission.validate_manifest(manifest)

    def test_rejects_legacy_overbroad_source_review_attestation_key(self) -> None:
        manifest = valid_manifest()
        attestations = manifest["participant_attestations"]
        attestations["all_submitted_code_reviewed_and_explainable"] = attestations.pop(
            "core_behavior_boundaries_artifacts_and_dependency_roles_reviewed_and_explainable"
        )
        with self.assertRaisesRegex(
            package_submission.GateError, "participant_attestations"
        ):
            package_submission.validate_manifest(manifest)

    def test_rejects_legacy_participant_only_owner_voice_attestation_key(self) -> None:
        manifest = valid_manifest()
        attestations = manifest["participant_attestations"]
        attestations["owner_voice_written_by_participant"] = attestations.pop(
            "owner_voice_ai_assistance_disclosed_and_participant_reviewed"
        )
        with self.assertRaisesRegex(
            package_submission.GateError, "participant_attestations"
        ):
            package_submission.validate_manifest(manifest)

    def test_rejects_non_nfc_submission_identity(self) -> None:
        manifest = valid_manifest()
        manifest["submission_identity"]["team_name"] = unicodedata.normalize(
            "NFD", "홍길동전"
        )
        with self.assertRaisesRegex(package_submission.GateError, "NFC"):
            package_submission.validate_manifest(manifest)

    def test_rejects_unsafe_submission_identity_components(self) -> None:
        for unsafe in ("../123", "123/456", "123\\456", "123\n456", " 123"):
            with self.subTest(unsafe=repr(unsafe)):
                manifest = valid_manifest()
                manifest["submission_identity"]["receipt_number"] = unsafe
                with self.assertRaises(package_submission.GateError):
                    package_submission.validate_manifest(manifest)

    def test_rejects_official_filename_over_portable_byte_limit(self) -> None:
        manifest = valid_manifest()
        manifest["submission_identity"]["team_name"] = "가" * 80
        with self.assertRaisesRegex(package_submission.GateError, "255 UTF-8 bytes"):
            package_submission.validate_manifest(manifest)

    def test_requires_exact_report_registration_identity_match(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        metadata = {
            "team_name": "홍길동전",
            "project_name": "RouteContract",
            "team_size": "1명",
            "division": "학생",
            "task_type": "자유과제",
        }
        package_submission.validate_submission_identity_matches_content(
            {"metadata": metadata}, manifest
        )
        for key in metadata:
            with self.subTest(key=key), self.assertRaisesRegex(
                package_submission.GateError, "exactly match"
            ):
                changed = dict(metadata)
                changed[key] += "-mismatch"
                package_submission.validate_submission_identity_matches_content(
                    {"metadata": changed}, manifest
                )

    def test_rejects_control_character_in_registration_identity(self) -> None:
        manifest = valid_manifest()
        manifest["submission_identity"]["registered_project_name"] = "Route\u200bContract"
        with self.assertRaisesRegex(package_submission.GateError, "control or formatting"):
            package_submission.validate_manifest(manifest)

    def test_rejects_required_duplicate_form_until_official_source_is_validated(self) -> None:
        manifest = valid_manifest()
        manifest["duplicate_benefit_confirmation"] = {
            "status": "required",
            "sha256": digest("8"),
        }
        with self.assertRaisesRegex(package_submission.GateError, "exact source form"):
            package_submission.validate_manifest(manifest)
        with self.assertRaisesRegex(package_submission.GateError, "exact source form"):
            package_submission.validate_duplicate_confirmation(Path("arbitrary.pdf"), manifest)

    def test_rejects_placeholder_before_packaging(self) -> None:
        manifest = valid_manifest()
        manifest["project"]["repository_url"] = "[[PUBLIC_REPOSITORY_URL]]"
        with self.assertRaisesRegex(package_submission.GateError, "unresolved"):
            package_submission.validate_manifest(manifest)

    def test_rejects_packaging_at_or_after_official_deadline(self) -> None:
        at_deadline = datetime(2026, 8, 27, 9, 0, 0, tzinfo=timezone.utc)
        with self.assertRaisesRegex(package_submission.GateError, "deadline has passed"):
            package_submission.validate_submission_deadline(at_deadline)

    def test_video_duration_window_and_caption_end_are_inclusive(self) -> None:
        for duration in (172.5, 175.0):
            manifest = valid_manifest()
            manifest["video"]["duration_seconds"] = duration
            with self.subTest(duration=duration):
                package_submission.validate_manifest(manifest)

        for duration in (169.9, 175.1):
            manifest = valid_manifest()
            manifest["video"]["duration_seconds"] = duration
            with self.subTest(duration=duration), self.assertRaisesRegex(
                package_submission.GateError, "170 through 175 seconds inclusive"
            ):
                package_submission.validate_manifest(manifest)

        for duration in (170.0, 172.499):
            manifest = valid_manifest()
            manifest["video"]["duration_seconds"] = duration
            with self.subTest(duration=duration), self.assertRaisesRegex(
                package_submission.GateError, "selected caption branch's final cue"
            ):
                package_submission.validate_manifest(manifest)

    def test_rejects_unsupported_or_mismatched_video_evidence_branch(self) -> None:
        for branch in ([], {}, None, True, 1, 1.0):
            manifest = valid_manifest()
            manifest["video"]["external_evidence_branch"] = branch
            with self.subTest(branch=branch), self.assertRaisesRegex(
                package_submission.GateError,
                r"^video\.external_evidence_branch must be rc_only or zero$",
            ):
                package_submission.validate_manifest(manifest)

        manifest = valid_manifest()
        manifest["video"]["external_evidence_branch"] = "final_stable"
        with self.assertRaisesRegex(
            package_submission.GateError,
            r"video\.external_evidence_branch must be rc_only or zero",
        ):
            package_submission.validate_manifest(manifest)

        manifest = package_submission.validate_manifest(valid_manifest())
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("zero"),
            manifest,
            current_utc=TEST_CURRENT_UTC,
        )
        with self.assertRaisesRegex(
            package_submission.GateError,
            r"must exactly match.*report external-evidence branch",
        ):
            package_submission.validate_video_external_evidence_branch_matches_content(
                content, manifest
            )

        manifest["video"]["external_evidence_branch"] = "zero"
        package_submission.validate_video_external_evidence_branch_matches_content(
            content, manifest
        )

    def test_rejects_release_candidate_tag_for_final_package(self) -> None:
        manifest = valid_manifest()
        manifest["project"]["tag"] = "v0.1.0-rc1"
        manifest["project"]["release_url"] += "-rc1"
        manifest["release_evidence"]["source_archive_filename"] = (
            "routecontract-0.1.0-rc1-source.zip"
        )
        with self.assertRaisesRegex(package_submission.GateError, "stable"):
            package_submission.validate_manifest(manifest)

    def test_rejects_inapplicable_duplicate_form_checksum(self) -> None:
        manifest = valid_manifest()
        manifest["duplicate_benefit_confirmation"]["sha256"] = digest("8")
        with self.assertRaisesRegex(package_submission.GateError, "must be null"):
            package_submission.validate_manifest(manifest)

    def test_rejects_unreviewed_participant_attestation(self) -> None:
        manifest = valid_manifest()
        manifest["participant_attestations"]["single_entry_per_participant_confirmed"] = False
        with self.assertRaisesRegex(package_submission.GateError, "attestations"):
            package_submission.validate_manifest(manifest)

    def test_human_review_attestations_require_literal_true(self) -> None:
        for key in (
            "core_behavior_boundaries_artifacts_and_dependency_roles_reviewed_and_explainable",
            "report_free_text_contains_no_external_evidence_claims",
            "report_free_text_privacy_reviewed",
            "five_year_public_repository_visibility_obligation_if_selected_accepted",
            "owner_voice_ai_assistance_disclosed_and_participant_reviewed",
            "maintenance_order_and_period_confirmed",
            "origin_and_prior_work_statement_confirmed",
        ):
            for value in (False, 1, "true", None):
                manifest = valid_manifest()
                manifest["participant_attestations"][key] = value
                with self.subTest(key=key, value=value), self.assertRaisesRegex(
                    package_submission.GateError, key
                ):
                    package_submission.validate_manifest(manifest)

    def test_rejects_unimplemented_release_signature_claim(self) -> None:
        manifest = valid_manifest()
        manifest["release_evidence"]["signature_filenames"] = [
            "routecontract-shardingsphere-5.5-0.1.0.jar.asc"
        ]
        with self.assertRaisesRegex(package_submission.GateError, "must be empty"):
            package_submission.validate_manifest(manifest)


class LocalVideoEvidenceTest(unittest.TestCase):
    VIDEO = Path("/private/final-video.mp4")

    def probe(self, payload: object) -> dict:
        with patch.object(
            package_submission.shutil,
            "which",
            side_effect=lambda name: "/usr/local/bin/ffprobe"
            if name == "ffprobe"
            else None,
        ), patch.object(
            package_submission,
            "run",
            return_value=json.dumps(payload),
        ) as run:
            result = package_submission.local_video_metadata(self.VIDEO)

        run.assert_called_once_with(
            [
                "/usr/local/bin/ffprobe",
                "-v",
                "error",
                "-f",
                "mov",
                "-protocol_whitelist",
                "file",
                "-show_entries",
                "format=duration:format_tags:stream=index,codec_type,width,height:stream_tags:stream_disposition=default,attached_pic,still_image:chapter=id:chapter_tags:program=id:program_tags",
                "-of",
                "json",
                str(self.VIDEO),
            ],
            failure_label="ffprobe video metadata probe",
        )
        return result

    def test_ffprobe_rejects_duplicate_nonfinite_and_deep_json_generically(self) -> None:
        canary = "CANARY_FFPROBE_PRIVATE_KEY"
        malformed = (
            '{"' + canary + '":1,"' + canary + '":2}',
            '{"value":NaN}',
            '{"value":1e999}',
            PublicJsonTransportTest.nested_json(
                package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
                + 1,
                "object",
            ).decode(),
            "[" * 10_000 + "0" + "]" * 10_000,
            " " * (package_submission.MAX_JSON_TOOL_OUTPUT_BYTES + 1),
        )
        for output in malformed:
            with self.subTest(prefix=output[:24]), patch.object(
                package_submission.shutil, "which", return_value="/usr/local/bin/ffprobe"
            ), patch.object(
                package_submission, "run", return_value=output
            ), self.assertRaisesRegex(
                package_submission.GateError, "incomplete video metadata"
            ) as caught:
                package_submission.local_video_metadata(self.VIDEO)
            self.assertIsNone(caught.exception.__cause__)
            self.assertNotIn(canary, str(caught.exception))

    def test_accepts_1080p_caption_first_video_and_normal_container_tags(self) -> None:
        result = self.probe(valid_ffprobe_video())

        self.assertEqual(173.0, result["duration_seconds"])
        self.assertEqual(1920, result["width"])
        self.assertEqual(1080, result["height"])
        self.assertEqual(1, result["video_stream_count"])
        self.assertEqual(0, result["audio_stream_count"])
        self.assertEqual("ffprobe", result["probe"])

    def test_accepts_larger_silent_video(self) -> None:
        payload = valid_ffprobe_video()
        payload["streams"][0]["width"] = 3840
        payload["streams"][0]["height"] = 2160

        result = self.probe(payload)

        self.assertEqual(3840, result["width"])
        self.assertEqual(2160, result["height"])
        self.assertEqual(0, result["audio_stream_count"])

    def test_uses_default_motion_video_instead_of_larger_alternate(self) -> None:
        payload = valid_ffprobe_video()
        alternate = dict(payload["streams"][0])
        alternate["index"] = 2
        alternate["width"] = 3840
        alternate["height"] = 2160
        alternate["disposition"] = dict(
            payload["streams"][0]["disposition"], default=0
        )
        payload["streams"].append(alternate)

        result = self.probe(payload)

        self.assertEqual(1920, result["width"])
        self.assertEqual(1080, result["height"])
        self.assertEqual(0, result["selected_video_stream_index"])
        self.assertTrue(result["selected_video_is_default"])

    def test_rejects_720p_default_even_with_4k_nondefault_alternate(self) -> None:
        payload = valid_ffprobe_video()
        payload["streams"][0]["width"] = 1280
        payload["streams"][0]["height"] = 720
        alternate = dict(payload["streams"][0])
        alternate["index"] = 2
        alternate["width"] = 3840
        alternate["height"] = 2160
        alternate["disposition"] = dict(
            payload["streams"][0]["disposition"], default=0
        )
        payload["streams"].append(alternate)

        with self.assertRaisesRegex(package_submission.GateError, "at least 1920x1080"):
            self.probe(payload)

    def test_rejects_cover_art_or_still_image_as_resolution_evidence(self) -> None:
        for excluded_disposition in ("attached_pic", "still_image"):
            payload = valid_ffprobe_video()
            payload["streams"][0]["width"] = 1280
            payload["streams"][0]["height"] = 720
            cover = dict(payload["streams"][0])
            cover["index"] = 2
            cover["width"] = 1920
            cover["height"] = 1080
            cover["disposition"] = {
                "default": 0,
                "attached_pic": 0,
                "still_image": 0,
                excluded_disposition: 1,
            }
            payload["streams"].append(cover)

            with self.subTest(disposition=excluded_disposition), self.assertRaisesRegex(
                package_submission.GateError, "at least 1920x1080"
            ):
                self.probe(payload)

    def test_without_default_uses_first_motion_stream_not_larger_secondary(self) -> None:
        payload = valid_ffprobe_video()
        payload["streams"][0]["width"] = 1280
        payload["streams"][0]["height"] = 720
        payload["streams"][0]["disposition"]["default"] = 0
        secondary = dict(payload["streams"][0])
        secondary["index"] = 2
        secondary["width"] = 3840
        secondary["height"] = 2160
        secondary["disposition"] = dict(
            payload["streams"][0]["disposition"]
        )
        payload["streams"].append(secondary)

        with self.assertRaisesRegex(package_submission.GateError, "at least 1920x1080"):
            self.probe(payload)

    def test_rejects_multiple_default_motion_video_streams(self) -> None:
        payload = valid_ffprobe_video()
        secondary = dict(payload["streams"][0])
        secondary["index"] = 2
        secondary["disposition"] = dict(
            payload["streams"][0]["disposition"], default=1
        )
        payload["streams"].append(secondary)

        with self.assertRaisesRegex(package_submission.GateError, "multiple default"):
            self.probe(payload)

    def test_rejects_missing_or_duplicate_motion_video_stream_index(self) -> None:
        payload = valid_ffprobe_video()
        del payload["streams"][0]["index"]
        with self.assertRaisesRegex(package_submission.GateError, "stream indexes"):
            self.probe(payload)

        payload = valid_ffprobe_video()
        secondary = dict(payload["streams"][0])
        secondary["disposition"] = dict(
            payload["streams"][0]["disposition"], default=0
        )
        payload["streams"].append(secondary)
        with self.assertRaisesRegex(package_submission.GateError, "stream indexes"):
            self.probe(payload)

    def test_rejects_only_cover_art_or_still_image_video(self) -> None:
        for excluded_disposition in ("attached_pic", "still_image"):
            payload = valid_ffprobe_video()
            payload["streams"][0]["disposition"][excluded_disposition] = 1
            with self.subTest(disposition=excluded_disposition), self.assertRaisesRegex(
                package_submission.GateError, "no playable motion video stream"
            ):
                self.probe(payload)

    def test_rejects_missing_or_malformed_video_disposition(self) -> None:
        for disposition in (None, {}, {"default": False, "attached_pic": 0, "still_image": 0}):
            payload = valid_ffprobe_video()
            if disposition is None:
                del payload["streams"][0]["disposition"]
            else:
                payload["streams"][0]["disposition"] = disposition
            with self.subTest(disposition=disposition), self.assertRaisesRegex(
                package_submission.GateError, "video stream disposition"
            ):
                self.probe(payload)

    def test_requires_ffprobe_instead_of_incomplete_mdls_fallback(self) -> None:
        with patch.object(package_submission.shutil, "which", return_value=None):
            with self.assertRaisesRegex(package_submission.GateError, "ffprobe is required"):
                package_submission.local_video_metadata(self.VIDEO)

    def test_rejects_invalid_or_incomplete_ffprobe_output(self) -> None:
        invalid_payloads = (
            "not-an-object",
            {},
            {"format": {}, "streams": []},
            {"format": {"duration": "NaN"}, "streams": []},
            {"format": {"duration": "179"}, "streams": "not-a-list"},
        )
        for payload in invalid_payloads:
            with self.subTest(payload=payload), self.assertRaisesRegex(
                package_submission.GateError, "ffprobe returned incomplete video metadata"
            ):
                self.probe(payload)

    def test_ffprobe_schema_diagnostics_never_echo_json_values(self) -> None:
        canary = "CANARY_FFPROBE_PRIVATE_VALUE"
        cases = []
        invalid_duration = valid_ffprobe_video()
        invalid_duration["format"]["duration"] = canary
        cases.append(invalid_duration)
        sensitive_stream_tag = valid_ffprobe_video()
        sensitive_stream_tag["streams"][0]["index"] = canary
        sensitive_stream_tag["streams"][0]["tags"]["author"] = "private"
        cases.append(sensitive_stream_tag)
        oversized_integer = valid_ffprobe_video()
        oversized_integer["format"]["duration"] = int("9" * 1_000)
        cases.append(oversized_integer)
        for payload in cases:
            with self.subTest(payload=payload), self.assertRaises(
                package_submission.GateError
            ) as caught:
                self.probe(payload)
            self.assertIsNone(caught.exception.__cause__)
            self.assertNotIn(canary, str(caught.exception))
            self.assertNotIn("9" * 32, str(caught.exception))

    def test_rejects_any_audio_stream(self) -> None:
        payload = valid_ffprobe_video()
        payload["streams"].append(
            {
                "index": 1,
                "codec_type": "audio",
                "tags": {
                    "language": "kor",
                    "handler_name": "SoundHandler",
                },
            }
        )

        with self.assertRaisesRegex(package_submission.GateError, "no audio streams"):
            self.probe(payload)

    def test_rejects_missing_video_stream(self) -> None:
        payload = valid_ffprobe_video()
        payload["streams"] = [
            {
                "index": 1,
                "codec_type": "audio",
                "tags": {
                    "language": "kor",
                    "handler_name": "SoundHandler",
                },
            }
        ]

        with self.assertRaisesRegex(package_submission.GateError, "no video stream"):
            self.probe(payload)

    def test_rejects_video_below_1920_by_1080(self) -> None:
        for width, height in ((1919, 1080), (1920, 1079), (1280, 720)):
            payload = valid_ffprobe_video()
            payload["streams"][0]["width"] = width
            payload["streams"][0]["height"] = height
            with self.subTest(width=width, height=height), self.assertRaisesRegex(
                package_submission.GateError, "at least 1920x1080"
            ):
                self.probe(payload)

    def test_rejects_malformed_video_dimensions(self) -> None:
        for width, height in (("1920", 1080), (1920, None), (True, 1080)):
            payload = valid_ffprobe_video()
            payload["streams"][0]["width"] = width
            payload["streams"][0]["height"] = height
            with self.subTest(width=width, height=height), self.assertRaisesRegex(
                package_submission.GateError, "dimensions"
            ):
                self.probe(payload)

    def test_rejects_every_explicit_sensitive_tag_at_format_and_stream_scope(
        self,
    ) -> None:
        for tag in package_submission.SENSITIVE_VIDEO_METADATA_TAGS:
            for scope in ("format", "stream"):
                payload = valid_ffprobe_video()
                target = (
                    payload["format"]["tags"]
                    if scope == "format"
                    else payload["streams"][0]["tags"]
                )
                target[f" {tag.swapcase()} "] = "private fixture value"
                with self.subTest(tag=tag, scope=scope), self.assertRaisesRegex(
                    package_submission.GateError, "sensitive metadata tag"
                ):
                    self.probe(payload)

    def test_rejects_normalized_gps_and_quicktime_location_namespace_tags(self) -> None:
        for tag in (
            " GPSCOORDINATES ",
            " com.apple.quicktime.location.name ",
            "COM.APPLE.QUICKTIME.LOCATION.BODY",
            "com.apple.quicktime.location.role",
        ):
            payload = valid_ffprobe_video()
            payload["format"]["tags"][tag] = "private location fixture"
            with self.subTest(tag=tag), self.assertRaisesRegex(
                package_submission.GateError, "sensitive metadata tag"
            ):
                self.probe(payload)

    def test_does_not_broadly_reject_non_location_namespace_tag_names(self) -> None:
        payload = valid_ffprobe_video()
        payload["format"]["tags"].update(
            {
                "application_name": "RouteContract",
                "encoder_location_hint": "portable encoder fixture",
                "com.apple.quicktime.displayname": "RouteContract demo",
            }
        )

        result = self.probe(payload)

        self.assertEqual(6, result["metadata_tag_count"])

    def test_rejects_sensitive_chapter_and_program_metadata_tags(self) -> None:
        for scope in ("chapters", "programs"):
            payload = valid_ffprobe_video()
            payload[scope] = [{"id": 1, "tags": {"author": "private fixture"}}]
            with self.subTest(scope=scope), self.assertRaisesRegex(
                package_submission.GateError, "sensitive metadata tag"
            ):
                self.probe(payload)

    def test_rejects_malformed_tag_objects(self) -> None:
        for scope in ("format", "stream"):
            payload = valid_ffprobe_video()
            if scope == "format":
                payload["format"]["tags"] = ["encoder"]
            else:
                payload["streams"][0]["tags"] = ["language"]
            with self.subTest(scope=scope), self.assertRaisesRegex(
                package_submission.GateError, "metadata tags"
            ):
                self.probe(payload)

    def test_rejects_private_path_or_email_in_otherwise_allowed_tag(self) -> None:
        for leaked_value in (
            "/Users/person/private/final.mov",
            "owner@example.com",
        ):
            payload = valid_ffprobe_video()
            payload["format"]["tags"]["encoder"] = leaked_value
            with self.subTest(value=leaked_value), self.assertRaisesRegex(
                package_submission.GateError, "private path/identity metadata"
            ):
                self.probe(payload)

    def test_full_decode_requires_ffmpeg(self) -> None:
        with patch.object(package_submission.shutil, "which", return_value=None), \
            self.assertRaisesRegex(package_submission.GateError, "ffmpeg is required"):
            package_submission.validate_full_motion_video_decode(self.VIDEO, 0, 173.0)

    def test_full_decode_maps_selected_stream_and_requires_completion(self) -> None:
        output = (
            "frame=1\nout_time_us=1000\nprogress=continue\n"
            "frame=5190\nout_time_us=173000000\nprogress=end\n"
        )
        with patch.object(
            package_submission.shutil,
            "which",
            return_value="/usr/local/bin/ffmpeg",
        ), patch.object(package_submission, "run", return_value=output) as run:
            result = package_submission.validate_full_motion_video_decode(
                self.VIDEO, 7, 173.0
            )

        run.assert_called_once_with(
            [
                "/usr/local/bin/ffmpeg",
                "-hide_banner",
                "-nostdin",
                "-v",
                "error",
                "-xerror",
                "-err_detect",
                "explode",
                "-f",
                "mov",
                "-protocol_whitelist",
                "file",
                "-progress",
                "pipe:1",
                "-nostats",
                "-i",
                str(self.VIDEO),
                "-map",
                "0:7",
                "-an",
                "-sn",
                "-dn",
                "-f",
                "null",
                "-",
            ],
            timeout_seconds=600,
            failure_label="ffmpeg full motion-video decode",
        )
        self.assertEqual(
            {
                "decode_probe": "ffmpeg",
                "decoded_frame_count": 5190,
                "decoded_duration_seconds": 173.0,
                "minimum_frame_count": 3460,
            },
            result,
        )

    def test_full_decode_rejects_zero_frames_or_incomplete_progress(self) -> None:
        for output in (
            "frame=0\nout_time_us=173000000\nprogress=end\n",
            "frame=10\nout_time_us=1000000\nprogress=continue\n",
            "progress=end\n",
        ):
            with self.subTest(output=output), patch.object(
                package_submission.shutil,
                "which",
                return_value="/usr/local/bin/ffmpeg",
            ), patch.object(
                package_submission, "run", return_value=output
            ), self.assertRaisesRegex(
                package_submission.GateError, "did not complete"
            ):
                package_submission.validate_full_motion_video_decode(
                    self.VIDEO, 0, 173.0
                )

    def test_full_decode_binds_duration_and_minimum_frame_count(self) -> None:
        cases = (
            (
                "frame=5190\nout_time_us=1000000\nprogress=end\n",
                173.0,
                "duration differs",
            ),
            (
                "frame=1\nout_time_us=173000000\nprogress=end\n",
                173.0,
                "frame count is too low",
            ),
            (
                "frame=3400\nout_time_us=170010000\nprogress=end\n",
                170.01,
                "frame count is too low",
            ),
        )
        for output, duration, message in cases:
            with self.subTest(duration=duration, message=message), patch.object(
                package_submission.shutil,
                "which",
                return_value="/usr/local/bin/ffmpeg",
            ), patch.object(
                package_submission, "run", return_value=output
            ), self.assertRaisesRegex(package_submission.GateError, message):
                package_submission.validate_full_motion_video_decode(
                    self.VIDEO, 0, duration
                )

    def test_validate_local_video_preserves_hash_and_duration_gates(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            video = Path(raw) / "final.mp4"
            video.write_bytes(b"final video fixture")
            manifest = valid_manifest()
            manifest["video"]["local_file_sha256"] = hashlib.sha256(
                video.read_bytes()
            ).hexdigest()
            manifest = package_submission.validate_manifest(manifest)

            with patch.object(
                package_submission,
                "local_video_metadata",
                return_value={
                    "duration_seconds": 173.0,
                    "width": 1920,
                    "height": 1080,
                    "video_stream_count": 1,
                    "audio_stream_count": 0,
                    "selected_video_stream_index": 0,
                    "probe": "ffprobe",
                },
            ), patch.object(
                package_submission,
                "validate_full_motion_video_decode",
                return_value={
                    "decode_probe": "ffmpeg",
                    "decoded_frame_count": 5190,
                    "decoded_duration_seconds": 173.0,
                    "minimum_frame_count": 3460,
                },
            ) as decode:
                accepted = package_submission.validate_local_video(video, manifest)
            decode.assert_called_once_with(video, 0, 173.0)
            self.assertEqual(manifest["video"]["local_file_sha256"], accepted["sha256"])
            self.assertEqual(5190, accepted["decoded_frame_count"])

            changed = package_submission.validate_manifest(valid_manifest())
            with self.assertRaisesRegex(package_submission.GateError, "SHA-256 mismatch"):
                package_submission.validate_local_video(video, changed)

            with patch.object(
                package_submission,
                "local_video_metadata",
                return_value={"duration_seconds": 175.1},
            ), self.assertRaisesRegex(
                package_submission.GateError, "170 through 175 seconds inclusive"
            ):
                package_submission.validate_local_video(video, manifest)

            with patch.object(
                package_submission,
                "local_video_metadata",
                return_value={"duration_seconds": 172.6},
            ), self.assertRaisesRegex(package_submission.GateError, "differs from manifest"):
                package_submission.validate_local_video(video, manifest)

    def test_validate_local_video_duration_window_and_caption_end_are_inclusive(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            video = Path(raw) / "final.mp4"
            video.write_bytes(b"final video fixture")
            expected_hash = hashlib.sha256(video.read_bytes()).hexdigest()

            for duration in (172.5, 175.0):
                manifest = package_submission.validate_manifest(valid_manifest())
                manifest["video"]["duration_seconds"] = duration
                manifest["video"]["local_file_sha256"] = expected_hash
                with self.subTest(duration=duration), patch.object(
                    package_submission,
                    "local_video_metadata",
                    return_value={
                        "duration_seconds": duration,
                        "selected_video_stream_index": 0,
                    },
                ), patch.object(
                    package_submission,
                    "validate_full_motion_video_decode",
                    return_value={
                        "decode_probe": "ffmpeg",
                        "decoded_frame_count": int(duration * 20),
                        "decoded_duration_seconds": duration,
                        "minimum_frame_count": int(duration * 20),
                    },
                ):
                    package_submission.validate_local_video(video, manifest)

            for duration in (169.9, 175.1):
                manifest = package_submission.validate_manifest(valid_manifest())
                manifest["video"]["duration_seconds"] = duration
                manifest["video"]["local_file_sha256"] = expected_hash
                with self.subTest(duration=duration), patch.object(
                    package_submission,
                    "local_video_metadata",
                    return_value={"duration_seconds": duration},
                ), self.assertRaisesRegex(
                    package_submission.GateError,
                    "170 through 175 seconds inclusive",
                ):
                    package_submission.validate_local_video(video, manifest)

            for duration in (170.0, 172.499):
                manifest = package_submission.validate_manifest(valid_manifest())
                manifest["video"]["duration_seconds"] = duration
                manifest["video"]["local_file_sha256"] = expected_hash
                with self.subTest(duration=duration), patch.object(
                    package_submission,
                    "local_video_metadata",
                    return_value={"duration_seconds": duration},
                ), self.assertRaisesRegex(
                    package_submission.GateError,
                    "selected caption branch's final cue",
                ):
                    package_submission.validate_local_video(video, manifest)

    def test_validate_local_video_rejects_change_during_decode(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        expected = manifest["video"]["local_file_sha256"]
        with patch.object(
            package_submission,
            "sha256",
            side_effect=[expected, "f" * 64],
        ), patch.object(
            package_submission,
            "local_video_metadata",
            return_value={
                "duration_seconds": 173.0,
                "selected_video_stream_index": 0,
            },
        ), patch.object(
            package_submission,
            "validate_full_motion_video_decode",
            return_value={
                "decode_probe": "ffmpeg",
                "decoded_frame_count": 3460,
                "decoded_duration_seconds": 173.0,
                "minimum_frame_count": 3460,
            },
        ), self.assertRaisesRegex(
            package_submission.GateError, "changed during validation"
        ):
            package_submission.validate_local_video(self.VIDEO, manifest)


class GitStateTest(unittest.TestCase):
    COMMIT = "1" * 40
    LOCAL_TAG_OBJECT = "a" * 40

    def validate_with_remote_refs(self, root: Path, remote_refs: str) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        (root / "build.gradle").write_text(
            "group = 'io.github.example-owner.routecontract'\n"
            "version = '0.1.0'\n",
            encoding="utf-8",
        )

        def git_run(
            command: list[str],
            *,
            cwd: Path | None = None,
            env: dict[str, str] | None = None,
        ) -> str:
            self.assertEqual(root, cwd)
            if command == ["git", "rev-parse", "--show-toplevel"]:
                return f"{root}\n"
            if command == ["git", "rev-parse", "HEAD"]:
                return f"{self.COMMIT}\n"
            if command == [
                "git",
                "status",
                "--porcelain=v1",
                "--untracked-files=all",
            ]:
                return ""
            if command == ["git", "remote", "get-url", "origin"]:
                return "https://github.com/example-owner/routecontract.git\n"
            if command == ["git", "cat-file", "-t", "refs/tags/v0.1.0"]:
                return "tag\n"
            if command == ["git", "rev-parse", "refs/tags/v0.1.0^{commit}"]:
                return f"{self.COMMIT}\n"
            if command == ["git", "rev-parse", "refs/tags/v0.1.0"]:
                return f"{self.LOCAL_TAG_OBJECT}\n"
            if command == [
                "git",
                "ls-remote",
                "--tags",
                "origin",
                "refs/tags/v0.1.0",
                "refs/tags/v0.1.0^{}",
            ]:
                self.assertEqual("0", env.get("GIT_TERMINAL_PROMPT") if env else None)
                return remote_refs
            self.fail(f"unexpected command: {command}")

        with patch.object(package_submission, "run", side_effect=git_run):
            package_submission.validate_git_state(root, manifest)

    def test_accepts_exact_remote_annotated_tag_object_and_peel(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw).resolve()
            self.validate_with_remote_refs(
                root,
                f"{self.LOCAL_TAG_OBJECT}\trefs/tags/v0.1.0\n"
                f"{self.COMMIT}\trefs/tags/v0.1.0^{{}}\n",
            )

    def test_rejects_remote_lightweight_tag_at_same_commit(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw).resolve()
            with self.assertRaisesRegex(package_submission.GateError, "exact annotated"):
                self.validate_with_remote_refs(
                    root,
                    f"{self.COMMIT}\trefs/tags/v0.1.0\n",
                )

    def test_rejects_different_remote_tag_object_with_same_peel(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw).resolve()
            with self.assertRaisesRegex(package_submission.GateError, "tag object"):
                self.validate_with_remote_refs(
                    root,
                    f"{'b' * 40}\trefs/tags/v0.1.0\n"
                    f"{self.COMMIT}\trefs/tags/v0.1.0^{{}}\n",
                )


class ReportContractTest(unittest.TestCase):
    BASE = (
        "2026년 오픈소스 개발자대회 결과보고서 "
        "개발 보조 AI OpenAI ChatGPT 및 Codex "
        f"{package_submission.NO_RUNTIME_AI_DISCLOSURE} "
        "붙임1 SBOM(소프트웨어 자재명세서)"
    )
    PORTRAIT = (595.0, 842.0)
    LANDSCAPE = (842.0, 595.0)
    SBOM_ROWS = [
        {"name": "Alpha", "version": "1.0"},
        {"name": "Omega", "version": "2.0"},
    ]

    def test_visible_release_identity_binds_full_coordinate_and_commit(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        coordinate = (
            "io.github.example-owner.routecontract:"
            "routecontract-shardingsphere-5.5:0.1.0"
        )
        commit = manifest["project"]["commit"]
        package_submission.validate_report_release_identity(
            f"설치 {coordinate} 최종 revision {commit}", manifest
        )

        cases = (
            (
                "artifact",
                f"설치 io.github.example-owner.routecontract:wrong:0.1.0 {commit}",
                "Maven install coordinate",
            ),
            (
                "version",
                "설치 io.github.example-owner.routecontract:"
                f"routecontract-shardingsphere-5.5:0.1.1 {commit}",
                "Maven install coordinate",
            ),
            (
                "commit",
                f"설치 {coordinate} {'2' * 40}",
                "final commit SHA",
            ),
        )
        for label, text, message in cases:
            with self.subTest(label=label), self.assertRaisesRegex(
                package_submission.GateError, message
            ):
                package_submission.validate_report_release_identity(text, manifest)

    def test_sbom_visible_text_uses_required_ten_point_body_font(self) -> None:
        self.assertEqual(10, build_official_report.BODY_FONT_PT)
        self.assertEqual(
            build_official_report.BODY_FONT_PT,
            build_official_report.SBOM_FONT_PT,
        )
        self.assertEqual(11.2, build_official_report.SBOM_LINE_SPACING_PT)

    def test_fill_sbom_applies_ten_point_visible_text_and_exact_line_spacing(
        self,
    ) -> None:
        class FakeFormat:
            pass

        class FakeRun:
            def __init__(self, paragraph, text: str) -> None:
                self.paragraph = paragraph
                self.text = text

        class FakeParagraph:
            def __init__(self) -> None:
                self.alignment = None
                self.paragraph_format = FakeFormat()
                self.runs = []

            def add_run(self, text: str = "") -> FakeRun:
                run = FakeRun(self, text)
                self.runs.append(run)
                return run

        class FakeCell:
            def __init__(self) -> None:
                self.text = ""
                self.paragraphs = [FakeParagraph()]
                self.vertical_alignment = None

        class FakeRow:
            def __init__(self) -> None:
                self.cells = [FakeCell() for _ in range(6)]

        class FakeTable:
            def __init__(self) -> None:
                self.rows = [FakeRow() for _ in range(11)]
                self.columns = [object() for _ in range(6)]

            def cell(self, row: int, column: int) -> FakeCell:
                return self.rows[row].cells[column]

        class FakeParagraphAlignment:
            LEFT = object()
            CENTER = object()

        class FakeVerticalAlignment:
            CENTER = object()

        rows = [
            {
                "name": f"component-{index}",
                "version": f"{index}.0",
                "license": "Apache-2.0",
                "url": f"https://example.test/component-{index}",
                "purpose": f"purpose-{index}",
            }
            for index in range(1, 11)
        ]
        table = FakeTable()
        # python-docx serializes a Pt length as exact line spacing.
        exact_11_2_pt = ("pt", 11.2)

        with (
            patch.object(build_official_report, "find_table", return_value=table),
            patch.object(build_official_report, "set_cell_margins"),
            patch.object(build_official_report, "mark_repeat_header"),
            patch.object(build_official_report, "mark_data_rows_cannot_split"),
            patch.object(build_official_report, "remove_fixed_row_heights"),
            patch.object(
                build_official_report,
                "WD_ALIGN_PARAGRAPH",
                new=FakeParagraphAlignment,
                create=True,
            ),
            patch.object(
                build_official_report,
                "WD_CELL_VERTICAL_ALIGNMENT",
                new=FakeVerticalAlignment,
                create=True,
            ),
            patch.object(
                build_official_report,
                "Pt",
                side_effect=lambda value: ("pt", value),
                create=True,
            ),
            patch.object(build_official_report, "set_run_font") as plain_font,
            patch.object(build_official_report, "append_hyperlink") as hyperlink,
        ):
            build_official_report.fill_sbom(object(), rows)

        data_paragraphs = {
            cell.paragraphs[0]
            for row in table.rows[1:]
            for cell in row.cells
        }
        self.assertEqual(60, len(data_paragraphs))
        self.assertEqual(50, plain_font.call_count)
        self.assertEqual(10, hyperlink.call_count)

        plain_paragraphs = set()
        for invocation in plain_font.call_args_list:
            run, size_pt = invocation.args
            self.assertEqual(10, size_pt)
            self.assertEqual({"bold": False}, invocation.kwargs)
            plain_paragraphs.add(run.paragraph)

        actual_hyperlinks = {
            tuple(invocation.args) for invocation in hyperlink.call_args_list
        }
        expected_hyperlinks = {
            (
                table.cell(index, 4).paragraphs[0],
                item["url"],
                item["url"],
                10,
            )
            for index, item in enumerate(rows, start=1)
        }
        self.assertEqual(expected_hyperlinks, actual_hyperlinks)

        hyperlink_paragraphs = {
            invocation.args[0] for invocation in hyperlink.call_args_list
        }
        self.assertFalse(plain_paragraphs & hyperlink_paragraphs)
        self.assertEqual(
            data_paragraphs,
            plain_paragraphs | hyperlink_paragraphs,
        )
        for paragraph in data_paragraphs:
            self.assertEqual(
                exact_11_2_pt,
                paragraph.paragraph_format.line_spacing,
            )

    def test_report_text_blocks_do_not_split_across_pages(self) -> None:
        class FakeFormat:
            pass

        class FakeRun:
            pass

        class FakeParagraph:
            def __init__(self) -> None:
                self.alignment = None
                self.paragraph_format = FakeFormat()

            def add_run(self, _text: str = "") -> FakeRun:
                return FakeRun()

        class FakeCell:
            def __init__(self) -> None:
                self.text = ""
                self.paragraphs = [FakeParagraph()]
                self.vertical_alignment = None

            def add_paragraph(self) -> FakeParagraph:
                paragraph = FakeParagraph()
                self.paragraphs.append(paragraph)
                return paragraph

        class FakeVerticalAlignment:
            TOP = object()

        cell = FakeCell()
        with (
            patch.object(
                build_official_report, "Pt", new=lambda value: value, create=True
            ),
            patch.object(
                build_official_report,
                "WD_CELL_VERTICAL_ALIGNMENT",
                new=FakeVerticalAlignment,
                create=True,
            ),
            patch.object(build_official_report, "set_run_font"),
            patch.object(build_official_report, "append_text_with_hyperlinks"),
        ):
            build_official_report.set_block_cell(
                cell,
                [
                    {
                        "lead": "첫 문단",
                        "text": "한 문단은 페이지 사이에서 분리되지 않는다.",
                    },
                    {
                        "lead": "둘째 문단",
                        "text": "다음 문단도 같은 규칙을 따른다.",
                    },
                ],
            )

        self.assertEqual(2, len(cell.paragraphs))
        for paragraph in cell.paragraphs:
            self.assertIs(True, paragraph.paragraph_format.keep_together)

    def test_report_other_continuation_row_is_kept_whole(self) -> None:
        class FakeTable:
            def __init__(self) -> None:
                self.rows = [object() for _ in range(12)]

            @staticmethod
            def cell(row: int, column: int) -> tuple[int, int]:
                return row, column

        table = FakeTable()
        blocks = [
            {"lead": "차별성", "text": "앞 행"},
            {"lead": "품질관리·발전 로드맵", "text": "앞 행"},
            {"lead": "오픈소스SW 조합", "text": "앞 행"},
            {"lead": "현재 한계", "text": "계속 행"},
            {"lead": "공개 증거 gate", "text": "계속 행"},
        ]

        def append_row(target, _row) -> None:
            target.rows.append(object())

        with (
            patch.object(build_official_report, "copy_row", side_effect=append_row),
            patch.object(build_official_report, "set_plain_cell") as plain,
            patch.object(build_official_report, "set_block_cell") as block,
            patch.object(
                build_official_report, "mark_data_rows_cannot_split"
            ) as marker,
        ):
            build_official_report.fill_other_report_rows(
                table,
                blocks,
                hyperlink_targets={"https://example.test/evidence"},
                hyperlink_aliases={"[증거]": "https://example.test/evidence"},
            )

        self.assertEqual(13, len(table.rows))
        self.assertEqual(blocks[:1], block.call_args_list[0].args[1])
        self.assertEqual(blocks[1:], block.call_args_list[1].args[1])
        plain.assert_called_once_with((12, 0), "")
        marker.assert_called_once_with([table.rows[12]])

    def test_report_other_continuation_requires_the_source_to_be_last(self) -> None:
        class FakeTable:
            def __init__(self) -> None:
                self.rows = [object() for _ in range(13)]

        with self.assertRaisesRegex(
            ValueError, "source row must be the final row"
        ):
            build_official_report.fill_other_report_rows(
                FakeTable(),
                [
                    {"lead": "차별성", "text": "앞 행"},
                    {"lead": "품질관리·발전 로드맵", "text": "앞 행"},
                    {"lead": "오픈소스SW 조합", "text": "앞 행"},
                    {"lead": "현재 한계", "text": "계속 행"},
                ],
                original_row_index=11,
            )

    def test_report_installation_has_an_unsplittable_continuation_row(self) -> None:
        from docx import Document

        table = Document().add_table(rows=12, cols=2)
        blocks = [
            {"lead": f"기능 {index}", "text": f"본문 {index}"}
            for index in range(6)
        ] + [{"lead": "설치·릴리스", "text": "검증된 stable 설치 경로"}]

        with (
            patch.object(build_official_report, "set_plain_cell") as plain,
            patch.object(build_official_report, "set_block_cell") as block,
            patch.object(
                build_official_report, "mark_data_rows_cannot_split"
            ) as marker,
        ):
            build_official_report.fill_feature_report_rows(
                table,
                blocks,
                image_path=Path("diagram.png"),
                image_caption="검증 흐름",
                hyperlink_targets={"https://example.test/evidence"},
            )

        self.assertEqual(13, len(table.rows))
        self.assertEqual(blocks[:6], block.call_args_list[0].args[1])
        self.assertEqual(blocks[6:], block.call_args_list[1].args[1])
        self.assertIs(plain.call_args.args[0]._tc, table.cell(10, 0)._tc)
        plain.assert_called_once_with(plain.call_args.args[0], "")
        marker.assert_called_once()
        self.assertIs(marker.call_args.args[0][0]._tr, table.rows[10]._tr)

    def test_report_effects_row_does_not_split_across_pages(self) -> None:
        class FakeTable:
            def __init__(self) -> None:
                self.rows = [object() for _ in range(12)]

        table = FakeTable()
        with patch.object(
            build_official_report, "mark_data_rows_cannot_split"
        ) as marker:
            build_official_report.mark_effects_row_cannot_split(table)

        marker.assert_called_once_with([table.rows[11]])

    def test_report_environment_row_does_not_split_across_pages(self) -> None:
        class FakeTable:
            def __init__(self) -> None:
                self.rows = [object() for _ in range(8)]

        table = FakeTable()
        with patch.object(
            build_official_report, "mark_data_rows_cannot_split"
        ) as marker:
            build_official_report.mark_environment_row_cannot_split(table)

        marker.assert_called_once_with([table.rows[7]])

    def test_report_main_table_structure_is_fail_closed(self) -> None:
        from docx import Document
        from docx.oxml import OxmlElement

        table = Document().add_table(rows=14, cols=2)
        for index, label in {
            9: "프로젝트 주요기능",
            10: "",
            11: "기대효과 및 활용분야",
            12: "기타",
            13: "",
        }.items():
            table.cell(index, 0).text = label
        table.cell(10, 1).text = "설치·릴리스: 검증된 stable 설치 경로"
        for index in (7, 10, 11, 13):
            table.rows[index]._tr.get_or_add_trPr().append(
                OxmlElement("w:cantSplit")
            )

        build_official_report.assert_main_report_table_structure(table)

        table.cell(11, 0).text = "잘못된 라벨"
        with self.assertRaisesRegex(ValueError, "continuation row labels changed"):
            build_official_report.assert_main_report_table_structure(table)

    def test_cloned_report_rows_drop_volatile_identity_attributes(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        table = Document().add_table(rows=1, cols=2)
        source = table.rows[0]
        source._tr.set(qn("w:rsidR"), "00112233")
        source.cells[0].paragraphs[0]._p.set(qn("w14:paraId"), "AABBCCDD")
        source.cells[0].paragraphs[0]._p.set(qn("w14:textId"), "EEFF0011")

        cloned = build_official_report.insert_row_after(table, source)

        volatile = []
        for element in cloned._tr.iter():
            volatile.extend(
                attribute
                for attribute in element.attrib
                if attribute.rsplit("}", 1)[-1].startswith("rsid")
                or attribute.rsplit("}", 1)[-1] in {"paraId", "textId"}
            )
        self.assertEqual([], volatile)

    def test_report_story_sanitizer_removes_revision_attributes_and_elements(self) -> None:
        build_official_report.load_document_dependencies()
        source = b'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
  <w:body><w:p w:rsidR="00112233" w14:paraId="AABBCCDD">
    <w:r w:rsidRPr="44556677"><w:t>visible</w:t></w:r>
    <w:rsids><w:rsidRoot w:val="00112233"/><w:rsid w:val="44556677"/></w:rsids>
  </w:p></w:body>
</w:document>'''

        sanitized = build_official_report.sanitize_story_part(source)

        self.assertNotIn(b"rsid", sanitized)
        self.assertIn(b"paraId", sanitized)
        self.assertIn(b"visible", sanitized)

    def test_report_privacy_part_selector_includes_settings_and_styles(self) -> None:
        for name in ("word/document.xml", "word/settings.xml", "word/styles.xml"):
            with self.subTest(name=name):
                self.assertTrue(
                    build_official_report.is_revision_identifier_part(name)
                )
        self.assertFalse(
            build_official_report.is_revision_identifier_part("word/numbering.xml")
        )

    def test_report_package_restore_scrubs_all_revision_identifier_shapes(self) -> None:
        build_official_report.load_document_dependencies()
        core = b'''<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties
 xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/">
  <dc:creator>template author</dc:creator>
  <cp:lastModifiedBy>template editor</cp:lastModifiedBy>
  <cp:revision>7</cp:revision>
</cp:coreProperties>'''
        document = b'''<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p w:rsidR="00112233"><w:r><w:t>visible</w:t></w:r></w:p></w:body>
</w:document>'''
        settings = b'''<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:zoom w:percent="100"/>
  <w:rsids><w:rsidRoot w:val="00112233"/><w:rsid w:val="44556677"/></w:rsids>
</w:settings>'''
        styles = b'''<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:rsidR="00112233">
    <w:name w:val="Normal"/><w:rsid w:val="44556677"/>
  </w:style>
</w:styles>'''
        parts = {
            "[Content_Types].xml": b"<Types/>",
            "docProps/core.xml": core,
            "word/document.xml": document,
            "word/settings.xml": settings,
            "word/styles.xml": styles,
        }

        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            template = root / "template.docx"
            output = root / "report.docx"
            for path in (template, output):
                with ZipFile(path, "w", ZIP_DEFLATED) as package:
                    for name, content in parts.items():
                        package.writestr(name, content)

            build_official_report.restore_preserve_only_package_parts(
                template, output
            )

            with ZipFile(output) as package:
                for name in (
                    "word/document.xml",
                    "word/settings.xml",
                    "word/styles.xml",
                ):
                    with self.subTest(name=name):
                        self.assertNotIn(b"rsid", package.read(name))
                self.assertIn(b"visible", package.read("word/document.xml"))
                self.assertIn(b"zoom", package.read("word/settings.xml"))
                self.assertIn(b"Normal", package.read("word/styles.xml"))
            package_submission.validate_docx_privacy(output)

    def test_local_package_metadata_uses_current_schema(self) -> None:
        self.assertEqual(4, package_submission.PACKAGE_METADATA_SCHEMA_VERSION)

    def test_direct_report_validator_rejects_unsupported_final_stable_branch(self) -> None:
        with self.assertRaisesRegex(
            package_submission.GateError, "only rc_only or zero"
        ):
            package_submission.validate_report(
                Path("unsupported.docx"),
                Path("unsupported.pdf"),
                valid_report_content("final_stable"),
                valid_manifest(),
            )

    @staticmethod
    def visible_binding_fixture(content: dict) -> tuple[str, str]:
        metadata_values = set(content["metadata"].values())
        captions = [
            asset["caption"] for asset in content["assets"].values()
        ]
        caption_values = set(captions)
        sbom_values = {
            value
            for row in content["sbom"]
            for value in row.values()
        }
        ordinary = [
            value
            for value in package_submission.visible_report_values(content)
            if value not in sbom_values
            and value not in metadata_values
            and value not in caption_values
        ]
        metadata_rows = [
            f"팀명 {content['metadata']['team_name']}",
            f"팀인원 {content['metadata']['team_size']}",
            f"참가부문 {content['metadata']['division']}",
            f"과제유형 {content['metadata']['task_type']}",
            f"프로젝트명 {content['metadata']['project_name']}",
            f"프로젝트등록 {content['metadata']['repository_url']}",
            f"시연영상 {content['metadata']['video_url']}",
        ]
        pdf_metadata_rows = [
            f"팀명 {content['metadata']['team_name']} 팀인원 {content['metadata']['team_size']}",
            f"참가부문 {content['metadata']['division']} 과제유형 {content['metadata']['task_type']}",
            f"프로젝트명 {content['metadata']['project_name']}",
            f"프로젝트등록 {content['metadata']['repository_url']}",
            f"시연영상 {content['metadata']['video_url']}",
        ]
        title = "붙임1 SBOM(소프트웨어 자재명세서)"
        header = (
            f"{'번호':<8}{'라이브러리명':<36}{'버전':<24}{'라이선스':<20}"
            "공식 저장소 URL(GitHub 등) 사용 목적 및 주요 기능"
        )
        rows = [
            (
                f"{index:<8}{row['name']:<36}{row['version']:<24}"
                f"{row['license']:<20}{row['url']} {row['purpose']}"
            )
            for index, row in enumerate(content["sbom"], start=1)
        ]
        # Both strings have the exact same visible characters.  Only whitespace
        # and, in individual tests, bounded table-cell interleaving may differ.
        docx_text = "\n".join(
            [*metadata_rows, *captions, *ordinary, title, header, *rows]
        )
        pdf_text = "\n".join(
            [*pdf_metadata_rows, *captions, *ordinary, title, header, *rows]
        )
        return docx_text, pdf_text

    @staticmethod
    def write_docx(path: Path, body_xml: str) -> None:
        document = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/'
            'wordprocessingml/2006/main">'
            f"<w:body>{body_xml}</w:body></w:document>"
        )
        content_types = (
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.document.main+xml"/>'
            "</Types>"
        )
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types)
            archive.writestr("word/document.xml", document)

    def test_docx_extractor_keeps_adjacent_runs_contiguous(self) -> None:
        prefix, _ = self.BASE.split("SBOM(", 1)
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "split.docx"
            self.write_docx(
                path,
                "<w:p>"
                f'<w:r><w:t xml:space="preserve">{prefix}</w:t></w:r>'
                "<w:r><w:t>SBOM(소프트웨어 자재명세서</w:t></w:r>"
                "<w:r><w:t>)</w:t></w:r>"
                "</w:p>",
            )
            extracted = package_submission.extract_docx_text(path)
            self.assertEqual(self.BASE, extracted)
            package_submission.validate_report_text_contract(extracted, self.BASE)

    def test_docx_run_split_cannot_hide_placeholder(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "placeholder.docx"
            self.write_docx(
                path,
                "<w:p><w:r><w:t>"
                f"{self.BASE}"
                "</w:t></w:r></w:p>"
                "<w:p><w:r><w:t>[[OWNER_</w:t></w:r>"
                "<w:r><w:t>VOICE]]</w:t></w:r></w:p>",
            )
            extracted = package_submission.extract_docx_text(path)
            self.assertIn("[[OWNER_VOICE]]", extracted)
            with self.assertRaisesRegex(package_submission.GateError, "unresolved gates"):
                package_submission.validate_report_text_contract(extracted, self.BASE)

    def test_final_visible_report_rejects_reader_facing_evidence_ids(self) -> None:
        markers = ("E09", "e09", "Ｅ０９", "Ｅ09", "E０9", "E\u200b09")
        for marker in markers:
            for label, docx_text, pdf_text in (
                ("DOCX", self.BASE + " " + marker, self.BASE),
                ("PDF", self.BASE, self.BASE + " " + marker),
            ):
                with self.subTest(label=label, marker=marker), self.assertRaisesRegex(
                    package_submission.GateError,
                    rf"{label} contains reader-facing audit evidence IDs \(count=1\)",
                ) as caught:
                    package_submission.validate_report_text_contract(
                        docx_text, pdf_text
                    )
                self.assertNotIn(marker, str(caught.exception))

    def test_docx_extractor_preserves_only_explicit_separators(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "separators.docx"
            self.write_docx(
                path,
                '<w:p><w:r><w:t xml:space="preserve">A</w:t><w:tab/>'
                "<w:t>B</w:t><w:br/><w:t>C</w:t></w:r></w:p>"
                "<w:p><w:r><w:t>D</w:t></w:r></w:p>",
            )
            self.assertEqual(
                "A\tB\nC\nD", package_submission.extract_docx_text(path)
            )

    def test_accepts_sbom_and_development_ai_without_attachment_2(self) -> None:
        package_submission.validate_report_text_contract(self.BASE, self.BASE)

    def test_accepts_canonical_report_content_ai_disclosure(self) -> None:
        content = json.loads(
            (REPOSITORY_ROOT / "submission" / "report-content.ko.json").read_text(
                encoding="utf-8"
            )
        )
        ai_row = next(
            row for row in content["environment"] if row["lead"] == "개발 보조 AI"
        )
        self.assertIn(package_submission.NO_RUNTIME_AI_DISCLOSURE, ai_row["text"])
        canonical_text = (
            f"{ai_row['lead']} {ai_row['text']} 붙임1 SBOM(소프트웨어 자재명세서)"
        )
        package_submission.validate_report_text_contract(
            canonical_text, canonical_text
        )

    def test_rejects_missing_canonical_no_runtime_ai_disclosure(self) -> None:
        text = self.BASE.replace(package_submission.NO_RUNTIME_AI_DISCLOSURE, "")
        with self.assertRaisesRegex(package_submission.GateError, "no-runtime-AI"):
            package_submission.validate_report_text_contract(text, text)

    def test_visible_binding_accepts_official_table_extraction_reordering(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        docx_text, pdf_text = self.visible_binding_fixture(content)
        package_submission.validate_report_visible_content(
            docx_text, pdf_text, content
        )

    def test_visible_binding_rejects_pdf_cell_deletion_or_substitution(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        text, _ = self.visible_binding_fixture(content)
        for label, mutated in (
            ("deletion", text.replace("Java 17", "Java 1", 1)),
            ("substitution", text.replace("Java 17", "Java X7", 1)),
        ):
            with self.subTest(label=label), self.assertRaisesRegex(
                package_submission.GateError, "visible-text inventory"
            ):
                package_submission.validate_report_visible_content(
                    text, mutated, content
                )

    def test_visible_binding_requires_exact_external_summary_in_pdf(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        text, _ = self.visible_binding_fixture(content)
        summary = package_submission.external_evidence_summary(content)
        same_inventory_without_sentence = text.replace(summary, summary[::-1], 1)
        with self.assertRaisesRegex(
            package_submission.GateError, "external-evidence summary"
        ):
            package_submission.validate_report_visible_content(
                text, same_inventory_without_sentence, content
            )

    def test_visible_binding_rejects_same_inventory_semantic_reordering(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        text, _ = self.visible_binding_fixture(content)
        target = content["background"][0]["text"]
        mutated = text.replace(target, target[::-1], 1)
        with self.assertRaisesRegex(
            package_submission.GateError, "ordered semantic text|lead-to-text row order"
        ):
            package_submission.validate_report_visible_content(text, mutated, content)

    def test_visible_binding_rejects_sbom_version_swap(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        docx_text, pdf_text = self.visible_binding_fixture(content)
        first = content["sbom"][0]["version"]
        second = content["sbom"][1]["version"]
        first_row = next(
            line for line in pdf_text.splitlines() if re.match(r"^\s*1\s+", line)
        )
        second_row = next(
            line for line in pdf_text.splitlines() if re.match(r"^\s*2\s+", line)
        )
        pdf_text = pdf_text.replace(first_row, first_row.replace(first, second, 1), 1)
        pdf_text = pdf_text.replace(second_row, second_row.replace(second, first, 1), 1)
        with self.assertRaisesRegex(package_submission.GateError, "SBOM row"):
            package_submission.validate_report_visible_content(
                docx_text, pdf_text, content
            )

    def test_visible_binding_rejects_wrapped_midpoint_version_swap(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        docx_text, pdf_text = self.visible_binding_fixture(content)
        rows = content["sbom"][:2]
        original_lines = [
            line
            for line in pdf_text.splitlines()
            if re.match(r"^\s*[12]\s+", line)
        ]
        first, second = rows
        wrapped = "\n".join(
            (
                f"{1:<8}{first['name']:<36}",
                f"{'':<44}{first['version']:<24}{first['license']:<20}"
                f"{first['url']} {first['purpose']}",
                f"{'':<44}{second['version']:<24}{second['license']:<20}"
                f"{second['url']} {second['purpose']}",
                f"{2:<8}{second['name']:<36}",
            )
        )
        pdf_text = pdf_text.replace("\n".join(original_lines), wrapped, 1)
        package_submission.validate_report_visible_content(
            docx_text, pdf_text, content
        )
        swapped_wrapped = wrapped.replace(first["version"], "[[VERSION]]", 1)
        swapped_wrapped = swapped_wrapped.replace(
            second["version"], first["version"], 1
        )
        swapped_wrapped = swapped_wrapped.replace(
            "[[VERSION]]", second["version"], 1
        )
        swapped = pdf_text.replace(wrapped, swapped_wrapped, 1)
        with self.assertRaisesRegex(
            package_submission.GateError, "version token is associated with another numbered row"
        ):
            package_submission.validate_report_visible_content(
                docx_text, swapped, content
            )

    def test_sbom_composite_version_cannot_straddle_neighbor_rows(self) -> None:
        rows = [
            {"name": "Alpha", "version": "1.0", "license": "MIT", "url": "u1", "purpose": "p1"},
            {"name": "Beta", "version": "2.0 / 2.1", "license": "MIT", "url": "u2", "purpose": "p2"},
            {"name": "Gamma", "version": "3.0", "license": "MIT", "url": "u3", "purpose": "p3"},
        ]
        header = (
            f"{'번호':<8}{'라이브러리명':<36}{'버전':<24}{'라이선스':<20}"
            "공식 저장소 URL(GitHub 등) 사용 목적 및 주요 기능"
        )
        text = "\n".join(
            (
                "붙임1 SBOM(소프트웨어 자재명세서)",
                header,
                f"{1:<8}{'Alpha':<36}{'1.0':<24}{'MIT':<20}u1 p1",
                f"{'':<44}{'2.0':<24}",
                "shared wrapped cells",
                f"{2:<8}{'Beta':<36}{'':<24}{'MIT':<20}u2 p2",
                "shared wrapped cells",
                f"{'':<44}{'2.1':<24}",
                f"{3:<8}{'Gamma':<36}{'3.0':<24}{'MIT':<20}u3 p3",
            )
        )
        with self.assertRaisesRegex(
            package_submission.GateError,
            "version token is associated with another numbered row",
        ):
            package_submission.validate_pdf_sbom_row_binding(text, rows)

    def test_sbom_accepts_junit_composite_version_split_around_own_row(self) -> None:
        rows = [
            {"name": "Alpha", "version": "1.0", "license": "MIT", "url": "u1", "purpose": "p1"},
            {
                "name": "JUnit Jupiter/Launcher",
                "version": "5.14.3 / 1.14.3",
                "license": "EPL-2.0",
                "url": "u2",
                "purpose": "p2",
            },
            {"name": "Gamma", "version": "3.0", "license": "MIT", "url": "u3", "purpose": "p3"},
        ]
        header = (
            f"{'번호':<8}{'라이브러리명':<36}{'버전':<24}{'라이선스':<20}"
            "공식 저장소 URL(GitHub 등) 사용 목적 및 주요 기능"
        )
        text = "\n".join(
            (
                "붙임1 SBOM(소프트웨어 자재명세서)",
                header,
                f"{1:<8}{'Alpha':<36}{'1.0':<24}{'MIT':<20}u1 p1",
                "",
                f"{'':<44}{'5.14.3 /':<24}",
                f"{2:<8}{'JUnit Jupiter/Launcher':<36}{'':<24}{'EPL-2.0':<20}u2 p2",
                f"{'':<44}{'1.14.3':<24}",
                "",
                f"{3:<8}{'Gamma':<36}{'3.0':<24}{'MIT':<20}u3 p3",
            )
        )
        package_submission.validate_pdf_sbom_row_binding(text, rows)

    def test_sbom_rejects_composite_token_tied_between_numbered_rows(self) -> None:
        rows = [
            {"name": "Alpha", "version": "1.0", "license": "MIT", "url": "u1", "purpose": "p1"},
            {
                "name": "Beta",
                "version": "2.0 / 2.1",
                "license": "MIT",
                "url": "u2",
                "purpose": "p2",
            },
            {"name": "Gamma", "version": "3.0", "license": "MIT", "url": "u3", "purpose": "p3"},
        ]
        header = (
            f"{'번호':<8}{'라이브러리명':<36}{'버전':<24}{'라이선스':<20}"
            "공식 저장소 URL(GitHub 등) 사용 목적 및 주요 기능"
        )
        text = "\n".join(
            (
                "붙임1 SBOM(소프트웨어 자재명세서)",
                header,
                f"{1:<8}{'Alpha':<36}{'1.0':<24}{'MIT':<20}u1 p1",
                "",
                f"{'':<44}{'2.0':<24}",
                "",
                f"{2:<8}{'Beta':<36}{'2.1':<24}{'MIT':<20}u2 p2",
                "",
                "",
                "",
                f"{3:<8}{'Gamma':<36}{'3.0':<24}{'MIT':<20}u3 p3",
            )
        )
        with self.assertRaisesRegex(
            package_submission.GateError,
            "version token is associated with another numbered row",
        ):
            package_submission.validate_pdf_sbom_row_binding(text, rows)

    def test_visible_binding_rejects_metadata_value_swap(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        docx_text, pdf_text = self.visible_binding_fixture(content)
        team = content["metadata"]["team_name"]
        project = content["metadata"]["project_name"]
        pdf_text = pdf_text.replace(team, "[[IDENTITY]]", 1)
        pdf_text = pdf_text.replace(project, team, 1).replace(
            "[[IDENTITY]]", project, 1
        )
        with self.assertRaisesRegex(package_submission.GateError, "metadata label"):
            package_submission.validate_report_visible_content(
                docx_text, pdf_text, content
            )

    def test_visible_binding_rejects_figure_caption_swap(self) -> None:
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            package_submission.validate_manifest(valid_manifest()),
            current_utc=TEST_CURRENT_UTC,
        )
        docx_text, pdf_text = self.visible_binding_fixture(content)
        captions = [asset["caption"] for asset in content["assets"].values()]
        pdf_text = pdf_text.replace(captions[0], "[[CAPTION]]", 1)
        pdf_text = pdf_text.replace(captions[1], captions[0], 1).replace(
            "[[CAPTION]]", captions[1], 1
        )
        with self.assertRaisesRegex(package_submission.GateError, "caption order"):
            package_submission.validate_report_visible_content(
                docx_text, pdf_text, content
            )

    def test_minimum_subsequence_extra_characters_preserves_order(self) -> None:
        self.assertEqual(
            2,
            package_submission.minimum_subsequence_extra_characters("aXbYc", "abc"),
        )
        self.assertIsNone(
            package_submission.minimum_subsequence_extra_characters("cba", "abc")
        )

    def test_report_hyperlink_allowlist_comes_only_from_structured_fields(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        content = package_submission.validate_and_materialize_report_content(
            valid_report_content("rc_only"),
            manifest,
            current_utc=TEST_CURRENT_UTC,
        )
        expected = {
            content["metadata"]["repository_url"],
            content["metadata"]["video_url"],
            manifest["project"]["ci_run_url"],
            manifest["project"]["release_url"],
            content["external_evidence"]["activation_record_url"],
            content["external_evidence"]["recruitment_record_url"],
            content["external_evidence"]["protocol_issue_url"],
            content["external_evidence"]["result_issue_url"],
            package_submission.UPSTREAM_ISSUE_38456_URL,
            *(row["url"] for row in content["sbom"]),
        }
        self.assertEqual(
            expected,
            package_submission.expected_report_hyperlink_targets(content, manifest),
        )
        content["other"][-1]["text"] += " https://evil.example/phish."
        self.assertEqual(
            expected,
            package_submission.expected_report_hyperlink_targets(content, manifest),
        )

    def test_report_hyperlink_errors_report_counts_without_private_values(self) -> None:
        canary = "PRIVATECANARYSECRETXYZ"
        content = valid_report_content("rc_only")
        manifest = valid_manifest()
        content["metadata"]["repository_url"] = (
            f"https://user:{canary}@example.test/private"
        )
        with self.assertRaisesRegex(
            package_submission.GateError, "unsafe structured hyperlink"
        ) as caught:
            package_submission.expected_report_hyperlink_targets(content, manifest)
        self.assertNotIn(canary, str(caught.exception))

        actual_url = f"https://actual.example/{canary}"
        expected_url = f"https://expected.example/{canary}"
        with patch.object(
            package_submission,
            "pdf_hyperlink_rows",
            return_value=[(1, actual_url)],
        ), patch.object(
            package_submission,
            "expected_report_hyperlink_targets",
            return_value={expected_url},
        ), self.assertRaisesRegex(
            package_submission.GateError,
            r"missing_count=1, extra_count=1",
        ) as caught:
            package_submission.validate_pdf_hyperlinks(
                Path("report.pdf"), content, manifest, 1
            )
        self.assertNotIn(canary, str(caught.exception))

        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "links.docx"
            document = (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<w:document xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main" '
                'xmlns:r="http://schemas.openxmlformats.org/officeDocument/'
                '2006/relationships"><w:body><w:p>'
                '<w:hyperlink r:id="rId1"><w:r><w:t>'
                f"{canary}</w:t></w:r></w:hyperlink>"
                "</w:p></w:body></w:document>"
            )
            relationships = (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/'
                'package/2006/relationships"><Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
                f'relationships/hyperlink" Target="{actual_url}" '
                'TargetMode="External"/></Relationships>'
            )
            with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
                archive.writestr("word/document.xml", document)
                archive.writestr("word/_rels/document.xml.rels", relationships)

            with patch.object(
                package_submission,
                "expected_report_hyperlink_targets",
                return_value={expected_url},
            ), self.assertRaisesRegex(
                package_submission.GateError,
                r"missing_count=1, extra_count=1",
            ) as caught:
                package_submission.validate_docx_hyperlinks(path, content, manifest)
            self.assertNotIn(canary, str(caught.exception))

            with patch.object(
                package_submission,
                "expected_report_hyperlink_targets",
                return_value={actual_url},
            ), patch.object(
                package_submission,
                "expected_report_hyperlink_bindings",
                return_value={("expected label", actual_url)},
            ), self.assertRaisesRegex(
                package_submission.GateError,
                r"missing_count=1, extra_count=1",
            ) as caught:
                package_submission.validate_docx_hyperlinks(path, content, manifest)
            self.assertNotIn(canary, str(caught.exception))

    def test_pdf_raster_binding_accepts_only_identical_page_pixels(self) -> None:
        docx = Path("canonical.docx")
        pdf = Path("supplied.pdf")
        with patch.object(
            package_submission,
            "export_canonical_report_pdf",
            return_value=Path("canonical.pdf"),
        ), patch.object(
            package_submission,
            "rasterize_pdf_pages",
            side_effect=(
                [b"page-1", b"page-2"],
                [b"page-1", b"page-2"],
                [b"page-1", b"page-2"],
            ),
        ), patch.object(
            package_submission,
            "extract_pdf_pages",
            side_effect=(
                (["canonical-1", "canonical-2"], []),
                (["canonical-1", "canonical-2"], []),
                (["canonical-1", "canonical-2"], []),
            ),
        ), patch.object(
            package_submission,
            "pdf_hyperlink_fragments",
            side_effect=(
                [(1, 10, 10, 20, 10, "https://example.test/evidence", "evidence")],
                [(1, 10, 10, 20, 10, "https://example.test/evidence", "evidence")],
                [(1, 10, 10, 20, 10, "https://example.test/evidence", "evidence")],
            ),
        ), patch.object(
            package_submission,
            "pdf_hyperlink_rows",
            side_effect=(
                [(1, "https://example.test/evidence")],
                [(1, "https://example.test/evidence")],
                [(1, "https://example.test/evidence")],
            ),
        ) as rasterized:
            package_submission.validate_pdf_raster_matches_docx(docx, pdf)
        self.assertEqual(3, rasterized.call_count)

        with patch.object(
            package_submission,
            "export_canonical_report_pdf",
            return_value=Path("canonical.pdf"),
        ), patch.object(
            package_submission,
            "rasterize_pdf_pages",
            side_effect=([b"visible"], [b"visible"], [b"all-white-mask"]),
        ), self.assertRaisesRegex(package_submission.GateError, "visual raster differs"):
            package_submission.validate_pdf_raster_matches_docx(docx, pdf)

    def test_pdf_raster_binding_rejects_same_pixels_with_alternate_text(self) -> None:
        with patch.object(
            package_submission,
            "export_canonical_report_pdf",
            return_value=Path("canonical.pdf"),
        ), patch.object(
            package_submission,
            "rasterize_pdf_pages",
            side_effect=([b"same-pixels"], [b"same-pixels"], [b"same-pixels"]),
        ), patch.object(
            package_submission,
            "extract_pdf_pages",
            side_effect=(
                (["ShardingSphere 5.5.3 Apache-2.0"], []),
                (["ShardingSphere 5.5.3 Apache-2.0"], []),
                (["ShardingSphere 5.5.3.0 Apache-2"], []),
            ),
        ), self.assertRaisesRegex(package_submission.GateError, "text layer differs"):
            package_submission.validate_pdf_raster_matches_docx(
                Path("canonical.docx"), Path("actualtext.pdf")
            )

    def test_pdf_raster_binding_rejects_link_annotation_drift(self) -> None:
        with patch.object(
            package_submission,
            "export_canonical_report_pdf",
            return_value=Path("canonical.pdf"),
        ), patch.object(
            package_submission,
            "rasterize_pdf_pages",
            side_effect=([b"same"], [b"same"], [b"same"]),
        ), patch.object(
            package_submission,
            "extract_pdf_pages",
            side_effect=((['same'], []), (['same'], []), (['same'], [])),
        ), patch.object(
            package_submission,
            "pdf_hyperlink_fragments",
            side_effect=(
                [(1, 10, 10, 20, 10, "https://example.test/reviewed", "reviewed")],
                [(1, 10, 10, 20, 10, "https://example.test/reviewed", "reviewed")],
                [(1, 10, 10, 20, 10, "https://evil.example/substituted", "reviewed")],
            ),
        ), patch.object(
            package_submission,
            "pdf_hyperlink_rows",
            side_effect=(
                [(1, "https://example.test/reviewed")],
                [(1, "https://example.test/reviewed")],
                [(1, "https://evil.example/substituted")],
            ),
        ), self.assertRaisesRegex(package_submission.GateError, "positioned links differ"):
            package_submission.validate_pdf_raster_matches_docx(
                Path("canonical.docx"), Path("links.pdf")
            )

    def test_pdf_raster_binding_rejects_page_count_drift(self) -> None:
        with patch.object(
            package_submission,
            "export_canonical_report_pdf",
            return_value=Path("canonical.pdf"),
        ), patch.object(
            package_submission,
            "rasterize_pdf_pages",
            side_effect=([b"one"], [b"one"], [b"one", b"two"]),
        ), self.assertRaisesRegex(package_submission.GateError, "page count"):
            package_submission.validate_pdf_raster_matches_docx(
                Path("canonical.docx"), Path("supplied.pdf")
            )

    def test_pdf_raster_binding_rejects_nondeterministic_canonical_exports(self) -> None:
        with patch.object(
            package_submission,
            "export_canonical_report_pdf",
            side_effect=(Path("first.pdf"), Path("second.pdf")),
        ), patch.object(
            package_submission,
            "rasterize_pdf_pages",
            side_effect=([b"first"], [b"second"], [b"first"]),
        ), self.assertRaisesRegex(package_submission.GateError, "two isolated canonical"):
            package_submission.validate_pdf_raster_matches_docx(
                Path("canonical.docx"), Path("supplied.pdf")
            )

    def test_rasterizer_rejects_unbounded_page_count_before_pdftoppm(self) -> None:
        with patch.object(
            package_submission.shutil, "which", side_effect=lambda name: f"/safe/{name}"
        ), patch.object(
            package_submission, "run", return_value="Pages:          8\n"
        ) as invoked, self.assertRaisesRegex(
            package_submission.GateError, "page count exceeds"
        ):
            package_submission.rasterize_pdf_pages(
                Path("oversized.pdf"), Path("pages"), "supplied report PDF"
            )
        invoked.assert_called_once_with(["/safe/pdfinfo", "oversized.pdf"])

    def test_rasterizer_rejects_unsafe_page_box_before_pdftoppm(self) -> None:
        with patch.object(
            package_submission.shutil, "which", side_effect=lambda name: f"/safe/{name}"
        ), patch.object(
            package_submission,
            "run",
            side_effect=(
                "Pages:          1\n",
                "Page 1 size: 99999 x 99999 pts\n",
            ),
        ) as invoked, self.assertRaisesRegex(
            package_submission.GateError, "unsafe PDF page box"
        ):
            package_submission.rasterize_pdf_pages(
                Path("oversized-box.pdf"), Path("pages"), "supplied report PDF"
            )
        self.assertEqual(2, invoked.call_count)
        self.assertEqual(
            [
                call(["/safe/pdfinfo", "oversized-box.pdf"]),
                call(
                    [
                        "/safe/pdfinfo",
                        "-f",
                        "1",
                        "-l",
                        "1",
                        "-box",
                        "oversized-box.pdf",
                    ]
                ),
            ],
            invoked.call_args_list,
        )

    def test_canonical_export_rejects_unreviewed_font_configuration(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            config = root / "fonts.conf"
            config.write_text("<fontconfig/>", encoding="utf-8")
            with patch.dict(os.environ, {"FONTCONFIG_FILE": str(config)}), patch.object(
                package_submission, "require_libreoffice_writer", return_value="/safe/soffice"
            ), self.assertRaisesRegex(
                package_submission.GateError, "reviewed report font configuration"
            ):
                package_submission.export_canonical_report_pdf(
                    root / "report.docx", root / "export"
                )

    def test_rejects_runtime_ai_attachment(self) -> None:
        text = self.BASE + " AI 모델 활용 및 라이선스 기술 명세서"
        with self.assertRaisesRegex(package_submission.GateError, "Attachment 2"):
            package_submission.validate_report_text_contract(text, text)

    def test_rejects_missing_sbom(self) -> None:
        text = self.BASE.replace("SBOM(소프트웨어 자재명세서)", "")
        with self.assertRaisesRegex(package_submission.GateError, "SBOM"):
            package_submission.validate_report_text_contract(text, text)

    def test_rejects_body_over_five_pages(self) -> None:
        pages = ["body"] * 6 + ["붙임1 SBOM(소프트웨어 자재명세서) Omega 2.0"]
        sizes = [self.PORTRAIT] * 6 + [self.LANDSCAPE]
        with self.assertRaisesRegex(package_submission.GateError, "maximum is 5"):
            package_submission.report_page_contract(pages, sizes, self.SBOM_ROWS)

    def test_accepts_two_trailing_attachment_1_pages(self) -> None:
        pages = [
            "body page 1",
            "body page 2",
            "붙임1 SBOM(소프트웨어 자재명세서) Alpha 1.0",
            "SBOM continued rows Omega 2.0",
        ]
        sizes = [self.PORTRAIT, self.PORTRAIT, self.LANDSCAPE, self.LANDSCAPE]
        self.assertEqual(
            {"body_pages": 2, "attachment_1_pages": 2, "total_pages": 4},
            package_submission.report_page_contract(pages, sizes, self.SBOM_ROWS),
        )

    def test_rejects_wrong_body_or_attachment_orientation(self) -> None:
        pages = ["body", "붙임1 SBOM(소프트웨어 자재명세서) Omega 2.0"]
        for sizes, message in (
            ([self.LANDSCAPE, self.LANDSCAPE], "body page 1 must be A4 portrait"),
            ([self.PORTRAIT, self.PORTRAIT], "Attachment 1 page 2 must be A4 landscape"),
        ):
            with self.subTest(sizes=sizes), self.assertRaisesRegex(
                package_submission.GateError, message
            ):
                package_submission.report_page_contract(pages, sizes, self.SBOM_ROWS)

    def test_rejects_letter_or_legal_page_sizes(self) -> None:
        pages = ["body", "붙임1 SBOM(소프트웨어 자재명세서) Omega 2.0"]
        for sizes, message in (
            ([(612.0, 792.0), self.LANDSCAPE], "A4 portrait"),
            ([self.PORTRAIT, (1008.0, 612.0)], "A4 landscape"),
        ):
            with self.subTest(sizes=sizes), self.assertRaisesRegex(
                package_submission.GateError, message
            ):
                package_submission.report_page_contract(pages, sizes, self.SBOM_ROWS)

    def test_parses_real_poppler_a4_page_size_format(self) -> None:
        info = (
            "Pages:           2\n"
            "Page    1 size:  595.304 x 841.89 pts (A4)\n"
            "Page    2 size:  841.89 x 595.304 pts (A4)\n"
            "File size:       12345 bytes\n"
        )
        self.assertEqual(
            [(595.304, 841.89), (841.89, 595.304)],
            package_submission.parse_pdf_page_sizes(info, 2),
        )

    def test_rejects_blank_or_unanchored_attachment_page(self) -> None:
        sizes = [self.PORTRAIT, self.LANDSCAPE, self.LANDSCAPE]
        for trailing, message in (
            ("", "blank pages"),
            ("continued without a declared component", "no declared SBOM row anchor"),
        ):
            pages = [
                "body",
                "붙임1 SBOM(소프트웨어 자재명세서) Alpha 1.0",
                trailing,
            ]
            with self.subTest(trailing=trailing), self.assertRaisesRegex(
                package_submission.GateError, message
            ):
                package_submission.report_page_contract(pages, sizes, self.SBOM_ROWS)

    def test_requires_last_declared_row_on_final_attachment_page(self) -> None:
        pages = [
            "body",
            "붙임1 SBOM(소프트웨어 자재명세서) Omega 2.0",
            "continued Alpha 1.0",
        ]
        sizes = [self.PORTRAIT, self.LANDSCAPE, self.LANDSCAPE]
        with self.assertRaisesRegex(package_submission.GateError, "last declared SBOM row"):
            package_submission.report_page_contract(pages, sizes, self.SBOM_ROWS)

    def test_recognizes_multiword_row_split_by_pdf_column_order(self) -> None:
        row = {"name": "CycloneDX Gradle Plugin", "version": "3.4.0"}
        page = "CycloneDX https://example.invalid 10 3.4.0 Gradle purpose Plugin"
        self.assertTrue(package_submission.page_contains_sbom_row(page, row))


class ReportContentSbomTest(unittest.TestCase):
    def test_report_traceability_ledger_matches_structured_public_evidence_contract(
        self,
    ) -> None:
        content = json.loads(
            (SCRIPT.parents[1] / "report-content.ko.json").read_text(encoding="utf-8")
        )
        matrix = (REPOSITORY_ROOT / "docs" / "evidence-matrix.md").read_text(
            encoding="utf-8"
        )

        evidence_row = next(
            row
            for row in content["other"]
            if row["lead"] == "오픈소스SW 조합"
        )
        evidence = evidence_row["text"]
        self.assertIn("SQLExecutionHook SPI", evidence)
        self.assertIn("JAR build에는 shading 설정이 없다", evidence)
        self.assertIn("부재나 의미적 출처를 증명하지 않는다", evidence)
        self.assertIn("SBOM·lock·checksum·NOTICE", evidence)
        self.assertIn("runtime POM에 없다", evidence)

        application_result = next(
            row["text"]
            for row in content["effects"]
            if row["lead"] == "재현한 적용 결과"
        )
        self.assertIn("same-checkout standalone consumer", application_result)
        self.assertIn("독립 외부 설치·채택 증거가 아니다", application_result)

        determinism = content["features"][3]
        self.assertEqual("4. 재현성·operation 격리", determinism["lead"])
        for measured_claim in ("8개", "각 20회", "20쌍", "혼합 0건"):
            self.assertIn(measured_claim, determinism["text"])

        self.assertEqual(
            "https://github.com/ym0506/routecontract",
            content["metadata"]["repository_url"],
        )
        self.assertIn(
            "[[PUBLIC_CI_RUN_URL_REQUIRED_BEFORE_SUBMISSION]]",
            content["features"][4]["text"],
        )
        self.assertEqual(
            "[[PUBLIC_V0_1_0_RELEASE_AND_INSTALL_COORDINATES]]",
            content["features"][6]["text"],
        )
        self.assertIn(
            "[[PUBLIC_REPOSITORY_REVISION_AND_RELEASE]]",
            content["other"][5]["text"],
        )

        for required in (
            "## Reader-facing report claim crosswalk",
            "The report uses descriptive Korean labels instead of inserting audit IDs",
            "The report's `공개 증거 gate` is an owner-supplied visible prose slot, not a structured evidence object",
            "the package manifest's structured public-evidence record",
            "visible exact public URLs in the report",
            "한 문장 소개; 사용자·검출 공백; 해결 방식 | E02, E04, E05",
            "검증된 효과; 실제 MySQL business-green / contract-red; 시연; "
            "재현한 적용 결과 | E04, E05, E06, E08",
            "later public main run [32440114569]",
            "설치·릴리스; stable 배포 후 4단계 적용 흐름",
            "승인 manifest와 structural manifest diff; 활용 경계",
            "외부 검증 | E10",
            "품질관리·발전 로드맵 | E13, E14",
            "`선행 작업 경계` and `개발 소감` are identity/disclosure/provenance/owner-voice blocks",
            "`ORIGIN_AND_PRIOR_WORK.md` is a participant provenance declaration, not independent proof",
        ):
            self.assertIn(required, matrix)
        self.assertNotRegex(json.dumps(content, ensure_ascii=False), r"\bE(?:0[1-9]|1[0-4])\b")
        self.assertNotIn("all numbers linked", matrix)

    def test_report_crosswalk_maps_each_evidentiary_block_exactly_once(self) -> None:
        content = json.loads(
            (SCRIPT.parents[1] / "report-content.ko.json").read_text(encoding="utf-8")
        )
        matrix = (REPOSITORY_ROOT / "docs" / "evidence-matrix.md").read_text(
            encoding="utf-8"
        )
        crosswalk = matrix.split("## Reader-facing report claim crosswalk", 1)[1].split(
            "## Claim promotion checklist", 1
        )[0]
        mapped: list[str] = []
        actual_groups: dict[str, tuple[str, ...]] = {}
        for line in crosswalk.splitlines():
            if not line.startswith("|"):
                continue
            cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
            if not cells or cells[0] in {"Reader-facing report block", "---"}:
                continue
            self.assertEqual(3, len(cells), line)
            self.assertNotIn(cells[0], actual_groups)
            ids = tuple(piece.strip() for piece in cells[1].split(",") if piece.strip())
            self.assertTrue(ids, line)
            self.assertEqual(len(ids), len(set(ids)), line)
            self.assertTrue(
                all(re.fullmatch(r"E(?:0[1-9]|1[0-4])", evidence_id) for evidence_id in ids),
                line,
            )
            self.assertEqual(", ".join(ids), cells[1])
            actual_groups[cells[0]] = ids
            mapped.extend(piece.strip() for piece in cells[0].split(";") if piece.strip())

        expected_groups = {
            "한 문장 소개; 사용자·검출 공백; 해결 방식": ("E02", "E04", "E05"),
            "검증된 효과; 실제 MySQL business-green / contract-red; 시연; 재현한 적용 결과": (
                "E04",
                "E05",
                "E06",
                "E08",
            ),
            "공개 문제 근거": ("E01", "E04"),
            "언어·대상; 빌드·검증; 재현과 패키징; 설치·릴리스; stable 배포 후 4단계 적용 흐름": (
                "E02",
                "E08",
                "E09",
                "E12",
            ),
            "관측·계약 흐름; 상관관계·Fail-closed; 정보 경계; Operation 단위 관측 계약; 승인 manifest와 structural manifest diff; 활용 경계": (
                "E02",
                "E03",
                "E05",
                "E07",
                "E08",
            ),
            "재현성·operation 격리": ("E06", "E07"),
            "차별성": ("E11",),
            "오픈소스SW 조합; 현재 한계; 라이선스 검토 상태": (
                "E02",
                "E08",
                "E11",
                "E12",
            ),
            "공개 증거 gate": ("E08", "E09", "E12", "E13", "E14"),
            "외부 검증": ("E10",),
            "품질관리·발전 로드맵": ("E13", "E14"),
        }
        self.assertEqual(expected_groups, actual_groups)
        self.assertEqual(
            {f"E{index:02d}" for index in range(1, 15)},
            {evidence_id for ids in actual_groups.values() for evidence_id in ids},
        )

        normalize = lambda lead: re.sub(r"^[1-9][0-9]*\.\s*", "", lead)
        excluded = {"개발 장비", "개발 보조 AI", "선행 작업 경계", "개발 소감"}
        expected = {
            normalize(row["lead"])
            for section in (
                "project_intro",
                "background",
                "environment",
                "architecture",
                "features",
                "effects",
                "other",
            )
            for row in content[section]
            if normalize(row["lead"]) not in excluded
        }
        normalized_mapped = [normalize(lead) for lead in mapped]
        self.assertEqual(expected, set(normalized_mapped))
        self.assertEqual(
            [],
            sorted(
                lead
                for lead in set(normalized_mapped)
                if normalized_mapped.count(lead) != 1
            ),
        )
        self.assertIn(
            "`개발 장비`, `개발 보조 AI`, `선행 작업 경계` and `개발 소감`",
            crosswalk,
        )

    def test_report_content_keeps_information_and_approval_boundaries_explicit(
        self,
    ) -> None:
        content = json.loads(
            (SCRIPT.parents[1] / "report-content.ko.json").read_text(encoding="utf-8")
        )
        self.assertEqual("사용자·검출 공백", content["background"][0]["lead"])
        self.assertEqual(
            "기능 테스트는 같은 한 행으로 통과했지만 실제 MySQL의 SQLExecutionHook이 "
            "보고한 물리 JDBC 실행 시도·data source는 1→2였다. 이를 성능 결함으로 "
            "단정하지 않고 merge 전 검토 대상으로 만든다. 대상은 ShardingSphere-JDBC "
            "5.5.3 Java 팀, 증거는 저장소 MySQL fixture다. "
            "지연·부하·비용은 측정하지 않았다.",
            content["background"][0]["text"],
        )
        self.assertEqual("검증된 효과", content["background"][1]["lead"])
        self.assertEqual(
            [
                {
                    "lead": "한 문장 소개",
                    "text": "ShardingSphere-JDBC는 한 SQL을 여러 DB로 나눠 실행할 수 "
                    "있다. RouteContract는 기능 결과가 같아도 보고된 JDBC 실행 구조가 "
                    "승인본과 달라지면 CI에서 검토·차단하는 Java 테스트 라이브러리다.",
                }
            ],
            content["project_intro"],
        )
        self.assertNotIn("SQLExecutionHook", content["project_intro"][0]["text"])
        self.assertNotIn("5.5.3", content["project_intro"][0]["text"])
        self.assertIn(
            "보고된 JDBC 실행 구조",
            content["project_intro"][0]["text"],
        )
        for boundary_text in (
            content["background"][0]["text"],
            content["background"][2]["text"],
            content["features"][0]["text"],
        ):
            self.assertIn(
                "SQLExecutionHook이 보고한 물리 JDBC 실행 시도",
                boundary_text,
            )
        serialized = json.dumps(content, ensure_ascii=False)
        for contradiction in (
            "complete route plan을 증명",
            "transaction commit을 증명",
            "stable Release는 이미 공개",
            "독립 외부 설치·채택 증거다",
            "candidate는 approved를 자동으로 덮는다",
        ):
            self.assertNotIn(contradiction, serialized)

        information = next(
            row["text"] for row in content["architecture"] if row["lead"] == "정보 경계"
        )
        self.assertEqual(
            "원문 SQL·bind 값·connection property·exception message는 저장하지 않는다. "
            "실제 data-source 이름은 메모리에 남고 manifest에서만 alias로 바꾼다. "
            "operationId·Java type·unsalted SQL fingerprint도 남는다. 비민감 식별자를 "
            "쓰고 snapshot을 외부 로그로 내보내지 않는다. 최소화는 익명화가 아니다.",
            information,
        )
        for required in (
            "실제 data-source 이름은 메모리에 남고",
            "operationId",
            "manifest에서만 alias",
            "Java type",
            "unsalted SQL fingerprint",
            "최소화는 익명화가 아니다",
            "원문 SQL·bind 값",
            "외부 로그로 내보내지 않는다",
        ):
            self.assertIn(required, information)

        ai_scope = next(
            row["text"] for row in content["environment"] if row["lead"] == "개발 보조 AI"
        )
        self.assertEqual(
            "OpenAI ChatGPT·Codex를 조사·설계·구현·테스트·문서(소감 초안 포함)·local "
            "command 보조에 사용하고 AI_ASSISTANCE.md에 공개했다. 참가자는 최종 "
            "diff·재현 테스트·소감 사실을 검토해 사실인 attestation만 true로 "
            "둔다. AI 출력은 증거가 아니며 runtime에는 AI 모델·데이터셋·외부 AI API가 없다.",
            ai_scope,
        )
        self.assertIn(package_submission.NO_RUNTIME_AI_DISCLOSURE, ai_scope)
        self.assertNotIn("코드는 독립 구현", json.dumps(content, ensure_ascii=False))

        build_validation = next(
            row["text"]
            for row in content["environment"]
            if row["lead"] == "빌드·검증"
        )
        self.assertEqual(
            "Gradle Wrapper 8.14.4, JUnit Jupiter 5.14.3, Testcontainers 1.21.4, "
            "Docker. 라이브러리는 Apache-2.0이며 CycloneDX 1.6 JSON/XML SBOM을 "
            "생성하고 dependency lock·checksum 검증을 수행한다.",
            build_validation,
        )

        development_process = next(
            row["text"]
            for row in content["features"]
            if row["lead"] == "5. 재현과 패키징"
        )
        self.assertEqual(
            "개발은 hook lifecycle 계약→실패경로 단위 테스트→MySQL corpus→"
            "결정성·격리→standalone consumer→release gate 순이었다. "
            "clean check·assemble·SBOM으로 52 tests와 배포물을 검증했다. "
            "[[PUBLIC_CI_RUN_URL_REQUIRED_BEFORE_SUBMISSION]]",
            development_process,
        )

        development_effect = next(
            row["text"]
            for row in content["effects"]
            if row["lead"] == "stable 배포 후 4단계 적용 흐름"
        )
        self.assertEqual(
            "stable Release 공개 후 package gate로 exact tag·checksum·attestation을 "
            "확인하고 ① 설치 ② v0.1 Java 테스트의 application operation을 "
            "capture(operationId)로 감싸 candidate 생성 ③ 사람이 approved diff 검토·승인 "
            "④ strict CI 순으로 적용한다. candidate는 "
            "approved를 자동 갱신하지 않는다.",
            development_effect,
        )
        self.assertIn("package gate", development_effect)
        self.assertIn("stable Release 공개 후", development_effect)
        self.assertIn("① 설치", development_effect)
        self.assertIn("② v0.1 Java 테스트", development_effect)
        self.assertIn("application operation", development_effect)
        self.assertIn("capture(operationId)", development_effect)
        self.assertIn("③ 사람이 approved diff 검토·승인", development_effect)
        self.assertIn("④ strict CI", development_effect)
        self.assertNotIn("의도한 baseline 변경", development_effect)

        reproducible_application = next(
            row["text"]
            for row in content["effects"]
            if row["lead"] == "재현한 적용 결과"
        )
        self.assertEqual(
            "동일 MySQL 한 행 assertion은 통과했지만 SQLExecutionHook이 보고한 "
            "물리 JDBC 실행 시도·data source 1→2를 RCM201·RCM202/non-zero exit으로 "
            "검출해 quickstart로 재현한다. same-checkout standalone consumer는 JAR "
            "SPI·MySQL 1건만 확인하므로 1→2 재현·독립 외부 설치·채택 증거가 아니다. 외부 결과는 "
            "cutoff 공개 사실만 보고한다.",
            reproducible_application,
        )

        application_boundary = next(
            row["text"] for row in content["effects"] if row["lead"] == "활용 경계"
        )
        self.assertEqual(
            "지원 범위에서 SQLExecutionHook이 보고한 물리 JDBC 실행 시도를 "
            "review·CI 계약으로 만든다. complete route plan·transaction commit·business "
            "success·성능을 증명하지 않는다. caller-supplied target universe 없이는 "
            "full-route detection을 주장하지 않는다. 기대 효과는 승인본과 다른 구조의 "
            "merge 전 노출이며 성능·비용 개선 보장은 없다. 공개 수요+fixture+"
            "real-MySQL CI를 gate로 확장한다. "
            "현재는 5.5.3/MySQL 8.4.11이다.",
            application_boundary,
        )
        for required in (
            "SQLExecutionHook이 보고한 물리 JDBC 실행 시도",
            "review·CI 계약",
            "complete route plan·transaction commit·business success·성능을 증명하지 않는다",
            "caller-supplied target universe 없이는 full-route detection을 주장하지 않는다",
        ):
            self.assertIn(required, application_boundary)

        community = next(
            row["text"]
            for row in content["other"]
            if row["lead"] == "품질관리·발전 로드맵"
        )
        self.assertEqual(
            "Issue #5→PR #6에서 Ubuntu CI checksum을 고쳐 Dependency Review·build를 "
            "통과시켰다. 1인이 Issue·PR·CI로 관리한다. 설치·문서→adapter·reporter 순이며 "
            "확장은 공개 수요+fixture+real-MySQL CI 뒤 한다. 외부 결과는 링크로만 "
            "보고한다.",
            community,
        )
        self.assertNotIn("v0.1.0", community)
        self.assertNotIn("동결", community)

        public_evidence_gate = next(
            row["text"]
            for row in content["other"]
            if row["lead"] == "공개 증거 gate"
        )
        self.assertEqual(
            "[[PUBLIC_REPOSITORY_REVISION_AND_RELEASE]] / "
            "[[PUBLIC_ISSUE_PR_COMMUNITY_FEEDBACK]]. 실제 링크 전에는 완료형으로 "
            "쓰지 않는다.",
            public_evidence_gate,
        )

        owner_voice = next(
            row["text"] for row in content["other"] if row["lead"] == "개발 소감"
        )
        self.assertEqual(
            "[[OWNER_VOICE: AI 보조 여부를 공개한 뒤 참가자가 사실·의미를 검토·확인한 "
            "1인칭 소감—실제로 가장 어려웠던 문제, 직접 한 분석/실험, 배운 점, 능동 "
            "유지보수 우선순위·순서와 의도한 기간]]",
            owner_voice,
        )

        differentiation = next(
            row["text"] for row in content["other"] if row["lead"] == "차별성"
        )
        self.assertEqual(
            "5.5.3 SQLExecutionHook이 보고한 물리 JDBC 시도를 operation별 최소 manifest→"
            "사람 승인→결정적 diff→RCM code→CI로 잇는다. 관측 도구를 대체하지 않으며 "
            "datasource-proxy도 data source 연결·상관관계·canonicalization·diff·assertion을 더하면 유사 "
            "검사가 가능하다.",
            differentiation,
        )

        evidence = next(
            row["text"]
            for row in content["other"]
            if row["lead"] == "오픈소스SW 조합"
        )
        self.assertEqual(
            "5.5.3 SQLExecutionHook SPI를 MySQL 8.4.11로 검증했다. JAR "
            "build에는 shading 설정이 없다. path·POM 검사는 renamed/copied bytes 부재나 "
            "의미적 출처를 증명하지 않는다. SBOM·lock·checksum·NOTICE로 추적한다. "
            "test-only datasource-proxy는 runtime POM에 없다. preflight는 전체 graph "
            "증명이 아니다.",
            evidence,
        )

        prior_work = next(
            row["text"] for row in content["other"] if row["lead"] == "선행 작업 경계"
        )
        self.assertEqual(
            "참가자 선언: ShardLens의 미구현 Route Guard 설계를 바탕으로 RouteContract "
            "라이브러리·manifest/diff·MySQL corpus·설치/CI를 새로 구현했고 애플리케이션 "
            "코드는 복사하지 않았다(ORIGIN_AND_PRIOR_WORK.md). 독립 증명은 아니다.",
            prior_work,
        )

        limitations = next(
            row["text"] for row in content["other"] if row["lead"] == "현재 한계"
        )
        self.assertEqual(
            "5.5.3 정상 반환·비-interrupt 동기식 non-batch PreparedStatement만 지원한다. "
            "Proxy·reactive·@Async·SQL Federation·route/table plan·commit·business "
            "success는 제외한다. test/example graph의 calcite-core 1.40.0 advisory는 "
            "2026-08-27까지 검토 예외이며 SBOM은 법률 검토가 아니다.",
            limitations,
        )
        license_disposition = next(
            row["text"]
            for row in content["other"]
            if row["lead"] == "라이선스 검토 상태"
        )
        self.assertEqual(
            "[[LICENSE_REVIEW_DISPOSITION: owner가 JTS/Mahout 비번들 배포경계와 MySQL "
            "OCI manual-review-required 유지 여부를 1차 자료·최종 payload 재검증 뒤 "
            "사실대로 작성]]",
            license_disposition,
        )

    def test_submission_readme_exactly_pins_privacy_and_owner_boundaries(self) -> None:
        readme = (SCRIPT.parents[1] / "README.md").read_text(encoding="utf-8")

        privacy_block = "The tracked report body is otherwise closed:" + readme.split(
            "The tracked report body is otherwise closed:", 1
        )[1].split("\n\nSet `submission_identity", 1)[0]
        self.assertEqual(
            "The tracked report body is otherwise closed: only the documented structured\n"
            "metadata and six owner free-text overlays may differ, and the external row is\n"
            "generated. The six `OWNER_FREE_TEXT_OVERLAY_STRING_PATHS` are `개발 장비`,\n"
            "`재현과 패키징`, `설치·릴리스`, `라이선스 검토 상태`, `공개 증거 gate`, and\n"
            "`개발 소감`; the high-confidence lexical privacy scanner applies only to those six values. It\n"
            "rejects common credential, contact, local-path, private-topology and raw-SQL\n"
            "leak forms, but it is a heuristic and cannot prove that arbitrary prose or\n"
            "topology is safe. The participant manually attests that those six free-text\n"
            "values do not introduce an external-result, adoption or stable-validation claim\n"
            "and sets `report_free_text_privacy_reviewed=true` only after manual privacy\n"
            "review. These are human attestations, not NLP classifiers. Structured\n"
            "registration identity is not checked by that lexical scanner: its six\n"
            "`submission_identity` values are exact-bound to the application, private\n"
            "manifest, report fields and official filenames where applicable, and the\n"
            "participant must manually review their exactness and disclosure. Reader-facing\n"
            "`E01`–`E14` audit IDs are rejected from every private overlay and from the final\n"
            "DOCX/PDF; the public crosswalk stays in\n"
            "`docs/evidence-matrix.md`. The check applies after Unicode compatibility and\n"
            "case normalization, so those tokens are reserved even when they would be part\n"
            "of a hardware model name; rewrite such a model descriptively in the private\n"
            "hardware disclosure. Stable Release metadata in its designated fields is allowed; a stable\n"
            "external-validation claim outside the generated row is not. Packaging verifies every canonical value in\n"
            "the DOCX, exact generated external text in DOCX/PDF, and the PDF's complete\n"
            "visible-character inventory and page/table anchors. A\n"
            "partial placeholder fill, altered generated-text marker, tracking Issue\n"
            "substituted for a participant result, mismatched tag, unsupported count, future\n"
            "cutoff, nonexistent public reference, or mixed branch fact is a hard failure.",
            privacy_block,
        )

        owner_boundary = "The `개발 소감` owner-voice block may use AI drafting or editing only when that" + readme.split(
            "The `개발 소감` owner-voice block may use AI drafting or editing only when that",
            1,
        )[1].split("\n\nDo not replace", 1)[0]
        self.assertEqual(
            "The `개발 소감` owner-voice block may use AI drafting or editing only when that\n"
            "assistance is disclosed. Before setting\n"
            "`owner_voice_ai_assistance_disclosed_and_participant_reviewed=true`, the\n"
            "participant must review and adopt the final text as an accurate first-person\n"
            "account, verify its concrete statements about the hardest problem, their own\n"
            "analysis or experiments, lessons learned, and active-maintenance priority, order,\n"
            "and intended period, and be able to explain it. This attestation does not claim\n"
            "participant-only authorship. Set `maintenance_order_and_period_confirmed=true`\n"
            "separately after confirming that concrete maintenance commitment. Contest rule Article\n"
            "10(3) requires a selected excellent or award-winning team to keep the public\n"
            "repository Public for five years from the award date. Set\n"
            "`five_year_public_repository_visibility_obligation_if_selected_accepted=true`\n"
            "only after understanding that conditional visibility obligation; it is not an\n"
            "active-maintenance promise. Confirm the tracked provenance statement with\n"
            "`origin_and_prior_work_statement_confirmed=true`; that is a participant\n"
            "self-attestation, not independent proof. The tool does not semantically decide\n"
            "whether the owner voice, its AI-assistance disclosure, or provenance statement is\n"
            "true or adequate.",
            owner_boundary,
        )

    def test_report_diagram_assets_are_self_contained_legible_and_pinned(self) -> None:
        assets = SCRIPT.parents[1] / "assets"
        expected = {
            "architecture": {
                "svg": "1d314f486fed80d5a1001d1d433530a874fa72506c3a4a654b06282ee7e4d3fe",
                "png": "c492713c535549fab2abb3571ba780e75e66d99fc2a24fdda978137c182da5f2",
                "text": (
                    "COMPARISON INPUTS",
                    "candidate canonical manifest",
                    "versioned approved manifest",
                    "deterministic structural manifest diff",
                    "stable RCM codes",
                    "value-minimized",
                    "no raw SQL / bind values",
                    "5.5.3 normal return",
                    "caller non-interrupted",
                    "synchronous non-batch PreparedStatement",
                ),
            },
            "baseline-candidate": {
                "svg": "d9d1abc3bab121c2541ca2ad892b36544fcaf9f0642944fc65397c170572ac94",
                "png": "2cbca82fb84267e092af919c6279290ba279b5350ba4a9f5b8ec0084b0c343a3",
                "text": (
                    "same one synthetic fixture row",
                    "approved baseline",
                    "candidate",
                    "structural manifest diff",
                    "1→2 attempts",
                    "1→2 aliases",
                    "RCM201",
                    "RCM202",
                    "Hook-reported attempts only",
                    "not route/table plan, commit, or business success",
                ),
            },
        }
        namespace = "{http://www.w3.org/2000/svg}"

        for name, contract in expected.items():
            with self.subTest(asset=name):
                svg_path = assets / f"{name}.svg"
                png_path = assets / f"{name}.png"
                svg_bytes = svg_path.read_bytes()
                png_bytes = png_path.read_bytes()
                self.assertEqual(contract["svg"], hashlib.sha256(svg_bytes).hexdigest())
                self.assertEqual(contract["png"], hashlib.sha256(png_bytes).hexdigest())

                root = ET.fromstring(svg_bytes)
                self.assertEqual(namespace + "svg", root.tag)
                self.assertEqual("1200", root.attrib.get("width"))
                self.assertEqual("675", root.attrib.get("height"))
                self.assertEqual("0 0 1200 675", root.attrib.get("viewBox"))
                title = root.find(namespace + "title")
                desc = root.find(namespace + "desc")
                self.assertIsNotNone(title)
                self.assertIsNotNone(desc)
                self.assertTrue((title.text or "").strip())
                self.assertTrue((desc.text or "").strip())

                normalized_text = " ".join(
                    fragment.strip()
                    for fragment in root.itertext()
                    if fragment.strip()
                )
                for required in contract["text"]:
                    self.assertIn(required, normalized_text)

                for element in root.iter():
                    self.assertNotIn(
                        element.tag.rsplit("}", 1)[-1],
                        {"foreignObject", "image", "script"},
                    )
                    for attribute, value in element.attrib.items():
                        if attribute.endswith("href"):
                            self.assertTrue(value.startswith("#"), value)
                        for reference in re.findall(r"url\(([^)]+)\)", value):
                            self.assertTrue(reference.startswith("#"), reference)
                        self.assertFalse(
                            re.match(r"(?i)^(?:https?:|file:|data:|//)", value),
                            value,
                        )

                font_sizes = [
                    int(size)
                    for size in re.findall(rb"font-size:\s*([0-9]+)px", svg_bytes)
                ]
                self.assertTrue(font_sizes)
                self.assertGreaterEqual(min(font_sizes), 28)

    def test_report_comparison_figure_excludes_fixture_identifiers(self) -> None:
        submission_root = SCRIPT.parents[1]
        svg = (submission_root / "assets" / "baseline-candidate.svg").read_text(
            encoding="utf-8"
        )
        content = (submission_root / "report-content.ko.json").read_text(
            encoding="utf-8"
        )
        combined = svg + "\n" + content

        for forbidden in ("user_id", "row 201", "= 3", "BETWEEN 3"):
            with self.subTest(forbidden=forbidden):
                self.assertNotIn(forbidden, combined)
        self.assertIn("one synthetic fixture row", svg)
        self.assertIn("structural", svg)
        self.assertIn("manifest diff", svg)
        self.assertNotIn("semantic diff", svg)
        self.assertIn("동일한 단일 fixture 행", content)

        png = (submission_root / "assets" / "baseline-candidate.png").read_bytes()
        self.assertEqual(
            "2cbca82fb84267e092af919c6279290ba279b5350ba4a9f5b8ec0084b0c343a3",
            hashlib.sha256(png).hexdigest(),
        )

    def test_current_content_declares_exactly_ten_prioritized_rows(self) -> None:
        content_path = SCRIPT.parents[1] / "report-content.ko.json"
        with content_path.open(encoding="utf-8") as stream:
            content = json.load(stream)

        expected_rows = [
            {
                "name": "OpenJDK standard-doclet assets",
                "version": "17.0.20.1+1",
                "license": "GPL-2.0-only WITH Classpath-exception-2.0",
                "url": "https://github.com/openjdk/jdk17u/tree/jdk-17.0.20.1%2B1/src/jdk.javadoc",
                "purpose": "stable Release Javadoc JAR에 CSS·JavaScript·PNG·legal 파일로 포함; main JAR·runtime 미포함",
            },
            {
                "name": "MySQL Connector/J",
                "version": "26.7.0",
                "license": "GPL-2.0-only WITH Universal-FOSS-exception-1.0",
                "url": "https://github.com/mysql/mysql-connector-j/tree/26.7.0",
                "purpose": "실제 MySQL 통합 테스트의 JDBC 드라이버로 사용(testRuntimeOnly); 배포 JAR 미포함",
            },
            {
                "name": "Apache ShardingSphere",
                "version": "5.5.3",
                "license": "Apache-2.0",
                "url": "https://github.com/apache/shardingsphere/tree/5.5.3",
                "purpose": "SQLExecutionHook SPI는 compileOnly, 실제 샤딩 fixture는 testImplementation/testRuntimeOnly로 사용; 배포 JAR 미내장",
            },
            {
                "name": "Alibaba TransmittableThreadLocal",
                "version": "2.14.2",
                "license": "Apache-2.0",
                "url": "https://github.com/alibaba/transmittable-thread-local/tree/v2.14.2",
                "purpose": "operation context 전달용 runtime 의존성",
            },
            {
                "name": "Jackson Core",
                "version": "3.1.5",
                "license": "Apache-2.0",
                "url": "https://github.com/FasterXML/jackson-core/tree/jackson-core-3.1.5",
                "purpose": "canonical JSON 생성·읽기용 runtime 의존성",
            },
            {
                "name": "Testcontainers (JUnit·MySQL)",
                "version": "1.21.4",
                "license": "MIT",
                "url": "https://github.com/testcontainers/testcontainers-java/tree/1.21.4",
                "purpose": "digest 고정 MySQL 통합 테스트 컨테이너 실행용 test 의존성",
            },
            {
                "name": "JUnit Jupiter / Platform Launcher",
                "version": "5.14.3 / 1.14.3",
                "license": "EPL-2.0",
                "url": "https://github.com/junit-team/junit-framework/tree/r5.14.3",
                "purpose": "단위·통합·실패경로 테스트와 launcher용 test 의존성",
            },
            {
                "name": "jQuery",
                "version": "3.7.1",
                "license": "MIT",
                "url": "https://github.com/jquery/jquery/tree/3.7.1",
                "purpose": "stable Release Javadoc JAR에 minified JavaScript 파일로 포함; main JAR·runtime 미포함",
            },
            {
                "name": "jQuery UI",
                "version": "1.14.1",
                "license": "MIT",
                "url": "https://github.com/jquery/jquery-ui/tree/1.14.1",
                "purpose": "stable Release Javadoc JAR에 JavaScript·CSS 파일로 포함; main JAR·runtime 미포함",
            },
            {
                "name": "Gradle Wrapper",
                "version": "8.14.4",
                "license": "Apache-2.0",
                "url": "https://github.com/gradle/gradle/tree/v8.14.4",
                "purpose": "repo에 Wrapper JAR로 포함; checksum 고정 build·test·package 도구; product runtime 미포함",
            },
        ]
        self.assertEqual(expected_rows, content["sbom"])
        self.assertTrue(all("GPL-" in row["license"] for row in content["sbom"][:2]))
        self.assertTrue(
            all("GPL-" not in row["license"] for row in content["sbom"][2:])
        )
        self.assertEqual(10, len({row["name"] for row in content["sbom"]}))
        self.assertEqual(10, len({row["url"] for row in content["sbom"]}))
        self.assertNotIn(
            "MySQL Community Server OCI image",
            {row["name"] for row in content["sbom"]},
        )
        self.assertNotIn("NOASSERTION", {row["license"] for row in content["sbom"]})
        for name in (
            "OpenJDK standard-doclet assets",
            "jQuery",
            "jQuery UI",
        ):
            row = next(item for item in content["sbom"] if item["name"] == name)
            self.assertIn("Javadoc JAR", row["purpose"])
            self.assertIn("main JAR·runtime 미포함", row["purpose"])

    def test_secondary_test_only_rows_remain_disclosed_outside_top_ten(self) -> None:
        content_path = SCRIPT.parents[1] / "report-content.ko.json"
        with content_path.open(encoding="utf-8") as stream:
            content = json.load(stream)
        names = {item["name"] for item in content["sbom"]}
        self.assertNotIn("HikariCP", names)
        self.assertNotIn("datasource-proxy", names)
        self.assertNotIn("MySQL Community Server OCI image", names)
        self.assertNotIn("CycloneDX Gradle Plugin", names)

        third_party = (REPOSITORY_ROOT / "THIRD_PARTY.md").read_text(
            encoding="utf-8"
        )
        self.assertIn("| HikariCP | 6.2.1 |", third_party)
        self.assertIn("| datasource-proxy | 1.11.0 |", third_party)
        self.assertIn("MySQL Community Server container image", third_party)
        self.assertIn("Image-wide conclusion not asserted", third_party)
        self.assertIn("manual package-level review required", third_party)
        self.assertIn(
            "https://github.com/docker-library/mysql/blob/"
            "01f90d87012e46cd174073bba02d64e9fc693ed3/8.4/Dockerfile.oracle",
            third_party,
        )
        self.assertIn(
            "image layers and installed programs are test-only", third_party
        )
        self.assertIn("| CycloneDX Gradle plugin | 3.4.0 |", third_party)

    def test_report_builder_requirements_pin_the_complete_python_closure(self) -> None:
        requirements_path = SCRIPT.parents[1] / "report-builder-requirements.txt"
        requirements = {
            line.strip()
            for line in requirements_path.read_text(encoding="utf-8").splitlines()
            if line.strip() and not line.startswith("#")
        }
        self.assertEqual(
            {
                "certifi==2026.7.22",
                "Pillow==12.3.0",
                "lxml==6.1.1",
                "python-docx==1.2.0",
                "typing_extensions==4.16.0",
            },
            requirements,
        )

    def test_ci_report_package_lock_pins_exact_linux_wheel_hashes(self) -> None:
        lock_path = SCRIPT.parents[1] / "report-package-ci-requirements.txt"
        self.assertEqual(
            "# CI-only wheel lock for CPython 3.12.14 on ubuntu-24.04 x86_64.\n"
            "# SHA-256 values are the official PyPI digests for the wheels selected on this platform.\n"
            "# Keep report-builder-requirements.txt for version-pinned cross-platform local setup.\n"
            "certifi==2026.7.22 \\\n"
            "    --hash=sha256:62f22742b58a1a33014a2b6b706588a8d7e2a88ae7bd1a6ebe8c992928483775\n"
            "Pillow==12.3.0 \\\n"
            "    --hash=sha256:78cb2c6865a35ab8ff8b75fd122f6033b92a62c82801110e48ddd6c936a45d91\n"
            "lxml==6.1.1 \\\n"
            "    --hash=sha256:ebe6af670449830d6d9b752c256a983291c766a1365ba5d5460048f9e33a7818\n"
            "python-docx==1.2.0 \\\n"
            "    --hash=sha256:3fd478f3250fbbbfd3b94fe1e985955737c145627498896a8a6bf81f4baf66c7\n"
            "typing_extensions==4.16.0 \\\n"
            "    --hash=sha256:481caa481374e813c1b176ada14e97f1f67a4539ce9cfeb3f350d78d6370c2e8\n",
            lock_path.read_text(encoding="utf-8"),
        )


class PrivacyAndPathTest(unittest.TestCase):
    PDF_INFO = """Author: RouteContract project
Creator: Writer
Producer: LibreOffice 26.2
Custom Metadata: no
UserProperties: no
Suspects: no
Form: none
JavaScript: no
Encrypted: no
Pages: 5
"""
    XMP = """<?xml version="1.0"?>
<x:xmpmeta xmlns:x="adobe:ns:meta/" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#">
  <rdf:RDF><rdf:Description><dc:creator><rdf:Seq><rdf:li>RouteContract project</rdf:li></rdf:Seq></dc:creator></rdf:Description></rdf:RDF>
</x:xmpmeta>"""

    def pdf_run(self, command: list[str]) -> str:
        if command[0].endswith("pdfdetach"):
            return "0 embedded files\n"
        if "-meta" in command:
            return self.XMP
        if "-js" in command:
            return ""
        return self.PDF_INFO

    def test_accepts_sanitized_pdf_without_attachments_or_javascript(self) -> None:
        with patch.object(package_submission.shutil, "which", side_effect=lambda name: f"/bin/{name}"), patch.object(
            package_submission, "run", side_effect=self.pdf_run
        ):
            package_submission.validate_pdf_privacy(Path("/tmp/report.pdf"))

    def test_rejects_private_path_in_pdf_metadata(self) -> None:
        def leaked_run(command: list[str]) -> str:
            value = self.pdf_run(command)
            if len(command) == 2 and command[0].endswith("pdfinfo"):
                return value + "Subject: /Users/person/private/report.docx\n"
            return value

        with patch.object(package_submission.shutil, "which", side_effect=lambda name: f"/bin/{name}"), patch.object(
            package_submission, "run", side_effect=leaked_run
        ):
            with self.assertRaisesRegex(package_submission.GateError, "private path"):
                package_submission.validate_pdf_privacy(Path("/tmp/report.pdf"))

    def test_rejects_docx_revision_session_identifiers(self) -> None:
        core = b'''<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties
 xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/">
  <dc:creator>RouteContract project</dc:creator>
  <cp:lastModifiedBy>RouteContract project</cp:lastModifiedBy>
</cp:coreProperties>'''
        document = b'''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p w:rsidR="00112233"><w:r><w:t>safe</w:t></w:r></w:p></w:body>
</w:document>'''
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "report.docx"
            with ZipFile(path, "w", ZIP_DEFLATED) as package:
                package.writestr("docProps/core.xml", core)
                package.writestr("word/document.xml", document)
            with self.assertRaisesRegex(
                package_submission.GateError, "revision session identifiers"
            ):
                package_submission.validate_docx_privacy(path)

    def test_accepts_docx_without_revision_identifiers_in_all_privacy_parts(self) -> None:
        core = b'''<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties
 xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/">
  <dc:creator>RouteContract project</dc:creator>
  <cp:lastModifiedBy>RouteContract project</cp:lastModifiedBy>
</cp:coreProperties>'''
        parts = {
            "word/document.xml": b'''<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>safe</w:t></w:r></w:p></w:body></w:document>''',
            "word/settings.xml": b'''<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:zoom w:percent="100"/></w:settings>''',
            "word/styles.xml": b'''<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:type="paragraph"><w:name w:val="Normal"/></w:style></w:styles>''',
        }
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "report.docx"
            with ZipFile(path, "w", ZIP_DEFLATED) as package:
                package.writestr("docProps/core.xml", core)
                for name, content in parts.items():
                    package.writestr(name, content)

            package_submission.validate_docx_privacy(path)

    def test_rejects_docx_revision_identifiers_in_settings_and_styles(self) -> None:
        core = b'''<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties
 xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/">
  <dc:creator>RouteContract project</dc:creator>
  <cp:lastModifiedBy>RouteContract project</cp:lastModifiedBy>
</cp:coreProperties>'''
        document = b'''<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body><w:p><w:r><w:t>safe</w:t></w:r></w:p></w:body></w:document>'''
        leaks = {
            "settings-container": (
                "word/settings.xml",
                b'''<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rsids><w:rsidRoot w:val="00112233"/><w:rsid w:val="44556677"/></w:rsids></w:settings>''',
            ),
            "styles-element": (
                "word/styles.xml",
                b'''<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:type="paragraph"><w:rsid w:val="00112233"/></w:style></w:styles>''',
            ),
            "styles-attribute": (
                "word/styles.xml",
                b'''<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:type="paragraph" w:rsidR="00112233"/></w:styles>''',
            ),
        }
        for label, (part_name, content) in leaks.items():
            with self.subTest(label=label), tempfile.TemporaryDirectory() as raw:
                path = Path(raw) / "report.docx"
                with ZipFile(path, "w", ZIP_DEFLATED) as package:
                    package.writestr("docProps/core.xml", core)
                    package.writestr("word/document.xml", document)
                    package.writestr(part_name, content)
                with self.assertRaisesRegex(
                    package_submission.GateError, "revision session identifiers"
                ):
                    package_submission.validate_docx_privacy(path)

    def test_report_metadata_errors_never_echo_private_values(self) -> None:
        canary = "PRIVATECANARYSECRETXYZ"
        with self.assertRaisesRegex(
            package_submission.GateError, "EMAIL_ADDRESS"
        ) as caught:
            package_submission.reject_sensitive_metadata(
                f"owner-{canary}@example.com", "report metadata"
            )
        self.assertNotIn(canary, str(caught.exception))
        self.assertNotIn("example.com", str(caught.exception))

        def run_with_pdf_info_canary(command: list[str]) -> str:
            value = self.pdf_run(command)
            if len(command) == 2 and command[0].endswith("pdfinfo"):
                return value.replace(
                    "Author: RouteContract project", f"Author: {canary}"
                )
            return value

        xmp_canary = self.XMP.replace("RouteContract project", canary)

        def run_with_xmp_canary(command: list[str]) -> str:
            if "-meta" in command:
                return xmp_canary
            return self.pdf_run(command)

        def run_with_attachment_canary(command: list[str]) -> str:
            if command[0].endswith("pdfdetach"):
                return f"1 embedded file\n1: {canary}.txt\n"
            return self.pdf_run(command)

        cases = (
            (run_with_pdf_info_canary, "field_count=1"),
            (run_with_xmp_canary, "field=creator"),
            (run_with_attachment_canary, "EMBEDDED_FILE"),
        )
        for side_effect, expected in cases:
            with self.subTest(expected=expected), patch.object(
                package_submission.shutil,
                "which",
                side_effect=lambda name: f"/bin/{name}",
            ), patch.object(
                package_submission, "run", side_effect=side_effect
            ), self.assertRaisesRegex(
                package_submission.GateError, expected
            ) as caught:
                package_submission.validate_pdf_privacy(Path("/tmp/report.pdf"))
            self.assertNotIn(canary, str(caught.exception))

        sensitive_key = f"author_{canary}"
        with self.assertRaisesRegex(
            package_submission.GateError, "index=0"
        ) as caught:
            with patch.object(
                package_submission,
                "SENSITIVE_VIDEO_METADATA_TAGS",
                frozenset({sensitive_key.casefold()}),
            ):
                package_submission.validate_video_metadata_tags(
                    {sensitive_key: "value"}, "format"
                )
        self.assertNotIn(canary, str(caught.exception))

    def test_rejects_symlink_in_input_path(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            target = root / "target.txt"
            target.write_text("value", encoding="utf-8")
            link = root / "link.txt"
            link.symlink_to(target)
            with self.assertRaisesRegex(package_submission.GateError, "symlink"):
                package_submission.require_file(link, "fixture")

    def test_rejects_corrupt_docx_as_gate_error(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "corrupt.docx"
            path.write_bytes(b"not-a-zip")
            for validator in (
                package_submission.extract_docx_text,
                package_submission.validate_docx_privacy,
            ):
                with self.subTest(validator=validator.__name__), self.assertRaises(
                    package_submission.GateError
                ):
                    validator(path)

    def test_requires_repository_output_to_be_gitignored(self) -> None:
        repository = SCRIPT.parents[2]
        package_submission.assert_ignored_if_inside_repository(
            repository / "submission" / "package" / "final",
            repository,
            "ignored output",
        )
        with self.assertRaisesRegex(package_submission.GateError, "not gitignored"):
            package_submission.assert_ignored_if_inside_repository(
                repository / "unexpected-final-output",
                repository,
                "public output",
            )

    def test_invalid_manifest_does_not_create_output_parent(self) -> None:
        repository = SCRIPT.parents[2]
        manifest = repository / "submission" / "package-manifest.example.json"
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw).resolve()
            missing_parent = root / "not-created"
            argv = [
                "package_submission.py",
                "--manifest",
                str(manifest),
                "--template",
                str(SCRIPT),
                "--content",
                str(SCRIPT),
                "--report-pdf",
                str(SCRIPT),
                "--video-file",
                str(SCRIPT),
                "--release-evidence-dir",
                str(root),
                "--release-evidence-artifact",
                str(SCRIPT),
                "--output",
                str(missing_parent / "final"),
                "--repository-root",
                str(repository),
            ]
            with patch.object(sys, "argv", argv), self.assertRaisesRegex(
                package_submission.GateError, "unresolved"
            ):
                package_submission.main()

            self.assertFalse(missing_parent.exists())


class ZipAllowlistTest(unittest.TestCase):
    def test_default_zip_has_exactly_two_deterministic_files(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            upload = root / "upload"
            upload.mkdir()
            filenames = package_submission.validate_manifest(valid_manifest())[
                "official_submission_filenames"
            ]
            names = [filenames["docx"], filenames["pdf"]]
            (upload / names[0]).write_bytes(b"docx")
            (upload / names[1]).write_bytes(b"pdf")
            first = root / "first.zip"
            second = root / "second.zip"
            package_submission.build_upload_zip(upload, first, names)
            package_submission.build_upload_zip(upload, second, names)
            self.assertEqual(hashlib.sha256(first.read_bytes()).digest(), hashlib.sha256(second.read_bytes()).digest())
            with ZipFile(first) as archive:
                self.assertEqual(names, archive.namelist())
                self.assertFalse(any(name.endswith((".json", ".xml", ".mp4", ".jar")) for name in archive.namelist()))

    def test_rejects_legacy_submission_filenames(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            upload = root / "upload"
            upload.mkdir()
            names = sorted(package_submission.LEGACY_SUBMISSION_FILENAMES)
            for name in names:
                (upload / name).write_bytes(b"legacy")
            with self.assertRaisesRegex(package_submission.GateError, "legacy"):
                package_submission.validate_upload_directory(upload, names)
            with self.assertRaisesRegex(package_submission.GateError, "legacy"):
                package_submission.build_upload_zip(upload, root / "legacy.zip", names)


class ChecksumManifestTest(unittest.TestCase):
    def test_accepts_flat_sha256sum_output(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "SHA256SUMS"
            path.write_text(f"{digest('8')}  artifact.jar\n", encoding="utf-8")
            self.assertEqual(
                {"artifact.jar": digest("8")},
                package_submission.parse_checksum_manifest(path),
            )

    def test_rejects_path_traversal(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "SHA256SUMS"
            path.write_text(f"{digest('9')}  ../secret\n", encoding="utf-8")
            with self.assertRaisesRegex(package_submission.GateError, "invalid"):
                package_submission.parse_checksum_manifest(path)


class SourceArchiveTest(unittest.TestCase):
    REQUIRED = ("README.md", "LICENSE", "NOTICE", "build.gradle", "gradlew")

    def write_archive(self, path: Path, extra: tuple[str, ...] = ()) -> None:
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
            for name in (*self.REQUIRED, *extra):
                archive.writestr(f"routecontract-0.1.0/{name}", name)

    def test_accepts_single_root_release_archive(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "source.zip"
            self.write_archive(path)
            self.assertEqual(
                "routecontract-0.1.0",
                package_submission.validate_source_archive(path, "routecontract-0.1.0"),
            )

    def test_rejects_corrupt_source_archive_as_gate_error(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "corrupt-source.zip"
            path.write_bytes(b"not-a-zip")
            with self.assertRaises(package_submission.GateError):
                package_submission.source_archive_members(path, "routecontract-0.1.0")

    def test_rejects_private_submission_material(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "source.zip"
            self.write_archive(path, ("submission/private/final.json",))
            with self.assertRaisesRegex(package_submission.GateError, "leaked"):
                package_submission.validate_source_archive(path, "routecontract-0.1.0")

    def test_matches_final_git_tree_and_rejects_changed_content(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            repository = Path(raw) / "repository"
            repository.mkdir()
            subprocess.run(["git", "init", "-q"], cwd=repository, check=True)
            for name in self.REQUIRED:
                path = repository / name
                path.write_text(name, encoding="utf-8")
            os.chmod(repository / "gradlew", 0o755)
            subprocess.run(["git", "add", "."], cwd=repository, check=True)
            subprocess.run(
                [
                    "git",
                    "-c",
                    "user.name=RouteContract test",
                    "-c",
                    "user.email=test@example.invalid",
                    "commit",
                    "-q",
                    "-m",
                    "fixture",
                ],
                cwd=repository,
                check=True,
            )
            commit = subprocess.run(
                ["git", "rev-parse", "HEAD"],
                cwd=repository,
                check=True,
                capture_output=True,
                text=True,
            ).stdout.strip()
            source = Path(raw) / "source.zip"
            subprocess.run(
                [
                    "git",
                    "archive",
                    "--format=zip",
                    "--prefix=routecontract-0.1.0/",
                    f"--output={source}",
                    commit,
                ],
                cwd=repository,
                check=True,
            )
            package_submission.validate_source_archive_identity(
                source, repository, commit, "routecontract-0.1.0"
            )

            changed = Path(raw) / "changed.zip"
            with ZipFile(source) as original, ZipFile(changed, "w") as rewritten:
                for info in original.infolist():
                    data = original.read(info.filename)
                    if info.filename.endswith("README.md"):
                        data = b"changed"
                    rewritten.writestr(info, data)
            with self.assertRaisesRegex(package_submission.GateError, "not content-identical"):
                package_submission.validate_source_archive_identity(
                    changed, repository, commit, "routecontract-0.1.0"
                )


class ReleaseEvidenceTest(unittest.TestCase):
    PRIVATE_EVIDENCE_NAMES = {
        "environment.txt",
        "mysql-image.txt",
        "routecontract-mysql-example-cyclonedx.json",
        "routecontract-mysql-example-cyclonedx.xml",
        "standalone-consumer.txt",
    }

    def write_public_checksums(self, root: Path) -> None:
        checksum_lines = [
            f"{package_submission.sha256(path)}  {path.name}"
            for path in sorted(root.iterdir())
            if path.name != "SHA256SUMS"
            and path.name not in self.PRIVATE_EVIDENCE_NAMES
        ]
        (root / "SHA256SUMS").write_text(
            "\n".join(checksum_lines) + "\n", encoding="utf-8"
        )

    def rebuild_artifact(self, root: Path, artifact: Path, manifest: dict) -> None:
        with ZipFile(artifact, "w", compression=ZIP_DEFLATED) as archive:
            for path in sorted(root.iterdir()):
                archive.write(path, path.name)
        manifest["release_evidence"]["workflow_artifact_sha256"] = (
            package_submission.sha256(artifact)
        )

    def build_evidence(self, root: Path, manifest: dict) -> Path:
        repository_root = root.parent
        security = repository_root / "security"
        security.mkdir()
        module = repository_root / "routecontract-shardingsphere-5.5"
        module.mkdir()
        (module / "gradle.lockfile").write_text(
            "example:runtime:1.0=runtimeClasspath\n", encoding="utf-8"
        )
        scanner_lock = {
            "database": {
                "ecosystem": "Maven",
                "generation": "1",
                "lastModified": "2026-08-09T03:03:50.782Z",
                "sha256": digest("b"),
                "size": 10,
                "url": "https://example.invalid/Maven/all.zip?generation=1",
            },
            "scanner": {
                "commit": "c" * 40,
                "name": "OSV-Scanner",
                "platforms": {
                    "linux-x86_64": {
                        "sha256": digest("d"),
                        "size": 10,
                        "url": "https://example.invalid/osv-scanner",
                    }
                },
                "scalibrVersion": "0.4.5",
                "version": "2.5.0",
            },
            "schemaVersion": 1,
        }
        (security / "osv-scanner.lock.json").write_text(
            json.dumps(scanner_lock, indent=2, sort_keys=True) + "\n", encoding="utf-8"
        )
        (security / "osv-scanner.toml").write_bytes(b"")
        license_reviews = valid_license_reviews()
        vulnerability_exceptions = valid_vulnerability_exceptions()
        policy = {
            "allowedLicenseIds": ["Apache-2.0"],
            "licenseExceptions": [],
            "licenseReviewExceptions": license_reviews,
            "schemaVersion": 3,
            "vulnerabilityExceptions": vulnerability_exceptions,
        }
        (security / "supply-chain-policy.json").write_text(
            json.dumps(policy, indent=2, sort_keys=True) + "\n", encoding="utf-8"
        )
        source_name = "routecontract-0.1.0-source.zip"
        with ZipFile(root / source_name, "w", compression=ZIP_DEFLATED) as archive:
            for name in ("README.md", "LICENSE", "NOTICE", "build.gradle", "gradlew"):
                archive.writestr(f"routecontract-0.1.0/{name}", name)
        files = {
            "environment.txt": f"revision={manifest['project']['commit']}\n",
            "mysql-image.txt": "image_id=sha256:test\n",
            "standalone-consumer.txt": (
                "ROUTECONTRACT_RELEASE_ASSET_CONSUMER "
                "coordinate=io.github.example-owner.routecontract:"
                "routecontract-shardingsphere-5.5:0.1.0 result=VERIFIED_MYSQL\n"
            ),
            "test-summary.txt": valid_test_summary(manifest["project"]["commit"]),
            "routecontract-shardingsphere-5.5.pom": (
                "<project><groupId>io.github.example-owner.routecontract</groupId>"
                "<artifactId>routecontract-shardingsphere-5.5</artifactId>"
                "<version>0.1.0</version></project>\n"
            ),
            "routecontract-shardingsphere-5.5-cyclonedx.json": "{}\n",
            "routecontract-shardingsphere-5.5-cyclonedx.xml": "<bom/>\n",
            "routecontract-aggregate-cyclonedx.json": "{}\n",
            "routecontract-aggregate-cyclonedx.xml": "<bom/>\n",
            "routecontract-mysql-example-cyclonedx.json": "{}\n",
            "routecontract-mysql-example-cyclonedx.xml": "<bom/>\n",
            "routecontract-shardingsphere-5.5-0.1.0.jar": "main\n",
            "routecontract-shardingsphere-5.5-0.1.0-sources.jar": "sources\n",
            "routecontract-shardingsphere-5.5-0.1.0-javadoc.jar": "javadoc\n",
        }
        for name, content in files.items():
            (root / name).write_text(content, encoding="utf-8")
        supply_chain = {
            "exampleProfile": {
                "componentLicenseCount": 2,
                "mavenPackageCount": 2,
                "resolvedProfileSha256": digest("1"),
                "sbomSha256": package_submission.sha256(
                    root / "routecontract-mysql-example-cyclonedx.json"
                ),
                "xmlComponentCount": 2,
                "xmlSha256": package_submission.sha256(
                    root / "routecontract-mysql-example-cyclonedx.xml"
                ),
            },
            "publishedModule": {
                "componentLicenseCount": 2,
                "dependencyLockSha256": package_submission.sha256(
                    module / "gradle.lockfile"
                ),
                "mavenPackageCount": 2,
                "pomDependencyCount": 1,
                "pomSha256": package_submission.sha256(
                    root / "routecontract-shardingsphere-5.5.pom"
                ),
                "resolvedProfileSha256": digest("4"),
                "runtimeClosureCount": 1,
                "runtimeClosureSha256": digest("5"),
                "sbomSha256": package_submission.sha256(
                    root / "routecontract-shardingsphere-5.5-cyclonedx.json"
                ),
                "xmlComponentCount": 2,
                "xmlSha256": package_submission.sha256(
                    root / "routecontract-shardingsphere-5.5-cyclonedx.xml"
                ),
            },
            "revision": manifest["project"]["commit"],
            "sbom": {
                "componentLicenseCount": 4,
                "inventorySha256": digest("6"),
                "licensePolicy": "passed",
                "licenseReviews": license_reviews,
                "mavenPackageCount": 4,
                "policySha256": package_submission.sha256(
                    security / "supply-chain-policy.json"
                ),
                "sha256": package_submission.sha256(
                    root / "routecontract-aggregate-cyclonedx.json"
                ),
                "unresolvedLicenseReviewCount": 2,
                "xmlComponentCount": 4,
                "xmlSha256": package_submission.sha256(
                    root / "routecontract-aggregate-cyclonedx.xml"
                ),
            },
            "scanner": {
                "binarySha256": scanner_lock["scanner"]["platforms"]["linux-x86_64"]["sha256"],
                "binarySize": 10,
                "binaryUrl": "https://example.invalid/osv-scanner",
                "commit": "c" * 40,
                "database": scanner_lock["database"],
                "name": "OSV-Scanner",
                "platform": "linux-x86_64",
                "scalibrVersion": "0.4.5",
                "scannerConfigSha256": hashlib.sha256(b"").hexdigest(),
                "scannerLockSha256": package_submission.sha256(
                    security / "osv-scanner.lock.json"
                ),
                "version": "2.5.0",
            },
            "schemaVersion": 1,
            "sourceTree": "e" * 40,
            "vulnerabilities": {
                "acceptedExceptionCount": 1,
                "findingCount": 1,
                "findings": findings_for(vulnerability_exceptions),
                "unreviewedCount": 0,
            },
        }
        (root / "supply-chain-evidence.json").write_text(
            json.dumps(supply_chain, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        source_hash = package_submission.sha256(root / source_name)
        json_hash = package_submission.sha256(root / "routecontract-aggregate-cyclonedx.json")
        xml_hash = package_submission.sha256(root / "routecontract-aggregate-cyclonedx.xml")
        manifest["release_evidence"].update(
            {
                "source_archive_sha256": source_hash,
                "aggregate_sbom_json_sha256": json_hash,
                "aggregate_sbom_xml_sha256": xml_hash,
            }
        )
        self.write_public_checksums(root)
        artifact = root.parent / "release-evidence.zip"
        self.rebuild_artifact(root, artifact, manifest)
        return artifact

    def test_rejects_corrupt_workflow_artifact_as_gate_error(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "corrupt-artifact.zip"
            path.write_bytes(b"not-a-zip")
            with self.assertRaises(package_submission.GateError):
                package_submission.zip_flat_file_metadata(path, "workflow artifact ZIP")

    def test_rejects_workflow_artifact_mutation_while_reading_members(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checked = package_submission.validate_manifest(manifest)
            original = package_submission.zip_flat_file_metadata

            def mutate_after_member_read(path: Path, label: str):
                result = original(path, label)
                with path.open("ab") as stream:
                    stream.write(b"changed")
                return result

            with patch.object(
                package_submission,
                "zip_flat_file_metadata",
                side_effect=mutate_after_member_read,
            ), self.assertRaisesRegex(
                package_submission.GateError,
                "workflow artifact ZIP changed during validation",
            ):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_accepts_complete_checksummed_release_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checked = package_submission.validate_manifest(manifest)
            commands: list[list[str]] = []

            def run_command(command: list[str], **_: object) -> str:
                if command[:2] == ["git", "rev-parse"]:
                    return "e" * 40 + "\n"
                commands.append(command)
                return ""

            with patch.object(
                package_submission,
                "run",
                side_effect=run_command,
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ):
                result = package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )
            self.assertEqual(16, result["release_evidence_file_count"])
            self.assertEqual(17, result["workflow_artifact_file_count"])
            self.assertEqual(12, len(result["public_release_assets"]))
            self.assertEqual(52, result["test_summary"]["test_count"])
            self.assertIn("SHA256SUMS", result["public_release_assets"])
            self.assertIn("test-summary.txt", result["public_release_assets"])
            self.assertIn("supply-chain-evidence.json", result["public_release_assets"])
            self.assertEqual(0, result["supply_chain"]["unreviewed_count"])
            self.assertEqual(2, result["supply_chain"]["unresolved_license_review_count"])
            self.assertEqual(1, result["supply_chain"]["finding_count"])
            self.assertNotIn("environment.txt", result["public_release_assets"])
            self.assertNotIn(
                "routecontract-mysql-example-cyclonedx.json",
                result["public_release_assets"],
            )
            self.assertEqual(2, len(commands))
            self.assertEqual(3, commands[0].count("--verify-pair"))
            self.assertEqual(3, commands[1].count("--pair"))
            self.assertTrue(
                commands[1][1].endswith("scripts/validate-official-cyclonedx.py")
            )
            input_root = commands[1].index("--input-root")
            self.assertEqual(str(root), commands[1][input_root + 1])

    def test_rejects_missing_workflow_only_example_pair_member(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            (root / "routecontract-mysql-example-cyclonedx.xml").unlink()
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)

            with self.assertRaisesRegex(package_submission.GateError, "exact allowlist"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_example_pair_in_public_sha256sums(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            example = root / "routecontract-mysql-example-cyclonedx.json"
            checksum_path = root / "SHA256SUMS"
            checksum_path.write_text(
                checksum_path.read_text(encoding="utf-8")
                + f"{package_submission.sha256(example)}  {example.name}\n",
                encoding="utf-8",
            )
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)

            with self.assertRaisesRegex(package_submission.GateError, "public SHA256SUMS"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_official_six_file_validation_failure_stops_packaging(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checked = package_submission.validate_manifest(manifest)

            def fail_official(command: list[str], **_: object) -> str:
                if command[1].endswith("scripts/validate-official-cyclonedx.py"):
                    raise package_submission.GateError("official validation failed")
                return ""

            with patch.object(
                package_submission, "run", side_effect=fail_official
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(
                package_submission.GateError, "official validation failed"
            ):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_sbom_mutation_after_workflow_artifact_binding(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checked = package_submission.validate_manifest(manifest)
            original = package_submission.validate_workflow_artifact_archive

            def mutate_after_binding(*args: object, **kwargs: object):
                result = original(*args, **kwargs)
                with (
                    root / "routecontract-mysql-example-cyclonedx.json"
                ).open("a", encoding="utf-8") as stream:
                    stream.write(" \n")
                return result

            with patch.object(
                package_submission,
                "validate_workflow_artifact_archive",
                side_effect=mutate_after_binding,
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(
                package_submission.GateError,
                "no longer matches the workflow artifact ZIP",
            ):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_mutation_between_semantic_and_official_sbom_validation(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checked = package_submission.validate_manifest(manifest)

            def mutate_after_semantic(command: list[str], **_: object) -> str:
                if command[1].endswith("scripts/finalize-sbom.py"):
                    with (
                        root / "routecontract-aggregate-cyclonedx.json"
                    ).open("a", encoding="utf-8") as stream:
                        stream.write("mutated after semantic validation\n")
                return ""

            with patch.object(
                package_submission, "run", side_effect=mutate_after_semantic
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(
                package_submission.GateError,
                "changed between semantic and official validation",
            ):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_boolean_supply_chain_vulnerability_count(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            evidence["vulnerabilities"]["unreviewedCount"] = False
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(
                package_submission.GateError, "must be a non-negative integer"
            ):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_non_integer_supply_chain_schema_version(self) -> None:
        for invalid_version in (True, 1.0):
            with self.subTest(value=invalid_version), tempfile.TemporaryDirectory() as raw:
                raw_root = Path(raw).resolve()
                root = raw_root / "evidence"
                root.mkdir()
                manifest = valid_manifest()
                artifact = self.build_evidence(root, manifest)
                evidence_path = root / "supply-chain-evidence.json"
                evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
                evidence["schemaVersion"] = invalid_version
                evidence_path.write_text(
                    json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                    encoding="utf-8",
                )
                self.write_public_checksums(root)
                self.rebuild_artifact(root, artifact, manifest)
                checked = package_submission.validate_manifest(manifest)
                with patch.object(
                    package_submission, "run", return_value="e" * 40 + "\n"
                ), patch.object(
                    package_submission, "validate_source_archive_identity", return_value=None
                ), self.assertRaisesRegex(package_submission.GateError, "schemaVersion"):
                    package_submission.validate_release_evidence(
                        root, artifact, checked, raw_root
                    )

    def test_accepts_reviewed_exception_and_rejects_non_boolean_reachability(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ):
                accepted = package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )
            self.assertEqual(1, accepted["supply_chain"]["finding_count"])

            evidence["vulnerabilities"]["findings"][0]["reachabilityEvidence"] = {
                "exampleProfile": 1,
                "publishedProfile": 0,
                "publishedRuntime": 0,
            }
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "flags must be booleans"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

            evidence["vulnerabilities"]["findings"][0]["reachabilityEvidence"] = {
                "exampleProfile": True,
                "publishedProfile": False,
                "publishedRuntime": False,
            }
            evidence["vulnerabilities"]["findings"][0]["exceptionId"] = "RC-TAMPERED"
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "differs from policy"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_supply_chain_evidence_for_another_revision(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            evidence["revision"] = "2" * 40
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "revision does not match"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_supply_chain_evidence_not_bound_to_public_sbom(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            evidence["sbom"]["sha256"] = "f" * 64
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "not source/artifact bound"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_supply_chain_evidence_not_bound_to_workflow_example_sbom(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            evidence["exampleProfile"]["sbomSha256"] = "f" * 64
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "workflow-artifact bound"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_incomplete_unresolved_license_review_set(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            evidence["sbom"]["licenseReviews"].pop()
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "exactly two license reviews"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_reversed_license_review_order(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            evidence_path = root / "supply-chain-evidence.json"
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
            evidence["sbom"]["licenseReviews"].reverse()
            evidence_path.write_text(
                json.dumps(evidence, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with patch.object(
                package_submission, "run", return_value="e" * 40 + "\n"
            ), patch.object(
                package_submission, "validate_source_archive_identity", return_value=None
            ), self.assertRaisesRegex(package_submission.GateError, "exact ordered policy"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_test_summary_for_another_revision(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            (root / "test-summary.txt").write_text(
                valid_test_summary("2" * 40), encoding="utf-8"
            )
            self.write_public_checksums(root)
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "7-suite/52-test"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_file_not_covered_by_sha256sums(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            (root / "unexpected.txt").write_text("not checksummed", encoding="utf-8")
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "byte-identical|unlisted"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_unexpected_file_even_when_artifact_and_checksums_match(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            (root / "private.log").write_text("unexpected", encoding="utf-8")
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "exact allowlist"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_workflow_only_log_in_public_sha256sums(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checksum_path = root / "SHA256SUMS"
            checksum_path.write_text(
                checksum_path.read_text(encoding="utf-8")
                + f"{package_submission.sha256(root / 'environment.txt')}  environment.txt\n",
                encoding="utf-8",
            )
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "public SHA256SUMS"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_missing_public_checksum_entry(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            checksum_path = root / "SHA256SUMS"
            lines = checksum_path.read_text(encoding="utf-8").splitlines()
            checksum_path.write_text(
                "\n".join(
                    line for line in lines if not line.endswith("  test-summary.txt")
                )
                + "\n",
                encoding="utf-8",
            )
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "public SHA256SUMS"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_tampered_public_payload(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            (root / "routecontract-shardingsphere-5.5-0.1.0.jar").write_text(
                "tampered\n", encoding="utf-8"
            )
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "checksum mismatch"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_legacy_source_published_consumer_marker(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            (root / "standalone-consumer.txt").write_text(
                "ROUTECONTRACT_STANDALONE artifact=published-jar\n",
                encoding="utf-8",
            )
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "final-asset"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )

    def test_rejects_duplicate_final_asset_consumer_marker(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            raw_root = Path(raw).resolve()
            root = raw_root / "evidence"
            root.mkdir()
            manifest = valid_manifest()
            artifact = self.build_evidence(root, manifest)
            standalone = root / "standalone-consumer.txt"
            marker = standalone.read_text(encoding="utf-8")
            standalone.write_text(marker + marker, encoding="utf-8")
            self.rebuild_artifact(root, artifact, manifest)
            checked = package_submission.validate_manifest(manifest)
            with self.assertRaisesRegex(package_submission.GateError, "marker once"):
                package_submission.validate_release_evidence(
                    root, artifact, checked, raw_root
                )


class PublicEvidenceTest(unittest.TestCase):
    def test_public_youtube_contract_preserves_title_and_duration_gates(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        local_video = {"duration_seconds": 173.0}
        youtube = {
            "title": "RouteContract demo",
            "duration_seconds": 172.5,
        }
        package_submission.validate_public_youtube_contract(
            manifest, local_video, youtube
        )

        changed_title = dict(youtube, title="Different title")
        with self.assertRaisesRegex(package_submission.GateError, "title mismatch"):
            package_submission.validate_public_youtube_contract(
                manifest, local_video, changed_title
            )

        over_limit = dict(youtube, duration_seconds=175.1)
        with self.assertRaisesRegex(
            package_submission.GateError, "170 through 175 seconds inclusive"
        ):
            package_submission.validate_public_youtube_contract(
                manifest, local_video, over_limit
            )

        different_upload = dict(youtube, duration_seconds=174.1)
        with self.assertRaisesRegex(
            package_submission.GateError, "checksummed local video"
        ):
            package_submission.validate_public_youtube_contract(
                manifest, local_video, different_upload
            )

    def test_public_youtube_contract_accepts_exact_one_second_rounding(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        package_submission.validate_public_youtube_contract(
            manifest,
            {"duration_seconds": 173.5},
            {"title": "RouteContract demo", "duration_seconds": 172.5},
        )

    def test_public_youtube_duration_window_and_caption_end_are_inclusive(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        for duration in (172.5, 175.0):
            with self.subTest(duration=duration):
                package_submission.validate_public_youtube_contract(
                    manifest,
                    {"duration_seconds": duration},
                    {"title": "RouteContract demo", "duration_seconds": duration},
                )

        for duration in (170.0, 172.499):
            with self.subTest(duration=duration), self.assertRaisesRegex(
                package_submission.GateError,
                "selected caption branch's final cue",
            ):
                package_submission.validate_public_youtube_contract(
                    manifest,
                    {"duration_seconds": duration},
                    {"title": "RouteContract demo", "duration_seconds": duration},
                )

        for duration in (169.9, 175.1):
            with self.subTest(duration=duration), self.assertRaisesRegex(
                package_submission.GateError,
                "170 through 175 seconds inclusive",
            ):
                package_submission.validate_public_youtube_contract(
                    manifest,
                    {"duration_seconds": duration},
                    {"title": "RouteContract demo", "duration_seconds": duration},
                )

    def test_tls_context_uses_the_pinned_certifi_bundle(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            ca_bundle = Path(raw) / "cacert.pem"
            ca_bundle.write_text("test CA bundle", encoding="utf-8")
            fake_certifi = type(
                "FakeCertifi",
                (),
                {"where": staticmethod(lambda: str(ca_bundle))},
            )
            sentinel = object()
            with patch.object(
                package_submission.importlib,
                "import_module",
                return_value=fake_certifi,
            ) as import_module, patch.object(
                package_submission.ssl,
                "create_default_context",
                return_value=sentinel,
            ) as create_context:
                context = package_submission.verified_tls_context()

            self.assertIs(sentinel, context)
            import_module.assert_called_once_with("certifi")
            create_context.assert_called_once_with(cafile=str(ca_bundle))

    def test_tls_context_fails_without_certifi(self) -> None:
        with patch.object(
            package_submission.importlib,
            "import_module",
            side_effect=ModuleNotFoundError("certifi"),
        ):
            with self.assertRaisesRegex(
                package_submission.GateError, "pinned certifi CA bundle"
            ):
                package_submission.verified_tls_context()

    def test_youtube_probe_ignores_local_config_and_playlists(self) -> None:
        url = "https://www.youtube.com/watch?v=abcdefghijk"
        commands: list[list[str]] = []

        def fake_run(command: list[str], **kwargs: object) -> str:
            commands.append(command)
            self.assertEqual(
                "yt-dlp public video metadata probe", kwargs.get("failure_label")
            )
            return json.dumps(valid_youtube_probe())

        with patch.object(
            package_submission,
            "request_json",
            return_value={"title": "RouteContract demo"},
        ), patch.object(
            package_submission.shutil, "which", return_value="/usr/local/bin/yt-dlp"
        ), patch.object(package_submission, "run", side_effect=fake_run):
            result = package_submission.public_youtube_metadata(url)

        self.assertEqual(173.0, result["duration_seconds"])
        self.assertEqual("public", result["availability"])
        self.assertEqual("not_live", result["live_status"])
        self.assertEqual(0, result["age_limit"])
        self.assertEqual(1080, result["max_video_height"])
        self.assertEqual(
            [
                "/usr/local/bin/yt-dlp",
                "--ignore-config",
                "--no-playlist",
                "--check-all-formats",
                "--dump-single-json",
                "--skip-download",
                "--",
                url,
            ],
            commands[0],
        )

    def test_youtube_probe_rejects_duplicate_nonfinite_and_deep_json_generically(
        self,
    ) -> None:
        canary = "CANARY_YTDLP_PRIVATE_KEY"
        malformed = (
            '{"' + canary + '":1,"' + canary + '":2}',
            '{"value":Infinity}',
            '{"value":1e999}',
            PublicJsonTransportTest.nested_json(
                package_submission.REPORT_CONTENT_CONTRACT.STRICT_JSON_MAX_CONTAINER_DEPTH
                + 1,
                "object",
            ).decode(),
            "[" * 10_000 + "0" + "]" * 10_000,
            " " * (package_submission.MAX_JSON_TOOL_OUTPUT_BYTES + 1),
        )
        url = "https://www.youtube.com/watch?v=abcdefghijk"
        for output in malformed:
            with self.subTest(prefix=output[:24]), patch.object(
                package_submission, "request_json", return_value={"title": "RouteContract demo"}
            ), patch.object(
                package_submission.shutil, "which", return_value="/usr/local/bin/yt-dlp"
            ), patch.object(
                package_submission, "run", return_value=output
            ), self.assertRaisesRegex(
                package_submission.GateError,
                "incomplete public video duration metadata",
            ) as caught:
                package_submission.public_youtube_metadata(url)
            self.assertIsNone(caught.exception.__cause__)
            self.assertNotIn(canary, str(caught.exception))
            self.assertNotIn("9" * 32, str(caught.exception))

    def youtube_probe(self, metadata: object) -> dict:
        url = "https://www.youtube.com/watch?v=abcdefghijk"
        with patch.object(
            package_submission,
            "request_json",
            return_value={"title": "RouteContract demo"},
        ), patch.object(
            package_submission.shutil, "which", return_value="/usr/local/bin/yt-dlp"
        ), patch.object(
            package_submission, "run", return_value=json.dumps(metadata)
        ):
            return package_submission.public_youtube_metadata(url)

    def test_youtube_probe_accepts_none_age_limit_and_4k_video(self) -> None:
        metadata = valid_youtube_probe()
        metadata["age_limit"] = None
        metadata["formats"][1]["height"] = 2160

        result = self.youtube_probe(metadata)

        self.assertIsNone(result["age_limit"])
        self.assertEqual(2160, result["max_video_height"])

        metadata["age_limit"] = 0.0
        result = self.youtube_probe(metadata)
        self.assertEqual(0.0, result["age_limit"])

    def test_youtube_probe_requires_yt_dlp_for_public_format_evidence(self) -> None:
        with patch.object(
            package_submission,
            "request_json",
            return_value={"title": "RouteContract demo"},
        ), patch.object(package_submission.shutil, "which", return_value=None):
            with self.assertRaisesRegex(package_submission.GateError, "yt-dlp is required"):
                package_submission.public_youtube_metadata(
                    "https://www.youtube.com/watch?v=abcdefghijk"
                )

    def test_youtube_probe_rejects_non_public_or_missing_availability(self) -> None:
        for availability in (None, "unlisted", "private", "needs_auth"):
            metadata = valid_youtube_probe()
            if availability is None:
                del metadata["availability"]
            else:
                metadata["availability"] = availability
            with self.subTest(availability=availability), self.assertRaisesRegex(
                package_submission.GateError, "availability must be public"
            ):
                self.youtube_probe(metadata)

    def test_youtube_probe_rejects_live_upcoming_or_unknown_status(self) -> None:
        for live_status in (
            None,
            "is_live",
            "is_upcoming",
            "post_live",
            "was_live",
        ):
            metadata = valid_youtube_probe()
            if live_status is None:
                del metadata["live_status"]
            else:
                metadata["live_status"] = live_status
            with self.subTest(live_status=live_status), self.assertRaisesRegex(
                package_submission.GateError, "must be a non-live upload"
            ):
                self.youtube_probe(metadata)

    def test_youtube_probe_rejects_restricted_or_malformed_age_limit(self) -> None:
        for age_limit in (18, 1, -1, "0", False):
            metadata = valid_youtube_probe()
            metadata["age_limit"] = age_limit
            with self.subTest(age_limit=age_limit), self.assertRaisesRegex(
                package_submission.GateError, "age_limit must be 0 or null"
            ):
                self.youtube_probe(metadata)

        missing = valid_youtube_probe()
        del missing["age_limit"]
        with self.assertRaisesRegex(
            package_submission.GateError, "age_limit must be 0 or null"
        ):
            self.youtube_probe(missing)

    def test_youtube_probe_rejects_missing_or_low_resolution_video_formats(
        self,
    ) -> None:
        invalid_formats = (
            [],
            [
                {
                    "format_id": "136",
                    "height": 720,
                    "vcodec": "avc1.4d401f",
                    "url": "https://example.invalid/720p",
                }
            ],
            [
                {
                    "format_id": "140",
                    "height": 1080,
                    "vcodec": "none",
                    "url": "https://example.invalid/audio",
                }
            ],
            [
                {
                    "format_id": "137",
                    "height": 1080,
                    "vcodec": "avc1.640028",
                    "url": "",
                    "has_drm": False,
                }
            ],
            [
                {
                    "format_id": "137",
                    "height": 1080,
                    "vcodec": "avc1.640028",
                    "url": "https://example.invalid/drm",
                    "has_drm": True,
                }
            ],
            [
                {
                    "format_id": "137",
                    "height": 1080,
                    "vcodec": "avc1.640028",
                    "url": "https://example.invalid/missing-drm-status",
                }
            ],
            [
                {
                    "format_id": "137",
                    "height": 1080,
                    "vcodec": "avc1.640028",
                    "url": "https://example.invalid/unknown-drm-status",
                    "has_drm": None,
                }
            ],
        )
        for formats in invalid_formats:
            metadata = valid_youtube_probe()
            metadata["formats"] = formats
            with self.subTest(formats=formats), self.assertRaisesRegex(
                package_submission.GateError, "downloadable video format at 1080p"
            ):
                self.youtube_probe(metadata)

    def test_youtube_probe_rejects_malformed_format_evidence(self) -> None:
        malformed_formats = (
            None,
            {},
            ["not-an-object"],
            [{"height": True}],
            [{"has_drm": "false"}],
        )
        for formats in malformed_formats:
            metadata = valid_youtube_probe()
            metadata["formats"] = formats
            with self.subTest(formats=formats), self.assertRaisesRegex(
                package_submission.GateError, "format metadata"
            ):
                self.youtube_probe(metadata)

    def test_youtube_probe_rejects_empty_or_none_like_video_codec(self) -> None:
        for vcodec in ("", "   ", "none", "NONE", " None "):
            metadata = valid_youtube_probe()
            metadata["formats"][1]["vcodec"] = vcodec
            with self.subTest(vcodec=vcodec), self.assertRaisesRegex(
                package_submission.GateError, "downloadable video format at 1080p"
            ):
                self.youtube_probe(metadata)

    def test_youtube_probe_preserves_exact_id_and_finite_duration(self) -> None:
        wrong_id = valid_youtube_probe()
        wrong_id["id"] = "zyxwvutsrqp"
        with self.assertRaisesRegex(package_submission.GateError, "different YouTube video ID"):
            self.youtube_probe(wrong_id)

        for duration in (None, True, 0, -1, float("nan"), float("inf")):
            metadata = valid_youtube_probe()
            if duration is None:
                del metadata["duration"]
            else:
                metadata["duration"] = duration
            with self.subTest(duration=duration), self.assertRaisesRegex(
                package_submission.GateError, "duration"
            ):
                self.youtube_probe(metadata)

        with self.assertRaisesRegex(
            package_submission.GateError, "incomplete public video duration metadata"
        ):
            self.youtube_probe("not-an-object")

    def test_youtube_schema_diagnostics_never_echo_json_values(self) -> None:
        canary = "CANARY_YTDLP_PRIVATE_VALUE"
        cases = []
        for field in ("duration", "availability", "live_status", "age_limit"):
            metadata = valid_youtube_probe()
            metadata[field] = canary
            cases.append((field, metadata))
        for field in ("duration", "age_limit"):
            metadata = valid_youtube_probe()
            metadata[field] = int("9" * 1_000)
            cases.append((f"oversized-{field}", metadata))
        oversized_height = valid_youtube_probe()
        oversized_height["formats"][1]["height"] = int("9" * 1_000)
        cases.append(("oversized-height", oversized_height))
        for field, metadata in cases:
            with self.subTest(field=field), self.assertRaises(
                package_submission.GateError
            ) as caught:
                self.youtube_probe(metadata)
            self.assertIsNone(caught.exception.__cause__)
            self.assertNotIn(canary, str(caught.exception))
            self.assertNotIn("9" * 32, str(caught.exception))

        manifest = package_submission.validate_manifest(valid_manifest())
        with self.assertRaises(package_submission.GateError) as caught:
            package_submission.validate_public_youtube_contract(
                manifest,
                {"duration_seconds": 173.0},
                {"title": canary, "duration_seconds": 172.5},
            )
        self.assertIsNone(caught.exception.__cause__)
        self.assertNotIn(canary, str(caught.exception))

    def test_accepts_matching_public_revision_release_and_video(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        project = manifest["project"]
        evidence = {
            "source_archive_filename": "routecontract-0.1.0-source.zip",
            "source_archive_sha256": digest("5"),
            "source_archive_size": 123,
            "workflow_artifact_size": 456,
            "public_release_assets": {
                "routecontract-0.1.0-source.zip": {"sha256": digest("5"), "size": 123}
            },
        }

        def fake_json(url: str) -> dict:
            if "/actions/runs/" in url:
                return {
                    "id": 123456,
                    "html_url": project["ci_run_url"],
                    "status": "completed",
                    "conclusion": "success",
                    "head_sha": project["commit"],
                    "head_branch": project["tag"],
                    "event": "push",
                    "path": ".github/workflows/release-evidence.yml",
                    "name": "Release evidence",
                    "repository": {"full_name": "example-owner/routecontract"},
                    "created_at": "2026-07-19T23:00:00Z",
                    "updated_at": "2026-07-19T23:30:00Z",
                }
            if "/actions/artifacts/" in url:
                return {
                    "id": 987654,
                    "url": (
                        "https://api.github.com/repos/example-owner/routecontract/"
                        "actions/artifacts/987654"
                    ),
                    "archive_download_url": (
                        "https://api.github.com/repos/example-owner/routecontract/"
                        "actions/artifacts/987654/zip"
                    ),
                    "name": f"routecontract-release-evidence-{project['commit']}",
                    "size_in_bytes": 456,
                    "expired": False,
                    "digest": f"sha256:{digest('a')}",
                    "created_at": "2026-07-19T23:10:00Z",
                    "updated_at": "2026-07-19T23:20:00Z",
                    "expires_at": "2026-08-19T23:20:00Z",
                    "workflow_run": {
                        "id": 123456,
                        "head_sha": project["commit"],
                        "head_branch": project["tag"],
                    },
                }
            if "/releases/tags/" in url:
                return {
                    "id": 7,
                    "created_at": "2026-07-20T00:00:00Z",
                    "published_at": "2026-07-20T00:01:00Z",
                    "updated_at": "2026-07-20T00:02:00Z",
                    "draft": False,
                    "prerelease": False,
                    "immutable": True,
                    "tag_name": project["tag"],
                    "html_url": project["release_url"],
                    "assets": [
                        {
                            "id": 8,
                            "name": evidence["source_archive_filename"],
                            "state": "uploaded",
                            "size": 123,
                            "digest": f"sha256:{digest('5')}",
                            "url": (
                                "https://api.github.com/repos/example-owner/"
                                "routecontract/releases/assets/8"
                            ),
                            "browser_download_url": (
                                "https://github.com/example-owner/routecontract/"
                                "releases/download/v0.1.0/routecontract-0.1.0-source.zip"
                            ),
                            "created_at": "2026-07-20T00:00:30Z",
                            "updated_at": "2026-07-20T00:01:30Z",
                        }
                    ],
                }
            if "/commits/" in url:
                return {
                    "sha": project["commit"],
                    "html_url": (
                        f"{project['repository_url']}/commit/{project['commit']}"
                    ),
                }
            return {
                "id": 1,
                "node_id": "R_example",
                "private": False,
                "archived": False,
                "full_name": "example-owner/routecontract",
                "html_url": project["repository_url"],
            }

        youtube = {
            "id": "abcdefghijk",
            "title": "RouteContract demo",
            "duration_seconds": 172.5,
            "availability": "public",
            "live_status": "not_live",
            "age_limit": 0,
            "max_video_height": 1080,
        }
        public_identity_events: list[str] = []
        with patch.object(
            package_submission, "request_json", side_effect=fake_json
        ), patch.object(
            package_submission, "public_youtube_metadata", return_value=youtube
        ), patch.object(
            package_submission,
            "verify_release_attestations",
            side_effect=lambda *_: public_identity_events.append("attestations"),
        ) as verify_attestations, patch.object(
            package_submission,
            "validate_remote_tag_identity",
            side_effect=lambda *_: public_identity_events.append("remote-tag-recheck"),
        ) as verify_remote_tag:
            result = package_submission.validate_public_evidence(
                manifest,
                {"duration_seconds": 173.0},
                evidence,
                Path("/evidence"),
                Path("/repository"),
            )
        self.assertEqual("success", result["ci_run"]["conclusion"])
        self.assertEqual("v0.1.0", result["release"]["tag_name"])
        self.assertIs(True, result["release"]["immutable"])
        self.assertEqual(8, result["release"]["assets"][0]["id"])
        self.assertEqual(456, result["workflow_artifact"]["size_in_bytes"])
        self.assertEqual("abcdefghijk", result["youtube_video_id"])
        self.assertEqual("RouteContract demo", result["youtube_title"])
        self.assertEqual(172.5, result["youtube_duration_seconds"])
        self.assertEqual("public", result["youtube_availability"])
        self.assertEqual("not_live", result["youtube_live_status"])
        self.assertEqual(0, result["youtube_age_limit"])
        self.assertEqual(1080, result["youtube_max_video_height"])
        verify_attestations.assert_called_once_with(
            manifest, evidence, Path("/evidence")
        )
        verify_remote_tag.assert_called_once_with(Path("/repository"), manifest)
        self.assertEqual(
            ["attestations", "remote-tag-recheck"], public_identity_events
        )

    def test_real_shaped_full_collector_requery_rejects_digest_or_time_drift(self) -> None:
        for drift in ("release_asset_digest", "run_updated_at"):
            manifest = package_submission.validate_manifest(valid_manifest())
            project = manifest["project"]
            asset_name = "routecontract-0.1.0-source.zip"
            evidence = {
                "source_archive_filename": asset_name,
                "source_archive_sha256": digest("5"),
                "source_archive_size": 123,
                "workflow_artifact_size": 456,
                "public_release_assets": {
                    asset_name: {"sha256": digest("5"), "size": 123}
                },
            }
            api_base = "https://api.github.com/repos/example-owner/routecontract"
            observation = 0

            def fake_json(url: str) -> dict:
                nonlocal observation
                if url == api_base:
                    observation += 1
                    return {
                        "id": 1,
                        "node_id": "R_example",
                        "private": False,
                        "archived": False,
                        "full_name": "example-owner/routecontract",
                        "html_url": project["repository_url"],
                    }
                if "/commits/" in url:
                    return {
                        "sha": project["commit"],
                        "html_url": f"{project['repository_url']}/commit/{project['commit']}",
                    }
                if "/actions/runs/" in url:
                    return {
                        "id": 123456,
                        "html_url": project["ci_run_url"],
                        "status": "completed",
                        "conclusion": "success",
                        "head_sha": project["commit"],
                        "head_branch": project["tag"],
                        "event": "push",
                        "path": ".github/workflows/release-evidence.yml",
                        "name": "Release evidence",
                        "repository": {"full_name": "example-owner/routecontract"},
                        "created_at": "2026-07-19T23:00:00Z",
                        "updated_at": (
                            "2026-07-19T23:31:00Z"
                            if drift == "run_updated_at" and observation == 2
                            else "2026-07-19T23:30:00Z"
                        ),
                    }
                if "/actions/artifacts/" in url:
                    return {
                        "id": 987654,
                        "url": f"{api_base}/actions/artifacts/987654",
                        "archive_download_url": f"{api_base}/actions/artifacts/987654/zip",
                        "name": f"routecontract-release-evidence-{project['commit']}",
                        "size_in_bytes": 456,
                        "expired": False,
                        "digest": f"sha256:{digest('a')}",
                        "created_at": "2026-07-19T23:10:00Z",
                        "updated_at": "2026-07-19T23:20:00Z",
                        "expires_at": "2026-08-19T23:20:00Z",
                        "workflow_run": {
                            "id": 123456,
                            "head_sha": project["commit"],
                            "head_branch": project["tag"],
                        },
                    }
                if "/releases/tags/" in url:
                    return {
                        "id": 7,
                        "created_at": "2026-07-20T00:00:00Z",
                        "published_at": "2026-07-20T00:01:00Z",
                        "updated_at": "2026-07-20T00:02:00Z",
                        "draft": False,
                        "prerelease": False,
                        "immutable": True,
                        "tag_name": project["tag"],
                        "html_url": project["release_url"],
                        "assets": [
                            {
                                "id": 8,
                                "name": asset_name,
                                "state": "uploaded",
                                "size": 123,
                                "digest": (
                                    f"sha256:{digest('5')}"
                                    if drift == "release_asset_digest" and observation == 2
                                    else None
                                ),
                                "url": f"{api_base}/releases/assets/8",
                                "browser_download_url": (
                                    f"{project['repository_url']}/releases/download/"
                                    f"{project['tag']}/{asset_name}"
                                ),
                                "created_at": "2026-07-20T00:00:30Z",
                                "updated_at": "2026-07-20T00:01:30Z",
                            }
                        ],
                    }
                self.fail(f"unexpected public collector URL: {url}")

            youtube = {
                "id": "abcdefghijk",
                "title": "RouteContract demo",
                "duration_seconds": 172.5,
                "availability": "public",
                "live_status": "not_live",
                "age_limit": 0,
                "max_video_height": 1080,
            }
            arguments = (
                manifest,
                {"duration_seconds": 173.0},
                evidence,
                Path("/evidence"),
                Path("/repository"),
            )
            with self.subTest(drift=drift), patch.object(
                package_submission, "request_json", side_effect=fake_json
            ), patch.object(
                package_submission, "public_youtube_metadata", return_value=youtube
            ), patch.object(
                package_submission, "hash_remote_file", return_value=digest("5")
            ), patch.object(
                package_submission, "verify_release_attestations"
            ), patch.object(
                package_submission, "validate_remote_tag_identity"
            ):
                initial = package_submission.validate_public_evidence(*arguments)
                with self.assertRaisesRegex(
                    package_submission.GateError,
                    "release/CI/video state changed between packaging observations",
                ):
                    package_submission.revalidate_public_evidence(initial, *arguments)
            self.assertEqual(2, observation)

    def test_release_attestation_verifier_checks_release_and_every_asset(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        evidence = {
            "public_release_assets": {
                "SHA256SUMS": {"sha256": digest("1"), "size": 1},
                "routecontract-0.1.0-source.zip": {
                    "sha256": digest("2"),
                    "size": 1,
                },
            }
        }
        with tempfile.TemporaryDirectory() as raw:
            evidence_dir = Path(raw)
            for asset_name in evidence["public_release_assets"]:
                (evidence_dir / asset_name).write_bytes(b"x")
            repository = "example-owner/routecontract"
            tag = manifest["project"]["tag"]
            expected_commands = [
                ["/usr/local/bin/gh", "version"],
                [
                    "/usr/local/bin/gh",
                    "release",
                    "verify",
                    tag,
                    "--repo",
                    repository,
                ],
                [
                    "/usr/local/bin/gh",
                    "release",
                    "verify-asset",
                    tag,
                    str(evidence_dir / "SHA256SUMS"),
                    "--repo",
                    repository,
                ],
                [
                    "/usr/local/bin/gh",
                    "release",
                    "verify-asset",
                    tag,
                    str(evidence_dir / "routecontract-0.1.0-source.zip"),
                    "--repo",
                    repository,
                ],
            ]
            for version_output in (
                (
                    "gh version 2.93.0 (2026-05-27)\n"
                    "https://github.com/cli/cli/releases/tag/v2.93.0\n"
                ),
                (
                    "gh version 2.97.0 (2026-07-31)\n"
                    "https://github.com/cli/cli/releases/tag/v2.97.0\n"
                ),
            ):
                with self.subTest(version_output=version_output):
                    commands: list[list[str]] = []

                    def fake_subprocess_run(
                        command: list[str], **_: object
                    ) -> subprocess.CompletedProcess[str]:
                        commands.append(command)
                        return subprocess.CompletedProcess(
                            command,
                            0,
                            stdout=version_output,
                            stderr="",
                        )

                    def fake_run(command: list[str]) -> str:
                        commands.append(command)
                        return ""

                    with patch.object(
                        package_submission.shutil,
                        "which",
                        return_value="/usr/local/bin/gh",
                    ), patch.object(
                        package_submission.subprocess,
                        "run",
                        side_effect=fake_subprocess_run,
                    ), patch.object(
                        package_submission, "run", side_effect=fake_run
                    ):
                        package_submission.verify_release_attestations(
                            manifest, evidence, evidence_dir
                        )

                    self.assertEqual(expected_commands, commands)

    def test_release_attestation_verifier_rejects_affected_github_cli_before_verify(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())

        for rendered_version in ("0.0.0", "2.87.3", "2.92.0"):
            with self.subTest(version=rendered_version):
                commands: list[list[str]] = []

                def fake_subprocess_run(
                    command: list[str], **_: object
                ) -> subprocess.CompletedProcess[str]:
                    commands.append(command)
                    return subprocess.CompletedProcess(
                        command,
                        0,
                        stdout=f"gh version {rendered_version}\n",
                        stderr="",
                    )

                with patch.object(
                    package_submission.shutil,
                    "which",
                    return_value="/usr/local/bin/gh",
                ), patch.object(
                    package_submission.subprocess,
                    "run",
                    side_effect=fake_subprocess_run,
                ), patch.object(package_submission, "run") as verify_command:
                    with self.assertRaises(package_submission.GateError) as caught:
                        package_submission.verify_release_attestations(
                            manifest,
                            {"public_release_assets": {}},
                            Path("/evidence"),
                        )

                message = str(caught.exception)
                self.assertIn("2.93.0 or newer", message)
                self.assertIn("GHSA-8xvp-7hj6-mcj9", message)
                self.assertIn(rendered_version, message)
                self.assertEqual([["/usr/local/bin/gh", "version"]], commands)
                verify_command.assert_not_called()

    def test_release_attestation_verifier_rejects_malformed_github_cli_version(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        malformed_outputs = (
            "",
            "GitHub CLI current\n",
            "gh version 02.093.000 (2026-05-27)\n",
            " gh version 2.93.0 (2026-05-27)\n",
            "gh version 2.93.0 (2026-05-27) \n",
            "gh version 2.93.0-rc.1 (2026-05-27)\n",
            "gh version 2.93.0\ngh version 2.97.0\n",
            (
                "gh version 2.93.0 (2026-05-27)\n"
                "https://github.com/cli/cli/releases/tag/v2.97.0\n"
            ),
            "gh version 2.93.0\nunexpected extra line\n",
        )

        for malformed_output in malformed_outputs:
            with self.subTest(output=malformed_output):
                commands: list[list[str]] = []

                def fake_subprocess_run(
                    command: list[str], **_: object
                ) -> subprocess.CompletedProcess[str]:
                    commands.append(command)
                    return subprocess.CompletedProcess(
                        command,
                        0,
                        stdout=malformed_output,
                        stderr="",
                    )

                with patch.object(
                    package_submission.shutil,
                    "which",
                    return_value="/usr/local/bin/gh",
                ), patch.object(
                    package_submission.subprocess,
                    "run",
                    side_effect=fake_subprocess_run,
                ), patch.object(package_submission, "run") as verify_command:
                    with self.assertRaisesRegex(
                        package_submission.GateError,
                        "not an unambiguous stable version",
                    ):
                        package_submission.verify_release_attestations(
                            manifest,
                            {"public_release_assets": {}},
                            Path("/evidence"),
                        )

                self.assertEqual([["/usr/local/bin/gh", "version"]], commands)
                verify_command.assert_not_called()

    def test_release_attestation_verifier_fails_without_or_with_failed_github_cli(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        with patch.object(
            package_submission.shutil, "which", return_value=None
        ), patch.object(package_submission.subprocess, "run") as version_command, patch.object(
            package_submission, "run"
        ) as verify_command:
            with self.assertRaisesRegex(package_submission.GateError, "GitHub CLI"):
                package_submission.verify_release_attestations(
                    manifest, {"public_release_assets": {}}, Path("/evidence")
                )
        version_command.assert_not_called()
        verify_command.assert_not_called()

        failed = subprocess.CompletedProcess(
            ["/usr/local/bin/gh", "version"],
            1,
            stdout="synthetic-sensitive-stdout",
            stderr="synthetic-sensitive-stderr",
        )
        with patch.object(
            package_submission.shutil,
            "which",
            return_value="/usr/local/bin/gh",
        ), patch.object(
            package_submission.subprocess, "run", return_value=failed
        ) as version_command, patch.object(
            package_submission, "run"
        ) as verify_command:
            with self.assertRaisesRegex(
                package_submission.GateError,
                "GitHub CLI version check failed",
            ) as caught:
                package_submission.verify_release_attestations(
                    manifest, {"public_release_assets": {}}, Path("/evidence")
                )
        self.assertNotIn("synthetic-sensitive", str(caught.exception))
        version_command.assert_called_once()
        verify_command.assert_not_called()

    def test_rejects_non_tag_push_release_evidence_run(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        project = manifest["project"]
        evidence = {"public_release_assets": {}}
        base_run = {
            "id": 123456,
            "html_url": project["ci_run_url"],
            "status": "completed",
            "conclusion": "success",
            "head_sha": project["commit"],
            "head_branch": project["tag"],
            "event": "push",
            "path": ".github/workflows/release-evidence.yml",
            "name": "Release evidence",
            "repository": {"full_name": "example-owner/routecontract"},
            "created_at": "2026-07-19T23:00:00Z",
            "updated_at": "2026-07-19T23:30:00Z",
        }
        cases = (
            {"event": "workflow_dispatch"},
            {"head_branch": "main"},
            {"path": ".github/workflows/ci.yml"},
            {"path": ".github/workflows/release-evidence.yml@evil.yaml"},
        )
        for override in cases:
            with self.subTest(override=override):

                def fake_json(url: str) -> dict:
                    if "/actions/runs/" in url:
                        return base_run | override
                    if "/commits/" in url:
                        return {
                            "sha": project["commit"],
                            "html_url": f"{project['repository_url']}/commit/{project['commit']}",
                        }
                    return {
                        "id": 1,
                        "node_id": "R_example",
                        "private": False,
                        "archived": False,
                        "full_name": "example-owner/routecontract",
                        "html_url": project["repository_url"],
                    }

                with patch.object(
                    package_submission, "request_json", side_effect=fake_json
                ):
                    with self.assertRaisesRegex(
                        package_submission.GateError, "expected tag-push workflow"
                    ):
                        package_submission.validate_public_evidence(
                            manifest,
                            {"duration_seconds": 173.0},
                            evidence,
                            Path("/evidence"),
                            Path("/repository"),
                        )

    def test_rejects_artifact_for_non_tag_branch(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        project = manifest["project"]
        evidence = {"public_release_assets": {}, "workflow_artifact_size": 456}

        def fake_json(url: str) -> dict:
            if "/actions/runs/" in url:
                return {
                    "id": 123456,
                    "html_url": project["ci_run_url"],
                    "status": "completed",
                    "conclusion": "success",
                    "head_sha": project["commit"],
                    "head_branch": project["tag"],
                    "event": "push",
                    "path": ".github/workflows/release-evidence.yml",
                    "name": "Release evidence",
                    "repository": {"full_name": "example-owner/routecontract"},
                    "created_at": "2026-07-19T23:00:00Z",
                    "updated_at": "2026-07-19T23:30:00Z",
                }
            if "/actions/artifacts/" in url:
                return {
                    "id": 987654,
                    "url": (
                        "https://api.github.com/repos/example-owner/routecontract/"
                        "actions/artifacts/987654"
                    ),
                    "archive_download_url": (
                        "https://api.github.com/repos/example-owner/routecontract/"
                        "actions/artifacts/987654/zip"
                    ),
                    "name": f"routecontract-release-evidence-{project['commit']}",
                    "size_in_bytes": 456,
                    "expired": False,
                    "digest": f"sha256:{digest('a')}",
                    "created_at": "2026-07-19T23:10:00Z",
                    "updated_at": "2026-07-19T23:20:00Z",
                    "expires_at": "2026-08-19T23:20:00Z",
                    "workflow_run": {
                        "id": 123456,
                        "head_sha": project["commit"],
                        "head_branch": "main",
                    },
                }
            if "/commits/" in url:
                return {
                    "sha": project["commit"],
                    "html_url": f"{project['repository_url']}/commit/{project['commit']}",
                }
            return {
                "id": 1,
                "node_id": "R_example",
                "private": False,
                "archived": False,
                "full_name": "example-owner/routecontract",
                "html_url": project["repository_url"],
            }

        with patch.object(package_submission, "request_json", side_effect=fake_json):
            with self.assertRaisesRegex(package_submission.GateError, "artifact ID/digest"):
                package_submission.validate_public_evidence(
                    manifest,
                    {"duration_seconds": 173.0},
                    evidence,
                    Path("/evidence"),
                    Path("/repository"),
                )

    def test_rejects_mutable_or_unreported_release_immutability(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        project = manifest["project"]
        evidence = {"public_release_assets": {}, "workflow_artifact_size": 456}
        for immutable in (False, None):
            with self.subTest(immutable=immutable):

                def fake_json(url: str) -> dict:
                    if "/actions/runs/" in url:
                        return {
                            "id": 123456,
                            "html_url": project["ci_run_url"],
                            "status": "completed",
                            "conclusion": "success",
                            "head_sha": project["commit"],
                            "head_branch": project["tag"],
                            "event": "push",
                            "path": ".github/workflows/release-evidence.yml",
                            "name": "Release evidence",
                            "repository": {
                                "full_name": "example-owner/routecontract"
                            },
                            "created_at": "2026-07-19T23:00:00Z",
                            "updated_at": "2026-07-19T23:30:00Z",
                        }
                    if "/actions/artifacts/" in url:
                        return {
                            "id": 987654,
                            "url": (
                                "https://api.github.com/repos/example-owner/routecontract/"
                                "actions/artifacts/987654"
                            ),
                            "archive_download_url": (
                                "https://api.github.com/repos/example-owner/routecontract/"
                                "actions/artifacts/987654/zip"
                            ),
                            "name": (
                                "routecontract-release-evidence-"
                                f"{project['commit']}"
                            ),
                            "expired": False,
                            "size_in_bytes": 456,
                            "digest": f"sha256:{digest('a')}",
                            "created_at": "2026-07-19T23:10:00Z",
                            "updated_at": "2026-07-19T23:20:00Z",
                            "expires_at": "2026-08-19T23:20:00Z",
                            "workflow_run": {
                                "id": 123456,
                                "head_sha": project["commit"],
                                "head_branch": project["tag"],
                            },
                        }
                    if "/releases/tags/" in url:
                        return {
                            "id": 7,
                            "created_at": "2026-07-20T00:00:00Z",
                            "published_at": "2026-07-20T00:01:00Z",
                            "updated_at": "2026-07-20T00:02:00Z",
                            "draft": False,
                            "prerelease": False,
                            "immutable": immutable,
                            "tag_name": project["tag"],
                            "html_url": project["release_url"],
                            "assets": [],
                        }
                    if "/commits/" in url:
                        return {
                            "sha": project["commit"],
                            "html_url": f"{project['repository_url']}/commit/{project['commit']}",
                        }
                    return {
                        "id": 1,
                        "node_id": "R_example",
                        "private": False,
                        "archived": False,
                        "full_name": "example-owner/routecontract",
                        "html_url": project["repository_url"],
                    }

                with patch.object(
                    package_submission, "request_json", side_effect=fake_json
                ):
                    with self.assertRaisesRegex(
                        package_submission.GateError, "not immutable"
                    ):
                        package_submission.validate_public_evidence(
                            manifest,
                            {"duration_seconds": 173.0},
                            evidence,
                            Path("/evidence"),
                            Path("/repository"),
                        )

    def test_rejects_release_run_for_different_commit(self) -> None:
        manifest = package_submission.validate_manifest(valid_manifest())
        project = manifest["project"]
        evidence = {
            "source_archive_filename": "routecontract-0.1.0-source.zip",
            "source_archive_sha256": digest("5"),
            "source_archive_size": 123,
            "workflow_artifact_size": 456,
            "public_release_assets": {
                "routecontract-0.1.0-source.zip": {"sha256": digest("5"), "size": 123}
            },
        }

        def fake_json(url: str) -> dict:
            if "/actions/runs/" in url:
                return {
                    "id": 123456,
                    "html_url": project["ci_run_url"],
                    "status": "completed",
                    "conclusion": "success",
                    "head_sha": "0" * 40,
                    "head_branch": project["tag"],
                    "event": "push",
                    "path": ".github/workflows/release-evidence.yml",
                    "name": "Release evidence",
                    "repository": {"full_name": "example-owner/routecontract"},
                    "created_at": "2026-07-19T23:00:00Z",
                    "updated_at": "2026-07-19T23:30:00Z",
                }
            if "/commits/" in url:
                return {
                    "sha": manifest["project"]["commit"],
                    "html_url": f"{project['repository_url']}/commit/{project['commit']}",
                }
            return {
                "id": 1,
                "node_id": "R_example",
                "private": False,
                "archived": False,
                "full_name": "example-owner/routecontract",
                "html_url": project["repository_url"],
            }

        with patch.object(package_submission, "request_json", side_effect=fake_json):
            with self.assertRaisesRegex(package_submission.GateError, "not green"):
                package_submission.validate_public_evidence(
                    manifest,
                    {"duration_seconds": 173.0},
                    evidence,
                    Path("/evidence"),
                    Path("/repository"),
                )


if __name__ == "__main__":
    unittest.main()
